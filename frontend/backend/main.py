"""
FastAPI Backend for QCA Renewable Energy Schedule Management Dashboard
"""
from fastapi import FastAPI, HTTPException, UploadFile, File, Query, Depends, Form, Header, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse, StreamingResponse, PlainTextResponse, HTMLResponse
from sqlalchemy.orm import Session
from sqlalchemy import inspect, text
from typing import Optional, List, Dict, Any, cast, Tuple
from pydantic import BaseModel, Field
import asyncio
import base64
import csv
import html
import io
import json
import math
import mimetypes
import random
import time
from decimal import Decimal
from datetime import datetime, date, timedelta, timezone
from zoneinfo import ZoneInfo
import os
import re
import hashlib
from urllib.parse import urlparse, quote
from urllib.request import urlopen
from urllib.error import HTTPError
from threading import Lock
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FutureTimeoutError
from xml.etree import ElementTree
from uuid import uuid4

# Manual changes ingest (local fallback for UI "Submit Changes").
from typing import Literal

from database import SessionLocal, engine, Base
from models import (
    Plant, Schedule, Forecast, Weather, Deviation, Report, Template, WhatsAppData, MeterData,
    ScheduleReadiness, ScheduleTrigger, ScheduleNotification, EmailSchedulerJob, EmailSendLog, EmailSchedulerSetting,
    EmailSchedulerSupportPreview, SiteMessageLog, WbesNotificationLog, DocumentationDocument, DocumentationHeading
)
from schemas import (
    PlantCreate, PlantUpdate, ScheduleCreate, ScheduleUpdate,
    ForecastCreate, WeatherCreate, DeviationCreate, ReportCreate, TemplateCreate,
    WhatsAppDataCreate, WhatsAppDataUpdate, MeterDataCreate, MeterDataUpdate,
    ScheduleReadinessResponse, ScheduleReadinessSummary, ScheduleTriggerResponse,
    ScheduleNotificationResponse, NotificationListResponse, TriggerCheckResult,
    ManualTriggerRequest, ContinueScheduleRequest, MarkReadyRequest,
    TemplateTransformRequest, TemplateTransformPreviewResponse, TemplateTransformGenerateResponse,
    ScheduleReadinessUploadTemplateRequest, ScheduleReadinessUploadTemplateResponse,
    ScheduleOverwriteRequest, ScheduleOverwriteResponse,
    ScheduleChangeLogRequest, ScheduleChangeLogResponse, ScheduleChangeLogEntry
)
from crud import (
    get_plants, get_plant, create_plant, update_plant, delete_plant,
    get_schedules, get_schedule, create_schedule, update_schedule, delete_schedule,
    get_forecasts, get_forecast, create_forecast,
    get_weather_data, create_weather,
    get_deviations, create_deviation,
    get_reports, get_report, create_report, update_report, delete_report,
    get_templates, get_template, create_template, delete_template,
    get_dashboard_stats as fetch_dashboard_stats,
    get_whatsapp_data, get_whatsapp_data_by_id, create_whatsapp_data, update_whatsapp_data, delete_whatsapp_data,
    get_meter_data, get_meter_data_by_id, get_latest_meter_data, create_meter_data, update_meter_data, delete_meter_data,
    get_schedule_readiness, get_schedule_readiness_by_plant, get_schedule_readiness_summary,
    get_schedule_triggers, create_schedule_trigger,
    get_schedule_notifications, get_schedule_notification_by_id, mark_notification_read, create_schedule_notification,
    update_schedule_readiness, create_schedule_readiness
)
from services.template_transform_service import (
    run_preview_pipeline, transform_rows, to_csv_bytes, publish_output_file,
    save_transform_audit_run, query_transform_history, load_pipeline_configs,
    get_active_template, get_template_mappings, get_plant_config,
    fetch_s3_text, parse_to_canonical_rows, validate_canonical_rows, compute_source_hash,
    list_schedule_files_for_date, get_transform_run_by_id, normalize_canonical_blocks,
    format_missing_blocks_summary
)
from services.template_transform_service import SCHEDULE_FILE_PREFIX
from services.email_scheduler_service import load_email_scheduler_metadata, normalize_day_ahead_body
from services.email_dispatch_service import send_email_smtp, EmailAttachment
from services.sldc_attachment_converter import ILIOS_PV_SITES, convert_ilios_pv_intraday_files_to_xlsx_bytes, maybe_convert_for_auto_email
from routers.all_plant_penalty import router as all_plant_penalty_router
from routers.utility_viewer import router as utility_viewer_router

app = FastAPI(
    title="QCA Renewable Energy Dashboard API",
    description="Backend API for Renewable Energy Schedule Management",
    version="1.0.0"
)
DAY_AHEAD_FILE_REGEX = re.compile(r"_DA0\.csv$", re.IGNORECASE)


class FrozenSchedulePersistRequest(BaseModel):
    plant_code: str
    schedule_date: str
    block: int
    status: str
    source_schedule_key: str
    freeze_time: Optional[str] = None
    reason: Optional[str] = None
    schedule_csv: Optional[str] = None
    edited_schedule_csv: Optional[str] = None
    system_schedule_csv: Optional[str] = None
    write_system_frozen: Optional[bool] = False
    summary: Optional[Dict[str, Any]] = None


class FrozenScheduleExclusionRequest(BaseModel):
    plant_code: str
    schedule_date: str
    source_schedule_key: str
    requested_by: Optional[str] = None


class ManualChangeItem(BaseModel):
    block: int
    mw: float


class ManualChangesIngestRequest(BaseModel):
    org_id: Optional[str] = None
    site_id: str
    schedule_date: str
    # Accept UI shorthands ("DA"/"ID") and normalize server-side.
    schedule_type: str
    # Optional: the base schedule CSV key that the operator edited (generated/.../schedule_from_XX.csv).
    # When provided, the API will write a full 96-block `edited_schedule.csv` by applying changes
    # on top of the base schedule. Without this, we only persist the sparse change list.
    source_file_key: Optional[str] = None
    request_id: Optional[str] = None
    submitted_at_ist: Optional[str] = None
    changes: List[ManualChangeItem]


class SiteMessageRequest(BaseModel):
    site_id: str
    site_id_raw: Optional[str] = None
    record_type: str = "site_event_message"
    source: str = "dashboard"
    event_date: str
    event_type: str
    raw_message: str
    start_time: Optional[str] = None
    end_time: Optional[str] = None
    date_context: Optional[str] = "today"
    mw: Optional[float] = None
    unit: Optional[str] = None
    reduction_type: Optional[str] = None
    status: Optional[str] = None
    delay_mode: Optional[str] = None
    minutes: Optional[int] = None
    description: Optional[str] = None


class WbesNotificationLogRequest(BaseModel):
    id: Optional[str] = None
    createdAt: Optional[str] = None
    utility: str
    date: date
    block: int
    interval: Optional[str] = None
    oa_remc: Optional[float] = None
    as_value: Optional[float] = Field(None, alias="as")
    total: Optional[float] = None
    status: Optional[str] = None
    fileName: Optional[str] = None
    message: str


class MultiGeneratorCapacity(BaseModel):
    ac_mw: Optional[float] = None
    dc_mw: Optional[float] = None


class MultiGeneratorAsset(BaseModel):
    asset_name: str
    capacity_ac_mw: Optional[float] = None
    capacity_dc_mw: Optional[float] = None
    meter_data_available: Optional[bool] = True


class MultiGeneratorBuyer(BaseModel):
    buyer_name: str
    schedule_capacity_mw: Optional[float] = None
    contract_id: Optional[str] = None
    approval_number: Optional[str] = None
    assets: List[MultiGeneratorAsset] = Field(default_factory=list)


class MultiGeneratorPlantRequest(BaseModel):
    plant_id: str
    plant_name: str
    state: str
    location: Optional[str] = None
    latitude: Optional[float] = None
    longitude: Optional[float] = None
    total_capacity: MultiGeneratorCapacity
    currently_scheduling_capacity: MultiGeneratorCapacity
    buyers: List[MultiGeneratorBuyer]
    penalty_config: Optional[Dict[str, Any]] = None
    template_config: Optional[Dict[str, Any]] = None

WEEK_AHEAD_SUPPORTED_PLANTS = {"BHUPALPALLY", "KOTHAGUDEM", "KASIPET", "OSEPL", "CME", "ZETRIC"}
WEEK_AHEAD_TELANGANA_PLANTS = {"BHUPALPALLY", "KOTHAGUDEM", "KASIPET"}
WEEK_AHEAD_TEMPLATE_PREFIX = os.getenv("WEEK_AHEAD_TEMPLATE_PREFIX", "templates/week-ahead").strip().strip("/")
WEEK_AHEAD_LOCAL_DIR = os.path.join(os.path.dirname(__file__), "uploads", "week_ahead_templates")

PLANTS_WITHOUT_S3_METER = {"CME", "KILAJ"}
MADHYA_PRADESH_EFFECTIVE_DELAY_PLANTS = {
    "ANJANGAON",
    "ANDAD",
    "BALAKWADA",
    "BAMKHAL",
    "GSNP",
    "GUGARIYAKHEDI",
    "NANDGAON",
    "SAWDA",
    "SIRMOUR",
}


def _is_madhya_pradesh_effective_delay_plant(plant_code: Any) -> bool:
    return _normalize_plant_code(str(plant_code or "").strip()) in MADHYA_PRADESH_EFFECTIVE_DELAY_PLANTS


def _effective_delay_blocks_for_plant(plant_code: Any) -> int:
    return 6 if _is_madhya_pradesh_effective_delay_plant(plant_code) else 3


def _derive_plant_code(name: str) -> str:
    """Derive a stable plant code from the stored plant name."""
    if not name:
        return ""
    if "(" in name and ")" in name:
        inside = re.search(r"\(([A-Za-z0-9_-]+)\)", name)
        if inside:
            return inside.group(1).upper()
    compact = re.sub(r"[^A-Za-z0-9]", "", name).upper()
    return compact


def _has_meter_data_in_s3(name: str) -> bool:
    code = _derive_plant_code(name)
    if not code:
        return True
    return code not in PLANTS_WITHOUT_S3_METER


def _manual_changes_sanitize(value: Any) -> str:
    text = str(value or "").strip()
    return "".join(ch for ch in text if ch.isalnum() or ch in {"-", "_", "."})


def _manual_changes_normalize(items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    seen: Dict[int, float] = {}
    for item in items or []:
        try:
            block = int(item.get("block") or 0)
            mw = float(item.get("mw") if item.get("mw") is not None else item.get("scheduled_mw"))
        except Exception:
            continue
        if 1 <= block <= 96:
            seen[block] = mw
    return [{"block": b, "mw": seen[b]} for b in sorted(seen.keys())]


def _manual_changes_to_csv(changes: List[Dict[str, Any]]) -> str:
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["block", "mw"])
    for item in changes or []:
        writer.writerow([item.get("block"), item.get("mw")])
    return buf.getvalue()


def _manual_changes_parse_base_schedule(csv_text: str) -> Dict[int, float]:
    """
    Parse a base schedule CSV into a block->MW mapping.

    We accept common formats produced by this app:
    - canonical "block,mw"
    - schedule_from_XX.csv with a leading Block column and some MW numeric column

    This is intentionally permissive (best-effort). Any missing blocks default to 0.
    """
    text = str(csv_text or "")
    if not text.strip():
        return {}
    lines = [l for l in re.split(r"\r?\n", text) if str(l).strip()]
    if not lines:
        return {}

    def split_csv_line(line: str) -> List[str]:
        out: List[str] = []
        cur = ""
        in_quotes = False
        for ch in str(line or ""):
            if ch == '"':
                in_quotes = not in_quotes
                continue
            if ch == "," and not in_quotes:
                out.append(cur.strip())
                cur = ""
                continue
            cur += ch
        out.append(cur.strip())
        return out

    header = split_csv_line(lines[0])
    norm = [re.sub(r"[^a-z0-9]+", "", str(h or "").strip().lower()) for h in header]
    block_idx = next((i for i, h in enumerate(norm) if h in {"block", "blockno", "blocknumber"} or h.startswith("block")), 0)

    # Prefer an explicit mw/schedule column; otherwise use the last numeric-ish column.
    preferred_cols = {"mw", "schedule", "stationschedule", "scheduledmw", "algoschedulemw", "algoschedule"}
    value_idx = next((i for i, h in enumerate(norm) if h in preferred_cols), -1)
    if value_idx < 0:
        value_idx = max(0, len(header) - 1)

    by_block: Dict[int, float] = {}
    for line in lines[1:]:
        cols = split_csv_line(line)
        if not cols:
            continue
        try:
            blk = int(str(cols[block_idx] if block_idx < len(cols) else cols[0]).strip())
        except Exception:
            continue
        if blk < 1 or blk > 96:
            continue
        raw_val = cols[value_idx] if value_idx < len(cols) else ""
        try:
            mw = float(str(raw_val or "").replace(",", "").strip() or 0.0)
        except Exception:
            mw = 0.0
        by_block[blk] = mw
    return by_block


def _manual_changes_apply_full_schedule(
    *,
    base_by_block: Dict[int, float],
    changes: List[Dict[str, Any]],
) -> str:
    """
    Build a full 96-block edited_schedule.csv in canonical format: block,mw
    """
    merged: Dict[int, float] = {int(b): float(v) for b, v in (base_by_block or {}).items() if 1 <= int(b) <= 96}
    for item in changes or []:
        try:
            blk = int(item.get("block") or 0)
            mw = float(item.get("mw") if item.get("mw") is not None else item.get("scheduled_mw"))
        except Exception:
            continue
        if 1 <= blk <= 96:
            merged[blk] = mw

    rows = [{"block": b, "mw": float(merged.get(b, 0.0))} for b in range(1, 97)]
    return _manual_changes_to_csv(rows)


def _manual_changes_pick_latest_generated_schedule_key(
    *,
    plant_code: str,
    schedule_date: str,
    schedule_type: str,
) -> str:
    """
    Best-effort fallback when UI doesn't pass `source_file_key`.

    Returns latest generated schedule_from_XX.csv key for the plant/date:
    - DAY_AHEAD: generated/vedanjay/<PLANT>/outputs/<DATE>/Day-ahead/schedule_from_*.csv
    - INTRADAY: generated/vedanjay/<PLANT>/outputs/<DATE>/schedule_from_*.csv
    """
    plant = _normalize_plant_code(plant_code)
    date_key = str(schedule_date or "").strip()
    if not plant or not re.fullmatch(r"\d{4}-\d{2}-\d{2}", date_key):
        return ""

    bucket = _derive_s3_bucket_name()
    region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"

    # Prefer boto3 if available; else fall back to public list API (handled by _list_s3_keys_safe).
    s3 = None
    try:
        import boto3  # type: ignore
        if bucket:
            s3 = boto3.client("s3", region_name=region)
    except Exception:
        s3 = None

    keys: List[str] = []
    schedule_type_key = "dayahead" if str(schedule_type or "").strip().upper() == "DAY_AHEAD" else "intraday"
    for prefix in _generated_schedule_prefixes_for_plant(plant, date_key, schedule_type_key):
        if not _s3_proxy_is_allowed_path(prefix):
            continue
        try:
            keys.extend(_list_s3_keys_safe(s3, bucket, prefix, max_keys=2000))
        except Exception:
            continue

    candidates = []
    for key in keys or []:
        k = str(key or "").strip()
        if not k.lower().endswith(".csv"):
            continue
        if not re.search(r"schedule_(?:free(?:z|ze)_)?from_\d+\.csv$", k, flags=re.IGNORECASE):
            continue
        candidates.append(k)
    if not candidates:
        return ""

    # Prefer highest revision token; if tie, lexical.
    def rev_key(k: str) -> int:
        r = _extract_schedule_revision_from_key(k)
        return int(r or -1)

    candidates.sort(key=lambda k: (rev_key(k), k), reverse=True)
    return candidates[0] if candidates else ""


def _manual_changes_pick_latest_manual_edited_schedule_key(
    *,
    org_id: str,
    plant_code: str,
    schedule_date: str,
    schedule_type: str,
    bucket: str,
) -> str:
    """
    Find the latest manual-edits edited_schedule.csv for the given plant/date/type.

    Expected keys:
      manual-edits/<org>/<PLANT>/<YYYY-MM-DD>/<DA|INTRADAY>/manual-<epoch>-<suffix>/edited_schedule.csv
    """
    org = _manual_changes_sanitize(org_id or "vedanjay") or "vedanjay"
    plant = _normalize_plant_code(plant_code)
    date_key = str(schedule_date or "").strip()
    if not org or not plant or not re.fullmatch(r"\d{4}-\d{2}-\d{2}", date_key):
        return ""

    schedule_type_norm = str(schedule_type or "").strip().upper().replace("-", "_")
    if schedule_type_norm in {"DA", "DAYAHEAD", "DAY_AHEAD"}:
        type_folder = "DA"
    else:
        type_folder = "INTRADAY"

    region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"

    s3 = None
    try:
        import boto3  # type: ignore
        if bucket:
            s3 = boto3.client("s3", region_name=region)
    except Exception:
        s3 = None

    keys: List[str] = []
    for folder in _special_s3_plant_folder_aliases(plant):
        prefix = f"manual-edits/{org}/{folder}/{date_key}/{type_folder}/"
        if not _s3_proxy_is_allowed_path(prefix):
            continue
        try:
            keys.extend(_list_s3_keys_safe(s3, bucket, prefix, max_keys=5000))
        except Exception:
            continue

    candidates: List[str] = []
    for key in keys or []:
        k = str(key or "").strip()
        if not k:
            continue
        if not k.lower().endswith("/edited_schedule.csv"):
            continue
        candidates.append(k)

    if not candidates:
        return ""

    def request_epoch(k: str) -> int:
        # Extract epoch from ".../<type_folder>/manual-<epoch>-.../edited_schedule.csv"
        m = re.search(r"/manual-(\d+)-[A-Za-z0-9]+/edited_schedule\.csv$", k, flags=re.IGNORECASE)
        if not m:
            return -1
        try:
            return int(m.group(1))
        except Exception:
            return -1

    candidates.sort(key=lambda k: (request_epoch(k), k), reverse=True)
    return candidates[0] if candidates else ""


def _load_seed_plants():
    config_path = os.path.join(
        os.path.dirname(__file__),
        "config",
        "template_pipeline",
        "plants.json"
    )
    try:
        with open(config_path, "r", encoding="utf-8") as handle:
            data = json.load(handle)
            return data if isinstance(data, list) else []
    except Exception as exc:
        print(f"Warning: Could not load plants seed file: {exc}")
        return []

def _ensure_plants_schema():
    try:
        if engine.dialect.name != "postgresql":
            return
        inspector = inspect(engine)
        columns = {col["name"] for col in inspector.get_columns("plants")}
        if "penalty_threshold_percent" not in columns:
            with engine.connect() as conn:
                conn.execute(text("ALTER TABLE plants ADD COLUMN penalty_threshold_percent FLOAT"))
                conn.commit()
            print("Added plants.penalty_threshold_percent column")
    except Exception as exc:
        print(f"Warning: Could not ensure plants schema: {exc}")

@app.on_event("startup")
async def startup_event():
    """Create database tables on startup"""
    try:
        Base.metadata.create_all(bind=engine)
        print("Database tables created/verified successfully")
        _ensure_plants_schema()

        # Ensure default plants required by schedule template conversion are present.
        db = SessionLocal()
        try:
            seed_plants = _load_seed_plants()
            existing = db.query(Plant).all()
            existing_by_key = {
                ((p.name or "").strip().lower(), (p.state or "").strip().lower()): p
                for p in existing
            }
            inserted = 0
            updated = 0
            for item in seed_plants:
                name = (item.get("name") or "").strip()
                state = (item.get("state") or "").strip()
                if not name or not state:
                    continue
                key = (name.lower(), state.lower())
                target = existing_by_key.get(key)
                payload = {
                    "name": name,
                    "type": item.get("type") or "Solar",
                    "capacity": item.get("capacity") if item.get("capacity") is not None else 0.0,
                    "state": state,
                    "status": item.get("status") or "Active",
                    "efficiency": item.get("efficiency") if item.get("efficiency") is not None else 0.0,
                    "penalty_threshold_percent": item.get("penalty_threshold_percent"),
                    "latitude": item.get("latitude"),
                    "longitude": item.get("longitude"),
                    "location_name": item.get("location_name") or item.get("location") or None,
                }
                if not target:
                    db.add(Plant(**payload))
                    inserted += 1
                    continue
                changed = False
                for field, value in payload.items():
                    if value is None:
                        continue
                    if getattr(target, field, None) != value:
                        setattr(target, field, value)
                        changed = True
                if changed:
                    updated += 1
            if inserted or updated:
                db.commit()
                print(f"Seeded plants (inserted={inserted}, updated={updated})")
        finally:
            db.close()
    except Exception as e:
        print(f"Warning: Could not create database tables: {e}")
        print("Tables may already exist or database may not be ready yet")

    # Auto-freeze daemon disabled; freezing happens only on SLDC confirmation.
    # (Preserve manual endpoints for explicit use.)

    # Backend auto-upload daemon (optional; enable with AUTO_UPLOAD_ENABLED=1).
    try:
        from services.auto_upload_worker import start_auto_upload_task
        start_auto_upload_task()
    except Exception as exc:
        print(f"Warning: auto-upload daemon not started: {exc}")

    # Backend Enercast frozen daemon (enabled by default; disable with ENERCAST_FROZEN_ENABLED=0).
    try:
        from services.enercast_frozen_worker import start_enercast_frozen_task
        start_enercast_frozen_task()
    except Exception as exc:
        print(f"Warning: Enercast frozen daemon not started: {exc}")


# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://localhost:80", "http://localhost", "http://frontend:80", "http://127.0.0.1:80", "http://127.0.0.1"],
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS", "PATCH"],
    allow_headers=["*"],
    expose_headers=["*"],
)

# Dependency to get DB session
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


def _create_operator_notification(
    db: Session,
    *,
    plant_id: int,
    plant_name: str,
    notification_type: str,
    title: str,
    message: str,
    priority: str = "NORMAL",
    action_required: bool = True,
) -> None:
    """Best-effort notification insert for operator bell + cross-user polling."""
    try:
        create_schedule_notification(
            db,
            {
                "plant_id": int(plant_id),
                "plant_name": str(plant_name or "").strip() or f"Plant {plant_id}",
                "notification_type": str(notification_type or "Schedule Alert"),
                "title": str(title or "Schedule alert"),
                "message": str(message or ""),
                "priority": str(priority or "NORMAL"),
                "action_required": bool(action_required),
            },
        )
    except Exception:
        # Do not fail core API actions because notification storage failed.
        pass


# ==================== ROOT ENDPOINTS ====================
@app.get("/api")
@app.get("/api/")
async def api_root():
    """API root endpoint - returns API information"""
    return {
        "name": "QCA Renewable Energy Dashboard API",
        "version": "1.0.0",
        "status": "running",
        "endpoints": {
            "dashboard": "/api/dashboard/stats",
            "plants": "/api/plants",
            "schedules": "/api/schedules",
            "forecasts": "/api/forecasts",
            "weather": "/api/weather",
            "deviations": "/api/deviations",
            "reports": "/api/reports",
            "templates": "/api/templates",
            "template_transform": "/api/template-transform/preview",
            "template_transform_source_files": "/api/template-transform/source-files",
            "template_transform_download": "/api/template-transform/download/{run_id}"
        }
    }


# ==================== DASHBOARD ENDPOINTS ====================
@app.get("/api/dashboard/stats")
async def get_dashboard_stats_endpoint(db: Session = Depends(get_db)):
    """Get dashboard statistics"""
    try:
        stats = fetch_dashboard_stats(db)
        return stats
    except Exception as e:
        import traceback
        traceback.print_exc()
        # Return a safe fallback instead of crashing
        return {
            "activePlants": 0,
            "totalCapacity": 0,
            "currentGeneration": 0,
            "efficiency": 0,
            "windCapacity": 0,
            "solarCapacity": 0,
            "schedules": {
                "total": 0,
                "pending": 0,
                "approved": 0,
                "revised": 0
            }
        }


@app.get("/api/dashboard/recent-activity")
async def get_recent_activity(
    limit: int = Query(5, ge=1, le=50),
    db: Session = Depends(get_db)
):
    """Get recent schedule activity"""
    try:
        schedules = get_schedules(db, limit=limit)
        return schedules
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ==================== PLANTS ENDPOINTS ====================
@app.get("/api/plants")
async def list_plants(
    search: Optional[str] = None,
    type: Optional[str] = None,
    state: Optional[str] = None,
    status: Optional[str] = None,
    db: Session = Depends(get_db)
):
    """List all plants with optional filtering"""
    try:
        filters = {}
        if search:
            filters['search'] = search
        if type and type != 'all' and type != 'All Types':
            filters['type'] = type
        if state and state != 'all' and state != 'All States':
            filters['state'] = state
        if status and status != 'all' and status != 'All':
            filters['status'] = status
        
        plants = get_plants(db, **filters)
        for plant in plants:
            try:
                plant.has_meter_data_in_s3 = _has_meter_data_in_s3(getattr(plant, "name", ""))
            except Exception:
                plant.has_meter_data_in_s3 = True
        # Return as list directly (FastAPI will serialize)
        return plants
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/plants/{plant_id}")
async def get_plant_by_id(plant_id: int, db: Session = Depends(get_db)):
    """Get a specific plant by ID"""
    try:
        plant = get_plant(db, plant_id)
        if not plant:
            raise HTTPException(status_code=404, detail="Plant not found")
        try:
            plant.has_meter_data_in_s3 = _has_meter_data_in_s3(getattr(plant, "name", ""))
        except Exception:
            plant.has_meter_data_in_s3 = True
        return plant
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/plants")
async def create_new_plant(plant: PlantCreate, db: Session = Depends(get_db)):
    """Create a new plant"""
    try:
        return create_plant(db, plant)
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.put("/api/plants/{plant_id}")
async def update_plant_by_id(
    plant_id: int,
    plant: PlantUpdate,
    db: Session = Depends(get_db)
):
    """Update an existing plant"""
    try:
        updated_plant = update_plant(db, plant_id, plant)
        if not updated_plant:
            raise HTTPException(status_code=404, detail="Plant not found")
        return updated_plant
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.delete("/api/plants/{plant_id}")
async def delete_plant_by_id(plant_id: int, db: Session = Depends(get_db)):
    """Delete a plant"""
    try:
        success = delete_plant(db, plant_id)
        if not success:
            raise HTTPException(status_code=404, detail="Plant not found")
        return {"message": "Plant deleted successfully"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ==================== SCHEDULES ENDPOINTS ====================
@app.get("/api/schedules")
async def list_schedules(
    type: Optional[str] = None,
    status: Optional[str] = None,
    plant: Optional[str] = None,
    startDate: Optional[str] = None,
    endDate: Optional[str] = None,
    limit: int = Query(10, ge=1, le=100),  # Allow limit from 1 to 100
    db: Session = Depends(get_db)
):
    """List all schedules with optional filtering"""
    try:
        filters = {}
        if type and type != 'all' and type != 'All':
            filters['type'] = type
        if status and status != 'all' and status != 'All':
            filters['status'] = status
        if plant and plant != 'all' and plant != 'All Plants' and plant != 'Select Plant':
            filters['plant'] = plant
        if startDate:
            filters['startDate'] = startDate
        if endDate:
            filters['endDate'] = endDate
        
        # Apply limit to schedules
        schedules = get_schedules(db, limit=limit, **filters)
        return schedules
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error fetching schedules: {str(e)}")


@app.get("/api/schedules/{schedule_id:int}")
async def get_schedule_by_id(schedule_id: int, db: Session = Depends(get_db)):
    """Get a specific schedule by ID"""
    try:
        schedule = get_schedule(db, schedule_id)
        if not schedule:
            raise HTTPException(status_code=404, detail="Schedule not found")
        return schedule
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/schedules")
async def create_new_schedule(
    schedule: ScheduleCreate,
    db: Session = Depends(get_db)
):
    """Create a new schedule"""
    try:
        return create_schedule(db, schedule)
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.put("/api/schedules/{schedule_id:int}")
async def update_schedule_by_id(
    schedule_id: int,
    schedule: ScheduleUpdate,
    db: Session = Depends(get_db)
):
    """Update an existing schedule"""
    try:
        updated_schedule = update_schedule(db, schedule_id, schedule)
        if not updated_schedule:
            raise HTTPException(status_code=404, detail="Schedule not found")
        return updated_schedule
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.delete("/api/schedules/{schedule_id:int}")
async def delete_schedule_by_id(
    schedule_id: int,
    db: Session = Depends(get_db)
):
    """Delete a schedule"""
    try:
        success = delete_schedule(db, schedule_id)
        if not success:
            raise HTTPException(status_code=404, detail="Schedule not found")
        return {"message": "Schedule deleted successfully"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/schedules/bulk-upload")
async def bulk_upload_schedules(
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    """Upload and import schedules from CSV file"""
    try:
        filename = str(getattr(file, "filename", "") or "")
        if not filename.lower().endswith(".csv"):
            raise HTTPException(status_code=400, detail="Only CSV files are supported")
        
        content = await file.read()
        csv_content = content.decode('utf-8')
        csv_reader = csv.DictReader(io.StringIO(csv_content))
        
        imported = 0
        failed = 0
        errors = []
        
        for row in csv_reader:
            try:
                # Parse scheduleDate - handle multiple formats
                schedule_date_str = row.get('scheduleDate', str(date.today()))
                try:
                    # Try ISO format first (YYYY-MM-DD)
                    if isinstance(schedule_date_str, str):
                        schedule_date = datetime.strptime(schedule_date_str, "%Y-%m-%d").date()
                    else:
                        schedule_date = date.today()
                except ValueError:
                    try:
                        # Try DD-MM-YYYY format
                        schedule_date = datetime.strptime(schedule_date_str, "%d-%m-%Y").date()
                    except ValueError:
                        # Default to today if parsing fails
                        schedule_date = date.today()
                
                schedule_data = ScheduleCreate(
                    plantName=row.get('plantName', ''),
                    type=row.get('type', 'Day-Ahead'),
                    scheduleDate=schedule_date,
                    capacity=float(row.get('capacity', 0)),
                    forecasted=float(row.get('forecasted', 0)),
                    actual=float(row.get('actual', 0)),
                    status=row.get('status', 'Pending')
                )
                create_schedule(db, schedule_data)
                imported += 1
            except Exception as e:
                failed += 1
                errors.append(f"Row {imported + failed}: {str(e)}")
        
        return {
            "success": True,
            "imported": imported,
            "failed": failed,
            "errors": errors[:10]  # Limit errors to first 10
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/api/schedules/upload-96-blocks")
async def upload_schedule_96_blocks(
    file: UploadFile = File(...),
    plant_name: str = Query(...),
    schedule_type: str = Query("Day-Ahead"),
    schedule_date: str = Query(..., description="Date in YYYY-MM-DD format"),
    db: Session = Depends(get_db)
):
    """Upload schedule data with 96 time blocks (15-min intervals) from CSV file"""
    try:
        filename = str(getattr(file, "filename", "") or "")
        if not filename.lower().endswith(".csv"):
            raise HTTPException(status_code=400, detail="Only CSV files are supported")
        
        content = await file.read()
        csv_content = content.decode('utf-8')
        csv_reader = csv.DictReader(io.StringIO(csv_content))
        
        rows = list(csv_reader)
        
        if len(rows) == 0:
            raise HTTPException(status_code=400, detail="CSV file is empty")
        
        # Parse date
        try:
            parsed_date = datetime.strptime(schedule_date, "%Y-%m-%d").date()
        except ValueError:
            raise HTTPException(status_code=400, detail="Invalid date format. Use YYYY-MM-DD")
        
        # Parse block data from CSV
        block_data = {}
        total_forecasted = 0
        total_actual = 0
        total_scheduled = 0
        valid_blocks = 0
        
        for idx, row in enumerate(rows):
            try:
                # Get block number (default to row index + 1)
                block_num = int(row.get('block', idx + 1))
                
                # Get time (default to calculated time)
                time_str = row.get('time', '')
                if not time_str:
                    time_str = f"{(idx * 15) // 60:02d}:{(idx * 15) % 60:02d}"
                
                # Parse values
                forecasted = float(row.get('forecasted', row.get('forecast', 0))) or 0
                actual = float(row.get('actual', 0)) or 0
                scheduled = float(row.get('scheduled', forecasted)) or forecasted
                
                block_key = f"block_{block_num}"
                block_data[block_key] = {
                    "block": block_num,
                    "time": time_str,
                    "forecasted": forecasted,
                    "actual": actual,
                    "scheduled": scheduled
                }
                
                total_forecasted += forecasted
                total_actual += actual
                total_scheduled += scheduled
                valid_blocks += 1
                
            except Exception as e:
                print(f"Warning: Could not parse row {idx}: {str(e)}")
                continue
        
        if valid_blocks == 0:
            raise HTTPException(status_code=400, detail="Could not parse any valid blocks from CSV")
        
        # Calculate capacity (average of scheduled values)
        capacity = total_scheduled / valid_blocks if valid_blocks > 0 else 0
        
        # Calculate deviation
        deviation = ((total_actual - total_forecasted) / total_forecasted * 100) if total_forecasted > 0 else 0
        
        # Create schedule with block data
        schedule_create = ScheduleCreate(
            plantName=plant_name,
            type=schedule_type,
            scheduleDate=parsed_date,
            capacity=round(capacity, 2),
            forecasted=round(total_forecasted, 2),
            actual=round(total_actual, 2),
            status="Pending",
            deviation=round(deviation, 2),
            blockData=block_data
        )
        
        created_schedule = create_schedule(db, schedule_create)

        # Shared notification for all operators when a schedule is generated.
        plant_record = db.query(Plant).filter(Plant.name == plant_name).first()
        resolved_plant_id = int(getattr(plant_record, "id", 0) or 0)
        _create_operator_notification(
            db,
            plant_id=resolved_plant_id,
            plant_name=plant_name,
            notification_type="Schedule Generated",
            title="Schedule generated",
            message=f"{plant_name}: {schedule_type} schedule generated for {parsed_date.isoformat()}",
            priority="NORMAL",
            action_required=False,
        )
        
        return {
            "success": True,
            "message": f"Schedule uploaded successfully with {valid_blocks} time blocks",
            "scheduleId": created_schedule.id,
            "plantName": plant_name,
            "scheduleDate": str(parsed_date),
            "type": schedule_type,
            "totalBlocks": valid_blocks,
            "totalForecasted": round(total_forecasted, 2),
            "totalActual": round(total_actual, 2),
            "deviation": round(deviation, 2)
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


# ==================== MANUAL CHANGES (UI SUBMIT) ====================
@app.post("/api/manual-changes")
async def ingest_manual_changes(request: ManualChangesIngestRequest):
    """
    Local/manual endpoint used by the Schedule Preparation UI to persist operator edits.

    In production this can point to an external API (e.g., API Gateway + Lambda),
    but for local/dev we accept the same payload and write to S3 when configured.
    """
    try:
        # NOTE: we store manual changes under `manual-edits/` so the Templates screen
        # can always find `edited_schedule.csv` in a single canonical location.
        site_id = _manual_changes_sanitize(request.site_id).upper()
        if not site_id:
            raise HTTPException(status_code=400, detail="site_id is required")

        schedule_date = str(request.schedule_date or "").strip()
        if not re.fullmatch(r"\d{4}-\d{2}-\d{2}", schedule_date):
            raise HTTPException(status_code=400, detail="schedule_date must be YYYY-MM-DD")

        schedule_type = str(request.schedule_type or "").strip().upper().replace("-", "_")
        if schedule_type in {"DA", "DAYAHEAD", "DAY_AHEAD"}:
            schedule_type = "DAY_AHEAD"
        if schedule_type in {"ID", "INTRADAY"}:
            schedule_type = "INTRADAY"
        if schedule_type not in {"DAY_AHEAD", "INTRADAY"}:
            raise HTTPException(status_code=400, detail="schedule_type must be DAY_AHEAD/INTRADAY (aliases: DA/ID)")

        normalized_changes = _manual_changes_normalize([c.model_dump() for c in request.changes])
        if not normalized_changes:
            raise HTTPException(status_code=400, detail="changes must include at least one valid block")

        # Canonical org folder for this deployment.
        org_id = _manual_changes_sanitize(request.org_id or "vedanjay") or "vedanjay"

        # Canonical manual request id folder format (matches existing UI expectations).
        raw_request_id = _manual_changes_sanitize(request.request_id or "")
        if raw_request_id and raw_request_id.lower().startswith("manual-"):
            request_id = raw_request_id
        else:
            request_id = f"manual-{int(time.time() * 1000)}-{uuid4().hex[:8]}"

        submitted_at_ist = _manual_changes_sanitize(request.submitted_at_ist or datetime.utcnow().replace(tzinfo=timezone.utc).isoformat())

        payload: Dict[str, Any] = {
            "org_id": org_id,
            "site_id": site_id,
            "schedule_date": schedule_date,
            "schedule_type": schedule_type,
            "source_file_key": str(request.source_file_key or "").strip() or None,
            "request_id": request_id,
            "submitted_at_ist": submitted_at_ist,
            "received_at_utc": datetime.utcnow().replace(tzinfo=timezone.utc).isoformat(),
            "changes": normalized_changes,
        }

        # Persist to S3 under the canonical manual-edits folder (NO `manual/changes/`).
        # Example:
        # manual-edits/vedanjay/OSEPL/2026-05-05/DA/manual-<id>/edited_schedule.csv
        type_folder = "DA" if schedule_type == "DAY_AHEAD" else "INTRADAY"
        site_folder = _special_s3_plant_folder(site_id)
        base = f"manual-edits/{org_id}/{site_folder}/{schedule_date}/{type_folder}/{request_id}"
        json_key = f"{base}/changes.json"
        csv_key = f"{base}/edited_schedule.csv"
        changes_csv_key = f"{base}/changes.csv"
        changes_csv_text = _manual_changes_to_csv(normalized_changes)

        bucket = str(os.getenv("MANUAL_CHANGES_BUCKET") or _derive_s3_bucket_name() or "").strip()
        region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"

        # Build the edited schedule CSV:
        # - if we have a base schedule file key, fetch it and apply changes to produce a full 96-block file
        # - otherwise, fall back to a full 96-block file using only the provided changes (missing blocks -> 0)
        base_key = str(request.source_file_key or "").strip()

        # Prefer latest manual-edits edited_schedule.csv as the base so successive manual saves preserve earlier edits.
        latest_manual_key = _manual_changes_pick_latest_manual_edited_schedule_key(
            org_id=org_id,
            plant_code=site_id,
            schedule_date=schedule_date,
            schedule_type=schedule_type,
            bucket=bucket,
        )

        if not base_key:
            base_key = latest_manual_key or ""

        if not base_key:
            # Final fallback: pick latest generated schedule for plant/date/type so edited_schedule.csv is complete.
            base_key = _manual_changes_pick_latest_generated_schedule_key(
                plant_code=site_id,
                schedule_date=schedule_date,
                schedule_type=schedule_type,
            )

        # Guardrail: don't allow INTRADAY edits to accidentally baseline off a day-ahead artifact (and vice-versa).
        if base_key:
            is_day_ahead_key = bool(re.search(r"(?:/day-ahead/|/dayahead/|/day_ahead/|_DA0\\.csv$)", base_key, re.IGNORECASE))
            if schedule_type == "INTRADAY" and is_day_ahead_key:
                base_key = _manual_changes_pick_latest_generated_schedule_key(
                    plant_code=site_id,
                    schedule_date=schedule_date,
                    schedule_type="INTRADAY",
                )
            if schedule_type == "DAY_AHEAD" and not is_day_ahead_key:
                da_candidate = _manual_changes_pick_latest_generated_schedule_key(
                    plant_code=site_id,
                    schedule_date=schedule_date,
                    schedule_type="DAY_AHEAD",
                )
                if da_candidate:
                    base_key = da_candidate

        # If the client passed a generated schedule file as the base, but we already have a manual-edited baseline,
        # prefer the manual baseline so previously edited blocks remain in the new edited_schedule.csv.
        if latest_manual_key and base_key and not str(base_key).startswith("manual-edits/"):
            base_key = latest_manual_key
        base_by_block: Dict[int, float] = {}
        if base_key:
            try:
                base_text = fetch_s3_text(base_key, DEFAULT_TEMPLATE_S3_BASE_URL)
                base_by_block = _manual_changes_parse_base_schedule(base_text)
            except Exception:
                base_by_block = {}
        csv_text = _manual_changes_apply_full_schedule(base_by_block=base_by_block, changes=normalized_changes)

        storage_mode = "local"
        json_location = json_key
        csv_location = csv_key
        error = ""

        if bucket:
            try:
                import boto3  # type: ignore

                s3 = boto3.client("s3", region_name=region)
                s3.put_object(
                    Bucket=bucket,
                    Key=json_key,
                    Body=(json.dumps(payload, ensure_ascii=True, separators=(",", ":")) + "\n").encode("utf-8"),
                    ContentType="application/json",
                )
                s3.put_object(
                    Bucket=bucket,
                    Key=csv_key,
                    Body=csv_text.encode("utf-8"),
                    ContentType="text/csv",
                )
                s3.put_object(
                    Bucket=bucket,
                    Key=changes_csv_key,
                    Body=changes_csv_text.encode("utf-8"),
                    ContentType="text/csv",
                )
                storage_mode = "s3"
            except Exception as exc:
                error = str(exc)

        if storage_mode != "s3":
            # Local fallback for dev environments.
            root = os.path.join(os.path.dirname(__file__), "storage", "manual_changes")
            local_json = os.path.join(root, json_key.replace("/", os.sep))
            local_csv = os.path.join(root, csv_key.replace("/", os.sep))
            local_changes_csv = os.path.join(root, changes_csv_key.replace("/", os.sep))
            os.makedirs(os.path.dirname(local_json), exist_ok=True)
            os.makedirs(os.path.dirname(local_csv), exist_ok=True)
            os.makedirs(os.path.dirname(local_changes_csv), exist_ok=True)
            with open(local_json, "w", encoding="utf-8", newline="") as handle:
                json.dump(payload, handle, ensure_ascii=False, indent=2)
            with open(local_csv, "w", encoding="utf-8", newline="") as handle:
                handle.write(csv_text)
            with open(local_changes_csv, "w", encoding="utf-8", newline="") as handle:
                handle.write(changes_csv_text)
            json_location = local_json
            csv_location = local_csv

        return {
            "ok": True,
            "message": "Manual schedule changes accepted",
            "request_id": request_id,
            "schedule_date": schedule_date,
            "site_id": site_id,
            "bucket": bucket if storage_mode == "s3" else None,
            "json_s3_key": json_key if storage_mode == "s3" else None,
            "csv_s3_key": csv_key if storage_mode == "s3" else None,
            "changes_csv_s3_key": changes_csv_key if storage_mode == "s3" else None,
            "json_location": json_location,
            "csv_location": csv_location,
            "storage_mode": storage_mode,
            "normalized_change_count": len(normalized_changes),
            "error": error or None,
        }
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc))


@app.post("/api/schedules/overwrite-latest", response_model=ScheduleOverwriteResponse)
async def overwrite_latest_schedule(
    request: ScheduleOverwriteRequest,
):
    """Overwrite latest schedule CSV in S3 (Option B)."""
    try:
        source_key = str(request.source_file_key or "").strip()
        if not source_key:
            raise HTTPException(status_code=400, detail="source_file_key is required")
        if not re.search(r"schedule_(?:free(?:z|ze)_)?from_\d+\.csv$", source_key, re.IGNORECASE):
            raise HTTPException(status_code=400, detail="source_file_key must be schedule_from_XX.csv or schedule_freeze_from_XX.csv")

        csv_text = str(request.csv_text or "")
        if not csv_text.strip():
            raise HTTPException(status_code=400, detail="csv_text is required")

        bucket = _derive_s3_bucket_name()
        if not bucket:
            raise HTTPException(status_code=500, detail="S3 bucket not configured")

        region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"
        output_file_key = source_key
        output_file_url = f"https://{bucket}.s3.{region}.amazonaws.com/{output_file_key}"
        uploaded_at = datetime.utcnow()

        try:
            import boto3  # type: ignore
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"boto3 not available: {e}")

        s3 = boto3.client("s3", region_name=region)
        s3.put_object(
            Bucket=bucket,
            Key=output_file_key,
            Body=csv_text.encode("utf-8"),
            ContentType="text/csv",
        )

        # Keep readiness upload history in sync for flows that overwrite schedule files directly.
        source_parts = [p for p in source_key.split("/") if p]
        inferred_plant = ""
        inferred_date = ""
        try:
            if "generated" in source_parts and "outputs" in source_parts:
                gen_idx = source_parts.index("generated")
                out_idx = source_parts.index("outputs")
                if out_idx >= gen_idx + 2:
                    inferred_plant = str(source_parts[gen_idx + 2]).strip().upper()
                if out_idx + 1 < len(source_parts):
                    inferred_date = str(source_parts[out_idx + 1]).strip()
        except Exception:
            inferred_plant = ""
            inferred_date = ""

        if not re.match(r"^\d{4}-\d{2}-\d{2}$", inferred_date):
            inferred_date = uploaded_at.date().isoformat()

        history_entry = {
            "id": f"{int(uploaded_at.timestamp() * 1000)}-{uuid4().hex[:8]}",
            "plant_code": inferred_plant or "UNKNOWN",
            "schedule_date": inferred_date,
            "template_file_name": os.path.basename(source_key),
            "source_file_key": source_key,
            "manual_request_id": None,
            "requested_by": str(request.requested_by or "").strip(),
            "bucket": bucket,
            "output_file_key": output_file_key,
            "output_file_url": output_file_url,
            "uploaded_at": uploaded_at.isoformat() + "Z",
            "storage_mode": "s3",
            "error": None,
            "csv_text": csv_text,
        }
        history_entry.update(_compute_submit_and_effective_blocks_from_iso(history_entry.get("uploaded_at", ""), plant_code=inferred_plant))
        try:
            _append_readiness_upload_history(history_entry)
            _readiness_dashboard_cache_clear_date(inferred_date)
        except Exception:
            # Do not fail overwrite flow if history persistence has an issue.
            pass

        return {
            "success": True,
            "message": "Latest schedule overwritten successfully",
            "bucket": bucket,
            "output_file_key": output_file_key,
            "output_file_url": output_file_url,
            "uploaded_at": uploaded_at,
            "error": None,
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/schedules/change-log", response_model=ScheduleChangeLogResponse)
def append_schedule_change_log(
    request: ScheduleChangeLogRequest,
):
    """Append a manual change log entry for a schedule (shared across users)."""
    try:
        plant_code = str(request.plant_code or "").strip().upper()
        if not plant_code:
            raise HTTPException(status_code=400, detail="plant_code is required")
        if plant_code in {"SHRIMOUR", "SHROMOUR"}:
            plant_code = "SIRMOUR"

        schedule_date = request.schedule_date
        source_key = str(request.source_file_key or "").strip()
        saved_at = request.saved_at or datetime.utcnow()
        requested_by = str(getattr(request, "requested_by", "") or "").strip()

        entry = {
            "block": int(request.block),
            "time": str(request.time or "").strip(),
            "old_value": str(request.old_value),
            "new_value": str(request.new_value),
            "saved_at": saved_at.isoformat(),
            "source_file_key": source_key,
            "requested_by": requested_by,
        }

        bucket = _derive_s3_bucket_name()
        region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"
        key = _schedule_change_log_s3_key(
            plant_code=plant_code,
            schedule_date=schedule_date,
            source_file_key=source_key,
        )
        output_file_url = f"https://{bucket}.s3.{region}.amazonaws.com/{key}" if bucket else ""

        with _CHANGE_LOG_LOCK:
            rows = []
            if bucket:
                try:
                    import boto3  # type: ignore
                    from botocore.config import Config  # type: ignore

                    s3 = boto3.client(
                        "s3",
                        region_name=region,
                        config=Config(
                            connect_timeout=2,
                            read_timeout=3,
                            retries={"max_attempts": 1, "mode": "standard"},
                        ),
                    )
                    try:
                        obj = s3.get_object(Bucket=bucket, Key=key)
                        text = obj["Body"].read().decode("utf-8")
                        rows = json.loads(text) if text else []
                    except Exception:
                        rows = []
                    if not isinstance(rows, list):
                        rows = []
                    rows.append(entry)
                    s3.put_object(
                        Bucket=bucket,
                        Key=key,
                        Body=json.dumps(rows, ensure_ascii=False, indent=2).encode("utf-8"),
                        ContentType="application/json",
                    )
                except Exception:
                    bucket = ""

            if not bucket:
                local_path = _schedule_change_log_local_path(
                    plant_code=plant_code,
                    schedule_date=schedule_date,
                    source_file_key=source_key,
                )
                rows = _load_change_log_local(local_path)
                if not isinstance(rows, list):
                    rows = []
                rows.append(entry)
                _save_change_log_local(local_path, rows)

        return {
            "success": True,
            "message": "Change log updated",
            "bucket": bucket or "LOCAL_FALLBACK",
            "output_file_key": key,
            "output_file_url": output_file_url,
            "uploaded_at": saved_at,
            "error": None,
            "items": rows,
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/schedules/change-log", response_model=ScheduleChangeLogResponse)
def get_schedule_change_log(
    plant_code: str = Query(...),
    schedule_date: date = Query(...),
    source_file_key: str = Query(""),
):
    """Fetch schedule change log entries."""
    try:
        plant_code = str(plant_code or "").strip().upper()
        if plant_code in {"SHRIMOUR", "SHROMOUR"}:
            plant_code = "SIRMOUR"

        bucket = _derive_s3_bucket_name()
        region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"
        key = _schedule_change_log_s3_key(
            plant_code=plant_code,
            schedule_date=schedule_date,
            source_file_key=source_file_key,
        )
        output_file_url = f"https://{bucket}.s3.{region}.amazonaws.com/{key}" if bucket else ""
        rows = []

        if bucket:
            try:
                import boto3  # type: ignore
                from botocore.config import Config  # type: ignore

                s3 = boto3.client(
                    "s3",
                    region_name=region,
                    config=Config(
                        connect_timeout=2,
                        read_timeout=3,
                        retries={"max_attempts": 1, "mode": "standard"},
                    ),
                )
                obj = s3.get_object(Bucket=bucket, Key=key)
                text = obj["Body"].read().decode("utf-8")
                rows = json.loads(text) if text else []
            except Exception:
                rows = []

        if not rows:
            local_path = _schedule_change_log_local_path(
                plant_code=plant_code,
                schedule_date=schedule_date,
                source_file_key=source_file_key,
            )
            rows = _load_change_log_local(local_path)

        if not isinstance(rows, list):
            rows = []

        return {
            "success": True,
            "message": "Change log loaded",
            "bucket": bucket or "LOCAL_FALLBACK",
            "output_file_key": key,
            "output_file_url": output_file_url,
            "uploaded_at": datetime.utcnow(),
            "error": None,
            "items": rows,
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/schedules/{schedule_id:int}/blocks")
async def get_schedule_blocks(
    schedule_id: int,
    db: Session = Depends(get_db)
):
    """Get schedule with 96-block data"""
    try:
        from crud import get_schedule_with_blocks
        schedule = get_schedule_with_blocks(db, schedule_id)
        if not schedule:
            raise HTTPException(status_code=404, detail="Schedule not found")
        return schedule
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ==================== FORECASTS ENDPOINTS ====================
@app.get("/api/forecasts")
async def list_forecasts(
    plantId: Optional[int] = None,
    db: Session = Depends(get_db)
):
    """List all forecasts"""
    try:
        filters = {}
        if plantId:
            filters['plantId'] = plantId
        forecasts = get_forecasts(db, **filters)
        return forecasts
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/forecasts/{plant_id}")
async def get_forecast_by_plant(plant_id: int, db: Session = Depends(get_db)):
    """Get forecast for a specific plant"""
    try:
        forecast = get_forecast(db, plant_id)
        if not forecast:
            raise HTTPException(status_code=404, detail="Forecast not found")
        return forecast
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/forecasts")
async def create_forecast_data(forecast: ForecastCreate, db: Session = Depends(get_db)):
    """Create a new forecast"""
    try:
        return create_forecast(db, forecast)
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.get("/api/forecasts/{plant_id}/data")
async def get_forecast_data_for_plant(
    plant_id: int,
    date: str = Query(..., description="Date in YYYY-MM-DD format"),
    db: Session = Depends(get_db)
):
    """Get forecast data for a specific plant and date (96 time blocks)"""
    try:
        # Try to get real forecast data first
        forecast = get_forecast(db, plant_id)
        if forecast:
            # Parse the hourlyData and return in expected format
            hourly_data = forecast.hourlyData
            if isinstance(hourly_data, str):
                hourly_data = json.loads(hourly_data)

            # Convert to dataPoints format
            data_points = []
            for hour in range(24):
                hour_data = hourly_data.get(str(hour), {}) if isinstance(hourly_data, dict) else {}
                for quarter in range(4):
                    minute = quarter * 15
                    time_str = f"{hour:02d}:{minute:02d}"

                    data_points.append({
                        "time": time_str,
                        "hour": hour,
                        "minute": minute,
                        "forecast": hour_data.get("forecast", 0),
                        "actual": hour_data.get("actual", 0),
                        "scheduled": hour_data.get("scheduled", 0)
                    })

            return {
                "date": forecast.forecastDate.isoformat() if forecast.forecastDate else date,
                "dataPoints": data_points,
                "totalForecast": sum(d["forecast"] for d in data_points),
                "totalActual": sum(d["actual"] for d in data_points),
                "createdAt": forecast.createdAt.isoformat() if forecast.createdAt else datetime.now().isoformat()
            }

        raise HTTPException(status_code=404, detail="Forecast data not found")

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ==================== WEATHER ENDPOINTS ====================
@app.get("/api/weather")
async def list_weather_data(db: Session = Depends(get_db)):
    """List all weather data"""
    try:
        weather = get_weather_data(db)
        return weather
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/weather/{location}")
async def get_weather_by_location(location: str, db: Session = Depends(get_db)):
    """Get weather data for a specific location"""
    try:
        weather = get_weather_data(db, location=location)
        if not weather:
            raise HTTPException(status_code=404, detail="Weather data not found")
        return weather
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ==================== DEVIATIONS ENDPOINTS ====================
@app.get("/api/deviations")
async def list_deviations(
    period: str = Query("hourly", regex="^(hourly|daily|weekly)$"),
    limit: int = Query(24, ge=1, le=1000),
    db: Session = Depends(get_db)
):
    """List deviations with period filtering"""
    try:
        deviations = get_deviations(db, period=period, limit=limit)
        return deviations
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ==================== REPORTS ENDPOINTS ====================
@app.get("/api/reports")
async def list_reports(
    type: Optional[str] = None,
    state: Optional[str] = None,
    skip: int = Query(0, ge=0),
    limit: int = Query(100, ge=1, le=1000),
    db: Session = Depends(get_db)
):
    """List all reports with optional filtering"""
    try:
        reports = get_reports(db, skip=skip, limit=limit, type=type, state=state)
        return reports
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/reports/generate")
async def generate_report(report: ReportCreate, db: Session = Depends(get_db)):
    """Track a new report in the database (PDF is generated client-side)"""
    try:
        # Validate required fields
        if not report.name or not report.name.strip():
            raise HTTPException(status_code=400, detail="Report name is required")
        if not report.type or not report.type.strip():
            raise HTTPException(status_code=400, detail="Report type is required")
        if not report.format or not report.format.strip():
            raise HTTPException(status_code=400, detail="Report format is required")
        
        # Create the report record (no PDF generation on server)
        created_report = create_report(db, report)
        
        # Report is tracked, client handles PDF generation
        # Update status to Ready since no file generation is needed
        created_report_id = int(getattr(created_report, "id", 0) or 0)
        update_report(db, created_report_id, status="Ready")
        
        # Refresh to get updated values
        db.refresh(created_report)
        
        return created_report
        
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=400, detail=f"Error tracking report: {str(e)}")


@app.get("/api/reports/{report_id}/download")
async def download_report(report_id: int, db: Session = Depends(get_db)):
    """Download a report PDF
    
    Note: Since PDF files are generated client-side, this endpoint
    returns an error message indicating the report file is not available
    on the server. The client should generate the PDF locally.
    """
    try:
        report = get_report(db, report_id)
        if not report:
            raise HTTPException(status_code=404, detail="Report not found")

        report_any = cast(Any, report)
        
        # Check if PDF file exists on server
        file_path = str(getattr(report_any, "filePath", "") or "")
        status = str(getattr(report_any, "status", "") or "")
        report_name = str(getattr(report_any, "name", "") or "report")
        if file_path and os.path.exists(file_path):
            # Update status to Ready if it was Generating
            if status == "Generating":
                update_report(db, report_id, status="Ready")
            
            # Return the actual PDF file
            return FileResponse(
                path=file_path,
                filename=f"{report_name.replace(' ', '_')}.pdf",
                media_type="application/pdf"
            )
        
        # No file exists on server - client-side PDF generation is expected
        raise HTTPException(
            status_code=410, 
            detail="Report file not available on server. Please generate the PDF using the report interface."
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/reports/{report_id}")
async def delete_report_by_id(report_id: int, db: Session = Depends(get_db)):
    """Delete a report"""
    try:
        success = delete_report(db, report_id)
        if not success:
            raise HTTPException(status_code=404, detail="Report not found")
        # Return proper success response format
        return {"success": True, "message": "Report deleted successfully", "id": report_id}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/reports/cleanup/generating")
async def cleanup_generating_reports(db: Session = Depends(get_db)):
    """Remove all reports with 'Generating' status from database"""
    try:
        from sqlalchemy import text
        # Delete reports with "Generating" status
        result = db.execute(
            text("DELETE FROM reports WHERE status = 'Generating'")
        )
        db.commit()
        deleted_count = result.rowcount
        return {
            "success": True, 
            "message": f"Cleaned up {deleted_count} report(s) with 'Generating' status"
        }
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))


# ==================== TEMPLATES ENDPOINTS ====================
@app.get("/api/templates")
async def list_templates(
    vendor: Optional[str] = None,
    type: Optional[str] = None,
    db: Session = Depends(get_db)
):
    """List all templates with optional filtering"""
    try:
        filters = {}
        if vendor and vendor != 'all':
            filters['vendor'] = vendor
        if type and type != 'all':
            filters['type'] = type
        templates = get_templates(db, **filters)
        return templates
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/templates")
async def create_new_template(
    template: TemplateCreate,
    db: Session = Depends(get_db)
):
    """Create a new template"""
    try:
        return create_template(db, template)
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.delete("/api/templates/{template_id}")
async def delete_template_by_id(
    template_id: int,
    db: Session = Depends(get_db)
):
    """Delete a template"""
    try:
        success = delete_template(db, template_id)
        if not success:
            raise HTTPException(status_code=404, detail="Template not found")
        return {"message": "Template deleted successfully"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ==================== WEEK-AHEAD TEMPLATE ENDPOINTS ====================
def _week_ahead_normalize_plant_code(value: Any) -> str:
    code = str(value or "").strip().upper()
    if code == "OSEL":
        code = "OSEPL"
    if code == "KOTHAGUDAM":
        code = "KOTHAGUDEM"
    return code


def _week_ahead_require_supported_plant(value: Any) -> str:
    code = _week_ahead_normalize_plant_code(value)
    if code not in WEEK_AHEAD_SUPPORTED_PLANTS:
        raise HTTPException(status_code=400, detail=f"Unsupported week-ahead plant: {code or value}")
    return code


def _week_ahead_safe_filename(value: str, fallback: str = "week_ahead_template.xlsx") -> str:
    raw = os.path.basename(str(value or "").strip().replace("\\", "/")) or fallback
    safe = re.sub(r"[^A-Za-z0-9._ -]+", "_", raw).strip(" .")
    return safe or fallback


def _week_ahead_template_key(plant_code: str, filename: str) -> str:
    return f"{WEEK_AHEAD_TEMPLATE_PREFIX}/{plant_code}/active/{_week_ahead_safe_filename(filename)}"


def _week_ahead_metadata_key(plant_code: str) -> str:
    return f"{WEEK_AHEAD_TEMPLATE_PREFIX}/{plant_code}/active/metadata.json"


def _week_ahead_local_template_path(plant_code: str, filename: str) -> str:
    return os.path.join(WEEK_AHEAD_LOCAL_DIR, plant_code, _week_ahead_safe_filename(filename))


def _week_ahead_local_metadata_path(plant_code: str) -> str:
    return os.path.join(WEEK_AHEAD_LOCAL_DIR, plant_code, "metadata.json")


def _week_ahead_s3_client() -> Tuple[Optional[Any], str, str]:
    bucket = _derive_s3_bucket_name()
    region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"
    if not bucket:
        return None, "", region
    try:
        import boto3  # type: ignore
        return boto3.client("s3", region_name=region), bucket, region
    except Exception:
        return None, "", region


def _week_ahead_read_metadata(plant_code: str) -> Dict[str, Any]:
    s3, bucket, _region = _week_ahead_s3_client()
    metadata_key = _week_ahead_metadata_key(plant_code)
    if s3 is not None and bucket:
        try:
            obj = s3.get_object(Bucket=bucket, Key=metadata_key)
            body = obj.get("Body")
            data = body.read() if body is not None else b"{}"
            parsed = json.loads(data.decode("utf-8", errors="replace"))
            if isinstance(parsed, dict):
                return parsed
        except Exception:
            pass

    local_path = _week_ahead_local_metadata_path(plant_code)
    try:
        if os.path.exists(local_path):
            with open(local_path, "r", encoding="utf-8") as f:
                parsed = json.load(f)
            if isinstance(parsed, dict):
                return parsed
    except Exception:
        pass
    return {}


def _week_ahead_fetch_template_bytes(metadata: Dict[str, Any]) -> bytes:
    storage = str(metadata.get("storage_mode") or "").strip().lower()
    if storage == "local":
        path = str(metadata.get("local_path") or "").strip()
        if path and os.path.exists(path):
            with open(path, "rb") as f:
                return f.read()
        raise HTTPException(status_code=404, detail="Week-ahead template not found in local storage")

    key = str(metadata.get("template_key") or "").strip()
    if not key:
        raise HTTPException(status_code=404, detail="Week-ahead template is not uploaded for this plant")
    s3, bucket, _region = _week_ahead_s3_client()
    if s3 is None or not bucket:
        raise HTTPException(status_code=500, detail="S3 bucket is not configured")
    try:
        obj = s3.get_object(Bucket=bucket, Key=key)
        body = obj.get("Body")
        return body.read() if body is not None else b""
    except Exception as exc:
        raise HTTPException(status_code=404, detail=f"Week-ahead template not found: {exc}") from exc


def _week_ahead_source_prefixes(plant_code: str, target_date: date) -> List[str]:
    if _normalize_plant_code(plant_code) == "ZETRIC":
        return [
            f"generated/vedanjay/multiple_generator/ZTRIC/{target_date.isoformat()}/Week-ahead/",
        ]
    return [f"raw/vedanjay/{plant_code}/{target_date.isoformat()}/enercast_data/week_ahead/"]


def _week_ahead_fetch_latest_source(plant_code: str, target_date: date) -> Tuple[str, bytes]:
    prefixes = _week_ahead_source_prefixes(plant_code, target_date)
    s3, bucket, _region = _week_ahead_s3_client()
    if s3 is None or not bucket:
        raise HTTPException(status_code=500, detail="S3 bucket is not configured")
    candidates = []
    list_errors = []
    for prefix in prefixes:
        try:
            response = s3.list_objects_v2(Bucket=bucket, Prefix=prefix, MaxKeys=1000)
        except Exception as exc:
            list_errors.append(str(exc))
            continue
        candidates.extend([
            item for item in (response.get("Contents", []) or [])
            if str(item.get("Key") or "").lower().endswith((".csv", ".xlsx"))
        ])
    if _normalize_plant_code(plant_code) == "ZETRIC":
        candidates = [
            item for item in candidates
            if re.search(r"/Week-ahead/schedule_weekahead.*\.csv$", str(item.get("Key") or ""), re.IGNORECASE)
        ]
    if not candidates:
        if list_errors:
            raise HTTPException(status_code=502, detail=f"Failed to list week-ahead source files: {list_errors[0]}")
        raise HTTPException(status_code=404, detail=f"No week-ahead file found under {', '.join(prefixes)}")
    candidates.sort(key=lambda item: item.get("LastModified") or datetime.min.replace(tzinfo=timezone.utc), reverse=True)
    key = str(candidates[0].get("Key") or "")
    try:
        obj = s3.get_object(Bucket=bucket, Key=key)
        body = obj.get("Body")
        return key, body.read() if body is not None else b""
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"Failed to fetch week-ahead source file: {exc}") from exc


def _week_ahead_parse_number(value: Any) -> Optional[float]:
    text = str(value if value is not None else "").strip().replace(",", "")
    if not text:
        return None
    match = re.search(r"-?\d+(?:\.\d+)?", text)
    if not match:
        return None
    try:
        return float(match.group(0))
    except Exception:
        return None


def _week_ahead_header_token(value: Any) -> str:
    return re.sub(r"[^a-z0-9]+", "", str(value or "").strip().lower())


def _week_ahead_pick_value_column(headers: List[Any], block_idx: int) -> int:
    tokens = [_week_ahead_header_token(h) for h in headers]
    preferred = [
        "weekaheadforecast", "weekahead", "forecastmw", "forecast", "schedulemw",
        "schedule", "generationmw", "powermw", "mw", "value",
    ]
    for needle in preferred:
        for idx, token in enumerate(tokens):
            if idx == block_idx:
                continue
            if needle in token and not any(skip in token for skip in ("date", "time", "block")):
                return idx
    for idx, token in enumerate(tokens):
        if idx != block_idx and not any(skip in token for skip in ("date", "time", "block")):
            return idx
    return max(0, block_idx + 1)


def _week_ahead_pick_named_column(headers: List[Any], names: List[str]) -> int:
    tokens = [_week_ahead_header_token(h) for h in headers]
    needles = [_week_ahead_header_token(name) for name in names]
    for needle in needles:
        for idx, token in enumerate(tokens):
            if token == needle:
                return idx
    for needle in needles:
        for idx, token in enumerate(tokens):
            if needle and needle in token:
                return idx
    return -1


def _week_ahead_parse_date_key(value: Any) -> str:
    if isinstance(value, datetime):
        return value.date().isoformat()
    if isinstance(value, date):
        return value.isoformat()
    text = str(value or "").strip()
    if not text:
        return ""
    match = re.search(r"\b(\d{4})-(\d{1,2})-(\d{1,2})\b", text)
    if match:
        try:
            return date(int(match.group(1)), int(match.group(2)), int(match.group(3))).isoformat()
        except Exception:
            return ""
    match = re.search(r"\b(\d{1,2})[-/](\d{1,2})[-/](\d{4})\b", text)
    if match:
        try:
            return date(int(match.group(3)), int(match.group(2)), int(match.group(1))).isoformat()
        except Exception:
            return ""
    return ""


def _week_ahead_parse_time_block(value: Any) -> Optional[int]:
    if isinstance(value, datetime):
        return min(96, max(1, (value.hour * 4) + (value.minute // 15) + 1))
    text = str(value or "").strip()
    match = re.search(r"\b([01]?\d|2[0-3]):([0-5]\d)\b", text)
    if not match:
        return None
    hour = int(match.group(1))
    minute = int(match.group(2))
    return min(96, max(1, (hour * 4) + (minute // 15) + 1))


def _week_ahead_parse_source_block(block_value: Any, from_value: Any = None) -> Optional[int]:
    time_block = _week_ahead_parse_time_block(from_value)
    if time_block is not None:
        return time_block
    raw_block = _week_ahead_parse_number(block_value)
    if raw_block is None:
        return None
    block = int(raw_block)
    if block < 1:
        return None
    return ((block - 1) % 96) + 1


def _week_ahead_parse_template_block_cell(value: Any) -> Optional[int]:
    if isinstance(value, bool):
        return None
    if isinstance(value, int):
        return value if 1 <= value <= 96 else None
    if isinstance(value, float) and math.isfinite(value) and value.is_integer():
        block = int(value)
        return block if 1 <= block <= 96 else None
    text = str(value if value is not None else "").strip()
    if not re.fullmatch(r"\d{1,2}", text):
        return None
    block = int(text)
    return block if 1 <= block <= 96 else None


def _week_ahead_parse_positive_int_cell(value: Any) -> Optional[int]:
    if isinstance(value, bool):
        return None
    if isinstance(value, int):
        return value if value >= 1 else None
    if isinstance(value, float) and math.isfinite(value) and value.is_integer():
        block = int(value)
        return block if block >= 1 else None
    text = str(value if value is not None else "").strip()
    if not re.fullmatch(r"\d+", text):
        return None
    block = int(text)
    return block if block >= 1 else None


def _week_ahead_template_date_columns_from_rows(rows: List[List[Any]], data_row_idx: int, block_col: int) -> Dict[str, int]:
    date_cols: Dict[str, int] = {}
    for header_row_idx in range(max(0, data_row_idx - 10), data_row_idx):
        row = rows[header_row_idx] if header_row_idx < len(rows) else []
        row_date_count = sum(1 for cell in row if _week_ahead_parse_date_key(cell))
        if row_date_count < 2:
            continue
        for col_idx in range(block_col + 1, len(row)):
            date_key = _week_ahead_parse_date_key(row[col_idx])
            header_tokens = []
            for check_row_idx in range(header_row_idx + 1, data_row_idx):
                check_row = rows[check_row_idx] if check_row_idx < len(rows) else []
                for check_col_idx in range(col_idx, min(len(check_row), col_idx + 3)):
                    header_tokens.append(_week_ahead_header_token(check_row[check_col_idx]))
            has_target_header = any(
                "avc" in token or "forecast" in token or "schedule" in token
                for token in header_tokens
            )
            if date_key and date_key not in date_cols:
                if has_target_header:
                    date_cols[date_key] = col_idx
    return date_cols


def _week_ahead_template_date_columns_from_sheet(sheet: Any, data_row_idx: int, block_col: int) -> Dict[str, int]:
    date_cols: Dict[str, int] = {}
    for header_row in range(max(1, data_row_idx - 10), data_row_idx):
        row_date_count = sum(
            1
            for col in range(1, (sheet.max_column or 1) + 1)
            if _week_ahead_parse_date_key(sheet.cell(header_row, col).value)
        )
        if row_date_count < 2:
            continue
        for col in range(block_col + 1, (sheet.max_column or 1) + 1):
            date_key = _week_ahead_parse_date_key(sheet.cell(header_row, col).value)
            header_tokens = []
            for check_row in range(header_row + 1, data_row_idx):
                for check_col in range(col, min(sheet.max_column or 1, col + 2) + 1):
                    header_tokens.append(_week_ahead_header_token(sheet.cell(check_row, check_col).value))
            has_target_header = any(
                "avc" in token or "forecast" in token or "schedule" in token
                for token in header_tokens
            )
            if date_key and date_key not in date_cols:
                if has_target_header:
                    date_cols[date_key] = col
    return date_cols


def _week_ahead_date_offset(date_key: str, date_cols: Dict[str, int], fallback_dates: List[str]) -> Optional[int]:
    if date_cols:
        ordered = [item[0] for item in sorted(date_cols.items(), key=lambda item: item[1])]
        return ordered.index(date_key) if date_key in ordered else None
    return fallback_dates.index(date_key) if date_key in fallback_dates else None


def _week_ahead_parse_csv_rows(content: bytes) -> List[List[str]]:
    text = content.decode("utf-8-sig", errors="replace")
    try:
        dialect = csv.Sniffer().sniff(text[:4096], delimiters=",;\t")
    except Exception:
        dialect = csv.excel
    return [[str(cell).strip() for cell in row] for row in csv.reader(io.StringIO(text), dialect)]


def _week_ahead_extract_values_from_rows(
    rows: List[List[Any]],
    use_telangana_mapping: bool = False,
    use_osepl_mapping: bool = False,
    mh_vedanjay_value_column_names: Optional[List[str]] = None,
) -> List[Any]:
    if not rows:
        return []
    header_idx = 0
    block_idx = -1
    for idx, row in enumerate(rows[:30]):
        tokens = [_week_ahead_header_token(cell) for cell in row]
        found = next((i for i, token in enumerate(tokens) if token in {"block", "blockno", "blk", "blkno", "blocknumber", "sno", "srno", "serialno"} or token.startswith("block") or token.startswith("blk")), -1)
        if found >= 0:
            header_idx = idx
            block_idx = found
            break
    if block_idx < 0:
        values: List[Any] = []
        for row in rows:
            for cell in row:
                num = _week_ahead_parse_number(cell)
                if num is not None:
                    values.append(num)
        if values:
            return values

    headers = rows[header_idx]
    from_idx = _week_ahead_pick_named_column(headers, ["From", "Start", "Start Time", "From Time", "Timestamp", "Date Time"])
    mh_value_names = mh_vedanjay_value_column_names or ["OSEPL"]
    osepl_idx = _week_ahead_pick_named_column(headers, mh_value_names) if use_osepl_mapping else -1
    availability_idx = _week_ahead_pick_named_column(headers, ["Availability Capacity"]) if use_osepl_mapping else -1
    if use_osepl_mapping and osepl_idx >= 0 and availability_idx >= 0:
        values = []
        for row in rows[header_idx + 1:]:
            if block_idx >= len(row):
                continue
            from_raw = row[from_idx] if from_idx >= 0 and from_idx < len(row) else ""
            block = _week_ahead_parse_source_block(row[block_idx], from_raw)
            date_key = _week_ahead_parse_date_key(from_raw)
            if block is None or not date_key:
                continue
            osepl_raw = row[osepl_idx] if osepl_idx < len(row) else ""
            availability_raw = row[availability_idx] if availability_idx < len(row) else ""
            osepl_num = _week_ahead_parse_number(osepl_raw)
            availability_num = _week_ahead_parse_number(availability_raw)
            osepl_value = osepl_num if osepl_num is not None else 0
            values.append({
                "date": date_key,
                "block": block,
                "declared_forecast": osepl_value,
                "inter_avc": availability_num if availability_num is not None else 0,
                "schedule": osepl_value,
            })
        if values:
            return values

    avc_idx = _week_ahead_pick_named_column(headers, ["AvC_MW", "AvC(MW)", "AvC MW", "AvC"]) if use_telangana_mapping else -1
    schedule_idx = _week_ahead_pick_named_column(headers, ["SCH_MW", "Schedule", "Schedule MW", "Schedule(MW)"]) if use_telangana_mapping else -1
    if use_telangana_mapping and avc_idx >= 0 and schedule_idx >= 0:
        values = []
        for row in rows[header_idx + 1:]:
            if block_idx >= len(row):
                continue
            from_raw = row[from_idx] if from_idx >= 0 and from_idx < len(row) else ""
            block = _week_ahead_parse_source_block(row[block_idx], from_raw)
            date_key = _week_ahead_parse_date_key(from_raw)
            if block is None or not date_key:
                continue
            avc_raw = row[avc_idx] if avc_idx < len(row) else ""
            schedule_raw = row[schedule_idx] if schedule_idx < len(row) else ""
            avc_num = _week_ahead_parse_number(avc_raw)
            schedule_num = _week_ahead_parse_number(schedule_raw)
            values.append({
                "date": date_key,
                "block": block,
                "avc": avc_num if avc_num is not None else 0,
                "schedule": schedule_num if schedule_num is not None else 0,
            })
        return values

    value_idx = _week_ahead_pick_value_column(headers, block_idx)
    values = []
    for row in rows[header_idx + 1:]:
        if block_idx >= len(row):
            continue
        block = _week_ahead_parse_number(row[block_idx])
        if block is None or int(block) < 1 or int(block) > 96:
            continue
        raw_value = row[value_idx] if value_idx < len(row) else ""
        num = _week_ahead_parse_number(raw_value)
        values.append(num if num is not None else 0)
    return values


def _week_ahead_extract_values(filename: str, content: bytes, plant_code: str = "") -> List[Any]:
    normalized_plant = _week_ahead_normalize_plant_code(plant_code)
    use_telangana_mapping = normalized_plant in WEEK_AHEAD_TELANGANA_PLANTS
    use_osepl_mapping = normalized_plant in {"OSEPL", "CME"}
    mh_value_column_names = ["CME"] if normalized_plant == "CME" else ["OSEPL", "OSEL"]
    if str(filename or "").lower().endswith(".xlsx"):
        try:
            from openpyxl import load_workbook  # type: ignore
        except Exception as exc:
            raise HTTPException(status_code=500, detail=f"openpyxl is not available: {exc}") from exc
        workbook = load_workbook(io.BytesIO(content), data_only=True, read_only=True)
        rows: List[List[Any]] = []
        for sheet in workbook.worksheets[:1]:
            for row in sheet.iter_rows(values_only=True):
                rows.append(list(row))
        return _week_ahead_extract_values_from_rows(
            rows,
            use_telangana_mapping=use_telangana_mapping,
            use_osepl_mapping=use_osepl_mapping,
            mh_vedanjay_value_column_names=mh_value_column_names,
        )
    return _week_ahead_extract_values_from_rows(
        _week_ahead_parse_csv_rows(content),
        use_telangana_mapping=use_telangana_mapping,
        use_osepl_mapping=use_osepl_mapping,
        mh_vedanjay_value_column_names=mh_value_column_names,
    )


def _week_ahead_format_value(value: Any) -> Any:
    if isinstance(value, (int, float)) and math.isfinite(float(value)):
        rounded = round(float(value), 4)
        return int(rounded) if rounded.is_integer() else rounded
    return value


def _week_ahead_normalize_cme_values(values: List[Any], target_date: date) -> List[Any]:
    normalized_items = [item for item in values if isinstance(item, dict)]

    start_date = target_date + timedelta(days=1)

    if not normalized_items:
        result: List[Dict[str, Any]] = []
        for idx, raw_value in enumerate(values):
            forecast = _week_ahead_parse_number(raw_value)
            forecast_value = forecast if forecast is not None else 0
            result.append({
                "date": (start_date + timedelta(days=idx // 96)).isoformat(),
                "block": (idx % 96) + 1,
                "declared_forecast": forecast_value,
                "inter_avc": 5 if abs(float(forecast_value)) > 1e-9 else 0,
                "schedule": forecast_value,
            })
        return result

    def sort_key(item: Dict[str, Any]) -> Tuple[int, int]:
        date_key = str(item.get("date") or "")
        parsed_date = _week_ahead_parse_date_key(date_key)
        day_index = 10_000
        if parsed_date:
            try:
                day_index = (date.fromisoformat(parsed_date) - start_date).days
            except ValueError:
                day_index = 10_000
        block = int(item.get("block") or 0)
        return day_index, block if block > 0 else 10_000

    ordered = sorted(normalized_items, key=sort_key)
    result: List[Dict[str, Any]] = []
    for idx, item in enumerate(ordered):
        forecast = _week_ahead_parse_number(item.get("declared_forecast"))
        if forecast is None:
            forecast = _week_ahead_parse_number(item.get("schedule"))
        forecast_value = forecast if forecast is not None else 0
        result.append({
            "date": (start_date + timedelta(days=idx // 96)).isoformat(),
            "block": (idx % 96) + 1,
            "declared_forecast": forecast_value,
            "inter_avc": 5 if abs(float(forecast_value)) > 1e-9 else 0,
            "schedule": forecast_value,
        })
    return result


def _week_ahead_normalize_osepl_values(values: List[Any]) -> List[Any]:
    normalized: List[Any] = []
    for item in values:
        if not isinstance(item, dict) or "inter_avc" not in item:
            normalized.append(item)
            continue
        forecast = _week_ahead_parse_number(item.get("declared_forecast"))
        schedule = _week_ahead_parse_number(item.get("schedule"))
        has_generation = any(
            value is not None and abs(float(value)) > 1e-9
            for value in (forecast, schedule)
        )
        next_item = dict(item)
        if not has_generation:
            next_item["inter_avc"] = 0
        normalized.append(next_item)
    return normalized


def _week_ahead_find_target_column(sheet: Any, row_idx: int, block_col: int) -> int:
    max_col = max(sheet.max_column or 1, block_col + 1)
    for header_row in range(max(1, row_idx - 5), row_idx):
        for col in range(1, max_col + 1):
            if col == block_col:
                continue
            token = _week_ahead_header_token(sheet.cell(header_row, col).value)
            if token and any(n in token for n in ("weekahead", "forecast", "schedule", "mw", "value")):
                return col
    for col in range(block_col + 1, max_col + 2):
        if sheet.cell(row_idx, col).value in (None, ""):
            return col
    return block_col + 1


def _week_ahead_find_paired_target_columns(sheet: Any, row_idx: int, block_col: int, day_offset: int = 0) -> Tuple[int, int]:
    max_col = max(sheet.max_column or 1, block_col + 2)
    pairs: List[Tuple[int, int]] = []
    for header_row in range(max(1, row_idx - 8), row_idx):
        for col in range(block_col + 1, max_col + 1):
            token = _week_ahead_header_token(sheet.cell(header_row, col).value)
            if "avc" not in token:
                continue
            for next_col in range(col + 1, min(max_col, col + 3) + 1):
                next_token = _week_ahead_header_token(sheet.cell(header_row, next_col).value)
                if "schedule" in next_token or "schmw" in next_token:
                    pairs.append((col, next_col))
                    break
    if pairs:
        return pairs[min(max(day_offset, 0), len(pairs) - 1)]
    return block_col + 1 + (day_offset * 2), block_col + 2 + (day_offset * 2)


def _week_ahead_find_osepl_target_columns(sheet: Any, row_idx: int, block_col: int, day_offset: int = 0) -> Tuple[int, int, int]:
    max_col = max(sheet.max_column or 1, block_col + 3)
    groups: List[Tuple[int, int, int]] = []
    for header_row in range(max(1, row_idx - 8), row_idx):
        for col in range(block_col + 1, max_col + 1):
            token = _week_ahead_header_token(sheet.cell(header_row, col).value)
            if "declaredforecast" not in token and token != "forecastmw" and token != "forecast":
                continue
            inter_col = -1
            schedule_col = -1
            for next_col in range(col + 1, min(max_col, col + 4) + 1):
                next_token = _week_ahead_header_token(sheet.cell(header_row, next_col).value)
                if inter_col < 0 and ("interavc" in next_token or "intraavc" in next_token or next_token == "avc"):
                    inter_col = next_col
                if "schedule" in next_token:
                    schedule_col = next_col
            if inter_col > 0 and schedule_col > 0:
                groups.append((col, inter_col, schedule_col))
                break
    if groups:
        return groups[min(max(day_offset, 0), len(groups) - 1)]
    return block_col + 1 + (day_offset * 3), block_col + 2 + (day_offset * 3), block_col + 3 + (day_offset * 3)


def _week_ahead_fill_xlsx(template_bytes: bytes, values: List[Any]) -> bytes:
    try:
        from openpyxl import load_workbook  # type: ignore
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"openpyxl is not available: {exc}") from exc
    workbook = load_workbook(io.BytesIO(template_bytes))
    candidates: List[Tuple[Any, int, int]] = []
    for sheet in workbook.worksheets:
        for row_idx in range(1, sheet.max_row + 1):
            for col_idx in range(1, sheet.max_column + 1):
                block_num = _week_ahead_parse_template_block_cell(sheet.cell(row_idx, col_idx).value)
                if block_num is not None:
                    candidates.append((sheet, row_idx, col_idx))
                    break
    start_idx = next((idx for idx, (sheet, row_idx, col_idx) in enumerate(candidates) if _week_ahead_parse_template_block_cell(sheet.cell(row_idx, col_idx).value) == 1), 0)
    block_rows = candidates[start_idx:start_idx + 96]
    if values and isinstance(values[0], dict) and block_rows:
        first_sheet, first_row_idx, first_block_col = block_rows[0]
        is_osepl_values = "inter_avc" in values[0]
        if is_osepl_values:
            for sheet in workbook.worksheets:
                for header_row in range(1, sheet.max_row + 1):
                    headers = [sheet.cell(header_row, col).value for col in range(1, (sheet.max_column or 1) + 1)]
                    tokens = [_week_ahead_header_token(cell) for cell in headers]
                    block_idx = next((idx for idx, token in enumerate(tokens) if token == "block"), -1)
                    forecast_idx = _week_ahead_pick_named_column(headers, ["Declared Forecast", "Declared F", "Forecast"])
                    inter_avc_idx = _week_ahead_pick_named_column(headers, ["Inter Avc", "Inter AVC", "AvC"])
                    schedule_idx = _week_ahead_pick_named_column(headers, ["Schedule"])
                    if min(block_idx, forecast_idx, inter_avc_idx, schedule_idx) < 0:
                        continue
                    max_block = 0
                    for row_idx in range(header_row + 1, min(sheet.max_row, header_row + len(values)) + 1):
                        block_num = _week_ahead_parse_positive_int_cell(sheet.cell(row_idx, block_idx + 1).value)
                        if block_num is not None:
                            max_block = max(max_block, block_num)
                    if max_block <= 96:
                        continue
                    for value_idx, item in enumerate(values):
                        row_idx = header_row + 1 + value_idx
                        sheet.cell(row_idx, block_idx + 1).value = value_idx + 1
                        sheet.cell(row_idx, forecast_idx + 1).value = _week_ahead_format_value(item.get("declared_forecast"))
                        sheet.cell(row_idx, inter_avc_idx + 1).value = _week_ahead_format_value(item.get("inter_avc"))
                        sheet.cell(row_idx, schedule_idx + 1).value = _week_ahead_format_value(item.get("schedule"))
                    output = io.BytesIO()
                    workbook.save(output)
                    return output.getvalue()

        date_cols = _week_ahead_template_date_columns_from_sheet(first_sheet, first_row_idx, first_block_col)
        fallback_dates = sorted({str(item.get("date") or "") for item in values if str(item.get("date") or "")})
        block_row_by_block: Dict[int, Tuple[Any, int, int]] = {}
        for sheet, row_idx, block_col in block_rows:
            block_num = int(_week_ahead_parse_template_block_cell(sheet.cell(row_idx, block_col).value) or 0)
            if 1 <= block_num <= 96:
                block_row_by_block[block_num] = (sheet, row_idx, block_col)
        for sheet, row_idx, _block_col in block_rows:
            for date_col in date_cols.values():
                width = 3 if is_osepl_values else 2
                for offset in range(width):
                    sheet.cell(row_idx, date_col + offset).value = None
        for value_idx, item in enumerate(values):
            date_key = str(item.get("date") or "")
            source_block = int(item.get("block") or 0)
            day_offset = _week_ahead_date_offset(date_key, date_cols, fallback_dates) if date_key else None
            if date_key and source_block:
                if day_offset is None or source_block not in block_row_by_block:
                    continue
                sheet, row_idx, block_col = block_row_by_block[source_block]
            else:
                block_pos = value_idx % 96
                day_offset = value_idx // 96
                sheet, row_idx, block_col = block_rows[block_pos]
            if "inter_avc" in item:
                date_col = date_cols.get(date_key) if date_key in date_cols else None
                if date_col is not None:
                    forecast_col, inter_avc_col, schedule_col = date_col, date_col + 1, date_col + 2
                else:
                    forecast_col, inter_avc_col, schedule_col = _week_ahead_find_osepl_target_columns(sheet, row_idx, block_col, day_offset)
                sheet.cell(row_idx, forecast_col).value = _week_ahead_format_value(item.get("declared_forecast"))
                sheet.cell(row_idx, inter_avc_col).value = _week_ahead_format_value(item.get("inter_avc"))
                sheet.cell(row_idx, schedule_col).value = _week_ahead_format_value(item.get("schedule"))
            else:
                date_col = date_cols.get(date_key) if date_key in date_cols else None
                if date_col is not None:
                    avc_col, schedule_col = date_col, date_col + 1
                else:
                    avc_col, schedule_col = _week_ahead_find_paired_target_columns(sheet, row_idx, block_col, day_offset)
                sheet.cell(row_idx, avc_col).value = _week_ahead_format_value(item.get("avc"))
                sheet.cell(row_idx, schedule_col).value = _week_ahead_format_value(item.get("schedule"))
        output = io.BytesIO()
        workbook.save(output)
        return output.getvalue()

    value_idx = 0
    expected_block = 1
    for sheet, row_idx, block_col in candidates[start_idx:]:
        if value_idx >= len(values):
            break
        block = int(_week_ahead_parse_template_block_cell(sheet.cell(row_idx, block_col).value) or 0)
        if block != expected_block:
            continue
        target_col = _week_ahead_find_target_column(sheet, row_idx, block_col)
        sheet.cell(row_idx, target_col).value = _week_ahead_format_value(values[value_idx])
        value_idx += 1
        expected_block = 1 if expected_block >= 96 else expected_block + 1
    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()


def _week_ahead_fill_csv(template_bytes: bytes, values: List[Any]) -> bytes:
    rows = _week_ahead_parse_csv_rows(template_bytes)
    candidates: List[Tuple[int, int]] = []
    for row_idx, row in enumerate(rows):
        for col_idx, cell in enumerate(row):
            block_num = _week_ahead_parse_template_block_cell(cell)
            if block_num is not None:
                candidates.append((row_idx, col_idx))
                break
    start_idx = next((idx for idx, (row_idx, col_idx) in enumerate(candidates) if _week_ahead_parse_template_block_cell(rows[row_idx][col_idx]) == 1), 0)
    output_rows: List[List[Any]] = [list(row) for row in rows]
    block_rows = candidates[start_idx:start_idx + 96]
    if values and isinstance(values[0], dict) and block_rows:
        first_row_idx, first_block_col = block_rows[0]
        is_osepl_values = "inter_avc" in values[0]
        if is_osepl_values:
            for header_idx, headers in enumerate(rows):
                block_idx = _week_ahead_pick_named_column(headers, ["Block"])
                forecast_idx = _week_ahead_pick_named_column(headers, ["Declared Forecast", "Declared F", "Forecast"])
                inter_avc_idx = _week_ahead_pick_named_column(headers, ["Inter Avc", "Inter AVC", "AvC"])
                schedule_idx = _week_ahead_pick_named_column(headers, ["Schedule"])
                if min(block_idx, forecast_idx, inter_avc_idx, schedule_idx) < 0:
                    continue
                max_block = 0
                for row in rows[header_idx + 1:header_idx + 1 + len(values)]:
                    block_num = _week_ahead_parse_positive_int_cell(row[block_idx] if block_idx < len(row) else "")
                    if block_num is not None:
                        max_block = max(max_block, block_num)
                if max_block <= 96:
                    continue
                max_target_col = max(block_idx, forecast_idx, inter_avc_idx, schedule_idx)
                for value_idx, item in enumerate(values):
                    row_idx = header_idx + 1 + value_idx
                    while len(output_rows) <= row_idx:
                        output_rows.append([])
                    while len(output_rows[row_idx]) <= max_target_col:
                        output_rows[row_idx].append("")
                    output_rows[row_idx][block_idx] = value_idx + 1
                    output_rows[row_idx][forecast_idx] = _week_ahead_format_value(item.get("declared_forecast"))
                    output_rows[row_idx][inter_avc_idx] = _week_ahead_format_value(item.get("inter_avc"))
                    output_rows[row_idx][schedule_idx] = _week_ahead_format_value(item.get("schedule"))
                buffer = io.StringIO()
                writer = csv.writer(buffer, lineterminator="\n")
                writer.writerows(output_rows)
                return buffer.getvalue().encode("utf-8")

        date_cols = _week_ahead_template_date_columns_from_rows(rows, first_row_idx, first_block_col)
        fallback_dates = sorted({str(item.get("date") or "") for item in values if str(item.get("date") or "")})
        block_row_by_block: Dict[int, Tuple[int, int]] = {}
        for row_idx, block_col in block_rows:
            block_num = int(_week_ahead_parse_template_block_cell(rows[row_idx][block_col]) or 0)
            if 1 <= block_num <= 96:
                block_row_by_block[block_num] = (row_idx, block_col)
        for row_idx, _block_col in block_rows:
            for date_col in date_cols.values():
                width = 3 if is_osepl_values else 2
                while len(output_rows[row_idx]) <= date_col + width - 1:
                    output_rows[row_idx].append("")
                for offset in range(width):
                    output_rows[row_idx][date_col + offset] = ""
        for value_idx, item in enumerate(values):
            date_key = str(item.get("date") or "")
            source_block = int(item.get("block") or 0)
            day_offset = _week_ahead_date_offset(date_key, date_cols, fallback_dates) if date_key else None
            if date_key and source_block:
                if day_offset is None or source_block not in block_row_by_block:
                    continue
                row_idx, block_col = block_row_by_block[source_block]
            else:
                block_pos = value_idx % 96
                day_offset = value_idx // 96
                row_idx, block_col = block_rows[block_pos]
            if "inter_avc" in item:
                date_col = date_cols.get(date_key) if date_key in date_cols else None
                forecast_col = date_col if date_col is not None else block_col + 1 + (day_offset * 3)
                inter_avc_col = forecast_col + 1
                schedule_col = forecast_col + 2
                while len(output_rows[row_idx]) <= schedule_col:
                    output_rows[row_idx].append("")
                output_rows[row_idx][forecast_col] = _week_ahead_format_value(item.get("declared_forecast"))
                output_rows[row_idx][inter_avc_col] = _week_ahead_format_value(item.get("inter_avc"))
                output_rows[row_idx][schedule_col] = _week_ahead_format_value(item.get("schedule"))
            else:
                date_col = date_cols.get(date_key) if date_key in date_cols else None
                avc_col = date_col if date_col is not None else block_col + 1 + (day_offset * 2)
                schedule_col = avc_col + 1
                while len(output_rows[row_idx]) <= schedule_col:
                    output_rows[row_idx].append("")
                output_rows[row_idx][avc_col] = _week_ahead_format_value(item.get("avc"))
                output_rows[row_idx][schedule_col] = _week_ahead_format_value(item.get("schedule"))
        buffer = io.StringIO()
        writer = csv.writer(buffer, lineterminator="\n")
        writer.writerows(output_rows)
        return buffer.getvalue().encode("utf-8")

    value_idx = 0
    expected_block = 1
    for row_idx, block_col in candidates[start_idx:]:
        if value_idx >= len(values):
            break
        block = int(_week_ahead_parse_template_block_cell(rows[row_idx][block_col]) or 0)
        if block != expected_block:
            continue
        target_col = block_col + 1
        while len(output_rows[row_idx]) <= target_col:
            output_rows[row_idx].append("")
        output_rows[row_idx][target_col] = _week_ahead_format_value(values[value_idx])
        value_idx += 1
        expected_block = 1 if expected_block >= 96 else expected_block + 1
    buffer = io.StringIO()
    writer = csv.writer(buffer, lineterminator="\n")
    writer.writerows(output_rows)
    return buffer.getvalue().encode("utf-8")


def _week_ahead_content_type(filename: str) -> str:
    if str(filename or "").lower().endswith(".csv"):
        return "text/csv"
    return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


@app.get("/api/week-ahead/templates/{plant_code}")
async def get_week_ahead_template_status(plant_code: str):
    plant = _week_ahead_require_supported_plant(plant_code)
    metadata = _week_ahead_read_metadata(plant)
    return {
        "plant_code": plant,
        "uploaded": bool(metadata.get("template_key") or metadata.get("local_path")),
        "filename": str(metadata.get("filename") or ""),
        "uploaded_at": str(metadata.get("uploaded_at") or ""),
        "uploaded_by": str(metadata.get("uploaded_by") or ""),
        "storage_mode": str(metadata.get("storage_mode") or ""),
    }


@app.post("/api/week-ahead/templates/upload")
async def upload_week_ahead_template(
    plant_code: str = Form(...),
    uploaded_by: str = Form(""),
    file: UploadFile = File(...),
    x_user_role: Optional[str] = Header(None),
):
    if str(x_user_role or "").strip().lower() not in {"admin", "intern", "employee", "member"}:
        raise HTTPException(status_code=403, detail="Only admin, intern, or employee can upload week-ahead templates")
    plant = _week_ahead_require_supported_plant(plant_code)
    filename = _week_ahead_safe_filename(file.filename or f"{plant}_week_ahead_template.xlsx")
    if not filename.lower().endswith((".xlsx", ".csv")):
        raise HTTPException(status_code=400, detail="Upload an .xlsx or .csv template")
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Uploaded template is empty")

    uploaded_at = datetime.utcnow().isoformat() + "Z"
    template_key = _week_ahead_template_key(plant, filename)
    metadata = {
        "plant_code": plant,
        "filename": filename,
        "template_key": template_key,
        "uploaded_at": uploaded_at,
        "uploaded_by": str(uploaded_by or "").strip()[:200],
        "storage_mode": "s3",
    }

    s3, bucket, _region = _week_ahead_s3_client()
    try:
        if s3 is None or not bucket:
            raise RuntimeError("S3 bucket is not configured")
        s3.put_object(
            Bucket=bucket,
            Key=template_key,
            Body=content,
            ContentType=_week_ahead_content_type(filename),
            Metadata={"plant_code": plant, "uploaded_by": str(uploaded_by or "")[:200]},
        )
        s3.put_object(
            Bucket=bucket,
            Key=_week_ahead_metadata_key(plant),
            Body=json.dumps(metadata, ensure_ascii=True).encode("utf-8"),
            ContentType="application/json",
        )
    except Exception as exc:
        local_path = _week_ahead_local_template_path(plant, filename)
        os.makedirs(os.path.dirname(local_path), exist_ok=True)
        with open(local_path, "wb") as f:
            f.write(content)
        metadata.update({
            "template_key": "",
            "local_path": local_path,
            "storage_mode": "local",
            "error": str(exc),
        })
        with open(_week_ahead_local_metadata_path(plant), "w", encoding="utf-8") as f:
            json.dump(metadata, f, ensure_ascii=True, indent=2)

    return {"success": True, "uploaded": True, **metadata}


@app.get("/api/week-ahead/download")
async def download_week_ahead_template(
    plant_code: str = Query(...),
    target_date: date = Query(...),
):
    plant = _week_ahead_require_supported_plant(plant_code)
    metadata = _week_ahead_read_metadata(plant)
    if not metadata:
        raise HTTPException(status_code=404, detail=f"No week-ahead template uploaded for {plant}")

    source_key, source_bytes = _week_ahead_fetch_latest_source(plant, target_date)
    values = _week_ahead_extract_values(source_key, source_bytes, plant_code=plant)
    if plant == "CME":
        values = _week_ahead_normalize_cme_values(values, target_date)
    if plant == "OSEPL":
        values = _week_ahead_normalize_osepl_values(values)
    if not values:
        raise HTTPException(status_code=400, detail=f"No week-ahead values found in {source_key}")

    template_bytes = _week_ahead_fetch_template_bytes(metadata)
    filename = str(metadata.get("filename") or f"{plant}_week_ahead_template.xlsx")
    if plant == "CME" or filename.lower().endswith(".csv"):
        output_bytes = _week_ahead_fill_csv(template_bytes, values)
        output_name = f"{plant}_{target_date.isoformat()}_week_ahead.csv"
    else:
        output_bytes = _week_ahead_fill_xlsx(template_bytes, values)
        output_name = f"{plant}_{target_date.isoformat()}_week_ahead.xlsx"

    headers = {
        "Content-Disposition": f'attachment; filename="{output_name}"',
        "X-Week-Ahead-Source-Key": source_key,
        "X-Week-Ahead-Value-Count": str(len(values)),
    }
    return StreamingResponse(
        io.BytesIO(output_bytes),
        media_type=_week_ahead_content_type(output_name),
        headers=headers,
    )


# ==================== FILE UPLOAD ENDPOINT ====================
@app.post("/api/upload")
async def upload_file(
    file: UploadFile = File(...),
    vendor: Optional[str] = None,
    type: Optional[str] = None
):
    """Upload a file"""
    try:
        # Create uploads directory if it doesn't exist
        os.makedirs("uploads", exist_ok=True)
        
        # Save file
        file_path = f"uploads/{datetime.now().timestamp()}-{file.filename}"
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        return {
            "message": "File uploaded successfully",
            "filename": file.filename,
            "size": len(content),
            "path": file_path
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


# ==================== DOCUMENTATION ENDPOINTS ====================
DOCUMENTATION_S3_PREFIX = "documentation/portal"
DOCUMENTATION_ACCESS_EVERYONE = "everyone"
DOCUMENTATION_ACCESS_ADMIN_INTERN = "admin_intern"
DOCUMENTATION_ACCESS_VALUES = {DOCUMENTATION_ACCESS_EVERYONE, DOCUMENTATION_ACCESS_ADMIN_INTERN}
DOCUMENTATION_DEFAULT_HEADING = "General"


def _documentation_ensure_schema(db: Session) -> None:
    try:
        inspector = inspect(db.bind)
        if not inspector.has_table("documentation_headings"):
            db.execute(
                text(
                    "CREATE TABLE documentation_headings ("
                    "id VARCHAR(64) PRIMARY KEY, "
                    "name VARCHAR(255) NOT NULL UNIQUE, "
                    "created_by VARCHAR(255), "
                    "created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP NOT NULL"
                    ")"
                )
            )
            db.commit()
        columns = {column.get("name") for column in inspector.get_columns("documentation_documents")}
        if "access_category" not in columns:
            db.execute(
                text(
                    "ALTER TABLE documentation_documents "
                    "ADD COLUMN access_category VARCHAR(50) DEFAULT 'everyone' NOT NULL"
                )
            )
            db.commit()
        if "heading" not in columns:
            db.execute(
                text(
                    "ALTER TABLE documentation_documents "
                    f"ADD COLUMN heading VARCHAR(255) DEFAULT '{DOCUMENTATION_DEFAULT_HEADING}' NOT NULL"
                )
            )
            db.commit()
    except Exception:
        db.rollback()


def _documentation_normalize_access_category(value: Optional[str]) -> str:
    raw = str(value or "").strip().lower().replace("-", "_").replace(" ", "_")
    if raw in {"adminintern", "admin_intern", "admin_and_intern", "admin_intern_only"}:
        return DOCUMENTATION_ACCESS_ADMIN_INTERN
    if raw in DOCUMENTATION_ACCESS_VALUES:
        return raw
    return DOCUMENTATION_ACCESS_EVERYONE


def _documentation_normalize_heading(value: Optional[str]) -> str:
    heading = re.sub(r"\s+", " ", str(value or "").strip())
    return heading[:255] if heading else DOCUMENTATION_DEFAULT_HEADING


def _documentation_heading_public_row(row: DocumentationHeading) -> Dict[str, Any]:
    return {
        "id": str(row.id or ""),
        "name": _documentation_normalize_heading(row.name),
        "created_by": str(row.created_by or ""),
        "created_at": row.created_at.isoformat() if row.created_at else "",
    }


def _documentation_find_heading_by_name(db: Session, name: str) -> Optional[DocumentationHeading]:
    normalized = _documentation_normalize_heading(name)
    rows = db.query(DocumentationHeading).all()
    for row in rows:
        if _documentation_normalize_heading(row.name).lower() == normalized.lower():
            return row
    return None


def _documentation_viewer_type(role: Optional[str], user_name: Optional[str] = None) -> str:
    role_value = str(role or "").strip().lower()
    user_value = str(user_name or "").strip().lower()
    if role_value == "admin":
        return "admin"
    if role_value == "intern" or user_value == "intern" or "intern" in user_value:
        return "intern"
    return "employee"


def _documentation_can_access(access_category: Optional[str], viewer_type: str) -> bool:
    category = _documentation_normalize_access_category(access_category)
    if category == DOCUMENTATION_ACCESS_ADMIN_INTERN:
        return viewer_type in {"admin", "intern"}
    return True


def _documentation_forbid_if_no_access(doc: Dict[str, Any], viewer_type: str) -> None:
    if not _documentation_can_access(str(doc.get("access_category") or ""), viewer_type):
        raise HTTPException(status_code=403, detail="You do not have access to this document")


def _documentation_safe_name(value: str) -> str:
    raw = os.path.basename(str(value or "document").strip()) or "document"
    return re.sub(r"[^A-Za-z0-9._ -]+", "_", raw).strip(" .") or "document"


def _documentation_s3_prefix() -> str:
    return str(os.getenv("DOCUMENTATION_S3_PREFIX") or DOCUMENTATION_S3_PREFIX).strip().strip("/") or DOCUMENTATION_S3_PREFIX


def _documentation_s3_metadata_key(document_id: str) -> str:
    return f"{_documentation_s3_prefix()}/{str(document_id or '').strip()}/metadata.json"


def _documentation_s3_content_key(document_id: str, filename: str) -> str:
    safe_name = _documentation_safe_name(filename or "document")
    return f"{_documentation_s3_prefix()}/{str(document_id or '').strip()}/content/{safe_name}"


def _documentation_s3_client() -> Tuple[Optional[Any], str]:
    bucket = _derive_s3_bucket_name()
    if not bucket:
        return None, ""
    try:
        import boto3  # type: ignore

        region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"
        return boto3.client("s3", region_name=region), bucket
    except Exception:
        return None, ""


def _documentation_normalize_link(value: str) -> str:
    raw = str(value or "").strip()
    if not raw:
        raise HTTPException(status_code=400, detail="Link URL is required")
    parsed = urlparse(raw)
    if parsed.scheme.lower() not in {"http", "https"} or not parsed.netloc:
        raise HTTPException(status_code=400, detail="Link must start with http:// or https://")
    return raw


def _documentation_is_link(row: DocumentationDocument) -> bool:
    return str(row.content_type or "").strip().lower() == "text/uri-list"


def _documentation_public_row(row: DocumentationDocument) -> Dict[str, Any]:
    doc_id = str(row.id or "").strip()
    uploaded_at = row.uploaded_at.isoformat() if row.uploaded_at else ""
    is_link = _documentation_is_link(row)
    link_url = ""
    if is_link:
        try:
            link_url = bytes(row.file_data or b"").decode("utf-8", errors="replace").strip()
        except Exception:
            link_url = ""
    return {
        "id": doc_id,
        "filename": str(row.filename or ""),
        "content_type": str(row.content_type or "application/octet-stream"),
        "size": int(row.size or 0),
        "uploaded_at": uploaded_at,
        "uploaded_by": str(row.uploaded_by or ""),
        "role": str(row.role or ""),
        "access_category": _documentation_normalize_access_category(getattr(row, "access_category", "")),
        "heading": _documentation_normalize_heading(getattr(row, "heading", "")),
        "document_type": "link" if is_link else "file",
        "link_url": link_url,
        "preview_url": f"/documentation/documents/{doc_id}/preview",
        "download_url": f"/documentation/documents/{doc_id}/download",
    }


def _documentation_public_s3_row(meta: Dict[str, Any]) -> Dict[str, Any]:
    doc_id = str(meta.get("id") or "").strip()
    is_link = str(meta.get("document_type") or "").strip().lower() == "link"
    content_type = str(meta.get("content_type") or ("text/uri-list" if is_link else "application/octet-stream"))
    return {
        "id": doc_id,
        "filename": str(meta.get("filename") or ""),
        "content_type": content_type,
        "size": int(meta.get("size") or 0),
        "uploaded_at": str(meta.get("uploaded_at") or ""),
        "uploaded_by": str(meta.get("uploaded_by") or ""),
        "role": str(meta.get("role") or ""),
        "access_category": _documentation_normalize_access_category(str(meta.get("access_category") or "")),
        "heading": _documentation_normalize_heading(str(meta.get("heading") or "")),
        "document_type": "link" if is_link else "file",
        "link_url": str(meta.get("link_url") or ""),
        "preview_url": f"/documentation/documents/{doc_id}/preview",
        "download_url": f"/documentation/documents/{doc_id}/download",
        "content_key": str(meta.get("content_key") or ""),
    }


def _documentation_s3_metadata_from_row(row: DocumentationDocument, *, content_key: str = "") -> Dict[str, Any]:
    public = _documentation_public_row(row)
    public["content_key"] = content_key
    return public


def _documentation_s3_put_json(s3: Any, bucket: str, key: str, payload: Dict[str, Any]) -> None:
    s3.put_object(
        Bucket=bucket,
        Key=key,
        Body=json.dumps(payload, ensure_ascii=True, separators=(",", ":")).encode("utf-8"),
        ContentType="application/json",
    )


def _documentation_mirror_row_to_s3(row: DocumentationDocument) -> None:
    s3, bucket = _documentation_s3_client()
    if not s3 or not bucket:
        return
    doc_id = str(row.id or "").strip()
    if not doc_id:
        return
    content_key = ""
    if not _documentation_is_link(row):
        content_key = _documentation_s3_content_key(doc_id, str(row.filename or "document"))
        s3.put_object(
            Bucket=bucket,
            Key=content_key,
            Body=bytes(row.file_data or b""),
            ContentType=str(row.content_type or "application/octet-stream"),
        )
    meta = _documentation_s3_metadata_from_row(row, content_key=content_key)
    _documentation_s3_put_json(s3, bucket, _documentation_s3_metadata_key(doc_id), meta)


def _documentation_list_s3_rows() -> List[Dict[str, Any]]:
    s3, bucket = _documentation_s3_client()
    if not s3 or not bucket:
        return []
    prefix = f"{_documentation_s3_prefix()}/"
    rows: List[Dict[str, Any]] = []
    token = None
    try:
        while True:
            kwargs = {"Bucket": bucket, "Prefix": prefix, "MaxKeys": 1000}
            if token:
                kwargs["ContinuationToken"] = token
            response = s3.list_objects_v2(**kwargs)
            for obj in response.get("Contents") or []:
                key = str(obj.get("Key") or "")
                if not key.endswith("/metadata.json"):
                    continue
                item = s3.get_object(Bucket=bucket, Key=key)
                body = item.get("Body")
                raw = body.read().decode("utf-8", errors="replace") if body is not None else "{}"
                meta = json.loads(raw or "{}")
                if isinstance(meta, dict) and str(meta.get("id") or "").strip():
                    rows.append(_documentation_public_s3_row(meta))
            if not response.get("IsTruncated"):
                break
            token = response.get("NextContinuationToken")
            if not token:
                break
    except Exception:
        return []
    return rows


def _documentation_get_s3_meta(document_id: str) -> Optional[Dict[str, Any]]:
    s3, bucket = _documentation_s3_client()
    if not s3 or not bucket:
        return None
    try:
        item = s3.get_object(Bucket=bucket, Key=_documentation_s3_metadata_key(document_id))
        body = item.get("Body")
        raw = body.read().decode("utf-8", errors="replace") if body is not None else "{}"
        meta = json.loads(raw or "{}")
        return _documentation_public_s3_row(meta) if isinstance(meta, dict) else None
    except Exception:
        return None


def _documentation_get_s3_content(meta: Dict[str, Any]) -> bytes:
    s3, bucket = _documentation_s3_client()
    if not s3 or not bucket:
        raise HTTPException(status_code=404, detail="Document not found")
    key = str(meta.get("content_key") or "").strip()
    if not key:
        raise HTTPException(status_code=404, detail="Document content not found")
    try:
        item = s3.get_object(Bucket=bucket, Key=key)
        body = item.get("Body")
        return body.read() if body is not None else b""
    except Exception as exc:
        raise HTTPException(status_code=404, detail="Document content not found") from exc


def _documentation_delete_s3_document(document_id: str, meta: Optional[Dict[str, Any]] = None) -> None:
    s3, bucket = _documentation_s3_client()
    if not s3 or not bucket:
        return
    doc_id = str(document_id or "").strip()
    if not doc_id:
        return
    keys = [_documentation_s3_metadata_key(doc_id)]
    content_key = str((meta or {}).get("content_key") or "").strip()
    if content_key:
        keys.append(content_key)
    for key in keys:
        try:
            s3.delete_object(Bucket=bucket, Key=key)
        except Exception:
            pass


def _documentation_find_or_404(document_id: str, db: Session) -> DocumentationDocument:
    safe_id = str(document_id or "").strip()
    row = db.query(DocumentationDocument).filter(DocumentationDocument.id == safe_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="Document not found")
    return row


def _documentation_extension(filename: str) -> str:
    _, ext = os.path.splitext(str(filename or "").strip().lower())
    return ext.lstrip(".")


def _documentation_html_page(title: str, body: str) -> HTMLResponse:
    safe_title = html.escape(str(title or "Document"))
    page = f"""<!doctype html>
<html>
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{safe_title}</title>
  <style>
    :root {{ color-scheme: light; }}
    body {{
      margin: 0;
      padding: 24px;
      color: #111827;
      background: #ffffff;
      font-family: Arial, Helvetica, sans-serif;
      font-size: 14px;
      line-height: 1.5;
    }}
    .doc-page {{ max-width: 980px; margin: 0 auto; }}
    .sheet {{ margin-bottom: 28px; }}
    h2 {{ margin: 0 0 12px; font-size: 18px; }}
    p {{ margin: 0 0 10px; white-space: pre-wrap; }}
    table {{ border-collapse: collapse; width: 100%; margin: 10px 0 18px; }}
    th, td {{ border: 1px solid #d1d5db; padding: 6px 8px; vertical-align: top; }}
    th {{ background: #f3f4f6; font-weight: 700; }}
    .empty {{ color: #64748b; text-align: center; margin-top: 96px; }}
  </style>
</head>
<body>
  <main class="doc-page">{body}</main>
</body>
</html>"""
    return HTMLResponse(page)


def _documentation_docx_preview(filename: str, content: bytes) -> HTMLResponse:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"python-docx is not available: {exc}") from exc

    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Unable to preview Word document: {exc}") from exc

    parts: List[str] = []
    for paragraph in document.paragraphs:
        text_value = paragraph.text.strip()
        if text_value:
            parts.append(f"<p>{html.escape(text_value)}</p>")

    for table in document.tables:
        rows: List[str] = []
        for row in table.rows:
            cells = "".join(f"<td>{html.escape(cell.text.strip())}</td>" for cell in row.cells)
            rows.append(f"<tr>{cells}</tr>")
        if rows:
            parts.append(f"<table>{''.join(rows)}</table>")

    body = "".join(parts) or '<div class="empty">No previewable text found in this Word document.</div>'
    return _documentation_html_page(filename, body)


def _documentation_xlsx_preview(filename: str, content: bytes) -> HTMLResponse:
    try:
        from openpyxl import load_workbook  # type: ignore
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"openpyxl is not available: {exc}") from exc

    try:
        workbook = load_workbook(io.BytesIO(content), data_only=True, read_only=True)
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Unable to preview Excel workbook: {exc}") from exc

    sheets: List[str] = []
    for worksheet in workbook.worksheets[:5]:
        rows: List[str] = []
        for row_index, row in enumerate(worksheet.iter_rows(max_row=200, max_col=50, values_only=True), start=1):
            values = ["" if value is None else str(value) for value in row]
            if not any(value.strip() for value in values):
                continue
            tag = "th" if row_index == 1 else "td"
            cells = "".join(f"<{tag}>{html.escape(value)}</{tag}>" for value in values)
            rows.append(f"<tr>{cells}</tr>")
        table = f"<table>{''.join(rows)}</table>" if rows else '<div class="empty">This sheet has no previewable data.</div>'
        sheets.append(f'<section class="sheet"><h2>{html.escape(worksheet.title)}</h2>{table}</section>')

    body = "".join(sheets) or '<div class="empty">No previewable sheets found in this workbook.</div>'
    return _documentation_html_page(filename, body)


@app.get("/api/documentation/documents")
async def list_documentation_documents(
    x_user_role: Optional[str] = Header(None, alias="X-User-Role"),
    x_user_name: Optional[str] = Header(None, alias="X-User-Name"),
    db: Session = Depends(get_db),
):
    _documentation_ensure_schema(db)
    viewer_type = _documentation_viewer_type(x_user_role, x_user_name)
    rows = (
        db.query(DocumentationDocument)
        .order_by(DocumentationDocument.uploaded_at.desc())
        .all()
    )
    for row in rows:
        try:
            _documentation_mirror_row_to_s3(row)
        except Exception:
            pass
    db_docs = [_documentation_public_row(row) for row in rows]
    seen_ids = {str(doc.get("id") or "") for doc in db_docs}
    s3_docs = [doc for doc in _documentation_list_s3_rows() if str(doc.get("id") or "") not in seen_ids]
    docs = [
        doc for doc in (db_docs + s3_docs)
        if _documentation_can_access(str(doc.get("access_category") or ""), viewer_type)
    ]
    docs.sort(key=lambda item: str(item.get("uploaded_at") or ""), reverse=True)
    return {"documents": docs}


@app.get("/api/documentation/headings")
async def list_documentation_headings(
    x_user_role: Optional[str] = Header(None, alias="X-User-Role"),
    x_user_name: Optional[str] = Header(None, alias="X-User-Name"),
    db: Session = Depends(get_db),
):
    _documentation_ensure_schema(db)
    viewer_type = _documentation_viewer_type(x_user_role, x_user_name)
    rows = db.query(DocumentationHeading).order_by(DocumentationHeading.created_at.asc()).all()
    headings = [_documentation_heading_public_row(row) for row in rows]
    if not any(item["name"].lower() == DOCUMENTATION_DEFAULT_HEADING.lower() for item in headings):
        headings.insert(0, {
            "id": "default",
            "name": DOCUMENTATION_DEFAULT_HEADING,
            "created_by": "",
            "created_at": "",
        })
    if viewer_type == "admin":
        return {"headings": headings}

    docs = [
        _documentation_public_row(row)
        for row in db.query(DocumentationDocument).all()
        if _documentation_can_access(getattr(row, "access_category", ""), viewer_type)
    ]
    visible_names = {_documentation_normalize_heading(doc.get("heading")).lower() for doc in docs}
    return {
        "headings": [
            item for item in headings
            if _documentation_normalize_heading(item.get("name")).lower() in visible_names
        ]
    }


@app.post("/api/documentation/headings")
async def create_documentation_heading(
    name: str = Form(...),
    created_by: Optional[str] = Form(""),
    x_user_role: Optional[str] = Header(None, alias="X-User-Role"),
    db: Session = Depends(get_db),
):
    if str(x_user_role or "").strip().lower() != "admin":
        raise HTTPException(status_code=403, detail="Only admin users can create documentation headings")
    _documentation_ensure_schema(db)
    heading_name = _documentation_normalize_heading(name)
    existing = _documentation_find_heading_by_name(db, heading_name)
    if existing:
        return {"ok": True, "heading": _documentation_heading_public_row(existing)}
    row = DocumentationHeading(
        id=uuid4().hex,
        name=heading_name,
        created_by=str(created_by or "").strip(),
        created_at=datetime.now(timezone.utc),
    )
    try:
        db.add(row)
        db.commit()
        db.refresh(row)
        return {"ok": True, "heading": _documentation_heading_public_row(row)}
    except Exception as exc:
        db.rollback()
        raise HTTPException(status_code=400, detail=f"Heading create failed: {exc}") from exc


@app.delete("/api/documentation/headings/{heading_id}")
async def delete_documentation_heading(
    heading_id: str,
    x_user_role: Optional[str] = Header(None, alias="X-User-Role"),
    db: Session = Depends(get_db),
):
    if str(x_user_role or "").strip().lower() != "admin":
        raise HTTPException(status_code=403, detail="Only admin users can delete documentation headings")
    if str(heading_id or "").strip() == "default":
        raise HTTPException(status_code=400, detail="Default heading cannot be deleted")
    _documentation_ensure_schema(db)
    row = db.query(DocumentationHeading).filter(DocumentationHeading.id == str(heading_id or "").strip()).first()
    if not row:
        raise HTTPException(status_code=404, detail="Heading not found")
    heading_name = _documentation_normalize_heading(row.name)
    try:
        docs = db.query(DocumentationDocument).filter(DocumentationDocument.heading == heading_name).all()
        for doc in docs:
            _documentation_delete_s3_document(str(doc.id or ""))
            db.delete(doc)
        db.delete(row)
        db.commit()
        return {"ok": True, "message": "Heading and documents deleted"}
    except Exception as exc:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Heading delete failed: {exc}") from exc


@app.post("/api/documentation/documents")
async def upload_documentation_document(
    file: UploadFile = File(...),
    uploaded_by: Optional[str] = Form(""),
    role: Optional[str] = Form(""),
    access_category: Optional[str] = Form(DOCUMENTATION_ACCESS_EVERYONE),
    heading: Optional[str] = Form(DOCUMENTATION_DEFAULT_HEADING),
    db: Session = Depends(get_db),
):
    try:
        _documentation_ensure_schema(db)
        original_name = _documentation_safe_name(file.filename or "document")
        content = await file.read()
        if not content:
            raise HTTPException(status_code=400, detail="Uploaded document is empty")

        content_type = (
            str(file.content_type or "").strip()
            or mimetypes.guess_type(original_name)[0]
            or "application/octet-stream"
        )
        row = DocumentationDocument(
            id=uuid4().hex,
            filename=original_name,
            content_type=content_type,
            size=len(content),
            file_data=content,
            uploaded_by=str(uploaded_by or "").strip(),
            role=str(role or "").strip(),
            access_category=_documentation_normalize_access_category(access_category),
            heading=_documentation_normalize_heading(heading),
            uploaded_at=datetime.now(timezone.utc),
        )
        db.add(row)
        db.commit()
        db.refresh(row)
        try:
            _documentation_mirror_row_to_s3(row)
        except Exception:
            pass
        return {"ok": True, "document": _documentation_public_row(row)}
    except HTTPException:
        raise
    except Exception as exc:
        db.rollback()
        raise HTTPException(status_code=400, detail=f"Document upload failed: {exc}") from exc


@app.post("/api/documentation/links")
async def create_documentation_link(
    title: str = Form(...),
    url: str = Form(...),
    uploaded_by: Optional[str] = Form(""),
    role: Optional[str] = Form(""),
    access_category: Optional[str] = Form(DOCUMENTATION_ACCESS_EVERYONE),
    heading: Optional[str] = Form(DOCUMENTATION_DEFAULT_HEADING),
    db: Session = Depends(get_db),
):
    try:
        _documentation_ensure_schema(db)
        link_url = _documentation_normalize_link(url)
        link_title = _documentation_safe_name(title or link_url)
        row = DocumentationDocument(
            id=uuid4().hex,
            filename=link_title,
            content_type="text/uri-list",
            size=0,
            file_data=link_url.encode("utf-8"),
            uploaded_by=str(uploaded_by or "").strip(),
            role=str(role or "").strip(),
            access_category=_documentation_normalize_access_category(access_category),
            heading=_documentation_normalize_heading(heading),
            uploaded_at=datetime.now(timezone.utc),
        )
        db.add(row)
        db.commit()
        db.refresh(row)
        try:
            _documentation_mirror_row_to_s3(row)
        except Exception:
            pass
        return {"ok": True, "document": _documentation_public_row(row)}
    except HTTPException:
        raise
    except Exception as exc:
        db.rollback()
        raise HTTPException(status_code=400, detail=f"Link upload failed: {exc}") from exc


@app.get("/api/documentation/documents/{document_id}/preview")
async def preview_documentation_document(
    document_id: str,
    role: Optional[str] = Query(None),
    user_name: Optional[str] = Query(None),
    x_user_role: Optional[str] = Header(None, alias="X-User-Role"),
    x_user_name: Optional[str] = Header(None, alias="X-User-Name"),
    db: Session = Depends(get_db),
):
    _documentation_ensure_schema(db)
    viewer_type = _documentation_viewer_type(x_user_role or role, x_user_name or user_name)
    safe_id = str(document_id or "").strip()
    row = db.query(DocumentationDocument).filter(DocumentationDocument.id == safe_id).first()
    s3_meta = None if row else _documentation_get_s3_meta(safe_id)
    if not row and not s3_meta:
        raise HTTPException(status_code=404, detail="Document not found")
    doc_meta = _documentation_public_row(row) if row else s3_meta
    _documentation_forbid_if_no_access(doc_meta, viewer_type)
    if row and _documentation_is_link(row):
        link_url = bytes(row.file_data or b"").decode("utf-8", errors="replace").strip()
        return JSONResponse({"url": link_url})
    if s3_meta and str(s3_meta.get("document_type") or "").lower() == "link":
        link_url = str(s3_meta.get("link_url") or "").strip()
        return JSONResponse({"url": link_url})
    filename = _documentation_safe_name((row.filename if row else s3_meta.get("filename")) or "document")
    content = (row.file_data or b"") if row else _documentation_get_s3_content(s3_meta)
    extension = _documentation_extension(filename)
    if extension == "docx":
        return _documentation_docx_preview(filename, content)
    if extension in {"xlsx", "xlsm", "xltx", "xltm"}:
        return _documentation_xlsx_preview(filename, content)
    return StreamingResponse(
        io.BytesIO(content),
        media_type=str((row.content_type if row else s3_meta.get("content_type")) or "application/octet-stream"),
        headers={"Content-Disposition": f'inline; filename="{filename}"'},
    )


@app.get("/api/documentation/documents/{document_id}/download")
async def download_documentation_document(
    document_id: str,
    role: Optional[str] = Query(None),
    user_name: Optional[str] = Query(None),
    x_user_role: Optional[str] = Header(None, alias="X-User-Role"),
    x_user_name: Optional[str] = Header(None, alias="X-User-Name"),
    db: Session = Depends(get_db),
):
    _documentation_ensure_schema(db)
    viewer_type = _documentation_viewer_type(x_user_role or role, x_user_name or user_name)
    safe_id = str(document_id or "").strip()
    row = db.query(DocumentationDocument).filter(DocumentationDocument.id == safe_id).first()
    s3_meta = None if row else _documentation_get_s3_meta(safe_id)
    if not row and not s3_meta:
        raise HTTPException(status_code=404, detail="Document not found")
    doc_meta = _documentation_public_row(row) if row else s3_meta
    _documentation_forbid_if_no_access(doc_meta, viewer_type)
    filename = _documentation_safe_name((row.filename if row else s3_meta.get("filename")) or "document")
    content = (row.file_data or b"") if row else _documentation_get_s3_content(s3_meta)
    return StreamingResponse(
        io.BytesIO(content),
        media_type=str((row.content_type if row else s3_meta.get("content_type")) or "application/octet-stream"),
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.delete("/api/documentation/documents/{document_id}")
async def delete_documentation_document(
    document_id: str,
    x_user_role: Optional[str] = Header(None, alias="X-User-Role"),
    db: Session = Depends(get_db),
):
    if str(x_user_role or "").strip().lower() != "admin":
        raise HTTPException(status_code=403, detail="Only admin users can delete documents")

    _documentation_ensure_schema(db)
    safe_id = str(document_id or "").strip()
    row = db.query(DocumentationDocument).filter(DocumentationDocument.id == safe_id).first()
    s3_meta = _documentation_get_s3_meta(safe_id)
    if not row and not s3_meta:
        raise HTTPException(status_code=404, detail="Document not found")
    try:
        if row:
            db.delete(row)
            db.commit()
        _documentation_delete_s3_document(safe_id, s3_meta)
    except Exception as exc:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Failed to delete document: {exc}") from exc
    return {"ok": True, "message": "Document deleted"}


# ==================== EXPORT ENDPOINTS ====================
@app.get("/api/export/schedules")
async def export_schedules(
    format: str = Query("csv", regex="^(csv|json)$"),
    db: Session = Depends(get_db)
):
    """Export schedules in CSV or JSON format"""
    try:
        schedules = get_schedules(db)
        
        if format == "csv":
            output = io.StringIO()
            if schedules:
                # Convert SQLAlchemy models to dicts
                schedule_dicts = [{
                    "id": s.id,
                    "plantName": s.plantName,
                    "type": s.type,
                    "scheduleDate": str(s.scheduleDate),
                    "capacity": s.capacity,
                    "forecasted": s.forecasted,
                    "actual": s.actual,
                    "status": s.status,
                    "deviation": s.deviation
                } for s in schedules]
                
                if schedule_dicts:
                    writer = csv.DictWriter(output, fieldnames=schedule_dicts[0].keys())
                    writer.writeheader()
                    writer.writerows(schedule_dicts)
            
            return StreamingResponse(
                io.BytesIO(output.getvalue().encode('utf-8')),
                media_type="text/csv",
                headers={"Content-Disposition": "attachment; filename=schedules.csv"}
            )
        else:  # JSON
            schedule_dicts = [{
                "id": s.id,
                "plantName": s.plantName,
                "type": s.type,
                "scheduleDate": str(s.scheduleDate),
                "capacity": s.capacity,
                "forecasted": s.forecasted,
                "actual": s.actual,
                "status": s.status,
                "deviation": s.deviation
            } for s in schedules]
            
            return JSONResponse(
                content=schedule_dicts,
                headers={"Content-Disposition": "attachment; filename=schedules.json"}
            )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/export/plants")
async def export_plants(
    format: str = Query("csv", regex="^(csv|json)$"),
    db: Session = Depends(get_db)
):
    """Export plants in CSV or JSON format"""
    try:
        plants = get_plants(db)
        
        if format == "csv":
            output = io.StringIO()
            if plants:
                # Convert SQLAlchemy models to dicts
                plant_dicts = [{
                    "id": p.id,
                    "name": p.name,
                    "type": p.type,
                    "capacity": p.capacity,
                    "state": p.state,
                    "status": p.status,
                    "efficiency": p.efficiency,
                    "penalty_threshold_percent": p.penalty_threshold_percent,
                    "lastUpdated": str(p.lastUpdated) if p.lastUpdated else ""
                } for p in plants]
                
                if plant_dicts:
                    writer = csv.DictWriter(output, fieldnames=plant_dicts[0].keys())
                    writer.writeheader()
                    writer.writerows(plant_dicts)
            
            return StreamingResponse(
                io.BytesIO(output.getvalue().encode('utf-8')),
                media_type="text/csv",
                headers={"Content-Disposition": "attachment; filename=plants.csv"}
            )
        else:  # JSON
            plant_dicts = [{
                "id": p.id,
                "name": p.name,
                "type": p.type,
                "capacity": p.capacity,
                "state": p.state,
                "status": p.status,
                "efficiency": p.efficiency,
                "penalty_threshold_percent": p.penalty_threshold_percent,
                "lastUpdated": str(p.lastUpdated) if p.lastUpdated else ""
            } for p in plants]
            
            return JSONResponse(
                content=plant_dicts,
                headers={"Content-Disposition": "attachment; filename=plants.json"}
            )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/export/deviations")
async def export_deviations(
    format: str = Query("csv", regex="^(csv|json)$"),
    db: Session = Depends(get_db)
):
    """Export deviations in CSV or JSON format"""
    try:
        deviations = get_deviations(db, period="hourly", limit=1000)
        
        if format == "csv":
            output = io.StringIO()
            if deviations:
                # deviations is already a list of dicts from get_deviations
                if deviations and isinstance(deviations[0], dict):
                    writer = csv.DictWriter(output, fieldnames=deviations[0].keys())
                    writer.writeheader()
                    writer.writerows(deviations)
            
            return StreamingResponse(
                io.BytesIO(output.getvalue().encode('utf-8')),
                media_type="text/csv",
                headers={"Content-Disposition": "attachment; filename=deviations.csv"}
            )
        else:  # JSON
            return JSONResponse(
                content=deviations,
                headers={"Content-Disposition": "attachment; filename=deviations.json"}
            )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ==================== WHATSAPP DATA ENDPOINTS ====================
@app.get("/api/whatsapp-data")
async def list_whatsapp_data(
    plant_id: Optional[int] = Query(None),
    date: Optional[str] = Query(None),
    status: Optional[str] = Query(None),
    skip: int = Query(0, ge=0),
    limit: int = Query(100, ge=1, le=1000),
    db: Session = Depends(get_db)
):
    """Get all WhatsApp data entries"""
    try:
        # Parse date with error handling - frontend sends YYYY-MM-DD
        parsed_date = None
        if date:
            try:
                # Try multiple date formats
                try:
                    parsed_date = datetime.strptime(date, "%Y-%m-%d").date()
                except ValueError:
                    try:
                        parsed_date = datetime.strptime(date, "%d-%m-%Y").date()
                    except ValueError:
                        try:
                            parsed_date = datetime.strptime(date, "%d/%m/%Y").date()
                        except ValueError:
                            pass  # Keep parsed_date as None if all formats fail
            except Exception:
                pass  # Keep parsed_date as None on any error
        
        whatsapp_data = get_whatsapp_data(db, skip=skip, limit=limit, plant_id=plant_id, date=parsed_date, status=status)
        # Return in format expected by frontend: { data: [...], total: X }
        return {"data": whatsapp_data, "total": len(whatsapp_data)}
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid date format. Use YYYY-MM-DD")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/whatsapp-instant")
async def get_whatsapp_instant_data(
    plant_id: Optional[str] = Query(None, min_length=1),
    since: Optional[str] = Query(None)
):
    """Get latest WhatsApp instant data from DynamoDB (single plant or updates feed)."""
    updates_mode = since is not None and str(since).strip() != ""
    single_mode = plant_id is not None and str(plant_id).strip() != ""
    if not updates_mode and not single_mode:
        return {"data": None}

    def _warn(exc: Exception) -> None:
        try:
            print(f"[whatsapp-instant] {type(exc).__name__}: {exc}")
        except Exception:
            pass

    try:
        table = _get_dynamodb_table("WHATSAPP_INSTANT_TABLE")
    except Exception as exc:
        _warn(exc if isinstance(exc, Exception) else Exception(str(exc)))
        return [] if updates_mode else {"data": None}

    if single_mode:
        requested_key = str(plant_id)
        item = _find_ddb_item_by_plant_id(table, requested_key)
        requested_site = str(requested_key or "").strip().upper()
        if not item and requested_site:
            # New schema: plant_id is constant (e.g. "vedanjay"), and per-site state is stored in site_states map.
            # If a site code is provided, try fetching the root item and then selecting that site.
            item = _find_ddb_item_by_plant_id(table, "vedanjay")
        if not item:
            return {"data": None}

        live_state: Dict[str, Any] = {}
        if requested_site:
            live_state = _extract_whatsapp_site_state(item, requested_site)
        message = str(live_state.get("last_message") or "").strip()
        if not message:
            message = str(
                item.get("last_message")
                or item.get("lastMessage")
                or item.get("message")
                or ""
            )
        parsed = _parse_whatsapp_message(message)
        # Ensure parsed plant status is available to the UI even if the DynamoDB record
        # stores the status only at the root item.
        if "plantStatus" not in parsed:
            fallback_status = str(live_state.get("plant_status") or item.get("plant_status") or item.get("status") or "").strip()
            if fallback_status:
                parsed["plantStatus"] = fallback_status.upper()
        if "curtailmentCapacity" not in parsed:
            capacity = live_state.get("curtailment_capacity") if live_state else item.get("curtailment_capacity")
            if capacity is not None:
                parsed["curtailmentCapacity"] = capacity
        if "curtailmentStatus" not in parsed:
            status_value = str(
                live_state.get("plant_status") if live_state else (item.get("plant_status") or item.get("status") or "")
            ).strip().lower()
            if status_value:
                parsed["curtailmentStatus"] = status_value == "curtailment"
        if "remarks" not in parsed and message:
            parsed["remarks"] = message
        windows = _load_whatsapp_windows_for_site(item.get("plant_id") or "vedanjay", requested_site, limit=50) if requested_site else []

        status_value = str(
            live_state.get("plant_status")
            or item.get("plant_status")
            or item.get("status")
            or ""
        ).strip()
        updated_value = str(
            live_state.get("updated_at")
            or item.get("updated_at")
            or item.get("updatedAt")
            or ""
        ).strip()
        return {
            "plantId": item.get("plant_id") or plant_id,
            "site": requested_site or item.get("site") or "",
            "message": message,
            "status": status_value,
            "updatedAt": updated_value,
            "parsed": parsed
            , "live": live_state
            , "windows": windows
        }

    if updates_mode:
        since_ms = _parse_ddb_timestamp(since) or 0
        results = []
        last_evaluated_key = None
        pages = 0
        while pages < 5:
            kwargs = {}
            if last_evaluated_key:
                kwargs["ExclusiveStartKey"] = last_evaluated_key
            try:
                response = table.scan(**kwargs)
            except Exception as exc:
                _warn(exc)
                return []
            raw_items = response.get("Items") or []
            for raw in raw_items:
                item = _normalize_ddb_item(raw)
                payload = _whatsapp_item_to_payload(item)
                ts = payload.get("timestamp_ms") or 0
                if ts > since_ms:
                    results.append(payload)
            last_evaluated_key = response.get("LastEvaluatedKey")
            pages += 1
            if not last_evaluated_key:
                break
        results.sort(key=lambda r: r.get("timestamp_ms") or 0)
        return results

    return {"data": None}


@app.get("/api/whatsapp-data/{whatsapp_id}")
async def get_whatsapp_data_by_id_endpoint(
    whatsapp_id: int,
    db: Session = Depends(get_db)
):
    """Get a single WhatsApp data entry"""
    try:
        whatsapp_data = get_whatsapp_data_by_id(db, whatsapp_id)
        if not whatsapp_data:
            raise HTTPException(status_code=404, detail="WhatsApp data not found")
        return whatsapp_data
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/whatsapp-data")
async def create_whatsapp_data_endpoint(
    whatsapp_data: WhatsAppDataCreate,
    db: Session = Depends(get_db)
):
    """Create a new WhatsApp data entry"""
    try:
        created = create_whatsapp_data(db, whatsapp_data)
        _mirror_whatsapp_sql_row_to_dynamodb(created)
        if bool(getattr(created, "curtailmentStatus", False)):
            reason = str(getattr(created, "curtailmentReason", "") or "").strip() or "Curtailment signal received"
            _create_operator_notification(
                db,
                plant_id=int(getattr(created, "plantId", 0) or 0),
                plant_name=str(getattr(created, "plantName", "") or ""),
                notification_type="Curtailment Alert",
                title="Curtailment message received",
                message=f"{created.plantName}: {reason}",
                priority="HIGH",
                action_required=True,
            )
        # Return the created record in a format the frontend expects
        return {
            "id": created.id,
            "plantId": created.plantId,
            "plantName": created.plantName,
            "state": created.state,
            "date": created.date,
            "time": created.time,
            "currentGeneration": created.currentGeneration,
            "expectedTrend": created.expectedTrend,
            "curtailmentStatus": created.curtailmentStatus,
            "curtailmentReason": created.curtailmentReason,
            "weatherCondition": created.weatherCondition,
            "inverterAvailability": created.inverterAvailability,
            "remarks": created.remarks,
            "status": created.status,
            "createdAt": created.createdAt.isoformat() if created.createdAt else datetime.now().isoformat()
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.put("/api/whatsapp-data/{whatsapp_id}")
async def update_whatsapp_data_endpoint(
    whatsapp_id: int,
    whatsapp_data: WhatsAppDataUpdate,
    db: Session = Depends(get_db)
):
    """Update a WhatsApp data entry"""
    try:
        updated = update_whatsapp_data(db, whatsapp_id, whatsapp_data)
        if not updated:
            raise HTTPException(status_code=404, detail="WhatsApp data not found")
        return updated
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.delete("/api/whatsapp-data/{whatsapp_id}")
async def delete_whatsapp_data_endpoint(
    whatsapp_id: int,
    db: Session = Depends(get_db)
):
    """Delete a WhatsApp data entry"""
    try:
        success = delete_whatsapp_data(db, whatsapp_id)
        if not success:
            raise HTTPException(status_code=404, detail="WhatsApp data not found")
        return {"message": "WhatsApp data deleted successfully"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ==================== METER DATA ENDPOINTS ====================
@app.get("/api/meter-data")
async def list_meter_data(
    plant_id: Optional[int] = Query(None),
    data_date: Optional[str] = Query(None),
    skip: int = Query(0, ge=0),
    limit: int = Query(100, ge=1, le=1000),
    db: Session = Depends(get_db)
):
    """Get all meter data entries"""
    try:
        parsed_date = datetime.strptime(data_date, "%Y-%m-%d").date() if data_date else None
        meter_data = get_meter_data(db, skip=skip, limit=limit, plant_id=plant_id, data_date=parsed_date)
        # Parse blockData JSON string back to dict for response
        result = []
        for md in meter_data:
            md_dict = {
                "id": md.id,
                "plantId": md.plantId,
                "plantName": md.plantName,
                "dataDate": md.dataDate,
                "blockData": json.loads(md.blockData) if isinstance(md.blockData, str) else md.blockData,
                "source": md.source,
                "lastReading": md.lastReading,
                "dataPoints": md.dataPoints,
                "delay": md.delay,
                "createdAt": md.createdAt,
                "updatedAt": md.updatedAt
            }
            result.append(md_dict)
        return {"data": result, "total": len(result)}
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid date format. Use YYYY-MM-DD")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/meter-data/{meter_id}")
async def get_meter_data_by_id_endpoint(
    meter_id: int,
    db: Session = Depends(get_db)
):
    """Get a single meter data entry"""
    try:
        meter_data = get_meter_data_by_id(db, meter_id)
        if not meter_data:
            raise HTTPException(status_code=404, detail="Meter data not found")
        # Parse blockData JSON string back to dict
        result = {
            "id": meter_data.id,
            "plantId": meter_data.plantId,
            "plantName": meter_data.plantName,
            "dataDate": meter_data.dataDate,
            "blockData": json.loads(meter_data.blockData) if isinstance(meter_data.blockData, str) else meter_data.blockData,
            "source": meter_data.source,
            "lastReading": meter_data.lastReading,
            "dataPoints": meter_data.dataPoints,
            "delay": meter_data.delay,
            "createdAt": meter_data.createdAt,
            "updatedAt": meter_data.updatedAt
        }
        return result
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/meter-data/plant/{plant_id}/latest")
async def get_latest_meter_data_endpoint(
    plant_id: int,
    db: Session = Depends(get_db)
):
    """Get the latest meter data for a plant"""
    try:
        meter_data = get_latest_meter_data(db, plant_id)
        if not meter_data:
            raise HTTPException(status_code=404, detail="Meter data not found")
        # Parse blockData JSON string back to dict
        result = {
            "id": meter_data.id,
            "plantId": meter_data.plantId,
            "plantName": meter_data.plantName,
            "dataDate": meter_data.dataDate,
            "blockData": json.loads(meter_data.blockData) if isinstance(meter_data.blockData, str) else meter_data.blockData,
            "source": meter_data.source,
            "lastReading": meter_data.lastReading,
            "dataPoints": meter_data.dataPoints,
            "delay": meter_data.delay,
            "createdAt": meter_data.createdAt,
            "updatedAt": meter_data.updatedAt
        }
        return result
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/meter-data")
async def create_meter_data_endpoint(
    meter_data: MeterDataCreate,
    db: Session = Depends(get_db)
):
    """Create a new meter data entry"""
    try:
        return create_meter_data(db, meter_data)
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/api/meter-data/upload-csv")
async def upload_meter_data_csv(
    file: UploadFile = File(...),
    plant_id: int = Query(...),
    plant_name: str = Query(...),
    data_date: str = Query(...),
    db: Session = Depends(get_db)
):
    """Upload meter data from CSV file"""
    try:
        # Parse date
        parsed_date = datetime.strptime(data_date, "%Y-%m-%d").date()
        
        # Read CSV file
        contents = await file.read()
        csv_content = contents.decode('utf-8')
        csv_reader = csv.DictReader(io.StringIO(csv_content))
        
        # Parse CSV and create block data
        block_data = {}
        rows = list(csv_reader)
        
        # Expected CSV format: Time Block, Generation (MW), etc.
        for idx, row in enumerate(rows):
            # Try to find time or block number
            time_key = None
            gen_key = None
            
            for key in row.keys():
                key_lower = key.lower()
                if 'time' in key_lower or 'block' in key_lower or 'blk' in key_lower:
                    time_key = key
                if 'generation' in key_lower or 'mw' in key_lower or 'actual' in key_lower:
                    gen_key = key
            
            if time_key and gen_key:
                block_num = idx + 1
                time_str = row[time_key].strip()
                gen_value = float(row[gen_key]) if row[gen_key] else 0.0
                block_data[f"block_{block_num}"] = {
                    "block": block_num,
                    "time": time_str,
                    "generation": gen_value
                }
            elif gen_key:
                # If no time key, use index
                block_num = idx + 1
                gen_value = float(row[gen_key]) if row[gen_key] else 0.0
                block_data[f"block_{block_num}"] = {
                    "block": block_num,
                    "time": f"{(block_num-1)*15:02d}:00",
                    "generation": gen_value
                }
        
        # Create meter data entry
        meter_data_create = MeterDataCreate(
            plantId=plant_id,
            plantName=plant_name,
            dataDate=parsed_date,
            blockData=block_data,
            source="Manual Upload",
            dataPoints=len(block_data),
            lastReading=datetime.now()
        )
        
        created = create_meter_data(db, meter_data_create)
        return {
            "message": "Meter data uploaded successfully",
            "data": {
                "id": created.id,
                "dataPoints": created.dataPoints,
                "blocks": len(block_data)
            }
        }
    except ValueError as e:
        raise HTTPException(status_code=400, detail=f"Invalid date format or CSV structure: {str(e)}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.put("/api/meter-data/{meter_id}")
async def update_meter_data_endpoint(
    meter_id: int,
    meter_data: MeterDataUpdate,
    db: Session = Depends(get_db)
):
    """Update a meter data entry"""
    try:
        updated = update_meter_data(db, meter_id, meter_data)
        if not updated:
            raise HTTPException(status_code=404, detail="Meter data not found")
        return updated
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.delete("/api/meter-data/{meter_id}")
async def delete_meter_data_endpoint(
    meter_id: int,
    db: Session = Depends(get_db)
):
    """Delete a meter data entry"""
    try:
        success = delete_meter_data(db, meter_id)
        if not success:
            raise HTTPException(status_code=404, detail="Meter data not found")
        return {"message": "Meter data deleted successfully"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/meter-data/plant/{plant_id}/data")
async def get_meter_data_points_for_plant(
    plant_id: int,
    date: str = Query(..., description="Date in YYYY-MM-DD format"),
    db: Session = Depends(get_db)
):
    """Get meter data points for a specific plant and date (96 time blocks)"""
    try:
        # Try to get real meter data first
        meter_data = get_latest_meter_data(db, plant_id)
        if meter_data:
            # Parse the blockData and return in expected format
            block_data = meter_data.blockData
            if isinstance(block_data, str):
                block_data = json.loads(block_data)

            # Convert to dataPoints format
            data_points = []
            for block_key, block_info in block_data.items():
                if isinstance(block_info, dict):
                    time_parts = block_info.get("time", "00:00").split(":")
                    hour = int(time_parts[0]) if len(time_parts) > 0 else 0
                    minute = int(time_parts[1]) if len(time_parts) > 1 else 0

                    data_points.append({
                        "time": block_info.get("time", "00:00"),
                        "hour": hour,
                        "minute": minute,
                        "generation": block_info.get("generation", 0),
                        "availableCapacity": block_info.get("availableCapacity", 95),
                        "availability": block_info.get("availability", 95)
                    })

            return {
                "date": meter_data.dataDate.isoformat() if meter_data.dataDate else date,
                "dataPoints": data_points,
                "totalGeneration": sum(d["generation"] for d in data_points),
                "lastReading": meter_data.lastReading.isoformat() if meter_data.lastReading else datetime.now().isoformat(),
                "source": meter_data.source or "SCADA",
                "status": "Live"
            }

        raise HTTPException(status_code=404, detail="Meter data not found")

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ==================== HEALTH CHECK ====================
@app.get("/api/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "ok", "message": "Server is running"}


@app.get("/")
async def root():
    """Root endpoint with API information"""
    return {
        "message": "Renewable Energy Dashboard API",
        "version": "1.0.0",
        "endpoints": {
            "dashboard": "/api/dashboard/stats",
            "plants": "/api/plants",
            "schedules": "/api/schedules",
            "forecasts": "/api/forecasts",
            "weather": "/api/weather",
            "deviations": "/api/deviations",
            "reports": "/api/reports",
            "templates": "/api/templates",
            "template_transform": "/api/template-transform/preview",
            "template_transform_source_files": "/api/template-transform/source-files",
            "template_transform_download": "/api/template-transform/download/{run_id}",
            "whatsapp-data": "/api/whatsapp-data",
            "meter-data": "/api/meter-data",
            "health": "/api/health"
        }
    }




# ==================== SCHEDULE READINESS ENDPOINTS ====================
@app.get("/api/schedule-readiness")
async def list_schedule_readiness(
    status: Optional[str] = Query(None, description="Filter by status: READY, PENDING, NO_ACTION"),
    db: Session = Depends(get_db)
):
    """List all site schedule readiness statuses with summary"""
    try:
        summary = get_schedule_readiness_summary(db)
        return summary
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/schedule-readiness/summary")
async def get_schedule_readiness_summary_endpoint(
    db: Session = Depends(get_db)
):
    """Get quick summary of all plant readiness statuses"""
    try:
        summary = get_schedule_readiness_summary(db)
        return {
            "total": summary["total_plants"],
            "ready": summary["ready_count"],
            "pending": summary["pending_count"],
            "no_action": summary["no_action_count"]
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ==================== SCHEDULE NOTIFICATIONS ENDPOINTS ====================
@app.get("/api/schedule-readiness/notifications")
async def get_notifications(
    unread_only: bool = Query(False, description="Show only unread notifications"),
    plant_id: Optional[int] = Query(None, description="Filter by plant ID"),
    limit: int = Query(50, ge=1, le=100),
    db: Session = Depends(get_db)
):
    """Get pending notifications"""
    try:
        notifications = get_schedule_notifications(db, plant_id=plant_id, unread_only=unread_only, limit=limit)
        unread_count = sum(1 for n in notifications if not bool(getattr(cast(Any, n), "read", False)))

        return {
            "notifications": notifications,
            "total": len(notifications),
            "unread_count": unread_count
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.put("/api/schedule-readiness/notifications/{notification_id}/read")
async def mark_notification_read_endpoint(
    notification_id: int,
    db: Session = Depends(get_db)
):
    """Mark a notification as read"""
    try:
        notification = mark_notification_read(db, notification_id)
        if not notification:
            raise HTTPException(status_code=404, detail="Notification not found")
        return {
            "success": True,
            "message": "Notification marked as read",
            "notification_id": notification_id
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ==================== FROZEN SCHEDULE AUTO-PERSIST ENDPOINT ====================
@app.post("/api/frozen-schedule/persist")
async def persist_frozen_schedule_artifacts(
    request: FrozenSchedulePersistRequest
):
    """Persist auto-frozen schedule CSV + audit log to S3 using naming convention."""
    try:
        plant_code = _normalize_plant_code(str(request.plant_code or "").strip())
        schedule_date = str(request.schedule_date or "").strip()
        if not plant_code or not re.fullmatch(r"\d{4}-\d{2}-\d{2}", schedule_date):
            raise HTTPException(status_code=400, detail="Invalid plant_code or schedule_date")

        block_value = max(1, min(96, int(request.block)))
        block_text = f"{block_value:02d}"
        freeze_time = str(request.freeze_time or datetime.utcnow().isoformat()).strip()

        frozen_folder = _special_s3_plant_folder(plant_code)
        frozen_prefix = f"frozenschedules/vedanjay/{frozen_folder}/{schedule_date}/"
        # Keep two overwriteable frozen CSV artifacts per plant/date.
        edited_schedule_key = f"{frozen_prefix}edited_frozen.csv"
        system_schedule_key = f"{frozen_prefix}system_frozen.csv"
        log_key = f"{frozen_prefix}{frozen_folder}_frozen.log"

        bucket = _derive_s3_bucket_name()
        if not bucket:
            raise HTTPException(status_code=500, detail="S3 bucket not configured")
        region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"

        try:
            import boto3  # type: ignore
            s3 = boto3.client("s3", region_name=region)
        except Exception as exc:
            raise HTTPException(status_code=500, detail=f"boto3 client unavailable: {exc}") from exc

        # Ensure marker exists for S3 console visibility.
        try:
            s3.put_object(Bucket=bucket, Key=frozen_prefix)
        except Exception:
            pass

        status_value = str(request.status or "").strip().lower()
        storage_mode = "s3"
        effective_bucket = bucket
        error_msg = None
        local_path = ""

        # Edited frozen is driven by user/manual confirmation payloads.
        edited_csv_text = request.edited_schedule_csv or request.schedule_csv
        # System frozen must be written only by the auto-upload pipeline.
        system_csv_text = request.system_schedule_csv

        if status_value in {"uploaded", "frozen"} and (edited_csv_text or (system_csv_text and request.write_system_frozen)):
            # Remove legacy per-block frozen files to keep only the consolidated artifact (Rule 3).
            try:
                resp = s3.list_objects_v2(Bucket=bucket, Prefix=frozen_prefix)
                legacy = [
                    it["Key"] for it in resp.get("Contents", [])
                    if re.search(r"schedule_free(?:z|ze)_from_\d+\.(?:csv|log)$", it.get("Key", ""), re.I)
                ]
                if legacy:
                    s3.delete_objects(Bucket=bucket, Delete={"Objects": [{"Key": k} for k in legacy]})
            except Exception:
                pass

            try:
                if edited_csv_text:
                    s3.put_object(
                        Bucket=bucket,
                        Key=edited_schedule_key,
                        Body=edited_csv_text.encode("utf-8"),
                        ContentType="text/csv",
                    )
                if system_csv_text and request.write_system_frozen:
                    s3.put_object(
                        Bucket=bucket,
                        Key=system_schedule_key,
                        Body=system_csv_text.encode("utf-8"),
                        ContentType="text/csv",
                    )
            except Exception as e:
                # Fallback to local storage if S3 fails
                storage_mode = "local"
                effective_bucket = "LOCAL_FALLBACK"
                error_msg = str(e)
                local_dir = os.path.join(READINESS_UPLOAD_LOCAL_DIR, "frozen", plant_code, schedule_date)
                os.makedirs(local_dir, exist_ok=True)
                if edited_csv_text:
                    local_path = os.path.join(local_dir, "edited_frozen.csv")
                    with open(local_path, "w", encoding="utf-8") as f:
                        f.write(edited_csv_text)
                if system_csv_text and request.write_system_frozen:
                    system_local_path = os.path.join(local_dir, "system_frozen.csv")
                    with open(system_local_path, "w", encoding="utf-8") as f:
                        f.write(system_csv_text)

        log_payload = {
            "plant_code": plant_code,
            "schedule_date": schedule_date,
            "block": block_value,
            "status": request.status,
            "source_schedule_key": request.source_schedule_key,
            "freeze_time": freeze_time,
            "reason": request.reason or "",
            "summary": request.summary or {},
            "stored_schedule_key": edited_schedule_key if (edited_csv_text and storage_mode == "s3" and status_value in {"uploaded", "frozen"}) else local_path,
            "stored_system_schedule_key": system_schedule_key if (system_csv_text and request.write_system_frozen and storage_mode == "s3" and status_value in {"uploaded", "frozen"}) else "",
            "stored_log_key": log_key,
            "created_at": datetime.utcnow().isoformat(),
            "storage_mode": storage_mode,
            "error": error_msg
        }
        
        try:
            s3.put_object(
                Bucket=bucket,
                Key=log_key,
                Body=json.dumps(log_payload, ensure_ascii=False, indent=2).encode("utf-8"),
                ContentType="application/json",
            )
        except Exception:
            # If even the log fails to upload to S3, we at least have the local fallback for the CSV.
            if storage_mode == "local" and local_path:
                with open(local_path.replace(".csv", ".log.json"), "w", encoding="utf-8") as f:
                    json.dump(log_payload, f, indent=2)

        return {
            "success": True,
            "bucket": effective_bucket,
            "schedule_key": edited_schedule_key if (edited_csv_text and storage_mode == "s3" and status_value in {"uploaded", "frozen"}) else local_path,
            "system_schedule_key": system_schedule_key if (system_csv_text and request.write_system_frozen and storage_mode == "s3" and status_value in {"uploaded", "frozen"}) else "",
            "log_key": log_key,
            "status": request.status,
            "freeze_time": freeze_time,
            "storage_mode": storage_mode,
            "error": error_msg
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


def _frozen_exclusions_key(plant_code: str, schedule_date: str) -> str:
    plant_code = str(plant_code or "").strip().upper()
    schedule_date = str(schedule_date or "").strip()
    plant_folder = _special_s3_plant_folder(plant_code)
    return f"frozenschedules/vedanjay/{plant_folder}/{schedule_date}/excluded_schedules.json"


def _normalize_s3_key(value: str) -> str:
    raw = str(value or "").strip()
    if not raw:
        return ""
    if raw.startswith("s3://"):
        return re.sub(r"^s3://[^/]+/?", "", raw)
    if re.match(r"^https?://", raw, flags=re.IGNORECASE):
        match = re.match(r"^https?://[^/]+/(.+)$", raw, flags=re.IGNORECASE)
        return match.group(1) if match and match.group(1) else raw
    return raw.lstrip("/")


def _is_allowed_schedule_key(key: str) -> bool:
    text = _normalize_s3_key(key)
    if not text:
        return False
    if len(text) > 1400:
        return False
    return bool(re.search(r"schedule_from_\d+(?:[_-][A-Za-z0-9]+)*\.csv$", text, flags=re.IGNORECASE))


@app.get("/api/frozen-schedule/exclusions")
async def list_frozen_schedule_exclusions(
    plant_code: str = Query(..., min_length=1, max_length=32),
    schedule_date: str = Query(..., min_length=10, max_length=10),
):
    """List schedule keys that should be excluded from frozen recomputation for plant/date."""
    plant_code = str(plant_code or "").strip().upper()
    schedule_date = str(schedule_date or "").strip()
    if not re.fullmatch(r"\d{4}-\d{2}-\d{2}", schedule_date):
        raise HTTPException(status_code=400, detail="Invalid schedule_date (expected YYYY-MM-DD)")

    bucket = _derive_s3_bucket_name()
    if not bucket:
        raise HTTPException(status_code=500, detail="S3 bucket not configured")
    region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"

    key = _frozen_exclusions_key(plant_code, schedule_date)
    try:
        import boto3  # type: ignore
        s3 = boto3.client("s3", region_name=region)
        try:
            obj = s3.get_object(Bucket=bucket, Key=key)
            body = obj.get("Body")
            text = body.read().decode("utf-8", errors="replace") if body is not None else ""
            parsed = json.loads(text) if text else []
        except Exception:
            parsed = []
        if isinstance(parsed, dict):
            parsed = parsed.get("items") or []
        if not isinstance(parsed, list):
            parsed = []
        normalized = []
        for item in parsed:
            item_key = _normalize_s3_key(str(item or ""))
            if _is_allowed_schedule_key(item_key):
                normalized.append(item_key)
        normalized = sorted(list(dict.fromkeys(normalized)))
        return {
            "plant_code": plant_code,
            "schedule_date": schedule_date,
            "items": normalized,
            "bucket": bucket,
            "key": key,
        }
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.post("/api/frozen-schedule/exclusions/add")
async def add_frozen_schedule_exclusion(
    request: FrozenScheduleExclusionRequest,
):
    """Add a schedule key to the exclusion list (so it won't be applied in frozen schedule)."""
    plant_code = str(request.plant_code or "").strip().upper()
    schedule_date = str(request.schedule_date or "").strip()
    source_key = _normalize_s3_key(str(request.source_schedule_key or ""))
    if not plant_code:
        raise HTTPException(status_code=400, detail="plant_code is required")
    if not re.fullmatch(r"\d{4}-\d{2}-\d{2}", schedule_date):
        raise HTTPException(status_code=400, detail="Invalid schedule_date (expected YYYY-MM-DD)")
    if not _is_allowed_schedule_key(source_key):
        raise HTTPException(status_code=400, detail="source_schedule_key must be schedule_from_XX.csv")

    bucket = _derive_s3_bucket_name()
    if not bucket:
        raise HTTPException(status_code=500, detail="S3 bucket not configured")
    region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"
    key = _frozen_exclusions_key(plant_code, schedule_date)
    frozen_folder = _special_s3_plant_folder(plant_code)
    frozen_prefix = f"frozenschedules/vedanjay/{frozen_folder}/{schedule_date}/"

    try:
        import boto3  # type: ignore
        s3 = boto3.client("s3", region_name=region)
        # Ensure frozen/ prefix exists in S3 console.
        try:
            s3.put_object(Bucket=bucket, Key=frozen_prefix)
        except Exception:
            # Non-fatal marker creation
            pass

        try:
            obj = s3.get_object(Bucket=bucket, Key=key)
            body = obj.get("Body")
            text = body.read().decode("utf-8", errors="replace") if body is not None else ""
            parsed = json.loads(text) if text else []
        except Exception:
            parsed = []

        if not isinstance(parsed, list):
            parsed = []

        normalized = [_normalize_s3_key(str(item or "")) for item in parsed]
        normalized = [k for k in normalized if _is_allowed_schedule_key(k)]
        normalized.append(source_key)
        normalized = sorted(list(dict.fromkeys(normalized)))

        payload = {
            "items": normalized,
            "updated_at": datetime.utcnow().isoformat(),
            "requested_by": str(request.requested_by or "").strip(),
        }
        s3.put_object(
            Bucket=bucket,
            Key=key,
            Body=json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8"),
            ContentType="application/json",
        )

        return {
            "success": True,
            "plant_code": plant_code,
            "schedule_date": schedule_date,
            "excluded_key": source_key,
            "items": normalized,
            "bucket": bucket,
            "key": key,
        }
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.post("/api/frozen-schedule/exclusions/remove")
async def remove_frozen_schedule_exclusion(
    request: FrozenScheduleExclusionRequest,
):
    """Remove a schedule key from the exclusion list."""
    plant_code = str(request.plant_code or "").strip().upper()
    schedule_date = str(request.schedule_date or "").strip()
    source_key = _normalize_s3_key(str(request.source_schedule_key or ""))
    if not plant_code:
        raise HTTPException(status_code=400, detail="plant_code is required")
    if not re.fullmatch(r"\d{4}-\d{2}-\d{2}", schedule_date):
        raise HTTPException(status_code=400, detail="Invalid schedule_date (expected YYYY-MM-DD)")
    if not _is_allowed_schedule_key(source_key):
        raise HTTPException(status_code=400, detail="source_schedule_key must be schedule_from_XX.csv")

    bucket = _derive_s3_bucket_name()
    if not bucket:
        raise HTTPException(status_code=500, detail="S3 bucket not configured")
    region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"
    key = _frozen_exclusions_key(plant_code, schedule_date)

    try:
        import boto3  # type: ignore
        s3 = boto3.client("s3", region_name=region)
        try:
            obj = s3.get_object(Bucket=bucket, Key=key)
            body = obj.get("Body")
            text = body.read().decode("utf-8", errors="replace") if body is not None else ""
            parsed = json.loads(text) if text else []
        except Exception:
            parsed = []

        if not isinstance(parsed, list):
            parsed = []

        normalized = [_normalize_s3_key(str(item or "")) for item in parsed]
        normalized = [k for k in normalized if _is_allowed_schedule_key(k)]
        normalized = [k for k in normalized if k != source_key]
        normalized = sorted(list(dict.fromkeys(normalized)))

        payload = {
            "items": normalized,
            "updated_at": datetime.utcnow().isoformat(),
            "requested_by": str(request.requested_by or "").strip(),
        }
        s3.put_object(
            Bucket=bucket,
            Key=key,
            Body=json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8"),
            ContentType="application/json",
        )

        return {
            "success": True,
            "plant_code": plant_code,
            "schedule_date": schedule_date,
            "removed_key": source_key,
            "items": normalized,
            "bucket": bucket,
            "key": key,
        }
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.post("/api/frozen-schedule/migrate-to-frozen-folder")
async def migrate_frozen_artifacts_to_frozen_folder(
    schedule_date: date = Query(..., description="Date in YYYY-MM-DD format"),
    plant_code: Optional[str] = Query(None, description="Optional single plant code, e.g. KOTHAGUDEM"),
    overwrite: bool = Query(False, description="Overwrite if target already exists"),
    delete_source: bool = Query(True, description="Delete old source file after copy"),
    dry_run: bool = Query(False, description="Preview only, do not copy/delete"),
):
    """Move old freeze files from outputs/<date>/ to outputs/<date>/frozen/."""
    try:
        bucket = _derive_s3_bucket_name()
        if not bucket:
            raise HTTPException(status_code=500, detail="S3 bucket not configured")
        region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"
        try:
            import boto3  # type: ignore
            s3 = boto3.client("s3", region_name=region)
        except Exception as exc:
            raise HTTPException(status_code=500, detail=f"boto3 client unavailable: {exc}") from exc

        target_date = schedule_date.isoformat()
        plants = [str(plant_code or "").strip().upper()] if plant_code else [
            "ANJANGAON", "ANDAD", "BALAKWADA", "BAMKHAL", "BHUPALPALLY", "CME", "GSNP", "GUGARIYAKHEDI", "KASIPET", "KILAJ", "KOTHAGUDEM", "NANDGAON", "OSEPL", "SAWDA", "SIRMOUR", "ZETRIC"
        ]
        plants = [p for p in plants if p]

        totals = {
            "plants_scanned": len(plants),
            "candidates": 0,
            "copied": 0,
            "deleted_sources": 0,
            "skipped_exists": 0,
            "errors": 0,
        }
        moved: List[Dict[str, str]] = []

        for plant in plants:
            output_prefix = f"generated/vedanjay/{plant}/outputs/{target_date}/"
            frozen_prefix = f"{output_prefix}frozen/"
            if not dry_run:
                # Create marker so "frozen" prefix is visible in S3 console.
                try:
                    s3.put_object(Bucket=bucket, Key=frozen_prefix)
                except Exception:
                    pass
            keys = _list_s3_keys_safe(s3, bucket, output_prefix)
            key_set = set(keys)

            for source_key in keys:
                key_text = str(source_key or "").strip()
                if not key_text or "/frozen/" in key_text:
                    continue
                match = re.search(r"schedule_free(?:z|ze)_from_(\d+)\.(csv|log)$", key_text, flags=re.IGNORECASE)
                if not match:
                    continue

                block = int(match.group(1))
                ext = str(match.group(2) or "").lower()
                target_key = f"{frozen_prefix}schedule_freeze_from_{block:02d}.{ext}"
                totals["candidates"] += 1

                if (target_key in key_set) and (not overwrite):
                    totals["skipped_exists"] += 1
                    continue

                try:
                    if not dry_run:
                        s3.copy_object(
                            Bucket=bucket,
                            CopySource={"Bucket": bucket, "Key": key_text},
                            Key=target_key,
                        )
                        totals["copied"] += 1
                        if delete_source and key_text != target_key:
                            s3.delete_object(Bucket=bucket, Key=key_text)
                            totals["deleted_sources"] += 1
                    moved.append({"from": key_text, "to": target_key})
                except Exception:
                    totals["errors"] += 1

        return {
            "success": True,
            "schedule_date": target_date,
            "dry_run": dry_run,
            "overwrite": overwrite,
            "delete_source": delete_source,
            "totals": totals,
            "moved": moved,
        }
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Migration failed: {exc}") from exc


# ==================== SCHEDULE TRIGGERS ENDPOINTS ====================
@app.get("/api/schedule-readiness/triggers")
async def get_schedule_triggers_endpoint(
    plant_id: Optional[int] = Query(None),
    trigger_type: Optional[str] = Query(None),
    processed: Optional[bool] = Query(None),
    limit: int = Query(50, ge=1, le=100),
    db: Session = Depends(get_db)
):
    """Get schedule trigger records"""
    try:
        triggers = get_schedule_triggers(db, plant_id=plant_id, trigger_type=trigger_type, processed=processed, limit=limit)
        return {"triggers": triggers, "total": len(triggers)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/schedule-readiness/{plant_id:int}")
async def get_plant_readiness(
    plant_id: int,
    db: Session = Depends(get_db)
):
    """Get specific plant's schedule readiness status"""
    try:
        readiness = get_schedule_readiness_by_plant(db, plant_id)
        if not readiness:
            # Get plant info to create readiness record
            plant = get_plant(db, plant_id)
            if not plant:
                raise HTTPException(status_code=404, detail="Plant not found")
            # Create new readiness record
            readiness_data = {
                "plant_id": plant_id,
                "plant_name": plant.name,
                "status": "NO_ACTION",
                "schedule_date": date.today()
            }
            readiness = create_schedule_readiness(db, readiness_data)
        return readiness
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/schedule-readiness/{plant_id:int}/trigger")
async def trigger_schedule_revision(
    plant_id: int,
    reason: str = Query(..., description="Reason for revision"),
    db: Session = Depends(get_db)
):
    """Manually trigger schedule revision for a plant"""
    try:
        # Check plant exists
        plant = get_plant(db, plant_id)
        if not plant:
            raise HTTPException(status_code=404, detail="Plant not found")
        
        # Use schedule service to trigger
        from services.schedule_service import ScheduleReadinessService
        service = ScheduleReadinessService(db)
        readiness = service.trigger_manual_revision(plant_id, reason)
        
        return {
            "success": True,
            "message": f"Schedule revision triggered for {plant.name}",
            "plant_id": plant_id,
            "status": readiness.status,
            "trigger_reason": reason
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/schedule-readiness/{plant_id:int}/continue")
async def continue_existing_schedule(
    plant_id: int,
    db: Session = Depends(get_db)
):
    """Continue with existing (day-ahead) schedule - clears triggers"""
    try:
        # Check plant exists
        plant = get_plant(db, plant_id)
        if not plant:
            raise HTTPException(status_code=404, detail="Plant not found")
        
        # Use schedule service to continue
        from services.schedule_service import ScheduleReadinessService
        service = ScheduleReadinessService(db)
        readiness = service.continue_existing_schedule(plant_id)
        
        return {
            "success": True,
            "message": f"Continuing existing schedule for {plant.name}",
            "plant_id": plant_id,
            "status": readiness.status if readiness else "NO_ACTION"
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/schedule-readiness/{plant_id:int}/mark-ready")
async def mark_schedule_ready(
    plant_id: int,
    upload_deadline: Optional[str] = Query(None, description="Upload deadline in ISO format"),
    db: Session = Depends(get_db)
):
    """Mark schedule as ready for upload"""
    try:
        # Check plant exists
        plant = get_plant(db, plant_id)
        if not plant:
            raise HTTPException(status_code=404, detail="Plant not found")
        
        # Parse deadline if provided
        deadline = None
        if upload_deadline:
            try:
                deadline = datetime.fromisoformat(upload_deadline)
            except ValueError:
                raise HTTPException(status_code=400, detail="Invalid date format. Use ISO format")
        
        # Use schedule service to mark ready
        from services.schedule_service import ScheduleReadinessService
        service = ScheduleReadinessService(db)
        readiness = service.mark_schedule_ready(plant_id, deadline)
        
        return {
            "success": True,
            "message": f"Schedule marked as ready for {plant.name}",
            "plant_id": plant_id,
            "status": readiness.status,
            "upload_deadline": readiness.upload_deadline.isoformat() if readiness.upload_deadline else None
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/schedule-readiness/check-triggers")
async def check_triggers_and_update_statuses(
    db: Session = Depends(get_db)
):
    """Run trigger check algorithm for all plants"""
    try:
        from services.schedule_service import ScheduleReadinessService
        service = ScheduleReadinessService(db)
        status_counts = service.check_all_plants()
        
        return {
            "success": True,
            "message": "Trigger check completed for all plants",
            "plants_checked": status_counts['READY'] + status_counts['PENDING'] + status_counts['NO_ACTION'],
            "ready_count": status_counts['READY'],
            "pending_count": status_counts['PENDING'],
            "no_action_count": status_counts['NO_ACTION']
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ==================== TEMPLATE TRANSFORM PIPELINE ENDPOINTS ====================
DEFAULT_TEMPLATE_S3_BASE_URL = os.getenv(
    "TEMPLATE_PIPELINE_S3_BASE_URL",
    "https://vedanjay-schedules-test-218708247175.s3.ap-south-1.amazonaws.com"
)
app.include_router(all_plant_penalty_router)
app.include_router(utility_viewer_router)
DEFAULT_TEMPLATE_S3_PREFIXES = os.getenv(
    "TEMPLATE_PIPELINE_S3_PREFIXES",
    "generated/vedanjay/BHUPALPALLY/outputs,generated/vedanjay/ANDAD/outputs,generated/vedanjay/BALAKWADA/outputs,generated/vedanjay/GUGARIYAKHEDI/outputs,generated/vedanjay/NANDGAON/outputs,generated/vedanjay/BAMKHAL/outputs,generated/vedanjay/SAWDA/outputs,generated/vedanjay/multiple_generator/ZTRIC,generated/vedanjay/CME/outputs,generated/vedanjay/GSNP/outputs,generated/vedanjay/KASIPET/outputs,generated/vedanjay/KILAJ/outputs,generated/vedanjay/KOTHAGUDEM/outputs,generated/vedanjay/OSEPL/outputs,generated/vedanjay/SIRMOUR/outputs,raw/vedanjay/BHUPALPALLY,raw/vedanjay/ANDAD,raw/vedanjay/BALAKWADA,raw/vedanjay/GUGARIYAKHEDI,raw/vedanjay/NANDGAON,raw/vedanjay/BAMKHAL,raw/vedanjay/SAWDA,raw/vedanjay/multiple_generator/ZTRIC,raw/vedanjay/CME,raw/vedanjay/GSNP,raw/vedanjay/KASIPET,raw/vedanjay/KILAJ,raw/vedanjay/KOTHAGUDEM,raw/vedanjay/OSEPL,raw/vedanjay/SIRMOUR,raw/GSNP/gsnp,generated/GSNP/gsnp/outputs,raw/Sirmour/sirmour,generated/Sirmour/sirmour/outputs,outputs"
)

DEFAULT_READINESS_UPLOAD_PREFIX = os.getenv(
    "READINESS_UPLOAD_PREFIX",
    "uploads/vedanjay"
).strip().strip("/")

READINESS_UPLOAD_LOCAL_DIR = os.path.join(os.path.dirname(__file__), "uploads", "readiness")
READINESS_UPLOAD_HISTORY_FILE = os.path.join(READINESS_UPLOAD_LOCAL_DIR, "upload_history.json")
_READINESS_UPLOAD_HISTORY_LOCK = Lock()

CHANGE_LOG_LOCAL_DIR = os.path.join(os.path.dirname(__file__), "uploads", "schedule_changes")


def _is_day_ahead_schedule_key(value: str) -> bool:
    text = str(value or "").strip().lower()
    if not text:
        return False
    return (
        "/day-ahead/" in text
        or "/dayahead/" in text
        or "/day_ahead/" in text
        or bool(re.search(r"_da\d*\.csv$", text, re.IGNORECASE))
    )


def _schedule_change_log_s3_key(*, plant_code: str, schedule_date: Any, source_file_key: str = "") -> str:
    suffix = "Day-ahead/" if _is_day_ahead_schedule_key(source_file_key) else ""
    if _normalize_plant_code(plant_code) == "ZETRIC":
        return f"generated/vedanjay/multiple_generator/ZTRIC/{schedule_date}/{suffix}schedule_changes.json"
    return f"generated/vedanjay/{plant_code}/outputs/{schedule_date}/{suffix}schedule_changes.json"


def _schedule_change_log_local_path(*, plant_code: str, schedule_date: Any, source_file_key: str = "") -> str:
    parts = [CHANGE_LOG_LOCAL_DIR, plant_code, str(schedule_date)]
    if _is_day_ahead_schedule_key(source_file_key):
        parts.append("Day-ahead")
    parts.append("schedule_changes.json")
    return os.path.join(*parts)
_CHANGE_LOG_LOCK = Lock()


def _ensure_change_log_dir(path: str) -> None:
    os.makedirs(path, exist_ok=True)


def _load_change_log_local(path: str) -> list:
    _ensure_change_log_dir(os.path.dirname(path))
    if not os.path.exists(path):
        return []
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, list) else []
    except Exception:
        return []


def _save_change_log_local(path: str, rows: list) -> None:
    _ensure_change_log_dir(os.path.dirname(path))
    with open(path, "w", encoding="utf-8") as f:
        json.dump(rows, f, ensure_ascii=False, indent=2, default=str)


def _ensure_readiness_upload_dirs() -> None:
    os.makedirs(READINESS_UPLOAD_LOCAL_DIR, exist_ok=True)


def _load_readiness_upload_history() -> list:
    _ensure_readiness_upload_dirs()
    if not os.path.exists(READINESS_UPLOAD_HISTORY_FILE):
        return []
    try:
        with open(READINESS_UPLOAD_HISTORY_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, list) else []
    except Exception:
        return []


def _save_readiness_upload_history(rows: list) -> None:
    _ensure_readiness_upload_dirs()
    with open(READINESS_UPLOAD_HISTORY_FILE, "w", encoding="utf-8") as f:
        json.dump(rows, f, ensure_ascii=False, indent=2, default=str)


def _append_readiness_upload_history(entry: dict) -> None:
    with _READINESS_UPLOAD_HISTORY_LOCK:
        rows = _load_readiness_upload_history()
        rows.append(entry)
        _save_readiness_upload_history(rows)


def _compute_submit_and_effective_blocks_from_iso(
    uploaded_at_iso: str,
    *,
    plant_code: Any = "",
    block_minutes: int = 15,
    total_blocks: int = 96,
    effective_delay_blocks: Optional[int] = None,
) -> Dict[str, Any]:
    """
    Compute submit/effective blocks for an upload timestamp.

    We interpret uploaded_at as UTC when timezone info is missing (matches frontend logic),
    then convert to IST (UTC+05:30) for block calculation so schedule windows match operator UI.
    """
    text = str(uploaded_at_iso or "").strip()
    if not text:
        return {"submit_block": None, "effective_start_block": None}

    dt: Optional[datetime] = None
    try:
        # Normalize ISO without timezone by assuming UTC.
        normalized = text
        if re.match(r"^\d{4}-\d{2}-\d{2}T\d{2}:\d{2}", normalized) and not re.search(r"[zZ]|[+-]\d{2}:\d{2}$", normalized):
            normalized = f"{normalized}Z"
        dt = datetime.fromisoformat(normalized.replace("Z", "+00:00"))
    except Exception:
        dt = None

    if not dt:
        return {"submit_block": None, "effective_start_block": None}

    try:
        # Ensure UTC-aware, then convert to IST fixed offset (+05:30).
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        ist = timezone(timedelta(hours=5, minutes=30))
        local_dt = dt.astimezone(ist)

        total_minutes = (local_dt.hour * 60) + local_dt.minute
        block_start = (total_minutes // int(block_minutes)) * int(block_minutes)
        submit_block = int(block_start // int(block_minutes)) + 1
        submit_block = max(1, min(int(total_blocks), submit_block))

        delay_blocks = int(effective_delay_blocks) if effective_delay_blocks is not None else _effective_delay_blocks_for_plant(plant_code)
        effective = submit_block + delay_blocks
        effective_start_block = effective if effective <= int(total_blocks) else None

        return {"submit_block": submit_block, "effective_start_block": effective_start_block}
    except Exception:
        return {"submit_block": None, "effective_start_block": None}


def _parse_sldc_template_schedule_map(csv_text: str) -> Dict[int, float]:
    """
    Parse an SLDC template (or generated schedule CSV) into a block->scheduled_mw mapping.

    Mirrors the frontend `parseSldcTemplateScheduleMap` heuristics:
    - Find the header row that contains "block"
    - Prefer "Station Schedule" / "Algo Schedule" columns when available
    - Otherwise fall back to a generic "schedule"/last-column choice
    """
    text = str(csv_text or "")
    if not text.strip():
        return {b: 0.0 for b in range(1, 97)}

    lines = [ln for ln in text.splitlines() if str(ln).strip()]
    if not lines:
        return {b: 0.0 for b in range(1, 97)}

    header_idx = 0
    for i, ln in enumerate(lines[:50]):
        if re.search(r"\bblock\b", ln, flags=re.IGNORECASE):
            header_idx = i
            break

    def _parse_csv_row(line: str) -> List[str]:
        try:
            return next(csv.reader([line]))
        except Exception:
            return [c.strip() for c in str(line).split(",")]

    headers = [h.strip().lstrip("\ufeff") for h in _parse_csv_row(lines[header_idx])]
    rows = [_parse_csv_row(ln) for ln in lines[header_idx + 1 :]]
    if not headers:
        return {b: 0.0 for b in range(1, 97)}

    def _norm(value: str) -> str:
        return re.sub(r"[^a-z0-9]+", "", re.sub(r"[\"']", "", str(value or "").lower()))

    normalized_headers = [_norm(h) for h in headers]

    def _find_col(needles: List[str]) -> int:
        for idx, h in enumerate(normalized_headers):
            for needle in needles:
                if needle in h:
                    return idx
        return -1

    block_idx = _find_col(["block", "blk", "blockno", "blocknumber"])
    station_schedule_idx = _find_col(["stationschedule"])
    schedule_idx = station_schedule_idx if station_schedule_idx != -1 else _find_col(["schedule"])
    forecast_idx = _find_col(["declaredforecast", "forecast"])
    algo_idx = _find_col(["algoschedulemw", "algoschedule", "algo"])
    base_idx = _find_col(["base"])

    value_idx = (
        algo_idx
        if algo_idx != -1
        else schedule_idx
        if schedule_idx != -1
        else base_idx
        if base_idx != -1
        else (forecast_idx if forecast_idx != -1 else max(0, len(headers) - 1))
    )

    def _to_num(value: Any) -> Optional[float]:
        try:
            raw = str(value if value is not None else "").replace(",", "").strip()
            if raw == "":
                return None
            parsed = float(raw)
            return parsed if math.isfinite(parsed) else None
        except Exception:
            return None

    out: Dict[int, float] = {}
    for idx, cols in enumerate(rows):
        try:
            block_raw = cols[block_idx] if block_idx != -1 and block_idx < len(cols) else (cols[0] if cols else "")
            block = int(str(block_raw or "").strip())
        except Exception:
            continue
        if block < 1 or block > 96:
            continue
        value = cols[value_idx] if value_idx < len(cols) else ""
        scheduled = _to_num(value)
        out[block] = float(scheduled) if isinstance(scheduled, float) else 0.0

    for b in range(1, 97):
        if b not in out:
            out[b] = 0.0
    return out


def _block_to_time_window(block: int) -> str:
    idx = max(1, min(96, int(block))) - 1
    start_minutes = idx * 15
    end_minutes = (idx + 1) * 15
    sh, sm = divmod(start_minutes, 60)
    eh, em = divmod(end_minutes, 60)
    return f"{sh:02d}:{sm:02d}-{eh:02d}:{em:02d}"


def _is_day_ahead_upload_history_row(row: Dict[str, Any]) -> bool:
    output_key = str(row.get("output_file_key") or "").strip()
    template_name = str(row.get("template_file_name") or "").strip()
    joined = " ".join([output_key, template_name]).strip().lower()
    return bool(
        re.search(r"/day-ahead/|/dayahead/|/day_ahead/", joined)
        or re.search(r"_da0\.csv$", template_name, flags=re.IGNORECASE)
        or re.search(r"_da0\.csv$", output_key, flags=re.IGNORECASE)
        or re.search(r"\bday[-\s_]*ahead\b", template_name, flags=re.IGNORECASE)
    )


def _load_latest_generated_day_ahead_baseline(
    *,
    s3_client: Any,
    bucket: str,
    plant_code: str,
    schedule_date: str,
) -> Optional[Tuple[Dict[int, float], str]]:
    if not s3_client or not bucket:
        return None

    code = _normalize_plant_code(plant_code)
    roots = _generated_schedule_prefixes_for_plant(code, schedule_date, "intraday")
    if code == "GSNP":
        roots.append(f"generated/GSNP/gsnp/outputs/{schedule_date}/")
    if code == "SIRMOUR":
        roots.append(f"generated/Sirmour/sirmour/outputs/{schedule_date}/")

    candidates: List[Dict[str, Any]] = []
    for root in dict.fromkeys(roots):
        for folder in ("Day-ahead", "day-ahead", "dayahead", "day_ahead"):
            prefix = f"{root}{folder}/"
            token: Optional[str] = None
            while True:
                payload: Dict[str, Any] = {
                    "Bucket": bucket,
                    "Prefix": prefix,
                    "MaxKeys": 1000,
                }
                if token:
                    payload["ContinuationToken"] = token
                try:
                    response = s3_client.list_objects_v2(**payload)
                except Exception:
                    break
                for item in response.get("Contents", []) or []:
                    key = str(item.get("Key") or "").strip()
                    if key.lower().endswith(".csv") and _extract_schedule_revision_from_key(key) is not None:
                        candidates.append(
                            {
                                "key": key,
                                "last_modified": item.get("LastModified"),
                                "revision": _extract_schedule_revision_from_key(key) or 0,
                            }
                        )
                if not response.get("IsTruncated"):
                    break
                token = response.get("NextContinuationToken")
                if not token:
                    break

    if not candidates:
        return None

    def _candidate_sort_key(item: Dict[str, Any]) -> Tuple[float, int, str]:
        last_modified = item.get("last_modified")
        try:
            modified_ts = float(last_modified.timestamp())
        except Exception:
            modified_ts = 0.0
        return (
            modified_ts,
            int(item.get("revision") or 0),
            str(item.get("key") or ""),
        )

    selected = max(candidates, key=_candidate_sort_key)
    selected_key = str(selected.get("key") or "").strip()
    if not selected_key:
        return None
    try:
        response = s3_client.get_object(Bucket=bucket, Key=selected_key)
        csv_text = response["Body"].read().decode("utf-8", errors="replace")
    except Exception:
        return None
    return _parse_sldc_template_schedule_map(csv_text), selected_key.split("/")[-1]


def _generate_edited_frozen_from_upload_history_rows(
    *,
    plant_code: str,
    schedule_date: str,
    rows: List[Dict[str, Any]],
    s3_client: Any = None,
    bucket: str = "",
) -> Optional[str]:
    """
    Generate a consolidated edited_frozen.csv from upload-history rows (SLDC-confirmed templates).

    Rules (aligned with frontend freezeRules):
    - Choose DA baseline: latest day-ahead upload strictly before the first intraday upload time,
      else the latest available day-ahead upload.
    - If no day-ahead upload exists, use the latest generated day-ahead schedule from S3.
    - Intraday uploads apply only from their effective_start_block onward (MP plants: 90 min; others: 45 min).
    - Later intraday uploads override earlier ones starting at their own effective block.
    """
    plant_code = str(plant_code or "").strip().upper()
    schedule_date = str(schedule_date or "").strip()
    if not plant_code or not re.fullmatch(r"\d{4}-\d{2}-\d{2}", schedule_date):
        return None

    filtered = [
        r for r in (rows or [])
        if _normalize_plant_code(str(r.get("plant_code") or "").strip()) == plant_code
        and str(r.get("schedule_date") or "").strip() == schedule_date
        and str(r.get("uploaded_at") or "").strip()
        and str(r.get("csv_text") or "").strip()
    ]
    if not filtered:
        return None

    def _ts(row: Dict[str, Any]) -> float:
        dt = None
        try:
            normalized = str(row.get("uploaded_at") or "").strip()
            if re.match(r"^\d{4}-\d{2}-\d{2}T\d{2}:\d{2}", normalized) and not re.search(r"[zZ]|[+-]\d{2}:\d{2}$", normalized):
                normalized = f"{normalized}Z"
            dt = datetime.fromisoformat(normalized.replace("Z", "+00:00"))
        except Exception:
            dt = None
        if not dt:
            return 0.0
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        return dt.timestamp()

    day_ahead_rows = [r for r in filtered if _is_day_ahead_upload_history_row(r)]
    intraday_rows = [r for r in filtered if not _is_day_ahead_upload_history_row(r)]

    intraday_rows_sorted = sorted(intraday_rows, key=_ts)
    first_intraday_time = _ts(intraday_rows_sorted[0]) if intraday_rows_sorted else float("inf")

    # Pick DA baseline before first intraday; else latest DA.
    da_candidates = sorted([r for r in day_ahead_rows if _ts(r) < first_intraday_time], key=_ts, reverse=True)
    baseline_row = da_candidates[0] if da_candidates else (sorted(day_ahead_rows, key=_ts, reverse=True)[0] if day_ahead_rows else None)

    baseline_map: Dict[int, float] = {b: 0.0 for b in range(1, 97)}
    baseline_label = "DA|day_ahead.csv"
    if baseline_row:
        baseline_map = _parse_sldc_template_schedule_map(str(baseline_row.get("csv_text") or ""))
        baseline_name = (
            str(baseline_row.get("source_file_key") or "").split("/")[-1].strip()
            or str(baseline_row.get("template_file_name") or "").strip()
            or "day_ahead.csv"
        )
        baseline_label = f"DA|{baseline_name}"
    else:
        generated_baseline = _load_latest_generated_day_ahead_baseline(
            s3_client=s3_client,
            bucket=bucket,
            plant_code=plant_code,
            schedule_date=schedule_date,
        )
        if generated_baseline:
            baseline_map, baseline_name = generated_baseline
            baseline_label = f"DA|{baseline_name}"

    def _display_intraday_source_name(row: Dict[str, Any]) -> str:
        source_key = str(row.get("source_file_key") or "").strip()
        base = source_key.split("/")[-1].strip() if source_key else ""
        fallback = str(row.get("template_file_name") or "").strip()
        manual_request_id = str(row.get("manual_request_id") or "").strip()

        name = base or fallback or "intraday.csv"
        low = name.lower()

        # Keep explicit schedule_from revision names as-is (best visibility per block).
        if re.search(r"schedule_from_\d+(?:[_-][A-Za-z0-9]+)*\.csv$", name, re.IGNORECASE):
            return name
        # Manual flow should be shown as edited schedule.
        if manual_request_id:
            return "edited_schedule.csv"
        # Zero-change/non-manual fallback when upstream key is generic.
        if low == "system_schedule.csv":
            return "schedule_from_XX.csv"
        return name

    # Prepare intraday layers with effective blocks.
    layers: List[Dict[str, Any]] = []
    for row in intraday_rows_sorted:
        uploaded_at = str(row.get("uploaded_at") or "").strip()
        computed = _compute_submit_and_effective_blocks_from_iso(uploaded_at, plant_code=plant_code)
        submit_block = computed.get("submit_block")
        effective_block = computed.get("effective_start_block")
        if not isinstance(effective_block, int) or effective_block < 1 or effective_block > 96:
            continue
        schedule_map = _parse_sldc_template_schedule_map(str(row.get("csv_text") or ""))
        name = _display_intraday_source_name(row)
        layers.append(
            {
                "effective_block": int(effective_block),
                "uploaded_at": uploaded_at,
                "timestamp": _ts(row),
                "submit_block": submit_block,
                "map": schedule_map,
                "name": name,
            }
        )

    # Sort layers by effective then time (matches frontend)
    layers.sort(key=lambda it: (int(it.get("effective_block") or 999), float(it.get("timestamp") or 0.0)))

    # Build final schedule + source labels.
    final_sched: Dict[int, float] = {b: float(baseline_map.get(b, 0.0)) for b in range(1, 97)}
    final_source: Dict[int, str] = {b: baseline_label for b in range(1, 97)}

    next_effective_by_idx: List[int] = [int(layers[i + 1]["effective_block"]) if i + 1 < len(layers) else 97 for i in range(len(layers))]
    for idx, layer in enumerate(layers):
        eff = int(layer["effective_block"])
        next_eff = int(next_effective_by_idx[idx])
        end_block = min(96, max(eff, next_eff - 1))
        src = f"ID-{idx + 1}|{str(layer.get('name') or '').strip()}"
        layer_map: Dict[int, float] = layer.get("map") or {}
        for b in range(eff, end_block + 1):
            final_sched[b] = float(layer_map.get(b, final_sched.get(b, 0.0)) or 0.0)
            final_source[b] = src

    headers = ["Block", "Time", "Scheduled MW", "Actual MW", "Deviation MW", "Deviation %", "Penalty Rs", "Source Schedule"]
    lines = [",".join(headers)]
    for b in range(1, 97):
        time_win = _block_to_time_window(b)
        mw = float(final_sched.get(b, 0.0) or 0.0)
        mw_text = str(int(mw)) if abs(mw - int(mw)) < 1e-9 else str(mw)
        lines.append(",".join([str(b), time_win, mw_text, "", "", "", "", str(final_source.get(b, ""))]))
    return "\n".join(lines)


def _extract_upload_path_parts_from_key(key: str) -> Optional[Dict[str, str]]:
    """
    Parse uploads key pattern:
    uploads/vedanjay/{plant_code}/{YYYY-MM-DD}/{file_name}
    """
    text = str(key or "").strip()
    if not text:
        return None
    normalized = text.replace("\\", "/")
    parts = [p for p in normalized.split("/") if p]
    if len(parts) < 5:
        return None
    if parts[0].lower() != "uploads" or parts[1].lower() != "vedanjay":
        return None
    plant_code = str(parts[2]).upper()
    schedule_date = str(parts[3]).strip()
    file_name = parts[-1]
    if not plant_code or not re.fullmatch(r"[A-Z0-9_-]{1,32}", plant_code):
        return None
    if not re.fullmatch(r"\d{4}-\d{2}-\d{2}", schedule_date):
        return None
    return {
        "plant_code": plant_code,
        "schedule_date": schedule_date,
        "template_file_name": file_name,
    }


def _list_s3_upload_objects_safe(
    *,
    s3_client: Any,
    bucket: str,
    prefix: str,
    max_items: int = 2000,
) -> List[Dict[str, str]]:
    objects: List[Dict[str, str]] = []
    max_items = int(max_items or 0)
    if max_items <= 0:
        max_items = 2000

    if s3_client is not None and bucket:
        continuation = None
        while True:
            remaining = max_items - len(objects)
            if remaining <= 0:
                break
            payload: Dict[str, Any] = {
                "Bucket": bucket,
                "Prefix": prefix,
                "MaxKeys": min(1000, remaining),
            }
            if continuation:
                payload["ContinuationToken"] = continuation
            try:
                response = s3_client.list_objects_v2(**payload)
            except Exception as e:
                print(f"Boto3 list error for {prefix}: {e}")
                break

            for item in response.get("Contents", []) or []:
                key = str(item.get("Key", "")).strip()
                if not key:
                    continue
                last_modified = item.get("LastModified")
                last_modified_text = ""
                try:
                    if last_modified is not None:
                        last_modified_text = last_modified.isoformat()
                except Exception:
                    last_modified_text = ""
                objects.append({"key": key, "last_modified": last_modified_text})
                if len(objects) >= max_items:
                    break

            if len(objects) >= max_items:
                break

            if response.get("IsTruncated"):
                continuation = response.get("NextContinuationToken")
                if not continuation:
                    break
            else:
                break

        return objects

    # Public/listable bucket fallback via XML listing endpoint.
    try:
        url = (
            f"{DEFAULT_TEMPLATE_S3_BASE_URL.rstrip('/')}/"
            f"?list-type=2&prefix={quote(prefix)}&max-keys={max_items}"
        )
        public_timeout = float(os.getenv("S3_PROXY_PUBLIC_LIST_TIMEOUT_SECONDS") or 5)
        with urlopen(url, timeout=max(2.0, min(public_timeout, 10.0))) as resp:
            xml = resp.read().decode("utf-8", errors="replace")
        root = ElementTree.fromstring(xml)
        for node in root.findall(".//{*}Contents"):
            key = node.findtext("{*}Key", default="")
            last_modified = node.findtext("{*}LastModified", default="")
            if key:
                objects.append({"key": key, "last_modified": last_modified})
                if len(objects) >= max_items:
                    break
    except Exception:
        return []

    return objects


def _load_s3_upload_history_rows(
    *,
    schedule_date: Optional[date],
    plant_code: Optional[str],
    candidate_plants: Optional[List[str]] = None,
    limit: int,
) -> List[Dict[str, Any]]:
    bucket = _derive_s3_bucket_name()
    region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"
    s3 = None
    try:
        import boto3  # type: ignore
        from botocore.config import Config  # type: ignore
        if bucket:
            s3 = boto3.client(
                "s3",
                region_name=region,
                config=Config(
                    connect_timeout=4,
                    read_timeout=10,
                    retries={"max_attempts": 2, "mode": "standard"},
                ),
            )
    except Exception:
        s3 = None

    date_values: List[str] = []
    if schedule_date is not None:
        date_values = [schedule_date.isoformat()]
    else:
        # Keep breadth bounded for endpoint performance.
        today_utc = datetime.utcnow().date()
        date_values = [(today_utc - timedelta(days=i)).isoformat() for i in range(0, 14)]

    plant_values: List[str] = []
    if plant_code:
        plant_values = [str(plant_code).strip().upper()]
    elif candidate_plants:
        plant_values = sorted({str(p or "").strip().upper() for p in candidate_plants if str(p or "").strip()})
    else:
        # When no plant filter is provided, avoid scanning all plants by default (slow + can 504).
        # Prefer plants present in persisted history when possible (passed via candidate_plants).
        # If we still have no candidates, fall back to a bounded subset from config.
        try:
            plants_path = os.path.join(os.path.dirname(__file__), "config", "template_pipeline", "plants.json")
            with open(plants_path, "r", encoding="utf-8") as f:
                plants_payload = json.load(f)
            plant_values = sorted(
                {
                    str((p or {}).get("code") or (p or {}).get("name") or "").strip().upper()
                    for p in (plants_payload or [])
                }
            )
            plant_values = [p for p in plant_values if p][:25]
        except Exception:
            plant_values = ["GSNP", "SIRMOUR"]

    discovered: List[Dict[str, Any]] = []
    for p in plant_values:
        for d in date_values:
            for folder in _special_s3_plant_folder_aliases(p):
                prefix = f"{DEFAULT_READINESS_UPLOAD_PREFIX}/{folder}/{d}/"
                for obj in _list_s3_upload_objects_safe(
                    s3_client=s3,
                    bucket=bucket,
                    prefix=prefix,
                    max_items=max(1, int(limit)) * 3,
                ):
                    key = str(obj.get("key", "")).strip()
                    if not key.lower().endswith(".csv"):
                        continue
                    parsed = _extract_upload_path_parts_from_key(key)
                    if not parsed:
                        continue
                    discovered.append(
                        {
                            "id": hashlib.md5(f"{key}|{str(obj.get('last_modified', '')).strip()}".encode("utf-8")).hexdigest(),
                            "plant_code": parsed["plant_code"],
                            "schedule_date": parsed["schedule_date"],
                            "template_file_name": parsed["template_file_name"],
                            "source_file_key": "",
                            "manual_request_id": "",
                            "requested_by": "",
                            "bucket": bucket or "UNKNOWN",
                            "output_file_key": key,
                            "output_file_url": f"https://{bucket}.s3.{region}.amazonaws.com/{key}" if bucket else "",
                            "uploaded_at": str(obj.get("last_modified", "")).strip(),
                            "storage_mode": "s3_discovered",
                            "error": None,
                            "csv_text": "",
                        }
                    )
                if len(discovered) >= max(1, int(limit)) * 3:
                    # Avoid unbounded S3 scans when a caller requests a large limit.
                    # We keep some headroom (x3) so the post-sort still returns enough rows.
                    break
            if len(discovered) >= max(1, int(limit)) * 3:
                break
        if len(discovered) >= max(1, int(limit)) * 3:
            break

    discovered = sorted(
        discovered,
        key=lambda r: str(r.get("uploaded_at", "")),
        reverse=True,
    )

    # Only enrich the top N rows to keep the endpoint fast.
    top_n = discovered[: max(1, int(limit))]
    if s3 is not None and bucket:
        for row in top_n:
            key = str(row.get("output_file_key", "") or "").strip()
            if not key:
                continue
            try:
                head = s3.head_object(Bucket=bucket, Key=key)
                meta = head.get("Metadata") or {}
                if isinstance(meta, dict):
                    requested_by = str(meta.get("requested_by") or "").strip()
                    source_file_key = str(meta.get("source_file_key") or "").strip()
                    manual_request_id = str(meta.get("manual_request_id") or "").strip()
                    freeze_time = str(meta.get("freeze_time") or "").strip()
                    trigger_reason = str(meta.get("trigger_reason") or "").strip()
                    slot_index = str(meta.get("slot_index") or "").strip()
                    if requested_by and not str(row.get("requested_by") or "").strip():
                        row["requested_by"] = requested_by
                    if source_file_key and not str(row.get("source_file_key") or "").strip():
                        row["source_file_key"] = source_file_key
                    if manual_request_id and not str(row.get("manual_request_id") or "").strip():
                        row["manual_request_id"] = manual_request_id
                    if freeze_time and not str(row.get("freeze_time") or "").strip():
                        row["freeze_time"] = freeze_time
                    if trigger_reason and not str(row.get("trigger_reason") or "").strip():
                        row["trigger_reason"] = trigger_reason
                    if slot_index and not str(row.get("slot_index") or "").strip():
                        row["slot_index"] = slot_index
            except Exception:
                continue

    # Attach computed submit/effective blocks so UI + frozen recompute can reuse stable values.
    for row in top_n:
        submit_source = str(row.get("freeze_time") or "").strip() or str(row.get("uploaded_at", "")).strip()
        computed = _compute_submit_and_effective_blocks_from_iso(submit_source, plant_code=row.get("plant_code") or row.get("plantCode") or row.get("site_code") or row.get("siteCode"))
        row.update(computed)
    return top_n


def _derive_s3_bucket_name() -> str:
    explicit_bucket = os.getenv("READINESS_UPLOAD_BUCKET", "").strip() or os.getenv("TEMPLATE_OUTPUT_BUCKET", "").strip()
    if explicit_bucket:
        return explicit_bucket
    parsed = urlparse(DEFAULT_TEMPLATE_S3_BASE_URL)
    host = parsed.netloc or ""
    if host:
        return host.split(".")[0]
    return ""


def _get_dynamodb_table(table_env_key: str) -> Any:
    table_name = os.getenv(table_env_key, "").strip()
    if not table_name and table_env_key == "WHATSAPP_INSTANT_TABLE":
        table_name = os.getenv("DDB_WHATSAPP_TABLE", "").strip()
    if not table_name and table_env_key == "WHATSAPP_WINDOWS_TABLE":
        table_name = os.getenv("DDB_WHATSAPP_WINDOWS_TABLE", "").strip() or "plant_control_windows_test"
    if not table_name:
        raise RuntimeError(f"{table_env_key} is not configured")
    region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"
    try:
        import boto3  # type: ignore
    except Exception as exc:
        raise RuntimeError(f"boto3 not available: {exc}") from exc
    try:
        from botocore.config import Config  # type: ignore
        dynamodb = boto3.resource(
            "dynamodb",
            region_name=region,
            config=Config(
                connect_timeout=4,
                read_timeout=8,
                retries={"max_attempts": 1},
            ),
        )
    except Exception:
        dynamodb = boto3.resource("dynamodb", region_name=region)
    return dynamodb.Table(table_name)


def _extract_whatsapp_site_state(state_item: Dict[str, Any], site_code: str) -> Dict[str, Any]:
    site_states = state_item.get("site_states") or state_item.get("siteStates") or {}
    if not isinstance(site_states, dict):
        return {}
    key = str(site_code or "").strip().upper()
    if not key:
        return {}
    raw = site_states.get(key) or site_states.get(key.title()) or site_states.get(key.lower()) or {}
    if not isinstance(raw, dict):
        return {}
    # Curtailement capacity may be missing/None.
    capacity = raw.get("curtailment_capacity")
    # Some producers store plant status under different keys.
    plant_status = raw.get("plant_status") or raw.get("plantStatus") or raw.get("status") or raw.get("plant_status_value") or ""
    return {
        "site": key,
        "plant_status": plant_status,
        "last_message": raw.get("last_message") or raw.get("lastMessage") or "",
        "updated_at": raw.get("updated_at") or raw.get("updatedAt") or "",
        "curtailment_capacity": capacity,
    }


def _load_whatsapp_windows_for_site(plant_id: str, site_code: str, limit: int = 50) -> List[Dict[str, Any]]:
    site = str(site_code or "").strip().upper()
    if not site:
        return []
    try:
        table = _get_dynamodb_table("WHATSAPP_WINDOWS_TABLE")
    except Exception:
        return []

    rows: List[Dict[str, Any]] = []
    try:
        # Prefer Query by partition key (plant_id) then filter by site.
        try:
            from boto3.dynamodb.conditions import Key, Attr  # type: ignore
            response = table.query(
                KeyConditionExpression=Key("plant_id").eq(str(plant_id)),
                FilterExpression=Attr("site").eq(site),
                Limit=max(1, min(int(limit), 200)),
            )
            items = response.get("Items") or []
        except Exception:
            # Fallback to scan (small pages).
            items = []
            last_evaluated_key = None
            pages = 0
            while pages < 5 and len(items) < limit:
                kwargs = {}
                if last_evaluated_key:
                    kwargs["ExclusiveStartKey"] = last_evaluated_key
                response = table.scan(**kwargs)
                raw = response.get("Items") or []
                items.extend(raw)
                last_evaluated_key = response.get("LastEvaluatedKey")
                pages += 1
                if not last_evaluated_key:
                    break

        for raw_item in items:
            item = _normalize_ddb_item(raw_item)
            if str(item.get("site") or "").strip().upper() != site:
                continue
            rows.append({
                "plant_id": item.get("plant_id"),
                "window_id": item.get("window_id"),
                "site": site,
                "plant_status": item.get("plant_status") or item.get("status"),
                "curtailment_capacity": item.get("curtailment_capacity"),
                "start_time": item.get("start_time"),
                "end_time": item.get("end_time"),
                "created_at": item.get("created_at"),
                "updated_at": item.get("updated_at"),
                "last_message": item.get("last_message") or item.get("lastMessage") or "",
            })
    except Exception:
        return []

    def _sort_key(row: Dict[str, Any]) -> str:
        # Prefer most-recent windows first.
        return str(row.get("updated_at") or row.get("created_at") or row.get("start_time") or "")
    rows.sort(key=_sort_key, reverse=True)
    return rows[:max(1, min(int(limit), 200))]


def _whatsapp_sql_site_code(row: Any) -> str:
    raw = str(getattr(row, "plantName", "") or "").strip()
    match = re.search(r"\(([A-Za-z0-9_-]+)\)", raw)
    if match:
        raw = match.group(1)
    code = re.sub(r"[^A-Za-z0-9_-]+", "", raw).upper()
    if code == "OSEL":
        code = "OSEPL"
    if code == "ANJANGAON":
        code = "ANJANGOAN"
    return code


def _whatsapp_sql_status(row: Any) -> str:
    if bool(getattr(row, "curtailmentStatus", False)):
        return "CURTAILMENT"
    haystack = " ".join(
        str(getattr(row, field, "") or "")
        for field in ("status", "curtailmentReason", "remarks", "expectedTrend")
    ).upper()
    if "SHUTDOWN" in haystack:
        return "SHUTDOWN"
    if "CURTAIL" in haystack:
        return "CURTAILMENT"
    return "NORMAL"


def _whatsapp_sql_row_start_iso(row: Any) -> str:
    day_value = getattr(row, "date", None)
    if isinstance(day_value, datetime):
        day = day_value.date()
    elif isinstance(day_value, date):
        day = day_value
    else:
        try:
            day = datetime.strptime(str(day_value or "").strip(), "%Y-%m-%d").date()
        except Exception:
            day = datetime.now(ZoneInfo("Asia/Kolkata")).date()

    time_text = str(getattr(row, "time", "") or "").strip()
    if not re.fullmatch(r"([01]\d|2[0-3]):[0-5]\d", time_text):
        time_text = datetime.now(ZoneInfo("Asia/Kolkata")).strftime("%H:%M")
    hour, minute = [int(part) for part in time_text.split(":")]
    return datetime(day.year, day.month, day.day, hour, minute, tzinfo=ZoneInfo("Asia/Kolkata")).isoformat()


def _whatsapp_sql_last_message(row: Any, plant_status: str) -> str:
    remarks = str(getattr(row, "remarks", "") or "").strip()
    if remarks:
        return remarks
    reason = str(getattr(row, "curtailmentReason", "") or "").strip()
    site = str(getattr(row, "plantName", "") or "").strip() or _whatsapp_sql_site_code(row)
    if reason:
        return f"{site} {plant_status.lower()}: {reason}"
    return f"{site} {plant_status.lower()}"


def _close_whatsapp_open_windows_for_site(
    table: Any,
    site: str,
    end_iso: str,
    updated_iso: str,
    last_message: str = "",
    source: str = "ui",
) -> None:
    try:
        from boto3.dynamodb.conditions import Key  # type: ignore
        response = table.query(KeyConditionExpression=Key("plant_id").eq(SITE_MESSAGES_PLANT_ID), Limit=500)
        items = response.get("Items") or []
    except Exception:
        try:
            response = table.scan(Limit=500)
            items = response.get("Items") or []
        except Exception:
            return

    for raw in items:
        item = _normalize_ddb_item(raw)
        if str(item.get("site") or "").strip().upper() != site:
            continue
        status = str(item.get("plant_status") or item.get("status") or "").strip().upper()
        if status not in {"SHUTDOWN", "CURTAILMENT"}:
            continue
        if item.get("active") is False:
            continue
        if item.get("is_open_ended") is False:
            continue
        window_id = str(item.get("window_id") or "").strip()
        if not window_id:
            continue
        try:
            table.update_item(
                Key={"plant_id": SITE_MESSAGES_PLANT_ID, "window_id": window_id},
                UpdateExpression=(
                    "SET active = :active, is_open_ended = :open, end_time = :end, "
                    "updated_at = :updated, last_message = :message, #src = :source"
                ),
                ExpressionAttributeNames={
                    "#src": "source",
                },
                ExpressionAttributeValues={
                    ":active": False,
                    ":open": False,
                    ":end": end_iso,
                    ":updated": updated_iso,
                    ":message": str(last_message or "").strip() or "normal/restored",
                    ":source": str(source or "").strip() or "ui",
                },
            )
        except Exception:
            continue


def _close_site_message_open_windows_for_normal(table: Any, item: Dict[str, Any]) -> None:
    if str(item.get("plant_status") or "").strip().upper() != "NORMAL":
        return
    site = str(item.get("site") or "").strip().upper()
    if not site:
        return
    close_iso = str(item.get("start_time") or item.get("end_time") or item.get("created_at") or item.get("updated_at") or "").strip()
    if not close_iso:
        return
    updated_iso = str(item.get("updated_at") or item.get("created_at") or close_iso).strip() or close_iso
    _close_whatsapp_open_windows_for_site(
        table,
        site,
        close_iso,
        updated_iso,
        str(item.get("last_message") or "").strip(),
        str(item.get("source") or "").strip() or "ui",
    )


def _mirror_whatsapp_sql_row_to_dynamodb(row: Any) -> None:
    try:
        site = _whatsapp_sql_site_code(row)
        if not site:
            return
        plant_status = _whatsapp_sql_status(row)
        if plant_status not in {"NORMAL", "SHUTDOWN", "CURTAILMENT"}:
            return

        table = _get_dynamodb_table("WHATSAPP_WINDOWS_TABLE")
        start_iso = _whatsapp_sql_row_start_iso(row)
        now_iso = datetime.now(timezone.utc).isoformat()
        end_iso = None if plant_status in {"SHUTDOWN", "CURTAILMENT"} else start_iso
        if plant_status == "NORMAL":
            _close_whatsapp_open_windows_for_site(
                table,
                site,
                start_iso,
                now_iso,
                _whatsapp_sql_last_message(row, plant_status),
                "ui",
            )

        end_token = "open" if end_iso is None else end_iso
        window_id = f"{site}#{plant_status}#{start_iso}#{end_token}"
        item: Dict[str, Any] = {
            "plant_id": SITE_MESSAGES_PLANT_ID,
            "window_id": window_id,
            "site": site,
            "plant_status": plant_status,
            "start_time": start_iso,
            "active": plant_status != "NORMAL",
            "is_open_ended": end_iso is None,
            "last_message": _whatsapp_sql_last_message(row, plant_status),
            "source": "ui",
            "created_at": now_iso,
            "updated_at": now_iso,
        }
        if end_iso is not None:
            item["end_time"] = end_iso
        table.put_item(Item=_site_message_decimalize(item))
    except Exception as exc:
        print(f"Warning: failed to mirror WhatsApp data to DynamoDB: {exc}")


def _unwrap_ddb_value(value: Any) -> Any:
    if not isinstance(value, dict) or len(value) != 1:
        return value
    if "S" in value:
        return value["S"]
    if "N" in value:
        try:
            return int(value["N"])
        except Exception:
            try:
                return float(value["N"])
            except Exception:
                return value["N"]
    if "BOOL" in value:
        return bool(value["BOOL"])
    if "M" in value:
        return {k: _unwrap_ddb_value(v) for k, v in value["M"].items()}
    if "L" in value:
        return [_unwrap_ddb_value(v) for v in value["L"]]
    return value


def _normalize_ddb_item(item: Any) -> Dict[str, Any]:
    if not isinstance(item, dict):
        return {}
    return {k: _unwrap_ddb_value(v) for k, v in item.items()}

def _parse_ddb_timestamp(value: Any) -> Optional[int]:
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return int(value)
    text = str(value).strip()
    if not text:
        return None
    if re.fullmatch(r"\d{10,}", text):
        try:
            return int(text)
        except Exception:
            return None
    try:
        cleaned = text.replace("Z", "+00:00") if "Z" in text else text
        return int(datetime.fromisoformat(cleaned).timestamp() * 1000)
    except Exception:
        return None

def _whatsapp_item_to_payload(item: Dict[str, Any]) -> Dict[str, Any]:
    message = str(
        item.get("last_message")
        or item.get("lastMessage")
        or item.get("message")
        or ""
    )
    plant = item.get("plant_id") or item.get("plant") or ""
    ts = (
        _parse_ddb_timestamp(item.get("updated_at"))
        or _parse_ddb_timestamp(item.get("updatedAt"))
        or _parse_ddb_timestamp(item.get("timestamp"))
    )
    ts = ts or 0
    msg_hash = hashlib.md5(message.encode("utf-8")).hexdigest()[:10] if message else "nomsg"
    item_id = f"{plant}:{ts}:{msg_hash}"
    return {
        "id": item_id,
        "plant": plant,
        "message": message,
        "templateType": "whatsapp",
        "timestamp": datetime.utcfromtimestamp(ts / 1000).isoformat() + "Z" if ts else "",
        "timestamp_ms": ts,
    }


SITE_MESSAGES_WINDOWS_TABLE_NAME = "plant_control_windows_test"
SITE_MESSAGES_PLANT_ID = "vedanjay"
SITE_MESSAGE_EVENT_STATUS = {
    "shutdown": "SHUTDOWN",
    "curtailment": "CURTAILMENT",
    "partial_shutdown": "SHUTDOWN",
    "normal": "NORMAL",
    "delay": "DELAY",
}
MULTI_GENERATOR_PLANT_TABLE_NAME = "multi_generator_plant"


def _site_message_ist_now() -> datetime:
    return datetime.now(ZoneInfo("Asia/Kolkata"))


def _site_message_parse_date(value: str) -> date:
    text = str(value or "").strip()
    try:
        return datetime.strptime(text, "%Y-%m-%d").date()
    except ValueError as exc:
        raise HTTPException(status_code=400, detail="event_date must use YYYY-MM-DD format") from exc


def _site_message_parse_time(value: Optional[str], field_name: str) -> Optional[str]:
    text = str(value or "").strip()
    if not text:
        return None
    if not re.fullmatch(r"([01]\d|2[0-3]):[0-5]\d", text):
        raise HTTPException(status_code=400, detail=f"{field_name} must use HH:mm format")
    return text


def _site_message_to_ist_iso(event_day: date, time_text: Optional[str]) -> Optional[str]:
    safe_time = _site_message_parse_time(time_text, "time")
    if not safe_time:
        return None
    hour, minute = [int(part) for part in safe_time.split(":")]
    return datetime(
        event_day.year,
        event_day.month,
        event_day.day,
        hour,
        minute,
        tzinfo=ZoneInfo("Asia/Kolkata"),
    ).isoformat()


def _site_message_decimalize(value: Any) -> Any:
    if isinstance(value, float):
        if not math.isfinite(value):
            return None
        return Decimal(str(value))
    if isinstance(value, dict):
        return {k: _site_message_decimalize(v) for k, v in value.items() if v is not None}
    if isinstance(value, list):
        return [_site_message_decimalize(v) for v in value if v is not None]
    return value


def _get_site_messages_windows_table() -> Any:
    configured = os.getenv("SITE_MESSAGES_WINDOWS_TABLE", SITE_MESSAGES_WINDOWS_TABLE_NAME).strip()
    region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"
    try:
        import boto3  # type: ignore
    except Exception as exc:
        raise RuntimeError(f"boto3 not available: {exc}") from exc
    dynamodb = boto3.resource("dynamodb", region_name=region)
    return dynamodb.Table(configured or SITE_MESSAGES_WINDOWS_TABLE_NAME)


def _multi_generator_ist_now() -> datetime:
    return datetime.now(ZoneInfo("Asia/Kolkata"))


def _get_multi_generator_plant_table() -> Any:
    table_name = os.getenv("MULTI_GENERATOR_PLANT_TABLE", MULTI_GENERATOR_PLANT_TABLE_NAME).strip()
    region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"
    try:
        import boto3  # type: ignore
        from botocore.exceptions import ClientError  # type: ignore
    except Exception as exc:
        raise RuntimeError(f"boto3 not available: {exc}") from exc

    dynamodb = boto3.resource("dynamodb", region_name=region)
    client = boto3.client("dynamodb", region_name=region)
    try:
        client.describe_table(TableName=table_name)
    except ClientError as exc:
        error_code = exc.response.get("Error", {}).get("Code")
        if error_code != "ResourceNotFoundException":
            raise
        client.create_table(
            TableName=table_name,
            AttributeDefinitions=[{"AttributeName": "plant_id", "AttributeType": "S"}],
            KeySchema=[{"AttributeName": "plant_id", "KeyType": "HASH"}],
            BillingMode="PAY_PER_REQUEST",
        )
        client.get_waiter("table_exists").wait(
            TableName=table_name,
            WaiterConfig={"Delay": 2, "MaxAttempts": 10},
        )
    return dynamodb.Table(table_name)


def _multi_generator_jsonable(value: Any) -> Any:
    if isinstance(value, Decimal):
        if value % 1 == 0:
            return int(value)
        return float(value)
    if isinstance(value, dict):
        return {k: _multi_generator_jsonable(v) for k, v in value.items()}
    if isinstance(value, list):
        return [_multi_generator_jsonable(v) for v in value]
    return value


def _build_multi_generator_item(
    payload: MultiGeneratorPlantRequest,
    *,
    user_name: str = "",
    existing: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    now_iso = _multi_generator_ist_now().isoformat()
    data = payload.model_dump(exclude_none=True)
    data["plant_id"] = str(payload.plant_id or "").strip() or "ZETRIC_SOLAR_PARK"
    data["plant_name"] = str(payload.plant_name or "").strip()
    data["record_type"] = "multi_generator_plant_config"
    data["created_at"] = str((existing or {}).get("created_at") or now_iso)
    data["updated_at"] = now_iso
    data["updated_by"] = str(user_name or "").strip() or "ui"
    return _site_message_decimalize(data)


def _site_message_storage_site_id(value: str) -> str:
    site = str(value or "").strip().upper()
    if site == "ANJANGAON":
        return "ANJANGOAN"
    return site


def _site_message_time_key(value: Any) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    if re.fullmatch(r"([01]\d|2[0-3]):[0-5]\d", text):
        return text
    try:
        normalized = text.replace("Z", "+00:00")
        parsed = datetime.fromisoformat(normalized)
        if parsed.tzinfo is None:
            parsed = parsed.replace(tzinfo=ZoneInfo("Asia/Kolkata"))
        return parsed.astimezone(ZoneInfo("Asia/Kolkata")).strftime("%Y-%m-%d %H:%M")
    except Exception:
        return text[:16]


def _site_message_float_key(value: Any) -> Optional[float]:
    if value is None:
        return None
    try:
        num = float(str(value).strip())
    except Exception:
        return None
    return round(num, 6) if math.isfinite(num) else None


def _site_message_event_date_key(item: Dict[str, Any]) -> str:
    raw_payload = item.get("raw_payload") if isinstance(item.get("raw_payload"), dict) else {}
    for value in (
        raw_payload.get("event_date") if isinstance(raw_payload, dict) else None,
        item.get("event_date"),
        item.get("eventDate"),
    ):
        text = str(value or "").strip()
        if not text:
            continue
        try:
            return _site_message_parse_date(text).isoformat()
        except Exception:
            pass
    created_at = str(item.get("created_at") or "").strip()
    if created_at:
        try:
            parsed = datetime.fromisoformat(created_at.replace("Z", "+00:00"))
            if parsed.tzinfo is None:
                parsed = parsed.replace(tzinfo=ZoneInfo("Asia/Kolkata"))
            return parsed.astimezone(ZoneInfo("Asia/Kolkata")).date().isoformat()
        except Exception:
            return created_at[:10]
    return ""


def _site_message_compare_key(item: Dict[str, Any]) -> Dict[str, Any]:
    status = str(item.get("plant_status") or item.get("plantStatus") or item.get("status") or "").strip().upper()
    control_mode = str(item.get("control_mode") or item.get("controlMode") or item.get("reduction_type") or "").strip().upper()
    if status == "PARTIAL_SHUTDOWN":
        status = "SHUTDOWN"
        if not control_mode:
            control_mode = "DC"
    if status == "CURTAILMENT" and not control_mode:
        control_mode = "AC"
    mw_value = None
    if status == "CURTAILMENT":
        mw_value = item.get("curtailment_capacity")
    elif status == "SHUTDOWN" and control_mode == "DC":
        mw_value = item.get("shutdown_reduction_mw") or item.get("curtailment_capacity")
    return {
        "site": _site_message_storage_site_id(str(item.get("site") or item.get("site_id") or "")),
        "plant_status": status,
        "control_mode": control_mode,
        "start_time": _site_message_time_key(item.get("start_time") or item.get("startTime")),
        "end_time": _site_message_time_key(item.get("end_time") or item.get("endTime")),
        "mw": _site_message_float_key(mw_value),
    }
    if status == "NORMAL":
        key["event_date"] = _site_message_event_date_key(item)
        key["message"] = re.sub(r"\s+", " ", str(item.get("last_message") or item.get("raw_message") or "").strip()).upper()
    return key


def _site_message_is_duplicate(table: Any, item: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    target = _site_message_compare_key(item)
    site = target.get("site")
    if not site:
        return None

    try:
        from boto3.dynamodb.conditions import Key  # type: ignore

        candidates = []
        last_evaluated_key = None
        pages = 0
        while pages < 10 and len(candidates) < 1000:
            kwargs: Dict[str, Any] = {
                "KeyConditionExpression": Key("plant_id").eq(str(item.get("plant_id") or SITE_MESSAGES_PLANT_ID)),
                "Limit": 200,
            }
            if last_evaluated_key:
                kwargs["ExclusiveStartKey"] = last_evaluated_key
            response = table.query(**kwargs)
            raw_items = response.get("Items") or []
            for raw in raw_items:
                if _site_message_storage_site_id(str(raw.get("site") or raw.get("site_id") or "")) == site:
                    candidates.append(raw)
            last_evaluated_key = response.get("LastEvaluatedKey")
            pages += 1
            if not last_evaluated_key:
                break
    except Exception:
        candidates = []
        last_evaluated_key = None
        pages = 0
        while pages < 5 and len(candidates) < 500:
            kwargs: Dict[str, Any] = {}
            if last_evaluated_key:
                kwargs["ExclusiveStartKey"] = last_evaluated_key
            response = table.scan(**kwargs)
            raw_items = response.get("Items") or []
            for raw in raw_items:
                if _site_message_storage_site_id(str(raw.get("site") or raw.get("site_id") or "")) == site:
                    candidates.append(raw)
            last_evaluated_key = response.get("LastEvaluatedKey")
            pages += 1
            if not last_evaluated_key:
                break

    for existing in candidates:
        existing_key = _site_message_compare_key(existing)
        if existing_key != target:
            continue
        existing_window_id = str(existing.get("window_id") or "").strip()
        if existing_window_id and existing_window_id == str(item.get("window_id") or "").strip():
            continue
        return existing
    return None


def _build_site_message_window_item(payload: SiteMessageRequest) -> Dict[str, Any]:
    site = _site_message_storage_site_id(str(payload.site_id or ""))
    if not site:
        raise HTTPException(status_code=400, detail="site_id is required")
    record_type = str(payload.record_type or "").strip()
    if record_type != "site_event_message":
        raise HTTPException(status_code=400, detail="record_type must be site_event_message")
    source = str(payload.source or "").strip().lower() or "ui"
    if source not in {"dashboard", "ui"}:
        raise HTTPException(status_code=400, detail="source must be ui or dashboard")
    stored_source = "ui"

    event_type = str(payload.event_type or "").strip().lower()
    plant_status = SITE_MESSAGE_EVENT_STATUS.get(event_type)
    if not plant_status:
        allowed = ", ".join(sorted(SITE_MESSAGE_EVENT_STATUS))
        raise HTTPException(status_code=400, detail=f"event_type must be one of: {allowed}")

    raw_message = str(payload.raw_message or "").strip()
    if not raw_message:
        raise HTTPException(status_code=400, detail="raw_message is required")

    event_day = _site_message_parse_date(payload.event_date)
    start_time = _site_message_parse_time(payload.start_time, "start_time")
    end_time = _site_message_parse_time(payload.end_time, "end_time")

    if event_type in {"shutdown", "curtailment", "partial_shutdown"} and not start_time:
        raise HTTPException(status_code=400, detail=f"start_time is required for {event_type}")
    if event_type in {"curtailment", "partial_shutdown"} and payload.mw is None:
        raise HTTPException(status_code=400, detail=f"mw is required for {event_type}")

    now_ist = _site_message_ist_now()
    now_iso = now_ist.isoformat()
    now_ms = int(now_ist.timestamp() * 1000)
    start_iso = _site_message_to_ist_iso(event_day, start_time)
    end_iso = _site_message_to_ist_iso(event_day, end_time)
    if event_type == "normal" and start_iso and not end_iso:
        end_iso = start_iso
    safe_time_token = (start_time or "none").replace(":", "-")
    window_id = f"{site}#{plant_status}#{event_day.isoformat()}#{safe_time_token}#{now_ms}"

    if event_type == "curtailment":
        control_mode = "AC"
    elif event_type == "partial_shutdown":
        control_mode = "DC"
    elif event_type == "shutdown":
        control_mode = "FULL"
    else:
        control_mode = None

    raw_payload = payload.model_dump(exclude_none=True)
    raw_payload["site_id"] = site
    raw_payload.setdefault("site_id_raw", payload.site_id_raw or site)
    raw_payload["source"] = stored_source
    raw_payload.setdefault("record_type", "site_event_message")
    raw_payload["plant_status"] = plant_status
    if control_mode:
        raw_payload["control_mode"] = control_mode
    if event_type == "partial_shutdown" and payload.mw is not None:
        raw_payload["shutdown_reduction_mw"] = payload.mw

    item: Dict[str, Any] = {
        "plant_id": SITE_MESSAGES_PLANT_ID,
        "window_id": window_id,
        "site": site,
        "source": stored_source,
        "plant_status": plant_status,
        "control_mode": control_mode,
        "last_message": raw_message,
        "start_time": start_iso,
        "end_time": end_iso,
        "created_at": now_iso,
        "updated_at": now_iso,
        "active": plant_status not in {"NORMAL"},
        "is_open_ended": end_iso is None,
        "record_type": "site_event_message",
        "raw_payload": raw_payload,
    }

    if event_type == "curtailment" and payload.mw is not None:
        item["curtailment_capacity"] = Decimal(str(payload.mw))
    if event_type == "partial_shutdown" and payload.mw is not None:
        item["shutdown_reduction_mw"] = Decimal(str(payload.mw))
    if payload.reduction_type:
        item["reduction_type"] = str(payload.reduction_type).strip().upper()
    if payload.unit:
        item["unit"] = str(payload.unit).strip().upper()
    if payload.description:
        item["description"] = str(payload.description).strip()
    if payload.status:
        item["status"] = str(payload.status).strip()
    if payload.delay_mode:
        item["delay_mode"] = str(payload.delay_mode).strip()
    if payload.minutes is not None:
        item["minutes"] = int(payload.minutes)

    return _site_message_decimalize(item)


def _site_message_normalize_user_role(role: Optional[str], user_name: Optional[str]) -> str:
    raw_role = str(role or "").strip().lower()
    raw_user = str(user_name or "").strip().lower()
    if raw_role == "admin":
        return "admin"
    if raw_role == "intern" or raw_user == "intern":
        return "intern"
    return raw_role or "employee"


@app.get("/api/multi-generator-plant/{plant_id}")
async def get_multi_generator_plant(plant_id: str):
    plant_key = str(plant_id or "").strip() or "ZETRIC_SOLAR_PARK"
    try:
        table = _get_multi_generator_plant_table()
        response = table.get_item(Key={"plant_id": plant_key})
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to load multi generator plant config: {exc}") from exc

    item = response.get("Item")
    if not item:
        return {
            "success": False,
            "message": "Multi generator plant config not found",
            "table": MULTI_GENERATOR_PLANT_TABLE_NAME,
            "plant_id": plant_key,
            "item": None,
        }

    return {
        "success": True,
        "table": MULTI_GENERATOR_PLANT_TABLE_NAME,
        "plant_id": plant_key,
        "item": _multi_generator_jsonable(_normalize_ddb_item(item)),
    }


@app.put("/api/multi-generator-plant/{plant_id}")
async def save_multi_generator_plant(
    plant_id: str,
    payload: MultiGeneratorPlantRequest,
    x_user_name: Optional[str] = Header(None, alias="X-User-Name"),
):
    plant_key = str(plant_id or "").strip() or "ZETRIC_SOLAR_PARK"
    if str(payload.plant_id or "").strip() and str(payload.plant_id).strip() != plant_key:
        raise HTTPException(status_code=400, detail="plant_id in path and payload must match")

    payload.plant_id = plant_key
    try:
        table = _get_multi_generator_plant_table()
        existing = table.get_item(Key={"plant_id": plant_key}).get("Item")
        item = _build_multi_generator_item(payload, user_name=x_user_name or "", existing=existing)
        table.put_item(Item=item)
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to save multi generator plant config: {exc}") from exc

    return {
        "success": True,
        "message": "Multi generator plant config saved",
        "table": MULTI_GENERATOR_PLANT_TABLE_NAME,
        "plant_id": plant_key,
        "item": _multi_generator_jsonable(_normalize_ddb_item(item)),
    }


def _site_message_log_row(
    *,
    db: Session,
    payload: SiteMessageRequest,
    event_day: date,
    username: str,
    user_role: str,
    status: str,
    window_id: str = "",
    error_message: str = "",
) -> SiteMessageLog:
    row = SiteMessageLog(
        username=str(username or "").strip()[:128] or None,
        user_role=str(user_role or "").strip()[:32] or None,
        site_id=str(payload.site_id or "").strip().upper(),
        site_id_raw=str(payload.site_id_raw or "").strip() or None,
        event_type=str(payload.event_type or "").strip().lower(),
        raw_message=str(payload.raw_message or "").strip(),
        event_date=event_day,
        start_time=str(payload.start_time or "").strip() or None,
        end_time=str(payload.end_time or "").strip() or None,
        mw=float(payload.mw) if payload.mw is not None else None,
        unit=str(payload.unit or "").strip().upper() or None,
        reduction_type=str(payload.reduction_type or "").strip().upper() or None,
        description=str(payload.description or "").strip() or None,
        dynamodb_table=SITE_MESSAGES_WINDOWS_TABLE_NAME,
        dynamodb_window_id=str(window_id or "").strip() or None,
        status=str(status or "").strip().upper() or "SUCCESS",
        error_message=str(error_message or "").strip() or None,
    )
    db.add(row)
    db.commit()
    db.refresh(row)
    return row


@app.post("/api/site-messages")
async def create_site_message(
    payload: SiteMessageRequest,
    x_user_role: Optional[str] = Header(None, alias="X-User-Role"),
    x_user_name: Optional[str] = Header(None, alias="X-User-Name"),
    db: Session = Depends(get_db),
):
    """Save a Site Messages row to the dedicated DynamoDB test windows table only."""
    item = _build_site_message_window_item(payload)
    event_day = _site_message_parse_date(payload.event_date)
    user_name = str(x_user_name or "").strip()
    user_role = _site_message_normalize_user_role(x_user_role, user_name)
    try:
        table = _get_site_messages_windows_table()
        duplicate_item = _site_message_is_duplicate(table, item)
        if duplicate_item:
            _close_site_message_open_windows_for_normal(table, item)
            log_row = _site_message_log_row(
                db=db,
                payload=payload,
                event_day=event_day,
                username=user_name,
                user_role=user_role,
                status="DUPLICATE",
                window_id=str(duplicate_item.get("window_id") or ""),
                error_message="Message has already been stored.",
            )
            return {
                "success": False,
                "duplicate": True,
                "message": "Message has already been stored.",
                "table": SITE_MESSAGES_WINDOWS_TABLE_NAME,
                "plant_id": duplicate_item.get("plant_id") or item.get("plant_id"),
                "window_id": duplicate_item.get("window_id"),
                "log_id": log_row.id,
            }
        _close_site_message_open_windows_for_normal(table, item)
        table.put_item(Item=item)
        log_row = _site_message_log_row(
            db=db,
            payload=payload,
            event_day=event_day,
            username=user_name,
            user_role=user_role,
            status="SUCCESS",
            window_id=str(item.get("window_id") or ""),
        )
    except HTTPException:
        raise
    except Exception as exc:
        try:
            _site_message_log_row(
                db=db,
                payload=payload,
                event_day=event_day,
                username=user_name,
                user_role=user_role,
                status="FAILED",
                error_message=str(exc),
            )
        except Exception:
            db.rollback()
        raise HTTPException(status_code=500, detail=f"Failed to write site message: {exc}") from exc

    return {
        "success": True,
        "message": "Site message saved",
        "table": SITE_MESSAGES_WINDOWS_TABLE_NAME,
        "plant_id": item.get("plant_id"),
        "window_id": item.get("window_id"),
        "log_id": log_row.id,
    }


@app.get("/api/site-messages/logs")
async def list_site_message_logs(
    date: str = Query(...),
    x_user_role: Optional[str] = Header(None, alias="X-User-Role"),
    db: Session = Depends(get_db),
):
    event_day = _site_message_parse_date(date)
    rows = (
        db.query(SiteMessageLog)
        .filter(SiteMessageLog.event_date == event_day)
        .order_by(SiteMessageLog.created_at.desc(), SiteMessageLog.id.desc())
        .limit(500)
        .all()
    )
    return {
        "date": event_day.isoformat(),
        "items": [
            {
                "id": row.id,
                "username": row.username,
                "user_role": row.user_role,
                "site_id": row.site_id,
                "site_id_raw": row.site_id_raw,
                "event_type": row.event_type,
                "raw_message": row.raw_message,
                "event_date": row.event_date.isoformat() if row.event_date else "",
                "start_time": row.start_time,
                "end_time": row.end_time,
                "mw": row.mw,
                "unit": row.unit,
                "reduction_type": row.reduction_type,
                "description": row.description,
                "dynamodb_table": row.dynamodb_table,
                "dynamodb_window_id": row.dynamodb_window_id,
                "status": row.status,
                "error_message": row.error_message,
                "created_at": row.created_at.isoformat() if row.created_at else "",
            }
            for row in rows
        ],
    }


def _parse_optional_datetime(value: Any) -> Optional[datetime]:
    raw = str(value or "").strip()
    if not raw:
        return None
    try:
        parsed = datetime.fromisoformat(raw.replace("Z", "+00:00"))
        if parsed.tzinfo is None:
            parsed = parsed.replace(tzinfo=timezone.utc)
        return parsed
    except Exception:
        return None


@app.post("/api/wbes-notification-logs")
async def create_wbes_notification_log(
    payload: WbesNotificationLogRequest,
    x_user_role: Optional[str] = Header(None, alias="X-User-Role"),
    x_user_name: Optional[str] = Header(None, alias="X-User-Name"),
    db: Session = Depends(get_db),
):
    utility = str(payload.utility or "").strip()
    message = str(payload.message or "").strip()
    notification_key = str(payload.id or "").strip()
    if not utility:
        raise HTTPException(status_code=400, detail="utility is required")
    if not message:
        raise HTTPException(status_code=400, detail="message is required")
    if payload.block < 1 or payload.block > 96:
        raise HTTPException(status_code=400, detail="block must be between 1 and 96")
    if not notification_key:
        notification_key = "|".join(
            [
                utility,
                payload.date.isoformat(),
                str(payload.block),
                str(payload.interval or ""),
                str(payload.fileName or ""),
                message,
            ]
        )

    existing = (
        db.query(WbesNotificationLog)
        .filter(WbesNotificationLog.notification_key == notification_key)
        .first()
    )
    if existing:
        return {"success": True, "duplicate": True, "log_id": existing.id}

    row = WbesNotificationLog(
        notification_key=notification_key[:1024],
        username=str(x_user_name or "").strip()[:128] or None,
        user_role=str(x_user_role or "").strip()[:32] or None,
        utility=utility[:128],
        notification_date=payload.date,
        block=int(payload.block),
        interval=str(payload.interval or "").strip()[:64] or None,
        oa_remc=payload.oa_remc,
        as_value=payload.as_value,
        total=payload.total,
        status=str(payload.status or "").strip()[:64] or None,
        file_name=str(payload.fileName or "").strip()[:500] or None,
        message=message,
        notification_created_at=_parse_optional_datetime(payload.createdAt),
    )
    try:
        db.add(row)
        db.commit()
        db.refresh(row)
    except Exception as exc:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Failed to store WBES notification log: {exc}") from exc

    return {"success": True, "duplicate": False, "log_id": row.id}


@app.get("/api/wbes-notification-logs")
async def list_wbes_notification_logs(
    notification_date: Optional[date] = Query(None),
    limit: int = Query(100, ge=1, le=500),
    db: Session = Depends(get_db),
):
    query = db.query(WbesNotificationLog)
    if notification_date is not None:
        query = query.filter(WbesNotificationLog.notification_date == notification_date)
    rows = (
        query
        .order_by(WbesNotificationLog.created_at.desc(), WbesNotificationLog.id.desc())
        .limit(limit)
        .all()
    )
    items = []
    for row in rows:
        items.append({
            "id": row.id,
            "notification_key": row.notification_key,
            "username": row.username,
            "user_role": row.user_role,
            "utility": row.utility,
            "date": row.notification_date.isoformat() if row.notification_date else "",
            "block": row.block,
            "interval": row.interval or "",
            "oa_remc": row.oa_remc,
            "as": row.as_value,
            "total": row.total,
            "status": row.status or "",
            "fileName": row.file_name or "",
            "message": row.message,
            "createdAt": row.notification_created_at.isoformat() if row.notification_created_at else "",
            "created_at": row.created_at.isoformat() if row.created_at else "",
        })
    return {"items": items, "total": len(items)}


def _normalize_plant_key(value: Any) -> str:
    return re.sub(r"[^a-z0-9]", "", str(value or "").strip().lower())

def _plant_id_candidates(raw: str) -> List[str]:
    base = str(raw or "").strip()
    if not base:
        return []
    no_space = re.sub(r"\s+", "", base)
    candidates = [
        base,
        base.upper(),
        base.lower(),
        base.title(),
        base.capitalize(),
        no_space,
        no_space.upper(),
        no_space.lower(),
        no_space.title(),
    ]
    seen = set()
    ordered = []
    for c in candidates:
        if c and c not in seen:
            ordered.append(c)
            seen.add(c)
    return ordered

def _find_ddb_item_by_plant_id(table: Any, plant_id: str) -> Dict[str, Any]:
    """Try exact and common-case variants; fall back to a small scan for case-insensitive match."""
    for candidate in _plant_id_candidates(plant_id):
        try:
            response = table.get_item(Key={"plant_id": candidate})
            item = _normalize_ddb_item(response.get("Item"))
            if item:
                return item
        except Exception:
            continue

    target_key = _normalize_plant_key(plant_id)
    if not target_key:
        return {}
    try:
        last_evaluated_key = None
        pages = 0
        while pages < 5:
            kwargs = {}
            if last_evaluated_key:
                kwargs["ExclusiveStartKey"] = last_evaluated_key
            response = table.scan(**kwargs)
            items = response.get("Items") or []
            for raw_item in items:
                item = _normalize_ddb_item(raw_item)
                if _normalize_plant_key(item.get("plant_id")) == target_key:
                    return item
            last_evaluated_key = response.get("LastEvaluatedKey")
            pages += 1
            if not last_evaluated_key:
                break
    except Exception:
        return {}
    return {}


def _parse_whatsapp_message(message: str) -> Dict[str, Any]:
    if not message:
        return {}
    parsed: Dict[str, Any] = {}

    lines = [line.strip() for line in re.split(r"[\r\n]+", message) if line.strip()]
    for line in lines:
        match = re.match(r"^([^:]+?)\s*[:\-]\s*(.+)$", line)
        if not match:
            continue
        raw_label = match.group(1).strip().lower()
        value = match.group(2).strip()

        label = re.sub(r"\s+", " ", raw_label)
        if "plant" in label and ("id" in label or "name" in label):
            parsed["plantName"] = value
        elif label.startswith("site"):
            parsed["site"] = value
        elif label.startswith("state"):
            parsed["state"] = value
        elif label.startswith("date"):
            parsed["date"] = value
        elif label.startswith("time"):
            parsed["time"] = value
        elif label.startswith("start"):
            parsed["startTime"] = value
        elif label.startswith("end"):
            parsed["endTime"] = value
        elif "current generation" in label:
            num_match = re.search(r"(\d+(?:\.\d+)?)", value)
            if num_match:
                try:
                    parsed["currentGeneration"] = float(num_match.group(1))
                except ValueError:
                    pass
        elif "expected" in label and "trend" in label:
            parsed["expectedTrend"] = value
        elif "curtailment status" in label:
            parsed["curtailmentStatus"] = value.strip().lower() in {"yes", "true", "1"}
        elif "curtailment reason" in label:
            parsed["curtailmentReason"] = value
        elif "weather" in label:
            parsed["weatherCondition"] = value
        elif "inverter" in label:
            num_match = re.search(r"(\d+(?:\.\d+)?)", value)
            if num_match:
                try:
                    parsed["inverterAvailability"] = float(num_match.group(1))
                except ValueError:
                    pass
        elif "remark" in label:
            parsed["remarks"] = value

    if "date" not in parsed:
        date_match = re.search(r"(\d{4}-\d{2}-\d{2}|\d{1,2}[-/]\d{1,2}[-/]\d{2,4})", message)
        if date_match:
            parsed["date"] = date_match.group(1)
    if "time" not in parsed:
        time_match = re.search(r"\b([01]?\d|2[0-3]):[0-5]\d\b", message)
        if time_match:
            parsed["time"] = time_match.group(0)
    if "currentGeneration" not in parsed:
        gen_match = re.search(r"(\d+(?:\.\d+)?)\s*MW\b", message, re.IGNORECASE)
        if gen_match:
            try:
                parsed["currentGeneration"] = float(gen_match.group(1))
            except ValueError:
                pass
    if "expectedTrend" not in parsed:
        trend_match = re.search(r"\b(increasing|decreasing|stable)\b", message, re.IGNORECASE)
        if trend_match:
            parsed["expectedTrend"] = trend_match.group(1).capitalize()
    if "remarks" not in parsed:
        remarks_match = re.search(r"\bremarks?\b[:\-]\s*(.+)$", message, re.IGNORECASE)
        if remarks_match:
            parsed["remarks"] = remarks_match.group(1).strip()

    # If start/end were not parsed from explicit labels, try best-effort extraction.
    if "startTime" not in parsed:
        m = re.search(r"\bstart\b\s*[:\-]\s*([0-9]{4}-[0-9]{2}-[0-9]{2}\s+[0-9]{1,2}:[0-9]{2})", message, re.IGNORECASE)
        if m:
            parsed["startTime"] = m.group(1).strip()
    if "endTime" not in parsed:
        m = re.search(r"\bend\b\s*[:\-]\s*([0-9]{4}-[0-9]{2}-[0-9]{2}\s+[0-9]{1,2}:[0-9]{2})", message, re.IGNORECASE)
        if m:
            parsed["endTime"] = m.group(1).strip()

    status_match = re.search(r"\b(curtaile?ment|normal)\b", message, re.IGNORECASE)
    if status_match and "curtailmentStatus" not in parsed:
        parsed["curtailmentStatus"] = status_match.group(1).lower().startswith("curtail")
    capacity_match = re.search(r"\bcurtaile?ment\s+(\d+(?:\.\d+)?)\b", message, re.IGNORECASE)
    if capacity_match and "curtailmentCapacity" not in parsed:
        try:
            parsed["curtailmentCapacity"] = float(capacity_match.group(1))
        except ValueError:
            pass

    # Plant status (used by Instant Data card). Prefer explicit "Plant status:" line, else keywords.
    if "plantStatus" not in parsed:
        status_line = None
        for line in lines:
            if str(line).lower().startswith("plant status"):
                status_line = str(line)
                break
        if status_line:
            m = re.search(r"plant\s*status\s*[:\-]\s*([A-Za-z_ -]+)", status_line, re.IGNORECASE)
            if m:
                token = m.group(1).strip().upper().split()[0]
                if token:
                    parsed["plantStatus"] = token
    if "plantStatus" not in parsed:
        if re.search(r"\bshutdown\b", message, re.IGNORECASE):
            parsed["plantStatus"] = "SHUTDOWN"
        elif re.search(r"\bcurtaile?ment\b", message, re.IGNORECASE):
            parsed["plantStatus"] = "CURTAILMENT"
        elif re.search(r"\bnormal\b", message, re.IGNORECASE):
            parsed["plantStatus"] = "NORMAL"

    return parsed


_TRIGGER_REASON_MAP = {
    "abrupt_weather": "ABRUPT_WEATHER",
    "curtailment": "CURTAILMENT",
    "dynamic_start": "DYNAMIC_START",
    "plant_status_change": "PLANT_STATUS_CHANGE",
    "intraday_revision": "INTRADAY_REVISION",
}


def _is_allowed_schedule_reason(value: Any) -> bool:
    text = str(value or "").strip().upper()
    if not text or text == "-":
        return False
    if text in set(_TRIGGER_REASON_MAP.values()):
        return True
    return bool(re.fullmatch(r"INTRADAY_REVISION(?:_R\d+)?", text))


def _sanitize_schedule_reason_plant(plant: str) -> str:
    value = str(plant or "").strip().upper()
    if not value or not re.fullmatch(r"[A-Z0-9_-]{1,32}", value):
        raise HTTPException(status_code=400, detail="Invalid plant")
    if value in {"SHRIMOUR", "SHROMOUR"}:
        return "SIRMOUR"
    if value == "ANJANGOAN":
        return "ANJANGAON"
    if value == "OSEL":
        return "OSEPL"
    return value


def _sanitize_schedule_reason_file_name(schedule_file: str) -> str:
    value = os.path.basename(str(schedule_file or "").strip())
    if not value:
        raise HTTPException(status_code=400, detail="schedule_file is required")
    if len(value) > 255:
        raise HTTPException(status_code=400, detail="Invalid schedule_file")
    if any(ch in value for ch in ["/", "\\", "\x00"]):
        raise HTTPException(status_code=400, detail="Invalid schedule_file")
    return value


def _sanitize_schedule_reason_date(date_str: str) -> str:
    value = str(date_str or "").strip()
    try:
        return datetime.strptime(value, "%Y-%m-%d").date().isoformat()
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid date; expected YYYY-MM-DD")


def _extract_schedule_id_from_name(file_name: str) -> Optional[str]:
    """
    Extract schedule revision/id from schedule-related filenames.

    Many filenames contain dates (e.g. 2026-04-13 / 2026_04_13) before the revision,
    so we must prefer explicit schedule tokens instead of the first number.
    """
    text = str(file_name or "")
    lower = text.lower()

    # Prefer explicit schedule tokens.
    patterns = [
        r"schedule_freeze_from_(\d+)",
        r"schedule_freez_from_(\d+)",
        r"schedule_from_(\d+)",
        r"schedule_(\d+)",
    ]
    for pat in patterns:
        m = re.search(pat, lower)
        if m:
            return m.group(1)

    # Fallback: take the last numeric token (avoid picking year-only tokens when possible).
    nums = re.findall(r"(\d+)", lower)
    if not nums:
        return None
    if len(nums) == 1:
        return nums[0]

    # Prefer the last token that is plausibly a revision (typically <= 3 digits).
    for token in reversed(nums):
        if 1 <= len(token) <= 3:
            return token
    return nums[-1]


def _expand_schedule_reason_prefix(template: str, plant: str, date_str: str) -> str:
    raw = str(template or "").strip()
    if not raw:
        return ""
    plant_text = str(plant or "").strip()
    expanded = (
        raw.replace("{plant}", plant_text)
        .replace("{plant_lower}", plant_text.lower())
        .replace("{date}", date_str)
        .strip()
    )
    return expanded.rstrip("/") + "/"


def _get_schedule_reason_log_prefixes(plant: str, date_str: str) -> List[str]:
    default_prefix = "generated/vedanjay/{plant}/logs/{date}/,generated/{plant}/{plant_lower}/logs/{date}/"
    raw = os.getenv("SCHEDULE_REASON_LOG_PREFIXES", default_prefix)
    # Allow comma or newline separated env values.
    parts = [p.strip() for p in re.split(r"[\r\n,]+", str(raw or "")) if p.strip()]
    plant_aliases = _generated_schedule_plant_folder_aliases(str(plant or "").strip().upper()) or [str(plant or "").strip().upper()]
    prefixes = [
        _expand_schedule_reason_prefix(p, alias, date_str)
        for alias in plant_aliases
        for p in parts
    ]
    prefixes = [p for p in prefixes if p]
    if not prefixes:
        prefixes = [_expand_schedule_reason_prefix(default_prefix, plant, date_str)]
    # Preserve order, dedupe.
    return list(dict.fromkeys(prefixes))


def _extract_trigger_reason_from_text(raw_text: str) -> str:
    text = str(raw_text or "")
    lower_text = text.lower()

    def _normalize_token(token: str) -> str:
        return re.sub(r"[\s-]+", "_", str(token or "").strip().lower())

    # 1) Prefer explicit reason markers from log lines.
    explicit_tokens: List[str] = []
    explicit_tokens.extend(
        re.findall(r"schedule\s+reason\s*:\s*([a-zA-Z_\-\s]+)", text, flags=re.IGNORECASE)
    )
    explicit_tokens.extend(
        re.findall(r"\breason\s*=\s*([a-zA-Z_\-\s]+)", text, flags=re.IGNORECASE)
    )
    for token in explicit_tokens:
        normalized = _normalize_token(token)
        if normalized in _TRIGGER_REASON_MAP:
            return _TRIGGER_REASON_MAP[normalized]
        # Handle common suffix variants like abrupt_weather_change.
        if normalized.startswith("abrupt_weather"):
            return _TRIGGER_REASON_MAP["abrupt_weather"]
        if normalized == "dynamic_start":
            return _TRIGGER_REASON_MAP["dynamic_start"]
        if normalized.startswith("plant_status_change"):
            return _TRIGGER_REASON_MAP["plant_status_change"]
        if normalized.startswith("curtailment"):
            return _TRIGGER_REASON_MAP["curtailment"]

    # 2) Fallback keyword scan.
    if re.search(r"\bplant[_\s-]?status[_\s-]?change(?:\b|[_-])", lower_text):
        return _TRIGGER_REASON_MAP["plant_status_change"]
    if re.search(r"\bdynamic[_\s-]?start\b(?![_\s-]?block)", lower_text):
        return _TRIGGER_REASON_MAP["dynamic_start"]
    if re.search(r"\bcurtailment(?:\b|[_-])", lower_text):
        return _TRIGGER_REASON_MAP["curtailment"]
    if re.search(r"\babrupt[_\s-]?weather(?:\b|[_-])", lower_text):
        return _TRIGGER_REASON_MAP["abrupt_weather"]
    return "-"

def _normalize_schedule_reason_token(token: Any) -> str:
    raw = str(token or "").strip()
    if not raw:
        return "-"
    normalized = re.sub(r"[\s-]+", "_", raw.strip().lower())
    if normalized in _TRIGGER_REASON_MAP:
        return _TRIGGER_REASON_MAP[normalized]
    # Accept arbitrary schedule_reason values from metadata.json
    # Example: intraday_revision -> INTRADAY_REVISION
    if re.fullmatch(r"[a-z0-9_]+", normalized):
        return normalized.upper()
    return "-"


def _extract_trigger_reason_from_metadata_value(value: Any) -> str:
    if isinstance(value, str):
        from_token = _normalize_schedule_reason_token(value)
        if from_token != "-":
            return from_token
        return _extract_trigger_reason_from_text(value)
    if isinstance(value, dict):
        plant_status_value = value.get("plant_status")
        if isinstance(plant_status_value, str) and plant_status_value.strip().upper() == "CURTAILMENT":
            return _TRIGGER_REASON_MAP["curtailment"]
        preferred_keys = [
            "reason",
            "trigger_reason",
            "schedule_reason",
            "schedule_reason_category",
            "triggerReason",
            "scheduleReason",
            "scheduleReasonCategory",
        ]
        for key in preferred_keys:
            if key in value:
                if key in {"schedule_reason", "schedule_reason_category", "scheduleReason", "scheduleReasonCategory"}:
                    explicit_schedule_reason = _normalize_schedule_reason_token(value.get(key))
                    if explicit_schedule_reason != "-":
                        return explicit_schedule_reason
                reason = _extract_trigger_reason_from_metadata_value(value.get(key))
                if reason != "-":
                    return reason
        for _, nested in value.items():
            reason = _extract_trigger_reason_from_metadata_value(nested)
            if reason != "-":
                return reason
    if isinstance(value, list):
        for item in value:
            reason = _extract_trigger_reason_from_metadata_value(item)
            if reason != "-":
                return reason
    return "-"

def _extract_importance_from_metadata_value(value: Any) -> str:
    if isinstance(value, str):
        text = str(value or "").strip().upper()
        if text in {"HIGH", "MEDIUM", "LOW"}:
            return text
        return "-"
    if isinstance(value, dict):
        preferred_keys = [
            "importance",
            "priority",
            "severity",
        ]
        for key in preferred_keys:
            if key in value:
                found = _extract_importance_from_metadata_value(value.get(key))
                if found != "-":
                    return found
        for _, nested in value.items():
            found = _extract_importance_from_metadata_value(nested)
            if found != "-":
                return found
    if isinstance(value, list):
        for item in value:
            found = _extract_importance_from_metadata_value(item)
            if found != "-":
                return found
    return "-"


def _read_s3_text_safe(s3_client: Any, bucket: str, key: str) -> Optional[str]:
    if s3_client is not None and bucket:
        try:
            obj = s3_client.get_object(Bucket=bucket, Key=key)
            return obj["Body"].read().decode("utf-8", errors="replace")
        except Exception:
            pass
    try:
        encoded_key = "/".join(quote(segment) for segment in str(key or "").split("/"))
        url = f"{DEFAULT_TEMPLATE_S3_BASE_URL.rstrip('/')}/{encoded_key}"
        with urlopen(url, timeout=30) as resp:
            return resp.read().decode("utf-8", errors="replace")
    except Exception:
        return None


def _list_s3_keys_safe(
    s3_client: Any,
    bucket: str,
    prefix: str,
    *,
    max_keys: int = 200,
) -> List[str]:
    max_keys = int(max_keys or 0)
    if max_keys <= 0:
        max_keys = 200

    if (s3_client is None or not bucket) and prefix:
        try:
            url = (
                f"{DEFAULT_TEMPLATE_S3_BASE_URL.rstrip('/')}/"
                f"?list-type=2&prefix={quote(prefix)}&max-keys={max_keys}"
            )
            with urlopen(url, timeout=20) as resp:
                xml = resp.read().decode("utf-8", errors="replace")
            root = ElementTree.fromstring(xml)
            keys: List[str] = []
            for node in root.findall(".//{*}Contents"):
                key = node.findtext("{*}Key", default="")
                if key:
                    keys.append(key)
                    if len(keys) >= max_keys:
                        break
            return keys
        except Exception:
            return []

    keys: List[str] = []
    continuation = None
    while True:
        payload: Dict[str, Any] = {"Bucket": bucket, "Prefix": prefix}
        payload["MaxKeys"] = max_keys
        if continuation:
            payload["ContinuationToken"] = continuation
        try:
            response = s3_client.list_objects_v2(**payload)
        except Exception:
            return keys

        for item in response.get("Contents", []) or []:
            key = str(item.get("Key", "")).strip()
            if key:
                keys.append(key)
                if len(keys) >= max_keys:
                    return keys

        if response.get("IsTruncated"):
            continuation = response.get("NextContinuationToken")
            if not continuation:
                break
        else:
            break
    return keys


def _find_trigger_reason_from_s3_logs(
    *,
    s3_client: Any,
    bucket: str,
    plant: str,
    schedule_id: str,
    date_str: str,
) -> str:
    id_regex = re.compile(rf"(?<!\d){re.escape(str(schedule_id))}(?!\d)")
    direct_name_templates = [
        "schedule from {id} block.log",
        "schedule from {id} block.log.txt",
        "schedule_from_{id}.log",
    ]

    for prefix in _get_schedule_reason_log_prefixes(plant, date_str):
        # Step 1: direct key match.
        for name_template in direct_name_templates:
            key = f"{prefix}{name_template.format(id=schedule_id)}"
            text = _read_s3_text_safe(s3_client, bucket, key)
            if text:
                reason = _extract_trigger_reason_from_text(text)
                if reason != "-":
                    return reason

        # Step 2: list and filter by schedule ID.
        candidate_keys = _list_s3_keys_safe(s3_client, bucket, prefix, max_keys=200)
        for key in candidate_keys:
            base = os.path.basename(str(key or ""))
            if not base:
                continue
            if not id_regex.search(base):
                continue
            text = _read_s3_text_safe(s3_client, bucket, key)
            if not text:
                continue
            reason = _extract_trigger_reason_from_text(text)
            if reason != "-":
                return reason

    return "-"


def _find_trigger_reason_from_metadata(
    *,
    s3_client: Any,
    bucket: str,
    plant: str,
    date_str: str,
    schedule_id: Optional[str] = None,
    schedule_file: Optional[str] = None,
) -> str:
    plant_code = str(plant or "").strip().upper()
    plant_aliases = _generated_schedule_plant_folder_aliases(plant_code) or [plant_code]
    plant_folder = None
    plant_lower = None
    if plant_code == "SIRMOUR":
        plant_folder = "Sirmour"
        plant_lower = "sirmour"
    elif plant_code == "GSNP":
        plant_folder = "GSNP"
        plant_lower = "gsnp"

    metadata_keys = []
    for alias in plant_aliases:
        metadata_keys.extend([
            f"generated/vedanjay/{alias}/outputs/{date_str}/metadata.json",
            f"generated/{alias}/outputs/{date_str}/metadata.json",
        ])
    metadata_keys.append(f"outputs/{date_str}/metadata.json")
    if plant_folder and plant_lower:
        metadata_keys.insert(1, f"generated/{plant_folder}/{plant_lower}/outputs/{date_str}/metadata.json")

    if schedule_id or schedule_file:
        schedule_metadata_names: List[str] = []
        # If we know the concrete schedule file name, try its adjacent meta file first.
        if schedule_file:
            base_name = os.path.basename(str(schedule_file or "").strip())
            if base_name.lower().endswith(".csv"):
                schedule_metadata_names.append(re.sub(r"\.csv$", ".meta.json", base_name, flags=re.IGNORECASE))

        if schedule_id:
            schedule_metadata_names.extend([
                f"schedule_freeze_from_{schedule_id}.meta.json",
                f"schedule_freez_from_{schedule_id}.meta.json",
                f"schedule_from_{schedule_id}.meta.json",
                f"schedule_{schedule_id}.meta.json",
            ])

        # Preserve order, dedupe.
        schedule_metadata_names = list(dict.fromkeys([n for n in schedule_metadata_names if n]))
        for name in schedule_metadata_names:
            alias_keys: List[str] = []
            for alias in plant_aliases:
                alias_keys.extend([
                    f"generated/vedanjay/{alias}/outputs/{date_str}/{name}",
                    f"generated/{alias}/outputs/{date_str}/{name}",
                ])
            for offset, key in enumerate(alias_keys):
                metadata_keys.insert(offset, key)
            if plant_folder and plant_lower:
                metadata_keys.insert(
                    len(alias_keys),
                    f"generated/{plant_folder}/{plant_lower}/outputs/{date_str}/{name}",
                )
            metadata_keys.insert(len(alias_keys), f"outputs/{date_str}/{name}")

    for key in metadata_keys:
        text = _read_s3_text_safe(s3_client, bucket, key)
        if not text:
            continue
        try:
            payload = json.loads(text)
        except Exception:
            continue
        reason = _extract_trigger_reason_from_metadata_value(payload)
        if reason != "-":
            return reason
    return "-"

def _normalize_plant_name(value: str) -> str:
    return "".join(ch for ch in str(value or "").strip().lower() if ch.isalnum())

_SCHEDULE_REASON_CACHE_LOCK = Lock()
_SCHEDULE_REASON_CACHE: Dict[str, Any] = {}


@app.get("/api/schedule/reason", response_class=PlainTextResponse)
def get_schedule_trigger_reason(
    plant: str = Query(..., description="Plant code, e.g. GSNP"),
    schedule_file: str = Query(..., description="Schedule file name, e.g. schedule_from_72.csv"),
    date: str = Query(..., description="Date in YYYY-MM-DD format"),
):
    """
    Resolve trigger reason for a schedule file by scanning S3 log files.
    Returns one of: CURTAILMENT, ABRUPT_WEATHER, DYNAMIC_START, or '-'.
    """
    safe_plant = _sanitize_schedule_reason_plant(plant)
    safe_file = _sanitize_schedule_reason_file_name(schedule_file)
    safe_date = _sanitize_schedule_reason_date(date)
    schedule_id = _extract_schedule_id_from_name(safe_file)
    if not schedule_id:
        return "-"

    cache_ttl_seconds = int(os.getenv("SCHEDULE_REASON_CACHE_TTL_SECONDS") or "300")
    if cache_ttl_seconds > 0:
        cache_key = f"{safe_plant}|{safe_file}|{safe_date}"
        now = time.time()
        with _SCHEDULE_REASON_CACHE_LOCK:
            cached = _SCHEDULE_REASON_CACHE.get(cache_key)
            if cached and isinstance(cached, tuple) and len(cached) == 2:
                cached_at, cached_value = cached
                try:
                    if (now - float(cached_at)) <= cache_ttl_seconds:
                        return str(cached_value)
                    _SCHEDULE_REASON_CACHE.pop(cache_key, None)
                except Exception:
                    pass

    s3 = None
    bucket = _derive_s3_bucket_name()
    region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"
    try:
        import boto3  # type: ignore
        if bucket:
            s3 = boto3.client("s3", region_name=region)
    except Exception:
        s3 = None

    metadata_reason = _find_trigger_reason_from_metadata(
        s3_client=s3,
        bucket=bucket,
        plant=safe_plant,
        date_str=safe_date,
        schedule_id=schedule_id,
        schedule_file=safe_file,
    )
    if _is_allowed_schedule_reason(metadata_reason):
        if cache_ttl_seconds > 0:
            with _SCHEDULE_REASON_CACHE_LOCK:
                if len(_SCHEDULE_REASON_CACHE) > 5000:
                    _SCHEDULE_REASON_CACHE.clear()
                _SCHEDULE_REASON_CACHE[cache_key] = (time.time(), metadata_reason)
        return metadata_reason

    reason = _find_trigger_reason_from_s3_logs(
        s3_client=s3,
        bucket=bucket,
        plant=safe_plant,
        schedule_id=schedule_id,
        date_str=safe_date,
    )
    if _is_allowed_schedule_reason(reason):
        if cache_ttl_seconds > 0:
            with _SCHEDULE_REASON_CACHE_LOCK:
                if len(_SCHEDULE_REASON_CACHE) > 5000:
                    _SCHEDULE_REASON_CACHE.clear()
                _SCHEDULE_REASON_CACHE[cache_key] = (time.time(), reason)
        return reason
    if cache_ttl_seconds > 0:
        with _SCHEDULE_REASON_CACHE_LOCK:
            if len(_SCHEDULE_REASON_CACHE) > 5000:
                _SCHEDULE_REASON_CACHE.clear()
            _SCHEDULE_REASON_CACHE[cache_key] = (time.time(), "-")
    return "-"


@app.get("/api/schedule/metadata")
def get_schedule_metadata(
    plant: str = Query(..., description="Plant code, e.g. GSNP"),
    schedule_file: str = Query(..., description="Schedule file name, e.g. schedule_from_72.csv"),
    date: str = Query(..., description="Date in YYYY-MM-DD format"),
):
    safe_plant = _sanitize_schedule_reason_plant(plant)
    safe_file = _sanitize_schedule_reason_file_name(schedule_file)
    safe_date = _sanitize_schedule_reason_date(date)
    schedule_id = _extract_schedule_id_from_name(safe_file)

    s3 = None
    bucket = _derive_s3_bucket_name()
    region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"
    try:
        import boto3  # type: ignore
        if bucket:
            s3 = boto3.client("s3", region_name=region)
    except Exception:
        s3 = None

    trigger_reason = _find_trigger_reason_from_metadata(
        s3_client=s3,
        bucket=bucket,
        plant=safe_plant,
        date_str=safe_date,
        schedule_id=schedule_id,
        schedule_file=safe_file,
    )

    importance = "-"
    plant_code = str(safe_plant or "").strip().upper()
    plant_aliases = _generated_schedule_plant_folder_aliases(plant_code) or [plant_code]
    metadata_keys: List[str] = []
    if safe_file:
        base_name = os.path.basename(str(safe_file or "").strip())
        if base_name.lower().endswith(".csv"):
            meta_name = re.sub(r"\.csv$", ".meta.json", base_name, flags=re.IGNORECASE)
            for alias in plant_aliases:
                metadata_keys.append(f"generated/vedanjay/{alias}/outputs/{safe_date}/{meta_name}")
                metadata_keys.append(f"generated/{alias}/outputs/{safe_date}/{meta_name}")
            metadata_keys.append(f"outputs/{safe_date}/{meta_name}")

    for key in metadata_keys:
        text = _read_s3_text_safe(s3, bucket, key)
        if not text:
            continue
        try:
            payload = json.loads(text)
        except Exception:
            continue
        importance = _extract_importance_from_metadata_value(payload)
        if importance != "-":
            break

    return {
        "trigger_reason": trigger_reason if trigger_reason and trigger_reason != "-" else "-",
        "importance": importance,
    }


def _resolve_pipeline_plant_id(requested_plant_id: int, db: Session) -> int:
    """
    Resolve runtime plant ID (from DB) to pipeline-config plant ID.
    Falls back to name-based match so pipeline configs stay stable across DB reseeds.
    """
    configs = load_pipeline_configs()
    try:
        get_plant_config(requested_plant_id, configs)
        return requested_plant_id
    except Exception:
        pass

    db_plant = get_plant(db, requested_plant_id)
    if db_plant:
        requested_name = _normalize_plant_name(getattr(db_plant, "name", ""))
        for plant in configs.get("plants", []):
            if _normalize_plant_name(plant.get("name", "")) == requested_name:
                return int(plant.get("plant_id"))

    raise ValueError(
        f"No template pipeline mapping found for plant_id={requested_plant_id}. "
        "Add/update backend/config/template_pipeline/plants.json."
    )


@app.get("/api/template-transform/active-plants")
async def list_template_transform_active_plants():
    """Return plant ids that have active template definitions configured."""
    configs = load_pipeline_configs()
    templates = configs.get("template_definitions", []) or []
    plants = configs.get("plants", []) or []
    plant_by_id = {
        int(p.get("plant_id")): p
        for p in plants
        if p.get("plant_id") is not None
    }
    active_templates = [t for t in templates if bool(t.get("is_active", False))]
    plant_ids = sorted({int(t.get("plant_id")) for t in active_templates if t.get("plant_id") is not None})
    plant_names = sorted({
        str(plant_by_id.get(int(pid), {}).get("name", "")).strip()
        for pid in plant_ids
        if str(plant_by_id.get(int(pid), {}).get("name", "")).strip()
    })
    return {
        "plant_ids": plant_ids,
        "plant_names": plant_names,
        "templates": [
            {
                "plant_id": int(t.get("plant_id")) if t.get("plant_id") is not None else None,
                "template_id": str(t.get("template_id", "")),
                "version": str(t.get("version", "")),
                "is_active": bool(t.get("is_active", False)),
                "name": str(t.get("name", "")),
                "plant_name": str(plant_by_id.get(int(t.get("plant_id", 0)), {}).get("name", "")).strip(),
            }
            for t in active_templates
        ],
    }


@app.get("/api/template-transform/source-files")
async def list_template_transform_source_files(
    plant_id: Optional[int] = Query(None),
    target_date: date = Query(..., description="Date in YYYY-MM-DD format"),
    db: Session = Depends(get_db),
):
    """List available schedule_from_*.csv files for a date across configured prefixes."""
    try:
        prefixes = [p.strip() for p in DEFAULT_TEMPLATE_S3_PREFIXES.split(",") if p.strip()]

        files: List[Dict[str, str]] = []
        bucket = _derive_s3_bucket_name()
        region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"
        s3_client = None
        try:
            import boto3  # type: ignore
            if bucket:
                s3_client = boto3.client("s3", region_name=region)
        except Exception:
            s3_client = None

        if s3_client is not None and bucket:
            date_str = target_date.isoformat()
            date_prefixes = [f"{prefix.rstrip('/')}/{date_str}/" for prefix in prefixes]
            objects: List[Dict[str, str]] = []
            for prefix in date_prefixes:
                continuation = None
                while True:
                    payload: Dict[str, Any] = {"Bucket": bucket, "Prefix": prefix}
                    if continuation:
                        payload["ContinuationToken"] = continuation
                    try:
                        response = s3_client.list_objects_v2(**payload)
                    except Exception:
                        break
                    for item in response.get("Contents", []) or []:
                        key = str(item.get("Key", "")).strip()
                        if not key:
                            continue
                        last_modified = item.get("LastModified")
                        last_modified_text = ""
                        try:
                            if last_modified is not None:
                                last_modified_text = last_modified.isoformat()
                        except Exception:
                            last_modified_text = ""
                        objects.append({"key": key, "last_modified": last_modified_text})
                    if response.get("IsTruncated"):
                        continuation = response.get("NextContinuationToken")
                        if not continuation:
                            break
                    else:
                        break

            unique = {obj["key"]: obj for obj in objects}
            files = [
                obj for obj in unique.values()
                if obj["key"].lower().endswith(".csv")
                and SCHEDULE_FILE_PREFIX in obj["key"].lower()
            ]
            files.sort(key=lambda item: item.get("last_modified", ""), reverse=True)

        if not files:
            files = list_schedule_files_for_date(
                target_date=target_date,
                s3_base_url=DEFAULT_TEMPLATE_S3_BASE_URL,
                prefixes=prefixes,
            )
        if plant_id is not None:
            try:
                configs = load_pipeline_configs()
                resolved_plant_id = _resolve_pipeline_plant_id(plant_id, db)
                plant = get_plant_config(resolved_plant_id, configs)
                plant_name = str(plant.get("name", "")).strip().lower()
                tokens = {
                    plant_name.replace(" ", ""),
                    plant_name.replace(" ", "_"),
                    plant_name.replace(" ", "-"),
                }
                code_match = re.search(r"\(([A-Za-z0-9_-]+)\)", plant_name, flags=re.IGNORECASE)
                if code_match:
                    tokens.add(code_match.group(1).strip().lower())
                if plant_name.isupper() and 2 <= len(plant_name) <= 6:
                    tokens.add(plant_name.lower())
                filtered = [f for f in files if any(token in f.get("key", "").lower().replace(" ", "") for token in tokens if token)]
                if filtered:
                    files = filtered
            except Exception:
                # Keep original file list if plant filter cannot be applied.
                pass
        return {"date": target_date, "files": files, "total": len(files)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/template-transform/preview", response_model=TemplateTransformPreviewResponse)
async def preview_template_transform(
    request: TemplateTransformRequest,
    db: Session = Depends(get_db),
):
    """
    Preview transformation:
    - Ingest source CSV from S3 key
    - Parse to canonical rows
    - Apply active plant template mapping
    - Validate + return preview rows
    """
    try:
        pipeline_plant_id = _resolve_pipeline_plant_id(request.plant_id, db)
        result = run_preview_pipeline(
            plant_id=pipeline_plant_id,
            target_date=request.date,
            source_file_key=request.source_file_key,
            s3_base_url=DEFAULT_TEMPLATE_S3_BASE_URL,
        )

        template = result["template"]
        validation = result["validation"]
        status = "PREVIEW_VALID" if validation.get("is_valid") else "PREVIEW_FAILED"

        save_transform_audit_run(
            db,
            plant_id=request.plant_id,
            source_file_key=request.source_file_key,
            source_hash=result["source_hash"],
            template_id=str(template.get("template_id", "")),
            template_version=str(template.get("version", "")),
            status=status,
            validation_errors=validation.get("errors", []),
            output_file_key=None,
            output_file_url=None,
            requested_by=request.requested_by,
            run_date=request.date,
        )

        return {
            "plant_id": request.plant_id,
            "template_id": str(template.get("template_id", "")),
            "template_version": str(template.get("version", "")),
            "source_file_key": request.source_file_key,
            "source_hash": result["source_hash"],
            "canonical_row_count": int(result["canonical_row_count"]),
            "validation": validation,
            "target_columns": result["target_columns"],
            "canonical_preview": result["canonical_preview"],
            "transformed_preview": result["transformed_preview"],
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/api/template-transform/generate", response_model=TemplateTransformGenerateResponse)
async def generate_template_transform(
    request: TemplateTransformRequest,
    db: Session = Depends(get_db),
):
    """
    Generate transformation output file.
    Generation is blocked on validation failure.
    """
    try:
        pipeline_plant_id = _resolve_pipeline_plant_id(request.plant_id, db)
        configs = load_pipeline_configs()
        template = get_active_template(pipeline_plant_id, configs)
        plant = get_plant_config(pipeline_plant_id, configs)
        mappings = get_template_mappings(str(template["template_id"]), configs)

        source_text = fetch_s3_text(request.source_file_key, DEFAULT_TEMPLATE_S3_BASE_URL)
        source_hash = compute_source_hash(source_text)
        canonical_rows = parse_to_canonical_rows(source_text)
        expected_blocks = int(template.get("expected_blocks", 96) or 96)
        auto_fill_missing = bool(template.get("auto_fill_missing_blocks", False))
        canonical_rows, missing_blocks = normalize_canonical_blocks(
            canonical_rows,
            expected_blocks=expected_blocks,
            auto_fill_missing=auto_fill_missing,
        )
        validation = validate_canonical_rows(canonical_rows, float(plant.get("capacity", 0)))
        if auto_fill_missing and missing_blocks:
            validation["warnings"].append(format_missing_blocks_summary(missing_blocks))

        if not validation.get("is_valid"):
            run = save_transform_audit_run(
                db,
                plant_id=request.plant_id,
                source_file_key=request.source_file_key,
                source_hash=source_hash,
                template_id=str(template.get("template_id", "")),
                template_version=str(template.get("version", "")),
                status="PREVIEW_FAILED",
                validation_errors=validation.get("errors", []),
                output_file_key=None,
                output_file_url=None,
                requested_by=request.requested_by,
                run_date=request.date,
            )
            raise HTTPException(
                status_code=400,
                detail={
                    "message": "Validation failed. Generation blocked.",
                    "run_id": run.id,
                    "errors": validation.get("errors", []),
                    "warnings": validation.get("warnings", []),
                },
            )

        target_columns, transformed_rows = transform_rows(canonical_rows, mappings)
        revision_source_key = str(request.revision_source_key or request.source_file_key or "").strip()
        schedule_type = "dayahead" if re.search(r"(?:day-ahead|day_ahead|dayahead|_da0\b)", revision_source_key, re.IGNORECASE) else "intraday"
        plant_code = _normalize_plant_code(str(plant.get("name") or plant.get("code") or ""))
        schedule_revision = _resolve_ordered_schedule_revision_number(
            plant_code=plant_code,
            schedule_date=request.date,
            schedule_type=schedule_type,
            source_key=revision_source_key,
        )
        payload = to_csv_bytes(
            target_columns,
            transformed_rows,
            template=template,
            plant=plant,
            target_date=request.date,
            schedule_type=schedule_type,
            schedule_revision=schedule_revision,
        )
        published = publish_output_file(
            payload,
            plant_id=request.plant_id,
            template_id=str(template.get("template_id", "")),
            run_ts=datetime.utcnow(),
        )

        run = save_transform_audit_run(
            db,
            plant_id=request.plant_id,
            source_file_key=request.source_file_key,
            source_hash=source_hash,
            template_id=str(template.get("template_id", "")),
            template_version=str(template.get("version", "")),
            status="GENERATED",
            validation_errors=[],
            output_file_key=published["output_file_key"],
            output_file_url=published["output_file_url"],
            requested_by=request.requested_by,
            run_date=request.date,
        )

        return {
            "run_id": run.id,
            "plant_id": request.plant_id,
            "template_id": str(template.get("template_id", "")),
            "template_version": str(template.get("version", "")),
            "source_file_key": request.source_file_key,
            "source_hash": source_hash,
            "output_file_key": published["output_file_key"],
            "output_file_url": published["output_file_url"],
            "status": "GENERATED",
            "validation": validation,
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.get("/api/template-transform/history")
async def get_template_transform_history(
    plant_id: Optional[int] = Query(None),
    run_date: Optional[date] = Query(None),
    status: Optional[str] = Query(None),
    limit: int = Query(100, ge=1, le=500),
    db: Session = Depends(get_db),
):
    """Get template transformation run history with optional filters."""
    try:
        rows = query_transform_history(
            db,
            plant_id=plant_id,
            run_date=run_date,
            status=status,
            limit=limit,
        )

        history = []
        for row in rows:
            parsed_errors = []
            if row.validation_errors:
                try:
                    parsed_errors = json.loads(row.validation_errors)
                except Exception:
                    parsed_errors = [str(row.validation_errors)]

            history.append(
                {
                    "id": row.id,
                    "plant_id": row.plant_id,
                    "run_date": row.run_date,
                    "source_file_key": row.source_file_key,
                    "source_hash": row.source_hash,
                    "template_id": row.template_id,
                    "template_version": row.template_version,
                    "status": row.status,
                    "validation_errors": parsed_errors,
                    "output_file_key": row.output_file_key,
                    "output_file_url": row.output_file_url,
                    "requested_by": row.requested_by,
                    "created_at": row.created_at,
                }
            )

        return {"items": history, "total": len(history)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/template-transform/download/{run_id}")
async def download_generated_template(
    run_id: int,
    db: Session = Depends(get_db),
):
    """Download generated template artifact for a given run."""
    try:
        run = get_transform_run_by_id(db, run_id)
        if not run:
            raise HTTPException(status_code=404, detail="Run not found")

        if run.status != "GENERATED":
            raise HTTPException(status_code=400, detail="Run is not in GENERATED state")

        local_path = (run.output_file_url or "").strip()
        if local_path and os.path.exists(local_path):
            return FileResponse(
                path=local_path,
                filename=os.path.basename(local_path),
                media_type="text/csv",
            )

        if run.output_file_url and str(run.output_file_url).startswith("http"):
            # Return URL for clients to redirect/open.
            return {"download_url": run.output_file_url}

        raise HTTPException(status_code=404, detail="Generated file not found")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/schedule-readiness/upload-template")
async def upload_schedule_readiness_template(
    request: ScheduleReadinessUploadTemplateRequest,
):
    """Upload confirmed SLDC template to S3 at uploads/vedanjay/{plant}/{date}/."""
    try:
        plant_code = str(request.plant_code or "").strip().upper()
        if not plant_code:
            raise HTTPException(status_code=400, detail="plant_code is required")
        if plant_code in {"SHRIMOUR", "SHROMOUR"}:
            plant_code = "SIRMOUR"
        allowed_codes = {"ANJANGAON", "ANDAD", "BALAKWADA", "BAMKHAL", "BHUPALPALLY", "CME", "GSNP", "GUGARIYAKHEDI", "KASIPET", "KILAJ", "KOTHAGUDEM", "NANDGAON", "OSEPL", "SAWDA", "SIRMOUR", "ZETRIC"}
        if plant_code not in allowed_codes:
            raise HTTPException(status_code=400, detail=f"Unsupported plant_code: {plant_code}")

        csv_text = str(request.csv_text or "")
        if not csv_text.strip():
            raise HTTPException(status_code=400, detail="csv_text is required")

        requested_by = str(request.requested_by or "").strip()
        source_file_key = str(request.source_file_key or "").strip()
        manual_request_id = str(getattr(request, "manual_request_id", "") or "").strip()

        if manual_request_id.startswith("combined-dayahead-download-"):
            for row in _load_readiness_upload_history():
                if not isinstance(row, dict):
                    continue
                if str(row.get("manual_request_id") or "").strip() != manual_request_id:
                    continue
                if str(row.get("plant_code") or "").strip().upper() != plant_code:
                    continue
                if str(row.get("schedule_date") or "").strip() != str(request.schedule_date):
                    continue
                if str(row.get("source_file_key") or "").strip() != source_file_key:
                    continue
                return {
                    "success": True,
                    "message": "Combined day-ahead download marker already stored",
                    "bucket": row.get("bucket") or "",
                    "output_file_key": row.get("output_file_key") or "",
                    "output_file_url": row.get("output_file_url") or "",
                    "uploaded_at": row.get("uploaded_at") or "",
                    "storage_mode": row.get("storage_mode") or "",
                    "error": row.get("error"),
                    "submit_block": row.get("submit_block"),
                    "effective_start_block": row.get("effective_start_block"),
                }

        raw_name = str(request.template_file_name or "").strip()
        safe_name = os.path.basename(raw_name).replace("\\", "_").replace("/", "_")
        if not safe_name:
            safe_name = f"{plant_code}_{request.schedule_date}_sldc_template.csv"
        if not safe_name.lower().endswith(".csv"):
            safe_name = f"{safe_name}.csv"

        bucket = _derive_s3_bucket_name()
        region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"
        uploaded_at = datetime.utcnow()
        upload_token = f"{uploaded_at.strftime('%Y%m%dT%H%M%S%fZ')}_{uuid4().hex[:10]}"
        plant_folder = _special_s3_plant_folder(plant_code)
        key = f"{DEFAULT_READINESS_UPLOAD_PREFIX}/{plant_folder}/{request.schedule_date}/{upload_token}_{safe_name}"
        output_file_key = key
        output_file_url = f"https://{bucket}.s3.{region}.amazonaws.com/{key}" if bucket else ""
        storage_mode = "s3"
        message = "Template uploaded to S3 successfully"
        upload_error = None
        effective_bucket = bucket or "UNKNOWN"

        try:
            if not bucket:
                raise RuntimeError("S3 bucket not configured for readiness uploads")
            try:
                import boto3  # type: ignore
            except Exception as e:
                raise RuntimeError(f"boto3 not available: {e}")

            s3 = boto3.client("s3", region_name=region)
            metadata: Dict[str, str] = {}
            if requested_by:
                metadata["requested_by"] = requested_by[:200]
            if source_file_key:
                metadata["source_file_key"] = source_file_key[:900]
            if manual_request_id:
                metadata["manual_request_id"] = manual_request_id[:200]
            s3.put_object(
                Bucket=bucket,
                Key=key,
                Body=csv_text.encode("utf-8"),
                ContentType="text/csv",
                Metadata=metadata,
            )
        except Exception as e:
            # Fallback: persist locally so upload flow does not fail when IAM creds are missing.
            storage_mode = "local"
            upload_error = str(e)
            effective_bucket = "LOCAL_FALLBACK"
            local_dir = os.path.join(
                READINESS_UPLOAD_LOCAL_DIR,
                plant_folder,
                str(request.schedule_date),
            )
            os.makedirs(local_dir, exist_ok=True)
            local_path = os.path.join(local_dir, safe_name)
            with open(local_path, "w", encoding="utf-8", newline="") as f:
                f.write(csv_text)
            output_file_key = f"local/readiness/{plant_folder}/{request.schedule_date}/{safe_name}"
            output_file_url = local_path
            message = "S3 upload unavailable; template stored in local fallback history"

        history_entry = {
            "id": f"{int(uploaded_at.timestamp() * 1000)}-{uuid4().hex[:8]}",
            "plant_code": plant_code,
            "schedule_date": str(request.schedule_date),
            "template_file_name": safe_name,
            "source_file_key": source_file_key,
            "manual_request_id": manual_request_id,
            "requested_by": requested_by,
            "bucket": effective_bucket,
            "output_file_key": output_file_key,
            "output_file_url": output_file_url,
            # Emit explicit UTC marker so all clients parse this consistently.
            "uploaded_at": uploaded_at.isoformat() + "Z",
            "storage_mode": storage_mode,
            "error": upload_error,
            "csv_text": csv_text,
        }
        history_entry.update(_compute_submit_and_effective_blocks_from_iso(history_entry.get("uploaded_at", ""), plant_code=plant_code))
        _append_readiness_upload_history(history_entry)
        _readiness_dashboard_cache_clear_date(request.schedule_date)

        # Auto-generate/update edited_frozen.csv server-side so it exists even if the browser
        # is closed or the operator doesn't click "Recompute Frozen".
        try:
            if bucket and storage_mode == "s3":
                frozen_prefix = f"frozenschedules/vedanjay/{plant_code}/{request.schedule_date}/"
                edited_key = f"{frozen_prefix}edited_frozen.csv"
                log_key = f"{frozen_prefix}{plant_code}_frozen.log"

                all_rows = _load_readiness_upload_history()
                frozen_csv = _generate_edited_frozen_from_upload_history_rows(
                    plant_code=plant_code,
                    schedule_date=str(request.schedule_date),
                    rows=[r for r in all_rows if isinstance(r, dict)],
                    s3_client=s3,
                    bucket=bucket,
                )
                if frozen_csv:
                    # Clean up legacy per-block frozen files; keep only consolidated artifacts.
                    try:
                        resp = s3.list_objects_v2(Bucket=bucket, Prefix=frozen_prefix)
                        legacy = [
                            it["Key"] for it in resp.get("Contents", [])
                            if re.search(r"schedule_free(?:z|ze)_from_\d+\.(?:csv|log)$", it.get("Key", ""), re.I)
                        ]
                        if legacy:
                            s3.delete_objects(Bucket=bucket, Delete={"Objects": [{"Key": k} for k in legacy]})
                    except Exception:
                        pass

                    # Ensure marker exists for S3 console visibility.
                    try:
                        s3.put_object(Bucket=bucket, Key=frozen_prefix)
                    except Exception:
                        pass

                    s3.put_object(
                        Bucket=bucket,
                        Key=edited_key,
                        Body=frozen_csv.encode("utf-8"),
                        ContentType="text/csv",
                    )

                    computed = _compute_submit_and_effective_blocks_from_iso(str(history_entry.get("uploaded_at") or ""), plant_code=plant_code)
                    log_payload = {
                        "plant_code": plant_code,
                        "schedule_date": str(request.schedule_date),
                        "status": "uploaded",
                        "source_schedule_key": source_file_key,
                        "freeze_time": str(history_entry.get("uploaded_at") or ""),
                        "reason": "AUTO_EDITED_FROZEN_FROM_UPLOAD_HISTORY",
                        "summary": {
                            "submit_block": computed.get("submit_block"),
                            "effective_start_block": computed.get("effective_start_block"),
                            "template_file_name": safe_name,
                            "output_file_key": output_file_key,
                        },
                        "stored_schedule_key": edited_key,
                        "stored_log_key": log_key,
                        "created_at": datetime.utcnow().isoformat(),
                        "storage_mode": "s3",
                        "error": "",
                    }
                    try:
                        s3.put_object(
                            Bucket=bucket,
                            Key=log_key,
                            Body=json.dumps(log_payload, ensure_ascii=False, indent=2).encode("utf-8"),
                            ContentType="application/json",
                        )
                    except Exception:
                        pass
        except Exception:
            # Do not fail the upload endpoint if frozen generation fails.
            pass

        computed_blocks = _compute_submit_and_effective_blocks_from_iso(history_entry.get("uploaded_at", ""), plant_code=plant_code)
        return {
            "success": True,
            "message": message,
            "bucket": effective_bucket,
            "output_file_key": output_file_key,
            "output_file_url": output_file_url,
            "uploaded_at": uploaded_at,
            "storage_mode": storage_mode,
            "error": upload_error,
            "submit_block": computed_blocks.get("submit_block"),
            "effective_start_block": computed_blocks.get("effective_start_block"),
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/schedule-readiness/upload-history")
@app.get("/api/schedule-readiness/uploads/history")
async def get_schedule_readiness_upload_history(
    schedule_date: Optional[date] = Query(None),
    plant_code: Optional[str] = Query(None),
    source_file_key: Optional[str] = Query(None),
    limit: int = Query(200, ge=1, le=2000),
    include_s3: bool = Query(True),
):
    """Get upload confirmation history (persisted even when S3 upload falls back locally)."""
    try:
        rows = _load_readiness_upload_history()
        candidate_plants: List[str] = []
        if schedule_date is not None:
            date_key = schedule_date.isoformat()
            candidate_plants = [
                str(r.get("plant_code", "")).strip().upper()
                for r in rows
                if str(r.get("schedule_date", "")).strip() == date_key
            ]
        else:
            # Use most-recent local rows as a proxy for "active" plants.
            candidate_plants = [
                str(r.get("plant_code", "")).strip().upper()
                for r in sorted(rows, key=lambda r: str(r.get("uploaded_at", "")), reverse=True)[:300]
            ]

        s3_rows: List[Dict[str, Any]] = []
        if include_s3:
            s3_rows = _load_s3_upload_history_rows(
                schedule_date=schedule_date,
                plant_code=plant_code,
                candidate_plants=candidate_plants,
                limit=limit,
            )

        # Merge local persisted history + (optional) S3 discovered rows.
        merged = rows + s3_rows
        deduped: Dict[str, Dict[str, Any]] = {}
        for r in merged:
            key = ""
            row_id = str(r.get("id", "")).strip()
            if row_id:
                key = f"id:{row_id}"
            if not key:
                output_key = str(r.get("output_file_key", "")).strip()
                uploaded_at = str(r.get("uploaded_at", "")).strip()
                if output_key:
                    key = f"{output_key}|{uploaded_at}"
            if not key:
                key = (
                    f"{str(r.get('plant_code','')).strip()}|"
                    f"{str(r.get('schedule_date','')).strip()}|"
                    f"{str(r.get('template_file_name','')).strip()}|"
                    f"{str(r.get('uploaded_at','')).strip()}"
                )
            prev = deduped.get(key)
            if prev is None:
                deduped[key] = r
                continue

            # Prefer rows that actually contain the template body + metadata (csv_text/source_file_key).
            # The "s3_discovered" rows are useful as fallback when we have no local record, but they often
            # contain empty csv_text and empty source_file_key. Do not let those overwrite richer rows.
            prev_text = str(prev.get("csv_text", "") or "").strip()
            curr_text = str(r.get("csv_text", "") or "").strip()
            prev_source = str(prev.get("source_file_key", "") or "").strip()
            curr_source = str(r.get("source_file_key", "") or "").strip()
            prev_requested_by = str(prev.get("requested_by", "") or "").strip()
            curr_requested_by = str(r.get("requested_by", "") or "").strip()
            prev_uploaded_by = str(prev.get("uploaded_by", "") or "").strip()
            curr_uploaded_by = str(r.get("uploaded_by", "") or "").strip()

            if prev_text and not curr_text:
                continue
            if curr_text and not prev_text:
                deduped[key] = r
                continue
            if prev_source and not curr_source:
                continue
            if curr_source and not prev_source:
                deduped[key] = r
                continue
            # Preserve uploader identity if one row has it and the other doesn't.
            if (prev_requested_by or prev_uploaded_by) and not (curr_requested_by or curr_uploaded_by):
                continue
            if (curr_requested_by or curr_uploaded_by) and not (prev_requested_by or prev_uploaded_by):
                deduped[key] = r
                continue

            prev_ts = str(prev.get("uploaded_at", ""))
            curr_ts = str(r.get("uploaded_at", ""))
            if curr_ts > prev_ts:
                deduped[key] = r

        filtered = list(deduped.values())

        if schedule_date is not None:
            d = schedule_date.isoformat()
            filtered = [r for r in filtered if str(r.get("schedule_date", "")).strip() == d]

        if plant_code:
            p = str(plant_code).strip().upper()
            filtered = [r for r in filtered if str(r.get("plant_code", "")).strip().upper() == p]

        if source_file_key:
            s = str(source_file_key).strip()
            filtered = [r for r in filtered if str(r.get("source_file_key", "")).strip() == s]

        filtered = sorted(
            filtered,
            key=lambda r: str(r.get("uploaded_at", "")),
            reverse=True,
        )[:limit]

        return {"items": filtered, "total": len(filtered)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/schedule-readiness/dashboard-summary")
def get_schedule_readiness_dashboard_summary(
    date: str = Query(..., min_length=10, max_length=10, description="YYYY-MM-DD"),
    plant_code: Optional[str] = Query(None, max_length=64),
    state: Optional[str] = Query(None, max_length=128),
    limit_per_plant: int = Query(20000, ge=1, le=20000),
):
    """
    Aggregate expensive Readiness screen S3 lookups into one short-cached response.
    The frontend still builds rows with its existing logic and falls back if this fails.
    """
    date_key = str(date or "").strip()
    if not re.fullmatch(r"\d{4}-\d{2}-\d{2}", date_key):
        raise HTTPException(status_code=400, detail="Invalid date format (expected YYYY-MM-DD)")

    bucket = _derive_s3_bucket_name() or str(os.getenv("S3_BUCKET") or "").strip() or "vedanjay-schedules-test-218708247175"
    region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"
    normalized_scope_plant = _normalize_plant_code(str(plant_code or "").strip()) if plant_code else ""
    normalized_scope_state = str(state or "").strip()
    cache_key = f"{bucket}|{region}|{date_key}|{int(limit_per_plant)}|{normalized_scope_plant}|{normalized_scope_state.lower()}"
    cache_ttl = _readiness_dashboard_cache_ttl(date_key)
    with _READINESS_DASHBOARD_CACHE_LOCK:
        cached = _cache_get(_READINESS_DASHBOARD_CACHE, key=cache_key, ttl_seconds=cache_ttl)
    if cached is not None:
        payload = dict(cached)
        payload["cache"] = {"hit": True, "ttl_seconds": cache_ttl}
        return payload

    if not bucket:
        raise HTTPException(status_code=500, detail="S3 bucket not configured")

    try:
        import boto3  # type: ignore
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"boto3 not available: {exc}") from exc

    s3 = boto3.client("s3", region_name=region)
    readiness_known_plant_states = {
        "ANJANGAON": "Madhya Pradesh",
        "ANDAD": "Madhya Pradesh",
        "BALAKWADA": "Madhya Pradesh",
        "BAMKHAL": "Madhya Pradesh",
        "BHUPALPALLY": "Telangana",
        "CME": "Maharashtra",
        "GSNP": "Madhya Pradesh",
        "GUGARIYAKHEDI": "Madhya Pradesh",
        "KASIPET": "Telangana",
        "KILAJ": "Maharashtra",
        "KOTHAGUDEM": "Telangana",
        "NANDGAON": "Madhya Pradesh",
        "OSEPL": "Maharashtra",
        "SAWDA": "Madhya Pradesh",
        "SIRMOUR": "Madhya Pradesh",
        "ZETRIC": "Maharashtra",
    }

    def _normalize_readiness_state(value: Any) -> str:
        raw = str(value or "").strip().lower()
        if raw in {"tl", "telangana"}:
            return "telangana"
        if raw in {"mh", "maharashtra"}:
            return "maharashtra"
        if raw in {"mp", "madhya pradesh", "madhyapradesh"}:
            return "madhya pradesh"
        return raw

    if normalized_scope_plant:
        discovered_codes = [normalized_scope_plant]
    elif normalized_scope_state:
        target_state = _normalize_readiness_state(normalized_scope_state)
        discovered_codes = [
            code for code, plant_state in readiness_known_plant_states.items()
            if _normalize_readiness_state(plant_state) == target_state
        ]
    else:
        try:
            discovered_codes = [
                _normalize_plant_code(code)
                for code in _list_generated_plants(s3_client=s3, bucket=bucket, max_plants=1000)
                if str(code or "").strip()
            ]
        except Exception:
            discovered_codes = []
    discovered_codes = sorted({code for code in discovered_codes if code})

    def _summary_schedule_items_for_code(code: str, schedule_type: str) -> List[Dict[str, Any]]:
        normalized_plant = _normalize_plant_code(code)
        normalized_type = str(schedule_type or "intraday").strip().lower()
        objects: List[Dict[str, str]] = []
        for prefix in _generated_schedule_prefixes_for_plant(normalized_plant, date_key, normalized_type):
            if not _s3_proxy_is_allowed_path(prefix):
                continue
            try:
                objects.extend(_list_s3_objects_paginated(
                    s3_client=s3,
                    bucket=bucket,
                    prefix=prefix,
                    max_items=int(limit_per_plant),
                ))
            except Exception:
                continue

        out: List[Dict[str, Any]] = []
        for obj in objects:
            key = str(obj.get("key") or "").strip()
            if not key.lower().endswith(".csv"):
                continue
            key_is_dayahead = bool(re.search(r"/day-ahead/|/dayahead/|/day_ahead/", key, re.IGNORECASE))
            if normalized_type == "dayahead" and not key_is_dayahead:
                continue
            if normalized_type == "intraday" and key_is_dayahead:
                continue
            revision = _extract_schedule_revision_from_key(key)
            if revision is None:
                continue
            out.append({
                "key": key,
                "last_modified": str(obj.get("last_modified") or "").strip(),
                "revision": int(revision),
            })
        out.sort(key=lambda item: (int(item.get("revision") or 0), str(item.get("last_modified") or ""), str(item.get("key") or "")), reverse=True)
        return out

    intraday_items: List[Dict[str, Any]] = []
    day_ahead_items: List[Dict[str, Any]] = []
    for code in discovered_codes:
        try:
            intraday_items.extend(_summary_schedule_items_for_code(code, "intraday"))
        except Exception:
            pass
        try:
            day_ahead_items.extend(_summary_schedule_items_for_code(code, "dayahead"))
        except Exception:
            pass

    def _plant_code_from_generated_key(key: Any) -> str:
        parts = str(key or "").strip().split("/")
        if len(parts) > 4 and parts[2].strip().lower() == "multiple_generator":
            return _normalize_plant_code(parts[3])
        return _normalize_plant_code(parts[2] if len(parts) > 2 else "")

    generated_plant_codes = sorted({
        _plant_code_from_generated_key(item.get("key"))
        for item in intraday_items
        if _plant_code_from_generated_key(item.get("key"))
    })
    generated_day_ahead_plant_codes = sorted({
        _plant_code_from_generated_key(item.get("key"))
        for item in day_ahead_items
        if _plant_code_from_generated_key(item.get("key"))
    })

    upload_plants = discovered_codes if (normalized_scope_plant or normalized_scope_state) else [
        "BHUPALPALLY", "CME", "GSNP", "KASIPET", "KILAJ", "KOTHAGUDEM",
        "OSEPL", "ANJANGAON", "ANJANGOAN", "SIRMOUR",
    ]
    upload_prefixes = [f"uploads/vedanjay/{plant}/{date_key}/" for plant in upload_plants]
    if "ANJANGAON" in upload_plants and "ANJANGOAN" not in upload_plants:
        upload_prefixes.append(f"uploads/vedanjay/ANJANGOAN/{date_key}/")
    uploaded_objects: List[Dict[str, Any]] = []
    for prefix in upload_prefixes:
        if not _s3_proxy_is_allowed_path(prefix):
            continue
        try:
            uploaded_objects.extend(_list_s3_objects_paginated(s3_client=s3, bucket=bucket, prefix=prefix, max_items=2000))
        except Exception:
            continue

    try:
        local_rows = [
            row for row in _load_readiness_upload_history()
            if str(row.get("schedule_date", "")).strip() == date_key
        ]
        if normalized_scope_plant:
            local_rows = [
                row for row in local_rows
                if _normalize_plant_code(str(row.get("plant_code") or "").strip()) == normalized_scope_plant
            ]
        elif normalized_scope_state:
            scoped_codes = set(discovered_codes)
            local_rows = [
                row for row in local_rows
                if _normalize_plant_code(str(row.get("plant_code") or "").strip()) in scoped_codes
            ]
        s3_rows = _load_s3_upload_history_rows(
            schedule_date=datetime.strptime(date_key, "%Y-%m-%d").date(),
            plant_code=normalized_scope_plant or None,
            candidate_plants=discovered_codes if (normalized_scope_state and not normalized_scope_plant) else [str(row.get("plant_code", "")).strip().upper() for row in local_rows],
            limit=500,
        )
        upload_history_items = sorted(
            [
                row for row in (local_rows + s3_rows)
                if str(row.get("output_file_key", "")).strip().lower().startswith("uploads/vedanjay/")
            ],
            key=lambda row: str(row.get("uploaded_at", "")),
            reverse=True,
        )[:500]
    except Exception:
        upload_history_items = []

    payload = {
        "success": True,
        "date": date_key,
        "bucket": bucket,
        "region": region,
        "generated_plant_codes": generated_plant_codes,
        "generated_day_ahead_plant_codes": generated_day_ahead_plant_codes,
        "schedule_files": intraday_items,
        "day_ahead_files": day_ahead_items,
        "uploaded_objects": uploaded_objects,
        "upload_history_items": upload_history_items,
        "cache": {"hit": False, "ttl_seconds": cache_ttl},
    }
    with _READINESS_DASHBOARD_CACHE_LOCK:
        _cache_set(_READINESS_DASHBOARD_CACHE, key=cache_key, value=payload)
    return payload


# ==================== S3 PROXY (avoids browser CORS on EC2/IP) ====================
class S3ProxyListRequest(BaseModel):
    prefixes: List[str] = []
    limit: int = 5000


class ScheduleListResponseItem(BaseModel):
    key: str
    last_modified: str = ""
    revision: Optional[int] = None


class ScheduleListResponse(BaseModel):
    plant_code: str
    date: str
    schedule_type: str
    prefix: str
    total: int
    items: List[ScheduleListResponseItem]


class ScheduleLatestFilesResponse(BaseModel):
    date: str
    schedule_type: str
    total: int
    items: List[ScheduleListResponseItem]


class SchedulePlantDiscoveryItem(BaseModel):
    plant_code: str
    latest_key: str = ""
    last_modified: str = ""
    revision: Optional[int] = None


class SchedulePlantDiscoveryResponse(BaseModel):
    date: str
    schedule_type: str
    total: int
    items: List[SchedulePlantDiscoveryItem]


def _s3_proxy_is_allowed_path(value: str) -> bool:
    text = str(value or "").strip()
    if not text or len(text) > 1024:
        return False
    # Limit to known app prefixes; prevents accidental exposure of unrelated bucket contents.
    return bool(
        re.match(
            r"^(raw/|generated/|outputs/|uploads/|manual-edits/|frozenschedules/|support-files/|Vedanjay SLDC Schedules/|\d{4}-\d{2}-\d{2}/meter/)",
            text,
            flags=re.IGNORECASE,
        )
    )


_S3_LIST_CACHE_LOCK = Lock()
_S3_LIST_CACHE: Dict[str, Dict[str, Any]] = {}
_S3_LIST_CACHE_TTL_SECONDS = int(os.getenv("S3_PROXY_LIST_CACHE_TTL_SECONDS") or 15)
_S3_LIST_EXECUTOR = ThreadPoolExecutor(max_workers=max(2, min(int(os.getenv("S3_PROXY_LIST_WORKERS") or 8), 32)))

_S3_TEXT_CACHE_LOCK = Lock()
_S3_TEXT_CACHE: Dict[str, Dict[str, Any]] = {}
_S3_TEXT_CACHE_TTL_SECONDS = int(os.getenv("S3_PROXY_TEXT_CACHE_TTL_SECONDS") or 30)
_S3_TEXT_CACHE_MAX_BYTES = int(os.getenv("S3_PROXY_TEXT_CACHE_MAX_BYTES") or (2 * 1024 * 1024))

_READINESS_DASHBOARD_CACHE_LOCK = Lock()
_READINESS_DASHBOARD_CACHE: Dict[str, Dict[str, Any]] = {}
_READINESS_DASHBOARD_CURRENT_TTL_SECONDS = int(os.getenv("READINESS_DASHBOARD_CURRENT_TTL_SECONDS") or 30)
_READINESS_DASHBOARD_PAST_TTL_SECONDS = int(os.getenv("READINESS_DASHBOARD_PAST_TTL_SECONDS") or 300)


def _now_ts() -> float:
    try:
        return datetime.utcnow().timestamp()
    except Exception:
        return 0.0


def _cache_get(cache: Dict[str, Dict[str, Any]], *, key: str, ttl_seconds: int) -> Optional[Any]:
    if not key or ttl_seconds <= 0:
        return None
    entry = cache.get(key)
    if not entry:
        return None
    created = float(entry.get("ts") or 0.0)
    if created <= 0.0:
        return None
    if _now_ts() - created > float(ttl_seconds):
        return None
    return entry.get("value")


def _cache_set(cache: Dict[str, Dict[str, Any]], *, key: str, value: Any) -> None:
    if not key:
        return
    cache[key] = {"ts": _now_ts(), "value": value}


def _readiness_dashboard_cache_ttl(date_key: str) -> int:
    today = datetime.now(ZoneInfo("Asia/Kolkata")).date().isoformat()
    return _READINESS_DASHBOARD_CURRENT_TTL_SECONDS if str(date_key or "").strip() == today else _READINESS_DASHBOARD_PAST_TTL_SECONDS


def _readiness_dashboard_cache_clear_date(date_key: Any) -> None:
    safe_date = str(date_key or "").strip()
    if not safe_date:
        return
    with _READINESS_DASHBOARD_CACHE_LOCK:
        for key in list(_READINESS_DASHBOARD_CACHE.keys()):
            if f"|{safe_date}|" in str(key):
                _READINESS_DASHBOARD_CACHE.pop(key, None)


def _s3_list_cache_key(*, bucket: str, region: str, prefixes: List[str], limit: int) -> str:
    normalized = sorted({str(p or "").strip() for p in (prefixes or []) if str(p or "").strip()})
    payload = {"bucket": bucket, "region": region, "limit": int(limit), "prefixes": normalized}
    encoded = json.dumps(payload, sort_keys=True, separators=(",", ":")).encode("utf-8")
    return hashlib.sha256(encoded).hexdigest()


def _s3_text_cache_key(*, bucket: str, region: str, key: str) -> str:
    payload = {"bucket": bucket, "region": region, "key": str(key or "").strip()}
    encoded = json.dumps(payload, sort_keys=True, separators=(",", ":")).encode("utf-8")
    return hashlib.sha256(encoded).hexdigest()


def _normalize_plant_code(value: str) -> str:
    code = str(value or "").strip().upper()
    if code == "ZTRIC":
        return "ZETRIC"
    if code in {"SHRIMOUR", "SHROMOUR"}:
        return "SIRMOUR"
    if code == "ANJANGOAN":
        return "ANJANGAON"
    if code == "OSEL":
        return "OSEPL"
    return code


def _special_s3_plant_folder(value: str) -> str:
    code = _normalize_plant_code(value)
    if code == "ANJANGAON":
        return "ANJANGOAN"
    return code


def _special_s3_plant_folder_aliases(value: str) -> List[str]:
    code = _normalize_plant_code(value)
    aliases: List[str] = []
    for item in (_special_s3_plant_folder(code), code):
        if item and item not in aliases:
            aliases.append(item)
    return aliases


def _generated_schedule_plant_folder_aliases(plant_code: str) -> List[str]:
    code = _normalize_plant_code(plant_code)
    if code == "ANJANGAON":
        return ["ANJANGAON", "ANJANGOAN"]
    return [code] if code else []


def _generated_schedule_prefixes_for_plant(plant_code: str, date_key: str, schedule_type: str = "intraday") -> List[str]:
    code = _normalize_plant_code(plant_code)
    date_text = str(date_key or "").strip()
    if not code or not date_text:
        return []
    if code == "ZETRIC":
        if str(schedule_type or "").strip().lower() == "dayahead":
            return [
                f"generated/vedanjay/multiple_generator/ZTRIC/{date_text}/Day-ahead/",
            ]
        return [f"generated/vedanjay/multiple_generator/ZTRIC/{date_text}/"]
    suffix = "Day-ahead/" if str(schedule_type or "").strip().lower() == "dayahead" else ""
    return [
        f"generated/vedanjay/{folder}/outputs/{date_text}/{suffix}"
        for folder in _generated_schedule_plant_folder_aliases(code)
    ]


def _raw_plant_folder_aliases(plant_code: str) -> List[str]:
    code = _normalize_plant_code(plant_code)
    if code == "ANJANGAON":
        return ["ANJANGAON", "ANJANGOAN"]
    return [code] if code else []


_SCHEDULE_FROM_RE = re.compile(r"schedule_(?:free(?:z|ze)_)?from_(\d+)(?:[_-][A-Za-z0-9]+)*\.csv$", re.IGNORECASE)


def _extract_schedule_revision_from_key(key: str) -> Optional[int]:
    name = os.path.basename(str(key or "").strip())
    match = _SCHEDULE_FROM_RE.search(name)
    if not match:
        return None
    try:
        rev = int(match.group(1))
    except Exception:
        return None
    return rev if 1 <= rev <= 96 else rev


def _list_generated_schedule_revision_items(
    *,
    plant_code: str,
    schedule_date: date,
    schedule_type: str,
    limit: int = 8000,
) -> List[Dict[str, Any]]:
    normalized_plant = _normalize_plant_code(plant_code)
    date_key = schedule_date.isoformat() if isinstance(schedule_date, date) else str(schedule_date or "").strip()
    normalized_type = str(schedule_type or "intraday").strip().lower()
    if normalized_type not in {"intraday", "dayahead"}:
        return []

    prefixes = [
        prefix
        for prefix in _generated_schedule_prefixes_for_plant(normalized_plant, date_key, normalized_type)
        if _s3_proxy_is_allowed_path(prefix)
    ]
    prefixes = list(dict.fromkeys(prefixes))
    if not prefixes:
        return []

    bucket = _derive_s3_bucket_name()
    region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"
    if not bucket:
        return []

    try:
        import boto3  # type: ignore
    except Exception:
        return []

    s3 = boto3.client("s3", region_name=region)
    objects: List[Dict[str, str]] = []
    for prefix in prefixes:
        objects.extend(_list_s3_objects_paginated(s3_client=s3, bucket=bucket, prefix=prefix, max_items=int(limit)))

    items: List[Dict[str, Any]] = []
    for obj in objects:
        key = str(obj.get("key") or "").strip()
        if not key.lower().endswith(".csv"):
            continue
        key_is_dayahead = bool(re.search(r"/day-ahead/|/dayahead/|/day_ahead/", key, re.IGNORECASE))
        if normalized_type == "dayahead" and not key_is_dayahead:
            continue
        if normalized_plant == "ZETRIC" and normalized_type == "dayahead":
            if not re.search(r"/Day-ahead/schedule_from_22\.csv$", key, re.IGNORECASE):
                continue
        if normalized_type == "intraday" and key_is_dayahead:
            continue
        revision = _extract_schedule_revision_from_key(key)
        if revision is None:
            continue
        items.append({
            "key": key,
            "last_modified": str(obj.get("last_modified") or "").strip(),
            "revision": int(revision),
        })
    return items


def _resolve_ordered_schedule_revision_number(
    *,
    plant_code: str,
    schedule_date: date,
    schedule_type: str,
    source_key: str,
) -> Optional[int]:
    items = _list_generated_schedule_revision_items(
        plant_code=plant_code,
        schedule_date=schedule_date,
        schedule_type=schedule_type,
    )
    ordered = sorted(
        items,
        key=lambda item: (
            int(item.get("revision") or 0),
            str(item.get("key") or ""),
        ),
    )

    revision_by_key: Dict[str, int] = {}
    current_position = 0
    last_revision_token: Optional[int] = None
    for item in ordered:
        revision_token = int(item.get("revision") or 0)
        if last_revision_token != revision_token:
            current_position += 1
            last_revision_token = revision_token
        revision_by_key[str(item.get("key") or "").strip()] = current_position

    lookup_key = str(source_key or "").strip()
    if lookup_key in revision_by_key:
        return revision_by_key[lookup_key]

    target_revision = _extract_schedule_revision_from_key(lookup_key)
    if target_revision is None:
        return None

    unique_tokens = sorted({int(item.get("revision") or 0) for item in ordered if item.get("revision") is not None})
    if target_revision in unique_tokens:
        return unique_tokens.index(target_revision) + 1
    return 1


def _pick_latest_csv(objects: List[Dict[str, str]], *, prefer_suffix: Optional[str] = None) -> Optional[Dict[str, str]]:
    """Pick the most recently modified CSV object from a list of {key,last_modified} dicts."""
    csvs = [o for o in (objects or []) if str(o.get("key") or "").lower().endswith(".csv")]
    if not csvs:
        return None

    def _parse_last_modified_ts(value: Any) -> float:
        """
        UI sorts S3 objects by `Date.parse(lastModified)`; emulate that here so we pick the
        same meter CSV as the frontend DSM preview.
        """
        raw = str(value or "").strip()
        if not raw:
            return 0.0
        # Handle "YYYY-MM-DDTHH:MM:SSZ" and the default boto3 string "YYYY-MM-DD HH:MM:SS+00:00"
        try:
            norm = raw.replace("Z", "+00:00")
            dt = datetime.fromisoformat(norm)
            if dt.tzinfo is None:
                dt = dt.replace(tzinfo=timezone.utc)
            return float(dt.timestamp())
        except Exception:
            return 0.0

    def sort_key(o: Dict[str, str]) -> Tuple[float, str]:
        return (_parse_last_modified_ts(o.get("last_modified")), str(o.get("key") or ""))

    ordered = sorted(csvs, key=sort_key, reverse=True)
    if prefer_suffix:
        suf = str(prefer_suffix or "").lower()
        for o in ordered:
            if str(o.get("key") or "").lower().endswith(suf):
                return o
    return ordered[0]


def _list_generated_plants(
    *,
    s3_client: Any,
    bucket: str,
    max_plants: int = 100,
) -> List[str]:
    """List plant codes under generated/vedanjay/<PLANT>/ using S3 delimiter prefixes."""
    if not s3_client or not bucket:
        return []
    prefix = "generated/vedanjay/"
    plants: List[str] = []
    token: Optional[str] = None
    while True:
        kwargs: Dict[str, Any] = {"Bucket": bucket, "Prefix": prefix, "Delimiter": "/", "MaxKeys": 1000}
        if token:
            kwargs["ContinuationToken"] = token
        resp = s3_client.list_objects_v2(**kwargs)  # type: ignore
        for common in (resp.get("CommonPrefixes") or []):
            pfx = str(common.get("Prefix") or "")
            if not pfx.lower().startswith(prefix.lower()):
                continue
            # expected: generated/vedanjay/<PLANT>/
            remainder = pfx[len(prefix):].strip("/")
            if not remainder:
                continue
            code = remainder.split("/", 1)[0].strip().upper()
            if code == "MULTIPLE_GENERATOR":
                code = "ZETRIC"
            if code and code not in plants:
                plants.append(code)
            if len(plants) >= max_plants:
                return plants
        if not resp.get("IsTruncated"):
            break
        token = resp.get("NextContinuationToken")
        if not token:
            break
    return plants


def _list_s3_objects_paginated(
    *,
    s3_client: Any,
    bucket: str,
    prefix: str,
    max_items: int = 5000,
) -> List[Dict[str, str]]:
    """
    List S3 objects with continuation tokens so we never miss schedule revisions.
    Returns list of {key,last_modified}.
    """
    if not s3_client or not bucket or not prefix:
        return []
    out: List[Dict[str, str]] = []
    token: Optional[str] = None
    while True:
        kwargs: Dict[str, Any] = {"Bucket": bucket, "Prefix": prefix, "MaxKeys": 1000}
        if token:
            kwargs["ContinuationToken"] = token
        resp = s3_client.list_objects_v2(**kwargs)  # type: ignore
        contents = resp.get("Contents") or []
        for obj in contents:
            key = str(obj.get("Key") or "").strip()
            if not key:
                continue
            if not _s3_proxy_is_allowed_path(key):
                continue
            last_modified = str(obj.get("LastModified") or "").strip()
            out.append({"key": key, "last_modified": last_modified})
            if len(out) >= max_items:
                return out
        if not resp.get("IsTruncated"):
            break
        token = resp.get("NextContinuationToken")
        if not token:
            break
    return out


@app.get("/api/schedules/list", response_model=ScheduleListResponse)
def list_generated_schedules(
    plant: str = Query(..., min_length=1, max_length=32, description="Plant code, e.g. KASIPET"),
    date: str = Query(..., min_length=10, max_length=10, description="YYYY-MM-DD"),
    type: str = Query("intraday", description="intraday | dayahead"),
    limit: int = Query(8000, ge=1, le=20000),
):
    """
    List generated schedule revisions from S3 for a single plant/date.
    This endpoint is revision-safe (uses S3 continuation tokens).

    Intraday: generated/vedanjay/<PLANT>/outputs/<DATE>/schedule_from_*.csv
    Day-ahead: generated/vedanjay/<PLANT>/outputs/<DATE>/Day-ahead/schedule_from_*.csv
    """
    plant_code = _normalize_plant_code(plant)
    date_key = str(date or "").strip()
    if not re.fullmatch(r"\d{4}-\d{2}-\d{2}", date_key):
        raise HTTPException(status_code=400, detail="Invalid date format (expected YYYY-MM-DD)")

    schedule_type = str(type or "intraday").strip().lower()
    if schedule_type not in {"intraday", "dayahead"}:
        raise HTTPException(status_code=400, detail="Invalid type (expected intraday or dayahead)")

    # Guardrail: prevent scanning arbitrary S3 prefixes.
    # We only allow "safe" plant codes and restrict scanning to `generated/vedanjay/<PLANT>/outputs/<DATE>/...`
    if not re.fullmatch(r"[A-Z0-9_-]{1,32}", plant_code):
        raise HTTPException(status_code=400, detail=f"Invalid plant code: {plant_code}")

    prefixes = _generated_schedule_prefixes_for_plant(plant_code, date_key, schedule_type)
    prefixes = [prefix for prefix in prefixes if _s3_proxy_is_allowed_path(prefix)]
    if not prefixes:
        raise HTTPException(status_code=400, detail="Prefix not allowed")

    bucket = _derive_s3_bucket_name()
    region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"

    try:
        import boto3  # type: ignore
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"boto3 not available: {exc}") from exc

    if not bucket:
        raise HTTPException(status_code=500, detail="S3 bucket not configured")

    raw_items = _list_generated_schedule_revision_items(
        plant_code=plant_code,
        schedule_date=date_key,
        schedule_type=schedule_type,
        limit=int(limit),
    )
    items = [
        ScheduleListResponseItem(
            key=str(item.get("key") or "").strip(),
            last_modified=str(item.get("last_modified") or "").strip(),
            revision=int(item.get("revision") or 0),
        )
        for item in raw_items
        if str(item.get("key") or "").strip()
    ]

    # Sort newest revision first, then last_modified, then key.
    items.sort(
        key=lambda r: (
            int(r.revision or 0),
            str(r.last_modified or ""),
            str(r.key or ""),
        ),
        reverse=True,
    )

    return ScheduleListResponse(
        plant_code=plant_code,
        date=date_key,
        schedule_type=schedule_type,
        prefix=prefixes[-1] if prefixes else "",
        total=len(items),
        items=items,
    )


@app.get("/api/schedules/latest-files", response_model=ScheduleLatestFilesResponse)
def list_latest_generated_schedule_files(
    plant: str = Query(..., min_length=1, max_length=512, description="Plant code or comma-separated plant codes"),
    date: str = Query(..., min_length=10, max_length=10, description="YYYY-MM-DD"),
    type: str = Query("intraday", description="intraday | dayahead"),
    limit_per_plant: int = Query(2000, ge=1, le=20000),
):
    started = time.monotonic()
    date_key = str(date or "").strip()
    if not re.fullmatch(r"\d{4}-\d{2}-\d{2}", date_key):
        raise HTTPException(status_code=400, detail="Invalid date format (expected YYYY-MM-DD)")

    schedule_type = str(type or "intraday").strip().lower()
    if schedule_type not in {"intraday", "dayahead"}:
        raise HTTPException(status_code=400, detail="Invalid type (expected intraday or dayahead)")

    plant_codes = [
        _normalize_plant_code(item)
        for item in re.split(r"[,|]", str(plant or ""))
        if str(item or "").strip()
    ]
    plant_codes = list(dict.fromkeys([code for code in plant_codes if re.fullmatch(r"[A-Z0-9_-]{1,32}", code)]))
    if not plant_codes:
        raise HTTPException(status_code=400, detail="Plant is required")

    items: List[ScheduleListResponseItem] = []
    for plant_code in plant_codes:
        raw_items = _list_generated_schedule_revision_items(
            plant_code=plant_code,
            schedule_date=date_key,
            schedule_type=schedule_type,
            limit=int(limit_per_plant),
        )
        raw_items.sort(
            key=lambda item: (
                int(item.get("revision") or 0),
                str(item.get("last_modified") or ""),
                str(item.get("key") or ""),
            ),
            reverse=True,
        )
        items.extend([
            ScheduleListResponseItem(
                key=str(item.get("key") or "").strip(),
                last_modified=str(item.get("last_modified") or "").strip(),
                revision=int(item.get("revision") or 0),
            )
            for item in raw_items
            if str(item.get("key") or "").strip()
        ])

    print(
        f"[timing] schedules.latest-files type={schedule_type} plants={len(plant_codes)} "
        f"items={len(items)} elapsed_ms={int((time.monotonic() - started) * 1000)}"
    )
    return ScheduleLatestFilesResponse(
        date=date_key,
        schedule_type=schedule_type,
        total=len(items),
        items=items,
    )


@app.get("/api/schedules/plants", response_model=SchedulePlantDiscoveryResponse)
def list_generated_schedule_plants(
    date: str = Query(..., min_length=10, max_length=10, description="YYYY-MM-DD"),
    type: str = Query("intraday", description="intraday | dayahead"),
    limit: int = Query(200, ge=1, le=1000),
):
    """
    Discover plant codes that have at least one generated schedule for a date.

    This is used by the Readiness UI so plants that exist in S3 but are missing
    from the DB seed list still show up.
    """
    try:
        date_key = str(date or "").strip()
        if not re.fullmatch(r"\d{4}-\d{2}-\d{2}", date_key):
            raise HTTPException(status_code=400, detail="Invalid date format (expected YYYY-MM-DD)")

        schedule_type = str(type or "intraday").strip().lower()
        if schedule_type not in {"intraday", "dayahead"}:
            raise HTTPException(status_code=400, detail="Invalid type (expected intraday or dayahead)")

        bucket = _derive_s3_bucket_name() or str(os.getenv("S3_BUCKET") or "").strip() or "vedanjay-schedules-test-218708247175"
        region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"
        if not bucket:
            raise HTTPException(status_code=500, detail="S3 bucket not configured")

        try:
            import boto3  # type: ignore
        except Exception as exc:
            raise HTTPException(status_code=500, detail=f"boto3 not available: {exc}") from exc

        try:
            from botocore.exceptions import ClientError  # type: ignore
        except Exception:
            ClientError = Exception  # type: ignore

        try:
            s3 = boto3.client("s3", region_name=region)
            plants = _list_generated_plants(s3_client=s3, bucket=bucket, max_plants=min(1000, int(limit)))
        except ClientError as exc:  # type: ignore[misc]
            msg = str(exc)
            code = ""
            try:
                err = (exc.response or {}).get("Error", {}) or {}  # type: ignore[attr-defined]
                code = str(err.get("Code") or "").strip()
            except Exception:
                code = ""
            if code in {"AccessDenied", "403"}:
                raise HTTPException(status_code=403, detail=f"S3 access denied for bucket {bucket}") from exc
            raise HTTPException(status_code=502, detail=f"Failed to list S3 prefixes: {msg}") from exc
        except Exception as exc:
            raise HTTPException(status_code=502, detail=f"Failed to list S3 prefixes: {exc}") from exc

        suffix = "Day-ahead/" if schedule_type == "dayahead" else ""

        # Aggregate by canonical plant code so aliases like SHRIMOUR/SHROMOUR
        # collapse into a single SIRMOUR entry for the UI.
        latest_by_plant: Dict[str, SchedulePlantDiscoveryItem] = {}
        for plant_code in plants:
            raw_plant_code = str(plant_code or "").strip().upper()
            canonical_plant_code = _normalize_plant_code(raw_plant_code)
            objects = []
            for prefix in _generated_schedule_prefixes_for_plant(canonical_plant_code, date_key, schedule_type):
                if not _s3_proxy_is_allowed_path(prefix):
                    continue
                try:
                    objects.extend(_list_s3_objects_paginated(s3_client=s3, bucket=bucket, prefix=prefix, max_items=5000))
                except Exception:
                    continue
            if not objects:
                continue

            # When listing intraday schedules we use the broader date prefix
            # `generated/.../outputs/<DATE>/` which also includes `Day-ahead/` children.
            # Filter to the requested schedule type so day-ahead files don't appear in intraday lists.
            filtered_objects: List[Dict[str, str]] = []
            for obj in objects or []:
                key = str(obj.get("key") or "").strip()
                lower = key.lower()
                is_day_ahead_path = (
                    "/day-ahead/" in lower or "/dayahead/" in lower or "/day_ahead/" in lower
                )
                if schedule_type == "dayahead":
                    if canonical_plant_code != "ZETRIC" and not is_day_ahead_path:
                        continue
                else:
                    if is_day_ahead_path:
                        continue
                filtered_objects.append(obj)

            if canonical_plant_code == "ZETRIC" and schedule_type == "dayahead":
                filtered_objects = [
                    item for item in filtered_objects
                    if re.search(r"/Day-ahead/schedule_from_22\.csv$", str(item.get("key") or ""), re.IGNORECASE)
                ]
                obj = _pick_latest_csv(filtered_objects)
            else:
                obj = _pick_latest_csv(filtered_objects, prefer_suffix="_da0.csv" if schedule_type == "dayahead" else None)
            if not obj:
                continue
            key = str(obj.get("key") or "").strip()
            if not key:
                continue
            rev = _extract_schedule_revision_from_key(key)
            candidate = SchedulePlantDiscoveryItem(
                plant_code=canonical_plant_code,
                latest_key=key,
                last_modified=str(obj.get("last_modified") or "").strip(),
                revision=rev,
            )
            existing = latest_by_plant.get(canonical_plant_code)
            if existing is None:
                latest_by_plant[canonical_plant_code] = candidate
                continue
            candidate_rev = int(candidate.revision or -1)
            existing_rev = int(existing.revision or -1)
            if candidate_rev > existing_rev:
                latest_by_plant[canonical_plant_code] = candidate
                continue
            if candidate_rev == existing_rev and str(candidate.last_modified or "") > str(existing.last_modified or ""):
                latest_by_plant[canonical_plant_code] = candidate

        items: List[SchedulePlantDiscoveryItem] = list(latest_by_plant.values())

        # Sort by plant code for stable UI.
        items.sort(key=lambda r: str(r.plant_code or ""))

        return SchedulePlantDiscoveryResponse(
            date=date_key,
            schedule_type=schedule_type,
            total=len(items),
            items=items,
        )
    except HTTPException:
        raise
    except Exception as exc:
        # Make debugging easier from the browser Network panel.
        raise HTTPException(status_code=500, detail=f"/api/schedules/plants failed: {exc}") from exc


VEDANJAY_SLDC_SCHEDULES_PREFIX = "Vedanjay SLDC Schedules"


def _normalize_vedanjay_sldc_plant_code(value: str) -> str:
    code = re.sub(r"[^A-Za-z0-9_-]", "", str(value or "").strip()).upper()
    if code == "OSEL":
        return "OSEPL"
    if code in {"SHRIMOUR", "SHROMOUR"}:
        return "SIRMOUR"
    if code == "ANJANGOAN":
        return "ANJANGAON"
    if code == "KASIPETH":
        return "KASIPET"
    if code == "KOTHAGUDAM":
        return "KOTHAGUDEM"
    return code


def _sanitize_vedanjay_sldc_filename(filename: str) -> str:
    base = os.path.basename(str(filename or "").strip()) or "schedule"
    base = re.sub(r"[^A-Za-z0-9._-]+", "_", base).strip("._")
    return base or "schedule"


def _vedanjay_sldc_prefix(plant_code: str, schedule_date: str) -> str:
    plant = _normalize_vedanjay_sldc_plant_code(plant_code)
    return f"{VEDANJAY_SLDC_SCHEDULES_PREFIX}/{plant}/{schedule_date}/"


def _vedanjay_sldc_latest_pointer_key(plant_code: str, schedule_date: str) -> str:
    return f"{_vedanjay_sldc_prefix(plant_code, schedule_date)}latest.json"


def _vedanjay_sldc_header_key(value: Any) -> str:
    return re.sub(r"[^a-z0-9]+", "", str(value or "").strip().lower())


def _vedanjay_sldc_parse_number(value: Any) -> Optional[float]:
    if value is None:
        return None
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        if math.isfinite(float(value)):
            return float(value)
        return None
    text = str(value).strip()
    if not text or text in {"-", "--", "NA", "N/A"}:
        return None
    text = text.replace(",", "")
    match = re.search(r"-?\d+(?:\.\d+)?", text)
    if not match:
        return None
    try:
        parsed = float(match.group(0))
    except Exception:
        return None
    return parsed if math.isfinite(parsed) else None


def _vedanjay_sldc_parse_block(value: Any) -> Optional[int]:
    if value is None:
        return None
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        number = int(value)
        return number if 1 <= number <= 96 else None

    text = str(value).strip()
    if not text:
        return None
    if re.fullmatch(r"\d{1,3}", text):
        number = int(text)
        return number if 1 <= number <= 96 else None

    time_match = re.search(r"(\d{1,2}):(\d{2})", text)
    if not time_match:
        return None
    hour = int(time_match.group(1))
    minute = int(time_match.group(2))
    if hour == 24 and minute == 0:
        return 96
    if not (0 <= hour <= 23 and 0 <= minute <= 59):
        return None
    block = (hour * 60 + minute) // 15 + 1
    return block if 1 <= block <= 96 else None


def _vedanjay_sldc_read_rows(filename: str, content: bytes) -> List[List[Any]]:
    ext = os.path.splitext(str(filename or ""))[1].lower()
    if ext == ".csv":
        text = content.decode("utf-8-sig", errors="replace")
        return [list(row) for row in csv.reader(io.StringIO(text))]
    if ext == ".xlsx":
        try:
            from openpyxl import load_workbook  # type: ignore
        except Exception as exc:
            raise HTTPException(status_code=500, detail=f"openpyxl is not available: {exc}") from exc
        workbook = load_workbook(io.BytesIO(content), data_only=True, read_only=True)
        sheet = workbook.active
        return [list(row) for row in sheet.iter_rows(values_only=True)]
    raise HTTPException(status_code=400, detail="Only .csv and .xlsx files are allowed")


def _vedanjay_sldc_preferred_schedule_column(keys: List[str], plant_code: str) -> Tuple[Optional[int], str]:
    plant = _normalize_vedanjay_sldc_plant_code(plant_code)
    if plant in {"SIRMOUR", "ANJANGAON", "BAMKHAL"}:
        match = next((idx for idx, key in enumerate(keys) if "forecast" in key and "block" not in key), None)
        return match, "Forecast"
    if plant in {"KASIPET", "KOTHAGUDEM", "BHUPALPALLY"}:
        match = next(
            (
                idx
                for idx, key in enumerate(keys)
                if "station" in key and "schedule" in key and "block" not in key
            ),
            None,
        )
        return match, "Station Schedule"
    if plant == "OSEPL":
        exact_headers = {"schedule", "schedulemw"}
        match = next((idx for idx, key in enumerate(keys) if key in exact_headers), None)
        if match is None:
            match = next(
                (
                    idx
                    for idx, key in enumerate(keys)
                    if "schedule" in key and "station" not in key and "block" not in key
                ),
                None,
            )
        return match, "Schedule"
    return None, ""


def _vedanjay_sldc_join_header_rows(primary: List[Any], secondary: Optional[List[Any]] = None) -> List[str]:
    width = max(len(primary or []), len(secondary or []))
    out: List[str] = []
    for idx in range(width):
        top = _vedanjay_sldc_header_key(primary[idx] if idx < len(primary or []) else "")
        bottom = _vedanjay_sldc_header_key(secondary[idx] if secondary and idx < len(secondary) else "")
        if top and bottom:
            out.append(f"{top}{bottom}")
        else:
            out.append(top or bottom)
    return out


def _vedanjay_sldc_find_columns(rows: List[List[Any]], plant_code: str) -> Tuple[int, int, int, int]:
    block_headers = {"block", "blockno", "blocknumber", "timeblock", "timeblockno", "srno", "sno"}
    mw_headers = {
        "mw",
        "scheduledmw",
        "schedulemw",
        "schedulemw",
        "implementedmw",
        "implementedsl dcmw".replace(" ", ""),
        "implementedsl dcschedule".replace(" ", ""),
        "sldcschedule",
        "sldcschedulemw",
        "generationmw",
        "forecastmw",
    }
    for row_index, row in enumerate(rows[:40]):
        next_row = rows[row_index + 1] if row_index + 1 < len(rows) else None
        key_sets: List[Tuple[List[str], int]] = [
            (_vedanjay_sldc_join_header_rows(row), row_index + 1),
        ]
        if next_row is not None:
            key_sets.append((_vedanjay_sldc_join_header_rows(row, next_row), row_index + 2))

        for keys, data_start_index in key_sets:
            block_col = next((idx for idx, key in enumerate(keys) if key in block_headers), None)
            mw_col, required_column_name = _vedanjay_sldc_preferred_schedule_column(keys, plant_code)
            if block_col is None:
                block_col = next((idx for idx, key in enumerate(keys) if "block" in key and "mw" not in key), None)
            if required_column_name and block_col is not None and mw_col is not None and block_col != mw_col:
                return row_index, block_col, mw_col, data_start_index
            if required_column_name:
                continue
            mw_col = next((idx for idx, key in enumerate(keys) if key in mw_headers), None)
            if mw_col is None:
                preferred = ("scheduled", "schedule", "implemented", "forecast", "mw")
                for token in preferred:
                    match = next((idx for idx, key in enumerate(keys) if token in key and "block" not in key), None)
                    if match is not None:
                        mw_col = match
                        break
            if block_col is not None and mw_col is not None and block_col != mw_col:
                return row_index, block_col, mw_col, data_start_index
    plant = _normalize_vedanjay_sldc_plant_code(plant_code)
    _, required_column_name = _vedanjay_sldc_preferred_schedule_column([], plant)
    if required_column_name:
        raise HTTPException(status_code=400, detail=f"Could not find Block/Time Block and required {required_column_name} column for {plant}")
    raise HTTPException(status_code=400, detail="Could not find Block/Time Block and MW/Schedule columns")


def _parse_vedanjay_sldc_schedule(filename: str, content: bytes, plant_code: str = "") -> List[Dict[str, Any]]:
    rows = _vedanjay_sldc_read_rows(filename, content)
    if not rows:
        raise HTTPException(status_code=400, detail="Uploaded file is empty")

    _, block_col, mw_col, data_start_index = _vedanjay_sldc_find_columns(rows, plant_code)
    by_block: Dict[int, float] = {}
    for row in rows[data_start_index:]:
        block_value = row[block_col] if block_col < len(row) else None
        mw_value = row[mw_col] if mw_col < len(row) else None
        block = _vedanjay_sldc_parse_block(block_value)
        mw = _vedanjay_sldc_parse_number(mw_value)
        if block is None or mw is None:
            continue
        by_block[block] = mw

    if len(by_block) != 96:
        missing = [str(block) for block in range(1, 97) if block not in by_block]
        detail = f"Expected 96 blocks, parsed {len(by_block)}"
        if missing:
            detail += f"; missing blocks: {', '.join(missing[:12])}{'...' if len(missing) > 12 else ''}"
        raise HTTPException(status_code=400, detail=detail)

    return [{"block": block, "mw": by_block[block]} for block in range(1, 97)]


_VEDANJAY_SLDC_S3_CLIENTS: Dict[str, Any] = {}
_VEDANJAY_SLDC_S3_CLIENT_LOCK = Lock()


def _get_vedanjay_sldc_s3_client() -> Any:
    try:
        import boto3  # type: ignore
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"boto3 not available: {exc}") from exc
    region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"
    with _VEDANJAY_SLDC_S3_CLIENT_LOCK:
        client = _VEDANJAY_SLDC_S3_CLIENTS.get(region)
        if client is None:
            client = boto3.client("s3", region_name=region)
            _VEDANJAY_SLDC_S3_CLIENTS[region] = client
        return client


def _vedanjay_sldc_validate_scope(plant_code: str, schedule_date: str) -> Tuple[str, str]:
    plant = _normalize_vedanjay_sldc_plant_code(plant_code)
    if not plant:
        raise HTTPException(status_code=400, detail="plant_code is required")
    date_text = str(schedule_date or "").strip()
    try:
        date.fromisoformat(date_text)
    except Exception as exc:
        raise HTTPException(status_code=400, detail="schedule_date must be YYYY-MM-DD") from exc
    return plant, date_text


def _validate_sldc_submission_time(value: str) -> str:
    time_text = str(value or "").strip()
    if not re.fullmatch(r"(?:[01]\d|2[0-3]):[0-5]\d", time_text):
        raise HTTPException(status_code=400, detail="sldc_submission_time must be HH:MM")
    return time_text


@app.post("/api/vedanjay-sldc-schedules/upload")
async def upload_vedanjay_sldc_schedule(
    file: UploadFile = File(...),
    plant_code: str = Form(...),
    plant_name: Optional[str] = Form(None),
    schedule_date: str = Form(...),
    state: str = Form(...),
    sldc_submission_time: str = Form(...),
    uploader: Optional[str] = Form(None),
    uploader_employee_id: Optional[str] = Form(None),
    uploader_name: Optional[str] = Form(None),
    uploader_role: Optional[str] = Form(None),
):
    plant, date_text = _vedanjay_sldc_validate_scope(plant_code, schedule_date)
    plant_name_text = str(plant_name or plant).strip()[:256]
    state_text = str(state or "").strip()
    if not state_text or state_text == "Select State":
        raise HTTPException(status_code=400, detail="state is required")
    submission_time = _validate_sldc_submission_time(sldc_submission_time)
    original_name = _sanitize_vedanjay_sldc_filename(file.filename or "schedule")
    ext = os.path.splitext(original_name)[1].lower()
    if ext not in {".csv", ".xlsx"}:
        raise HTTPException(status_code=400, detail="Only .csv and .xlsx files are allowed")

    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Uploaded file is empty")

    parsed_rows = _parse_vedanjay_sldc_schedule(original_name, content, plant)
    bucket = _derive_s3_bucket_name()
    if not bucket:
        raise HTTPException(status_code=500, detail="S3 bucket is not configured")

    uploaded_at_dt = datetime.now(timezone.utc)
    timestamp = uploaded_at_dt.strftime("%Y%m%d_%H%M%S")
    stored_filename = f"{timestamp}_{original_name}"
    key = f"{_vedanjay_sldc_prefix(plant, date_text)}{stored_filename}"
    log_key = f"{key}.metadata.json"
    content_type = "text/csv" if ext == ".csv" else "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    upload_id = str(uuid4())
    uploader_label = str(uploader or "").strip()[:256]
    uploader_details = {
        "employee_id": str(uploader_employee_id or "").strip()[:128],
        "name": str(uploader_name or "").strip()[:128],
        "role": str(uploader_role or "").strip()[:64],
        "label": uploader_label,
    }
    audit_log = {
        "upload_id": upload_id,
        "status": "successful",
        "state": state_text,
        "plant_name": plant_name_text,
        "plant_code": plant,
        "schedule_date": date_text,
        "original_filename": original_name,
        "stored_filename": stored_filename,
        "s3_key": key,
        "sldc_submission_time": submission_time,
        "timezone": "Asia/Kolkata",
        "portal_uploaded_at": uploaded_at_dt.isoformat(),
        "uploaded_by": uploader_details,
        "file_size_bytes": len(content),
        "file_extension": ext,
        "content_type": content_type,
        "file_checksum_sha256": hashlib.sha256(content).hexdigest(),
        "parsed_blocks": len(parsed_rows),
    }

    s3 = _get_vedanjay_sldc_s3_client()
    s3.put_object(
        Bucket=bucket,
        Key=key,
        Body=content,
        ContentType=content_type,
        Metadata={
            "plant_code": plant,
            "schedule_date": date_text,
            "original_filename": original_name,
            "uploader": uploader_label[:128],
            "sldc_submission_time": submission_time,
            "upload_id": upload_id,
        },
    )
    try:
        s3.put_object(
            Bucket=bucket,
            Key=log_key,
            Body=json.dumps(audit_log, ensure_ascii=False, indent=2).encode("utf-8"),
            ContentType="application/json",
        )
    except Exception:
        # Do not leave an unlogged schedule behind when audit-log creation fails.
        try:
            s3.delete_object(Bucket=bucket, Key=key)
        except Exception:
            pass
        raise
    try:
        latest_pointer = {
            "plant_code": plant,
            "schedule_date": date_text,
            "s3_key": key,
            "log_key": log_key,
            "filename": original_name,
            "stored_filename": stored_filename,
            "uploaded_at": uploaded_at_dt.isoformat(),
            "sldc_submission_time": submission_time,
            "uploader": uploader_label,
            "uploaded_by": uploader_details,
            "upload_id": upload_id,
            "bucket": bucket,
        }
        s3.put_object(
            Bucket=bucket,
            Key=_vedanjay_sldc_latest_pointer_key(plant, date_text),
            Body=json.dumps(latest_pointer, ensure_ascii=False, indent=2).encode("utf-8"),
            ContentType="application/json",
        )
    except Exception:
        pass

    return {
        "success": True,
        "found": True,
        "plant_code": plant,
        "schedule_date": date_text,
        "filename": original_name,
        "stored_filename": stored_filename,
        "uploaded_at": uploaded_at_dt.isoformat(),
        "sldc_submission_time": submission_time,
        "uploader": uploader_label,
        "uploaded_by": uploader_details,
        "upload_id": upload_id,
        "s3_key": key,
        "log_key": log_key,
        "bucket": bucket,
        "data": parsed_rows,
        "rows": parsed_rows,
    }


@app.get("/api/vedanjay-sldc-schedules/latest")
def get_latest_vedanjay_sldc_schedule(
    plant_code: str = Query(..., min_length=1),
    schedule_date: str = Query(..., min_length=10, max_length=10),
):
    plant, date_text = _vedanjay_sldc_validate_scope(plant_code, schedule_date)
    bucket = _derive_s3_bucket_name()
    if not bucket:
        raise HTTPException(status_code=500, detail="S3 bucket is not configured")

    prefix = _vedanjay_sldc_prefix(plant, date_text)
    s3 = _get_vedanjay_sldc_s3_client()
    try:
        pointer_obj = s3.get_object(Bucket=bucket, Key=_vedanjay_sldc_latest_pointer_key(plant, date_text))
        pointer_body = pointer_obj.get("Body")
        pointer_content = pointer_body.read() if pointer_body is not None else b""
        pointer = json.loads(pointer_content.decode("utf-8")) if pointer_content else {}
        if isinstance(pointer, dict):
            key = str(pointer.get("s3_key") or "").strip()
            if key and key.startswith(prefix) and os.path.splitext(key)[1].lower() in {".csv", ".xlsx"}:
                obj = s3.get_object(Bucket=bucket, Key=key)
                body = obj.get("Body")
                content = body.read() if body is not None else b""
                filename = key.rsplit("/", 1)[-1]
                metadata = obj.get("Metadata") or {}
                original_filename = str(pointer.get("filename") or metadata.get("original_filename") or "").strip()
                log_key = str(pointer.get("log_key") or f"{key}.metadata.json").strip()
                audit_log: Dict[str, Any] = {}
                try:
                    log_obj = s3.get_object(Bucket=bucket, Key=log_key)
                    log_body = log_obj.get("Body")
                    log_content = log_body.read() if log_body is not None else b""
                    loaded_log = json.loads(log_content.decode("utf-8")) if log_content else {}
                    if isinstance(loaded_log, dict):
                        audit_log = loaded_log
                except Exception:
                    audit_log = {}
                parsed_rows = _parse_vedanjay_sldc_schedule(filename, content, plant)
                uploaded_at = str(pointer.get("uploaded_at") or "").strip()
                if not uploaded_at:
                    last_modified = obj.get("LastModified")
                    uploaded_at = last_modified.isoformat() if hasattr(last_modified, "isoformat") else ""

                return {
                    "success": True,
                    "found": True,
                    "plant_code": plant,
                    "schedule_date": date_text,
                    "filename": original_filename or filename,
                    "stored_filename": str(pointer.get("stored_filename") or filename),
                    "uploaded_at": uploaded_at,
                    "sldc_submission_time": audit_log.get("sldc_submission_time") or pointer.get("sldc_submission_time") or metadata.get("sldc_submission_time") or "",
                    "uploader": (audit_log.get("uploaded_by") or {}).get("label") or pointer.get("uploader") or metadata.get("uploader") or "",
                    "uploaded_by": audit_log.get("uploaded_by") or pointer.get("uploaded_by") or {},
                    "upload_id": audit_log.get("upload_id") or pointer.get("upload_id") or metadata.get("upload_id") or "",
                    "s3_key": key,
                    "log_key": log_key if audit_log else "",
                    "bucket": bucket,
                    "data": parsed_rows,
                    "rows": parsed_rows,
                }
    except Exception:
        pass

    items: List[Dict[str, Any]] = []
    continuation_token: Optional[str] = None
    while True:
        kwargs: Dict[str, Any] = {"Bucket": bucket, "Prefix": prefix, "MaxKeys": 1000}
        if continuation_token:
            kwargs["ContinuationToken"] = continuation_token
        response = s3.list_objects_v2(**kwargs)
        for obj in response.get("Contents") or []:
            key = str(obj.get("Key") or "")
            if os.path.splitext(key)[1].lower() in {".csv", ".xlsx"}:
                items.append(obj)
        if not response.get("IsTruncated"):
            break
        continuation_token = response.get("NextContinuationToken")
        if not continuation_token:
            break

    if not items:
        return {
            "success": True,
            "found": False,
            "plant_code": plant,
            "schedule_date": date_text,
            "data": [],
            "rows": [],
        }

    latest = sorted(
        items,
        key=lambda obj: (
            obj.get("LastModified") or datetime.min.replace(tzinfo=timezone.utc),
            str(obj.get("Key") or ""),
        ),
        reverse=True,
    )[0]
    key = str(latest.get("Key") or "")
    obj = s3.get_object(Bucket=bucket, Key=key)
    body = obj.get("Body")
    content = body.read() if body is not None else b""
    filename = key.rsplit("/", 1)[-1]
    metadata = obj.get("Metadata") or {}
    original_filename = str(metadata.get("original_filename") or "").strip()
    log_key = f"{key}.metadata.json"
    audit_log: Dict[str, Any] = {}
    try:
        log_obj = s3.get_object(Bucket=bucket, Key=log_key)
        log_body = log_obj.get("Body")
        log_content = log_body.read() if log_body is not None else b""
        loaded_log = json.loads(log_content.decode("utf-8")) if log_content else {}
        if isinstance(loaded_log, dict):
            audit_log = loaded_log
    except Exception:
        # Backward compatibility for schedules uploaded before sidecar logs existed.
        audit_log = {}
    parsed_rows = _parse_vedanjay_sldc_schedule(filename, content, plant)
    last_modified = latest.get("LastModified")
    uploaded_at = last_modified.isoformat() if hasattr(last_modified, "isoformat") else ""

    return {
        "success": True,
        "found": True,
        "plant_code": plant,
        "schedule_date": date_text,
        "filename": original_filename or filename,
        "stored_filename": filename,
        "uploaded_at": uploaded_at,
        "sldc_submission_time": audit_log.get("sldc_submission_time") or metadata.get("sldc_submission_time") or "",
        "uploader": (audit_log.get("uploaded_by") or {}).get("label") or metadata.get("uploader") or "",
        "uploaded_by": audit_log.get("uploaded_by") or {},
        "upload_id": audit_log.get("upload_id") or metadata.get("upload_id") or "",
        "s3_key": key,
        "log_key": log_key if audit_log else "",
        "bucket": bucket,
        "data": parsed_rows,
        "rows": parsed_rows,
    }


@app.post("/api/s3/list")
def s3_proxy_list_objects(payload: S3ProxyListRequest):
    """List S3 objects across prefixes via backend (works even when S3 CORS blocks browser)."""
    prefixes = [str(p or "").strip() for p in (payload.prefixes or [])]
    prefixes = [p for p in prefixes if _s3_proxy_is_allowed_path(p)]
    if not prefixes:
        return {"items": []}

    limit = max(1, min(int(payload.limit or 5000), 8000))
    bucket = _derive_s3_bucket_name()
    region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"

    cache_key = _s3_list_cache_key(bucket=bucket, region=region, prefixes=prefixes, limit=limit)
    with _S3_LIST_CACHE_LOCK:
        cached = _cache_get(_S3_LIST_CACHE, key=cache_key, ttl_seconds=_S3_LIST_CACHE_TTL_SECONDS)
    if cached is not None:
        return cached

    s3 = None
    try:
        import boto3  # type: ignore
        if bucket:
            try:
                from botocore.config import Config  # type: ignore
                s3 = boto3.client(
                    "s3",
                    region_name=region,
                    config=Config(
                        connect_timeout=3,
                        read_timeout=5,
                        retries={"max_attempts": 1},
                    ),
                )
            except Exception:
                s3 = None
    except Exception:
        s3 = None

    if s3 is None or not bucket:
        return {
            "items": [],
            "bucket": bucket,
            "region": region,
            "partial": True,
            "scanned_prefixes": 0,
            "error": "s3_client_unavailable",
        }

    merged: Dict[str, Dict[str, str]] = {}
    max_prefixes = max(1, min(int(os.getenv("S3_PROXY_LIST_MAX_PREFIXES", "30") or 30), 80))
    max_items_per_prefix = max(
        50,
        min(limit, int(os.getenv("S3_PROXY_LIST_MAX_ITEMS_PER_PREFIX", "500") or 500), 2000),
    )
    try:
        time_budget_seconds = float(os.getenv("S3_PROXY_LIST_TIME_BUDGET_SECONDS", "8") or 8)
    except Exception:
        time_budget_seconds = 8.0
    try:
        per_prefix_timeout_seconds = float(os.getenv("S3_PROXY_LIST_PREFIX_TIMEOUT_SECONDS", "4") or 4)
    except Exception:
        per_prefix_timeout_seconds = 4.0
    deadline = time.monotonic() + max(2.0, min(time_budget_seconds, 120.0))
    partial = False
    scanned_prefixes = 0

    for prefix in prefixes[:max_prefixes]:
        if time.monotonic() >= deadline:
            partial = True
            break
        scanned_prefixes += 1
        remaining_budget = max(0.1, deadline - time.monotonic())
        prefix_timeout = max(0.1, min(per_prefix_timeout_seconds, remaining_budget))
        future = _S3_LIST_EXECUTOR.submit(
            _list_s3_upload_objects_safe,
            s3_client=s3,
            bucket=bucket,
            prefix=prefix,
            max_items=max_items_per_prefix,
        )
        try:
            prefix_objects = future.result(timeout=prefix_timeout)
        except FutureTimeoutError:
            future.cancel()
            partial = True
            break
        except Exception as exc:
            print(f"S3 proxy list error for {prefix}: {exc}")
            prefix_objects = []
        for obj in prefix_objects:
            if time.monotonic() >= deadline:
                partial = True
                break
            key = str(obj.get("key", "")).strip()
            if not key:
                continue
            if not _s3_proxy_is_allowed_path(key):
                continue
            last_modified = str(obj.get("last_modified", "")).strip()
            prev = merged.get(key)
            if prev is None or last_modified > str(prev.get("last_modified", "")):
                merged[key] = {"key": key, "last_modified": last_modified}
            if len(merged) >= limit:
                break
        if partial or len(merged) >= limit:
            break

    items = list(merged.values())
    items.sort(key=lambda r: str(r.get("last_modified", "")), reverse=True)
    response = {
        "items": items[:limit],
        "bucket": bucket,
        "region": region,
        "partial": partial,
        "scanned_prefixes": scanned_prefixes,
    }
    if not partial:
        with _S3_LIST_CACHE_LOCK:
            # Best-effort cleanup to keep memory bounded.
            if len(_S3_LIST_CACHE) > 2000:
                _S3_LIST_CACHE.clear()
            _cache_set(_S3_LIST_CACHE, key=cache_key, value=response)
    return response


@app.get("/api/s3/text")
async def s3_proxy_get_text(key: str = Query(..., min_length=1, max_length=1024)):
    """Fetch an S3 object as plain text via backend (works even when S3 CORS blocks browser)."""
    key = str(key or "").strip()
    if not _s3_proxy_is_allowed_path(key):
        raise HTTPException(status_code=400, detail="Key not allowed")

    bucket = _derive_s3_bucket_name()
    region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"

    cache_key = _s3_text_cache_key(bucket=bucket, region=region, key=key)
    with _S3_TEXT_CACHE_LOCK:
        cached = _cache_get(_S3_TEXT_CACHE, key=cache_key, ttl_seconds=_S3_TEXT_CACHE_TTL_SECONDS)
    if isinstance(cached, str):
        return PlainTextResponse(content=cached)

    # Prefer boto3 (private buckets / IAM role). Fallback to public HTTPS GET.
    try:
        import boto3  # type: ignore
        from botocore.exceptions import ClientError  # type: ignore
        if bucket:
            s3 = boto3.client("s3", region_name=region)
            try:
                obj = s3.get_object(Bucket=bucket, Key=key)
            except ClientError as exc:
                err = (exc.response or {}).get("Error", {}) or {}
                code = str(err.get("Code", "")).strip()
                if code in {"NoSuchKey", "NotFound", "404"}:
                    raise HTTPException(status_code=404, detail="S3 object not found") from exc
                if code in {"AccessDenied", "403"}:
                    raise HTTPException(status_code=403, detail="S3 access denied") from exc
                raise
            body = obj.get("Body")
            data = body.read() if body is not None else b""
            text = data.decode("utf-8", errors="replace")
            if _S3_TEXT_CACHE_TTL_SECONDS > 0 and len(data) <= _S3_TEXT_CACHE_MAX_BYTES:
                with _S3_TEXT_CACHE_LOCK:
                    if len(_S3_TEXT_CACHE) > 2000:
                        _S3_TEXT_CACHE.clear()
                    _cache_set(_S3_TEXT_CACHE, key=cache_key, value=text)
            return PlainTextResponse(content=text)
    except HTTPException:
        raise
    except Exception:
        pass

    try:
        url = f"{DEFAULT_TEMPLATE_S3_BASE_URL.rstrip('/')}/{quote(key)}"
        with urlopen(url, timeout=30) as resp:
            data = resp.read()
        text = data.decode("utf-8", errors="replace")
        if _S3_TEXT_CACHE_TTL_SECONDS > 0 and len(data) <= _S3_TEXT_CACHE_MAX_BYTES:
            with _S3_TEXT_CACHE_LOCK:
                if len(_S3_TEXT_CACHE) > 2000:
                    _S3_TEXT_CACHE.clear()
                _cache_set(_S3_TEXT_CACHE, key=cache_key, value=text)
        return PlainTextResponse(content=text)
    except HTTPError as e:
        if getattr(e, "code", None) == 404:
            raise HTTPException(status_code=404, detail="S3 object not found") from e
        if getattr(e, "code", None) == 403:
            raise HTTPException(status_code=403, detail="S3 access denied") from e
        raise HTTPException(status_code=502, detail=f"Failed to fetch S3 object: HTTP {e.code}") from e
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Failed to fetch S3 object: {e}")


# ==================== EMAIL SCHEDULER (FASTAPI) ====================

def _email_scheduler_normalize_role(role: Optional[str]) -> str:
    r = str(role or "").strip().lower()
    return "admin" if r == "admin" else "testing"


def _email_scheduler_normalize_user(user: Optional[str]) -> str:
    return str(user or "").strip()[:128]


def _email_scheduler_parse_scheduled_at_utc(*, date_str: str, time_str: str, am_pm: str = "AM") -> datetime:
    date_key = str(date_str or "").strip()
    if not re.fullmatch(r"\d{4}-\d{2}-\d{2}", date_key):
        raise HTTPException(status_code=400, detail="Invalid date (expected YYYY-MM-DD)")

    time_key = str(time_str or "").strip()
    if not re.fullmatch(r"\d{2}:\d{2}", time_key):
        raise HTTPException(status_code=400, detail="Invalid time (expected HH:MM)")

    hour = int(time_key.split(":")[0])
    minute = int(time_key.split(":")[1])
    if hour < 0 or hour > 23 or minute < 0 or minute > 59:
        raise HTTPException(status_code=400, detail="Invalid time value")

    # If the UI passed AM/PM with a 12h clock string, normalize (best-effort).
    # Note: browser <input type="time"> produces 24h time, so we primarily treat HH:MM as 24h.
    ampm = str(am_pm or "").strip().upper()
    if ampm in {"AM", "PM"} and 1 <= hour <= 12:
        if hour == 12 and ampm == "AM":
            hour = 0
        elif hour != 12 and ampm == "PM":
            hour = hour + 12

    ist = ZoneInfo("Asia/Kolkata")
    local_dt = datetime(int(date_key[0:4]), int(date_key[5:7]), int(date_key[8:10]), hour, minute, tzinfo=ist)
    return local_dt.astimezone(timezone.utc)


def _email_scheduler_template_category(template_id: str) -> str:
    t = str(template_id or "").strip().lower()
    if "portal" in t:
        return "Portal Issue"
    if "dsm" in t:
        return "DSM"
    # IMPORTANT: "intraday" contains the substring "day", so intraday must be checked first.
    if "intra" in t or "intraday" in t or "id" in t:
        return "Intraday"
    if "day" in t or "dayahead" in t or "da" in t:
        return "Day-Ahead"
    return "Custom"


EMAIL_SCHEDULER_SYSTEM_USER_DEFAULT = "code.vedanjay"
EMAIL_SCHEDULER_SIGNATURE_NAME_DEFAULT = "Vedanjay Team"
EMAIL_SCHEDULER_DEFAULT_FROM_EMAIL = "forecasting.vppl@gmail.com"


def _email_scheduler_system_user() -> str:
    return str(os.getenv("EMAIL_SCHEDULER_SYSTEM_USER") or EMAIL_SCHEDULER_SYSTEM_USER_DEFAULT).strip() or EMAIL_SCHEDULER_SYSTEM_USER_DEFAULT


def _email_scheduler_signature_name() -> str:
    return str(os.getenv("EMAIL_SCHEDULER_SIGNATURE_NAME") or EMAIL_SCHEDULER_SIGNATURE_NAME_DEFAULT).strip() or EMAIL_SCHEDULER_SIGNATURE_NAME_DEFAULT


def _email_scheduler_normalize_signature_name(value: Any) -> str:
    name = str(value or "").strip()
    if not name or name.upper() == "SYSTEM_CRON" or name == _email_scheduler_system_user():
        return _email_scheduler_signature_name()
    return name


def _email_scheduler_pick_template_for_plant(
    *,
    templates_for_plant: List[Dict[str, Any]],
    template_id: str,
) -> Tuple[Optional[Dict[str, Any]], str]:
    """
    Resolve a template config for a given plant from scheduler metadata.

    Supports "global" selectors like "DA0"/"DA1"/"DSM" by mapping them to the plant-specific
    template ids (e.g. "kothagudem_da0"). This keeps cron auto-emails aligned with the Email
    Scheduler UI defaults (subject/body/to/cc).
    """
    requested = str(template_id or "").strip()
    if not requested:
        return None, ""
    selector = requested.lower()
    templates = list(templates_for_plant or [])
    if not templates:
        return None, requested

    # 1) Exact id match (case-insensitive).
    for tpl in templates:
        tpl_id = str((tpl or {}).get("id") or "").strip()
        if tpl_id and tpl_id.lower() == selector:
            return tpl, tpl_id

    # 2) Map generic selectors to plant-specific ids/labels.
    # Examples:
    # - selector "da0" -> "<plant>_da0" or label startswith "DA0"
    # - selector "da1" -> "<plant>_da1" or label startswith "DA1"
    # - selector "dsm" -> "<plant>_dsm" or label contains "DSM"
    if selector in {"da0", "da1", "dsm"}:
        for tpl in templates:
            tpl_id = str((tpl or {}).get("id") or "").strip()
            tpl_label = str((tpl or {}).get("label") or "").strip().lower()
            if not tpl_id:
                continue
            tid = tpl_id.lower()
            if tid.endswith(f"_{selector}"):
                return tpl, tpl_id
            if selector in {"da0", "da1"} and tpl_label.startswith(selector):
                return tpl, tpl_id
            if selector == "dsm" and "dsm" in (tid + " " + tpl_label):
                return tpl, tpl_id

    # Intraday (common shorthand: "intraday", "id").
    if selector in {"intraday", "intra", "id"} or "intraday" in selector:
        for tpl in templates:
            tpl_id = str((tpl or {}).get("id") or "").strip()
            tpl_label = str((tpl or {}).get("label") or "").strip().lower()
            if not tpl_id:
                continue
            tid = tpl_id.lower()
            if "intra" in tid or tid.endswith("_id") or tid.endswith("_intraday"):
                return tpl, tpl_id
            if "intra" in tpl_label:
                return tpl, tpl_id

    # 3) Label match (e.g. UI passes "DA0 Schedule").
    for tpl in templates:
        tpl_label = str((tpl or {}).get("label") or "").strip().lower()
        tpl_id = str((tpl or {}).get("id") or "").strip()
        if tpl_label and tpl_label == selector and tpl_id:
            return tpl, tpl_id

    return None, requested


def _email_scheduler_log_event(
    *,
    requested_by: str,
    employee_name: str,
    role: str,
    template_id: str,
    plant_code: str,
    mode: str,
    status: str,
    from_email: str,
    to_email: str,
    cc_email: str,
    subject: str,
    scheduled_at: Optional[datetime],
    sent_at: Optional[datetime],
    error_message: Optional[str] = None,
) -> None:
    try:
        db = SessionLocal()
        try:
            entry = EmailSendLog(
                requested_by=str(requested_by or "").strip() or None,
                employee_name=str(employee_name or "").strip() or None,
                role=str(role or "").strip() or None,
                template_id=str(template_id or "").strip() or None,
                plant_code=str(plant_code or "").strip() or None,
                category=_email_scheduler_template_category(template_id),
                mode=str(mode or "").strip() or None,
                status=str(status or "").strip().upper() or "UNKNOWN",
                from_email=str(from_email or "").strip() or None,
                to_email=str(to_email or "").strip() or None,
                cc_email=str(cc_email or "").strip() or None,
                subject=str(subject or "").strip() or None,
                scheduled_at=scheduled_at,
                sent_at=sent_at,
                error_message=str(error_message or "").strip() or None,
            )
            db.add(entry)
            db.commit()
        finally:
            db.close()
    except Exception:
        # Logging must never break email flows.
        return


@app.get("/api/email-scheduler/metadata")
def email_scheduler_metadata(
    x_user_role: Optional[str] = Header(None, alias="X-User-Role"),
    x_user_name: Optional[str] = Header(None, alias="X-User-Name"),
):
    plants, templates_by_plant, meta = load_email_scheduler_metadata()

    role = _email_scheduler_normalize_role(x_user_role)
    _user = _email_scheduler_normalize_user(x_user_name)

    # Flatten templates and group by category so React can show grouped dropdowns.
    groups: Dict[str, List[Dict[str, Any]]] = {}
    for plant_code, templates in (templates_by_plant or {}).items():
        for tpl in (templates or []):
            tpl_id = str(tpl.get("id") or "").strip()
            if not tpl_id:
                continue
            cat = _email_scheduler_template_category(tpl_id)
            # Optional: hide some categories for testing if desired later.
            if role != "admin":
                pass
            groups.setdefault(cat, []).append(
                {
                    "id": tpl_id,
                    "label": str(tpl.get("label") or tpl_id),
                    "category": cat,
                    "plant_code": str(plant_code or "").strip(),
                    "timing_hint": str(tpl.get("timing_hint") or ""),
                    "time_24h": str(tpl.get("time_24h") or ""),
                    "am_pm": str(tpl.get("am_pm") or ""),
                    "subject": str(tpl.get("subject") or ""),
                    "body": str(tpl.get("body") or ""),
                    "default_to": str(tpl.get("default_to") or ""),
                    "default_cc": str(tpl.get("default_cc") or ""),
                    "active": bool(tpl.get("active", True)),
                }
            )

    # Stable ordering for UI.
    for key in list(groups.keys()):
        groups[key].sort(key=lambda r: (str(r.get("plant_code") or ""), str(r.get("label") or ""), str(r.get("id") or "")))

    source = str((meta or {}).get("source") or "").strip()
    source_url = f"email-scheduler:{source}" if source else ""
    return {"plants": plants, "templates": groups, "source_url": source_url}


class EmailSchedulerResolveAttachmentRequest(BaseModel):
    plant_name: str
    template_id: str
    date: str


class EmailSchedulerSupportPreviewSaveRequest(BaseModel):
    plant_code: str
    report_date: str
    file_name: Optional[str] = None
    source_type: Optional[str] = None
    payload: Dict[str, Any]
    xlsx_base64: Optional[str] = None
    xlsx_file_name: Optional[str] = None


def _email_scheduler_support_preview_key(plant_code: str, report_date: date) -> str:
    plant = _normalize_plant_code(plant_code)
    if not plant or not re.fullmatch(r"[A-Z0-9_-]{1,64}", plant):
        raise HTTPException(status_code=400, detail="Invalid plant_code")
    return f"support-files/vedanjay/{plant}/{report_date.isoformat()}/support_preview.json"


def _email_scheduler_support_preview_xlsx_key(plant_code: str, report_date: date, file_name: str) -> str:
    plant = _normalize_plant_code(plant_code)
    if not plant or not re.fullmatch(r"[A-Z0-9_-]{1,64}", plant):
        raise HTTPException(status_code=400, detail="Invalid plant_code")
    safe_name = os.path.basename(str(file_name or "").replace("\\", "/")).strip()
    safe_name = re.sub(r"[^A-Za-z0-9._ -]+", "_", safe_name).strip(" .")
    if not safe_name:
        safe_name = "support_preview.xlsx"
    if not safe_name.lower().endswith(".xlsx"):
        safe_name = re.sub(r"\.[^.]+$", "", safe_name) + ".xlsx"
    return f"support-files/vedanjay/{plant}/{report_date.isoformat()}/{safe_name[:180]}"


def _email_scheduler_s3_client() -> Tuple[Any, str]:
    bucket = _derive_s3_bucket_name()
    region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"
    if not bucket:
        raise HTTPException(status_code=500, detail="S3 bucket not configured")
    try:
        import boto3  # type: ignore
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"boto3 not available: {exc}") from exc
    return boto3.client("s3", region_name=region), bucket


def _email_scheduler_sldc_schedule_prefix(plant_code: str, date_key: str) -> str:
    return _vedanjay_sldc_prefix(plant_code, str(date_key or "").strip())


def _email_scheduler_pick_latest_sldc_schedule(objects: List[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
    candidates = [
        obj
        for obj in (objects or [])
        if re.search(r"\.(csv|xlsx|xlsm)$", str(obj.get("key") or ""), flags=re.IGNORECASE)
    ]
    if not candidates:
        return None

    def sort_key(obj: Dict[str, Any]) -> Tuple[float, str]:
        raw = str(obj.get("last_modified") or obj.get("lastModified") or "").strip()
        timestamp = 0.0
        if raw:
            try:
                parsed = datetime.fromisoformat(raw.replace("Z", "+00:00"))
                if parsed.tzinfo is None:
                    parsed = parsed.replace(tzinfo=timezone.utc)
                timestamp = float(parsed.timestamp())
            except Exception:
                timestamp = 0.0
        return (timestamp, str(obj.get("key") or ""))

    return sorted(candidates, key=sort_key)[-1]


def _email_scheduler_ilios_pv_intraday_attachment_data(
    *,
    s3_client: Any,
    bucket: str,
    date_key: str,
) -> Dict[str, Any]:
    site_files: Dict[str, Tuple[str, bytes]] = {}
    picked_keys: List[str] = []
    revisions: List[int] = []
    missing_sites: List[str] = []

    for site_code, _label, _capacity in ILIOS_PV_SITES:
        prefix = _email_scheduler_sldc_schedule_prefix(site_code, date_key)
        if not prefix or not _s3_proxy_is_allowed_path(prefix):
            missing_sites.append(site_code)
            continue
        objects = _list_s3_objects_paginated(s3_client=s3_client, bucket=bucket, prefix=prefix, max_items=5000)
        pick = _email_scheduler_pick_latest_sldc_schedule(objects)
        if not pick:
            missing_sites.append(site_code)
            continue
        s3_key = str(pick.get("key") or "").strip()
        if not s3_key:
            missing_sites.append(site_code)
            continue
        try:
            obj = s3_client.get_object(Bucket=bucket, Key=s3_key)
            body = obj.get("Body")
            data = body.read() if body is not None else b""
        except Exception as exc:
            raise HTTPException(status_code=502, detail=f"Failed to fetch ILIOS_PV {site_code} attachment: {exc}") from exc
        if not data:
            missing_sites.append(site_code)
            continue
        site_files[site_code] = (os.path.basename(s3_key) or f"{site_code}_{date_key}.csv", data)
        picked_keys.append(s3_key)
        revision = _extract_schedule_revision_from_key(s3_key)
        if revision is not None:
            revisions.append(int(revision))

    if missing_sites:
        raise HTTPException(
            status_code=404,
            detail=f"Missing ILIOS_PV intraday SLDC upload for: {', '.join(missing_sites)}.",
        )

    revision_label = str(max(revisions) if revisions else 1)
    file_bytes = convert_ilios_pv_intraday_files_to_xlsx_bytes(
        site_files,
        report_date=date_key,
        revision=revision_label,
    )
    file_name = f"Final_Scheule_Ilios_PV_{date_key}_{revision_label}.xlsx"
    return {
        "ok": True,
        "plant_code": "ILIOS_PV",
        "file_name": file_name,
        "file_bytes": file_bytes,
        "schedule_type": "intraday",
        "lookup_date": date_key,
        "s3_key": picked_keys[0] if picked_keys else "",
        "attachment_revision_source_key": "|".join(picked_keys),
    }


def _email_scheduler_schedule_bytes_to_csv_text(file_name: str, file_bytes: bytes) -> str:
    lower = str(file_name or "").strip().lower()
    if lower.endswith(".csv"):
        for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
            try:
                return bytes(file_bytes or b"").decode(encoding)
            except UnicodeDecodeError:
                continue
        return bytes(file_bytes or b"").decode("utf-8", errors="replace")
    if lower.endswith((".xlsx", ".xlsm")):
        try:
            from openpyxl import load_workbook  # type: ignore
        except Exception as exc:
            raise HTTPException(status_code=500, detail=f"openpyxl is not available: {exc}") from exc
        workbook = load_workbook(io.BytesIO(file_bytes or b""), data_only=True, read_only=True)
        sheet = workbook.active
        output = io.StringIO()
        writer = csv.writer(output)
        for row in sheet.iter_rows(values_only=True):
            writer.writerow(["" if value is None else value for value in row])
        return output.getvalue()
    return bytes(file_bytes or b"").decode("utf-8", errors="replace")


def _email_scheduler_resolve_schedule_attachment_data(
    *,
    plant_name: str,
    template_id: str,
    date_key: str,
) -> Dict[str, Any]:
    plant_raw = str(plant_name or "").strip()
    if not plant_raw:
        raise HTTPException(status_code=400, detail="Missing plant_name")

    # Accept "Plant Name (CODE)" and "CODE" formats.
    paren = re.search(r"\(([A-Za-z0-9_-]{1,32})\)", plant_raw)
    plant_code = _normalize_plant_code(paren.group(1) if paren else plant_raw)
    if not re.fullmatch(r"[A-Z0-9_-]{1,32}", plant_code):
        raise HTTPException(status_code=400, detail=f"Invalid plant code: {plant_code}")

    template_id = str(template_id or "").strip()
    if not template_id:
        raise HTTPException(status_code=400, detail="Missing template_id")

    date_key = str(date_key or "").strip()
    if not re.fullmatch(r"\d{4}-\d{2}-\d{2}", date_key):
        raise HTTPException(status_code=400, detail="Invalid date format (expected YYYY-MM-DD)")

    cat = _email_scheduler_template_category(template_id).lower()
    if "intra" in cat:
        schedule_type = "intraday"
        suffix = ""
    elif "day" in cat:
        schedule_type = "dayahead"
        suffix = "Day-ahead/"
    else:
        raise HTTPException(status_code=400, detail="This template does not require schedule CSV attachment")

    lookup_date = date_key
    # Day-ahead schedules are generated for the next day (D+1).
    if schedule_type == "dayahead":
        try:
            dt = datetime.strptime(date_key, "%Y-%m-%d").date()
            lookup_date = (dt + timedelta(days=1)).strftime("%Y-%m-%d")
        except Exception:
            lookup_date = date_key
    if schedule_type == "intraday":
        prefix = _email_scheduler_sldc_schedule_prefix(plant_code, lookup_date)
        suffix = ""
    else:
        prefixes = _generated_schedule_prefixes_for_plant(plant_code, lookup_date, "dayahead")
        prefix = next((item for item in prefixes if _s3_proxy_is_allowed_path(item)), "")
    if not prefix or not _s3_proxy_is_allowed_path(prefix):
        raise HTTPException(status_code=400, detail="Prefix not allowed")

    bucket = _derive_s3_bucket_name()
    region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"
    if not bucket:
        raise HTTPException(status_code=500, detail="S3 bucket not configured")

    try:
        import boto3  # type: ignore
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"boto3 not available: {exc}") from exc

    s3 = boto3.client("s3", region_name=region)
    if schedule_type == "intraday" and plant_code == "ILIOS_PV":
        return _email_scheduler_ilios_pv_intraday_attachment_data(
            s3_client=s3,
            bucket=bucket,
            date_key=lookup_date,
        )

    pick = None
    attachment_revision_source_key = ""
    template_key = str(template_id or "").strip().lower()
    if (
        schedule_type == "dayahead"
        and plant_code in {"BHUPALPALLY", "KASIPET", "KOTHAGUDEM"}
        and ("da0" in template_key or "da1" in template_key)
    ):
        manual_key = ""
        manual_lookup_dates = []
        for candidate_date in (lookup_date, date_key):
            if candidate_date and candidate_date not in manual_lookup_dates:
                manual_lookup_dates.append(candidate_date)
        for manual_date in manual_lookup_dates:
            manual_key = _manual_changes_pick_latest_manual_edited_schedule_key(
                org_id="vedanjay",
                plant_code=plant_code,
                schedule_date=manual_date,
                schedule_type="DAY_AHEAD",
                bucket=bucket,
            )
            if manual_key:
                break
        if manual_key:
            pick = {"key": manual_key}
            attachment_revision_source_key = manual_key

    objects = _list_s3_objects_paginated(s3_client=s3, bucket=bucket, prefix=prefix, max_items=5000)
    if schedule_type == "intraday":
        pick = _email_scheduler_pick_latest_sldc_schedule(objects)
    # For day-ahead mail types, pick the expected schedule revision (no fallback).
    # DA0 -> schedule_from_22.csv, DA1 -> schedule_from_88.csv
    if not pick and schedule_type == "dayahead":
        preferred_file = None
        if "da0" in template_key:
            preferred_file = "schedule_from_22.csv"
        elif "da1" in template_key:
            preferred_file = "schedule_from_88.csv"
        if preferred_file:
            pref_lower = preferred_file.lower()
            pick = next(
                (
                    o
                    for o in (objects or [])
                    if str(o.get("key") or "").lower().endswith(f"/{pref_lower}")
                ),
                None,
            )
            if not pick:
                raise HTTPException(
                    status_code=404,
                    detail=f"{preferred_file} not present in S3 for {plant_code} on {lookup_date} (dayahead).",
                )
    if not pick and schedule_type != "intraday":
        # Keep latest-file fallback only for non-DA0/DA1 templates.
        pick = _pick_latest_csv(objects, prefer_suffix=".csv")
    if not pick:
        raise HTTPException(status_code=404, detail=f"No schedule CSV found for {plant_code} on {lookup_date} ({schedule_type}).")

    s3_key = str(pick.get("key") or "").strip()
    if not s3_key:
        raise HTTPException(status_code=404, detail="No schedule CSV key found")

    # Fetch as bytes via boto3 (private bucket).
    try:
        obj = s3.get_object(Bucket=bucket, Key=s3_key)
        body = obj.get("Body")
        data = body.read() if body is not None else b""
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"Failed to fetch S3 attachment: {exc}") from exc

    display_source_key = attachment_revision_source_key or s3_key
    original_name = os.path.basename(display_source_key) or f"{plant_code}_{lookup_date}.csv"
    file_name = _email_scheduler_attachment_display_name(
        plant_code=plant_code,
        template_id=template_id,
        schedule_type=schedule_type,
        source_key=display_source_key,
        original_name=original_name,
        report_date=lookup_date,
        date_already_day_ahead=schedule_type == "dayahead",
    )
    return {
        "ok": True,
        "plant_code": plant_code,
        "file_name": file_name,
        "file_bytes": data or b"",
        "schedule_type": schedule_type,
        "lookup_date": lookup_date,
        "s3_key": s3_key,
        "attachment_revision_source_key": attachment_revision_source_key,
    }


@app.post("/email-scheduler/resolve-s3-schedule-attachment")
async def email_scheduler_resolve_s3_schedule_attachment(
    payload: EmailSchedulerResolveAttachmentRequest,
    x_user_role: Optional[str] = Header(None, alias="X-User-Role"),
    x_user_name: Optional[str] = Header(None, alias="X-User-Name"),
):
    role = _email_scheduler_normalize_role(x_user_role)
    _user = _email_scheduler_normalize_user(x_user_name)
    resolved = _email_scheduler_resolve_schedule_attachment_data(
        plant_name=payload.plant_name,
        template_id=payload.template_id,
        date_key=payload.date,
    )
    file_name = str(resolved.get("file_name") or "schedule.csv")
    file_bytes = bytes(resolved.get("file_bytes") or b"")
    csv_text = _email_scheduler_schedule_bytes_to_csv_text(file_name, file_bytes)
    return {
        "ok": True,
        "file_name": file_name,
        "csv_text": csv_text,
        "file_base64": base64.b64encode(file_bytes).decode("ascii") if file_bytes else "",
        "schedule_type": str(resolved.get("schedule_type") or ""),
        "lookup_date": str(resolved.get("lookup_date") or ""),
        "s3_key": str(resolved.get("s3_key") or ""),
        "role": role,
    }


@app.get("/email-scheduler/support-preview")
def email_scheduler_get_support_preview(
    plant_code: str = Query(..., min_length=1, max_length=64),
    report_date: str = Query(..., min_length=10, max_length=10),
    x_user_role: Optional[str] = Header(None, alias="X-User-Role"),
    x_user_name: Optional[str] = Header(None, alias="X-User-Name"),
):
    role = _email_scheduler_normalize_role(x_user_role)
    _user = _email_scheduler_normalize_user(x_user_name)
    _email_scheduler_role_guard(role=role, admin_only=False)
    plant = _normalize_plant_code(plant_code)
    try:
        parsed_date = date.fromisoformat(str(report_date or "").strip())
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid report_date") from None

    key = _email_scheduler_support_preview_key(plant, parsed_date)
    if not _s3_proxy_is_allowed_path(key):
        raise HTTPException(status_code=400, detail="Support preview key not allowed")
    try:
        s3, bucket = _email_scheduler_s3_client()
        obj = s3.get_object(Bucket=bucket, Key=key)
        body = obj.get("Body")
        content = body.read() if body is not None else b""
    except Exception as exc:
        code = str(getattr(exc, "response", {}).get("Error", {}).get("Code", "")).strip()
        if code in {"NoSuchKey", "NotFound", "404"}:
            return {"ok": True, "found": False}
        raise HTTPException(status_code=502, detail=f"Failed to fetch support preview from S3: {exc}") from exc

    try:
        stored = json.loads(content.decode("utf-8")) if content else {}
    except Exception:
        stored = {}
    payload = stored.get("payload") if isinstance(stored, dict) else {}
    if not isinstance(payload, dict):
        payload = {}
    return {
        "ok": True,
        "found": True,
        "plant_code": str(stored.get("plant_code") or plant) if isinstance(stored, dict) else plant,
        "report_date": str(stored.get("report_date") or parsed_date.isoformat()) if isinstance(stored, dict) else parsed_date.isoformat(),
        "file_name": str(stored.get("file_name") or "") if isinstance(stored, dict) else "",
        "source_type": str(stored.get("source_type") or "") if isinstance(stored, dict) else "",
        "payload": payload,
        "xlsx_file_name": str(stored.get("xlsx_file_name") or "") if isinstance(stored, dict) else "",
        "xlsx_key": str(stored.get("xlsx_key") or "") if isinstance(stored, dict) else "",
        "updated_at": str(stored.get("updated_at") or "") if isinstance(stored, dict) else "",
        "s3_key": key,
    }


@app.post("/email-scheduler/support-preview")
def email_scheduler_save_support_preview(
    payload: EmailSchedulerSupportPreviewSaveRequest,
    x_user_role: Optional[str] = Header(None, alias="X-User-Role"),
    x_user_name: Optional[str] = Header(None, alias="X-User-Name"),
):
    role = _email_scheduler_normalize_role(x_user_role)
    user = _email_scheduler_normalize_user(x_user_name)
    _email_scheduler_role_guard(role=role, admin_only=False)
    plant = _normalize_plant_code(payload.plant_code)
    try:
        parsed_date = date.fromisoformat(str(payload.report_date or "").strip())
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid report_date") from None
    if not plant:
        raise HTTPException(status_code=400, detail="plant_code is required")
    try:
        payload_json = json.dumps(payload.payload or {}, ensure_ascii=True, separators=(",", ":"))
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid support preview payload") from None

    key = _email_scheduler_support_preview_key(plant, parsed_date)
    if not _s3_proxy_is_allowed_path(key):
        raise HTTPException(status_code=400, detail="Support preview key not allowed")
    xlsx_bytes = b""
    xlsx_key = ""
    xlsx_file_name = ""
    raw_xlsx = str(payload.xlsx_base64 or "").strip()
    if raw_xlsx:
        try:
            xlsx_bytes = base64.b64decode(raw_xlsx, validate=True)
        except Exception:
            raise HTTPException(status_code=400, detail="Invalid support preview XLSX payload") from None
        if xlsx_bytes:
            xlsx_file_name = str(payload.xlsx_file_name or payload.file_name or "support_preview.xlsx").strip()
            xlsx_key = _email_scheduler_support_preview_xlsx_key(plant, parsed_date, xlsx_file_name)
            xlsx_file_name = os.path.basename(xlsx_key)
            if not _s3_proxy_is_allowed_path(xlsx_key):
                raise HTTPException(status_code=400, detail="Support preview XLSX key not allowed")
    stored = {
        "plant_code": plant,
        "report_date": parsed_date.isoformat(),
        "file_name": str(payload.file_name or "").strip()[:500],
        "source_type": str(payload.source_type or "").strip()[:64],
        "payload": payload.payload or {},
        "xlsx_file_name": xlsx_file_name,
        "xlsx_key": xlsx_key,
        "updated_by": user or "",
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }
    try:
        s3, bucket = _email_scheduler_s3_client()
        s3.put_object(
            Bucket=bucket,
            Key=key,
            Body=json.dumps(stored, ensure_ascii=True, separators=(",", ":")).encode("utf-8"),
            ContentType="application/json; charset=utf-8",
        )
        if xlsx_bytes and xlsx_key:
            s3.put_object(
                Bucket=bucket,
                Key=xlsx_key,
                Body=xlsx_bytes,
                ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        with _S3_TEXT_CACHE_LOCK:
            _S3_TEXT_CACHE.pop(_s3_text_cache_key(bucket=bucket, region=os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1", key=key), None)
        return {
            "ok": True,
            "plant_code": plant,
            "report_date": parsed_date.isoformat(),
            "s3_key": key,
            "xlsx_key": xlsx_key,
        }
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"Failed to store support preview in S3: {exc}") from exc


def _email_scheduler_role_guard(*, role: str, admin_only: bool = False) -> None:
    if admin_only and role != "admin":
        raise HTTPException(status_code=403, detail="Admin-only action")


async def _email_scheduler_read_upload_bytes(file: Optional[UploadFile]) -> Tuple[Optional[str], Optional[bytes], Optional[str]]:
    if file is None:
        return None, None, None
    data = await file.read()
    name = str(getattr(file, "filename", "") or "").strip() or "attachment.bin"
    ctype = str(getattr(file, "content_type", "") or "").strip() or "application/octet-stream"
    return name, data, ctype


def _email_scheduler_parse_json_payload(text_value: Optional[str]) -> Optional[Dict[str, Any]]:
    raw = str(text_value or "").strip()
    if not raw:
        return None
    try:
        parsed = json.loads(raw)
    except Exception:
        return None
    return parsed if isinstance(parsed, dict) else None


def _email_scheduler_build_dsm_payload_from_s3_for_email(
    *,
    plant_code: str,
    report_date: str,
) -> Optional[Dict[str, Any]]:
    """
    Build the DSM table payload directly from S3 so the client does not need to send
    a large JSON blob for S3-backed DSM mail sends.
    """
    pcode = _normalize_plant_code(plant_code)
    day = str(report_date or "").strip()
    if not pcode or not re.fullmatch(r"\d{4}-\d{2}-\d{2}", day):
        return None

    bucket = _derive_s3_bucket_name()
    region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"
    s3 = None
    if bucket:
        try:
            import boto3  # type: ignore

            s3 = boto3.client("s3", region_name=region)
        except Exception:
            s3 = None

    def _build_single(code: str) -> Optional[Dict[str, Any]]:
        row = None
        if s3 and bucket:
            row = _email_scheduler_build_daily_dsm_row_from_s3(
                s3_client=s3,
                bucket=bucket,
                plant_code=code,
                plant_name=code,
                report_date=day,
            )
        if row:
            columns = _email_scheduler_public_dsm_columns(row)
            variant = "default"
            if code == "OSEPL":
                variant = "osepl"
            elif code == "SIRMOUR":
                variant = "sirmour"
            elif code in _EMAIL_SCHEDULER_TELANGANA_DSM_CODES:
                variant = "multi"
            return {"variant": variant, "columns": columns, "rows": [row]}

        simple = _email_scheduler_build_simple_daily_dsm_table_payload(
            plant_code=code,
            plant_name=code,
            report_date=day,
        )
        return simple if isinstance(simple, dict) else None

    if pcode == "TELANGANA" or pcode in _EMAIL_SCHEDULER_TELANGANA_DSM_CODES:
        rows: List[Dict[str, Any]] = []
        columns: List[str] = []
        for code in _EMAIL_SCHEDULER_TELANGANA_DSM_CODES:
            row = None
            if s3 and bucket:
                row = _email_scheduler_build_daily_dsm_row_from_s3(
                    s3_client=s3,
                    bucket=bucket,
                    plant_code=code,
                    plant_name=code,
                    report_date=day,
                )
            if not row:
                continue
            if not columns:
                columns = _email_scheduler_public_dsm_columns(row)
            rows.append(row)

        if rows:
            return {
                "variant": "multi",
                "columns": columns or (_email_scheduler_public_dsm_columns(rows[0]) if isinstance(rows[0], dict) else list(rows[0].keys())),
                "rows": rows,
            }

        return None

    return _build_single(pcode)


def _email_scheduler_send_now(
    *,
    template_id: str,
    role: str,
    from_email: str,
    to_email: str,
    cc_email: str,
    subject: str,
    body: str,
    employee_name: str,
    dsm_payload: Optional[Dict[str, Any]],
    schedule_attachment: Optional[Tuple[str, bytes]],
    attachment: Optional[Tuple[str, bytes, str]],
    include_employee_mobile: bool = True,
) -> None:
    def _guess_attachment_content_type(name: str) -> str:
        lower = str(name or "").strip().lower()
        if lower.endswith(".xlsx"):
            return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        if lower.endswith(".xls"):
            return "application/vnd.ms-excel"
        if lower.endswith(".csv"):
            return "text/csv"
        if lower.endswith(".pdf"):
            return "application/pdf"
        if lower.endswith(".doc"):
            return "application/msword"
        if lower.endswith(".docx"):
            return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if lower.endswith(".png"):
            return "image/png"
        if lower.endswith(".jpg") or lower.endswith(".jpeg"):
            return "image/jpeg"
        if lower.endswith(".gif"):
            return "image/gif"
        if lower.endswith(".bmp"):
            return "image/bmp"
        if lower.endswith(".webp"):
            return "image/webp"
        return "application/octet-stream"

    attachments: List[EmailAttachment] = []
    if schedule_attachment and schedule_attachment[1]:
        attachments.append(
            EmailAttachment(
                filename=schedule_attachment[0],
                content_bytes=schedule_attachment[1],
                content_type=_guess_attachment_content_type(schedule_attachment[0]),
            )
        )
    if attachment and attachment[1]:
        supplied = str(attachment[2] or "").strip()
        ctype = supplied if supplied and supplied != "application/octet-stream" else _guess_attachment_content_type(attachment[0])
        attachments.append(EmailAttachment(filename=attachment[0], content_bytes=attachment[1], content_type=ctype))

    body = normalize_day_ahead_body(body, template_id)
    ok, msg = send_email_smtp(
        from_email=from_email,
        to_email=to_email,
        cc_email=cc_email,
        subject=subject,
        body_text=body,
        employee_name=_email_scheduler_normalize_signature_name(employee_name),
        include_employee_mobile=include_employee_mobile,
        dsm_payload=dsm_payload,
        attachments=attachments,
        smtp_profile="testing" if str(role or "").strip().lower() != "admin" else "default",
    )
    if not ok:
        raise HTTPException(status_code=502, detail=f"Email send failed: {msg}")


@app.post("/email-scheduler/send-report-now")
async def email_scheduler_send_report_now(
    template_id: str = Form(...),
    plant_code: str = Form(...),
    date: str = Form(...),
    time: str = Form(...),
    am_pm: str = Form("AM"),
    from_email: str = Form(...),
    to_email: str = Form(...),
    cc_email: str = Form(""),
    employee_name: str = Form(""),
    subject: str = Form(...),
    body: str = Form(...),
    auto_send: str = Form("0"),
    portal_issue: str = Form("0"),
    portal_issue_plants: str = Form(""),
    dsm_summary_payload: str = Form(""),
    day_ahead_date_already_adjusted: str = Form("0"),
    schedule_attachment: Optional[UploadFile] = File(None),
    attachment: Optional[UploadFile] = File(None),
    x_user_role: Optional[str] = Header(None, alias="X-User-Role"),
    x_user_name: Optional[str] = Header(None, alias="X-User-Name"),
):
    role = _email_scheduler_normalize_role(x_user_role)
    user = _email_scheduler_normalize_user(x_user_name)

    _email_scheduler_role_guard(role=role, admin_only=False)
    to_email = str(to_email or "").strip()
    cc_email = str(cc_email or "").strip()

    plant = _normalize_plant_code(plant_code)
    scheduled_at = _email_scheduler_parse_scheduled_at_utc(date_str=date, time_str=time, am_pm=am_pm)

    schedule_name, schedule_bytes, _ = await _email_scheduler_read_upload_bytes(schedule_attachment)
    att_name, att_bytes, att_type = await _email_scheduler_read_upload_bytes(attachment)
    normalized_plant_code = _normalize_plant_code(plant_code)
    is_dsm_template = "dsm" in _email_scheduler_template_category(str(template_id or "")).lower()
    portal_issue_flag = str(portal_issue or "").strip().lower() in {"1", "true", "yes", "y"}
    day_ahead_date_already_adjusted_flag = str(day_ahead_date_already_adjusted or "").strip().lower() in {"1", "true", "yes", "y"}

    db = SessionLocal()
    try:
        to_email, cc_email = _email_scheduler_apply_saved_recipients(
            db,
            plant_code=normalized_plant_code,
            template_id=str(template_id or "").strip(),
            to_email=to_email,
            cc_email=cc_email,
        )
    finally:
        db.close()

    # For INTRADAY templates, auto-resolve schedule CSV from S3 if caller did not upload one.
    if (not schedule_bytes) and ("intra" in _email_scheduler_template_category(str(template_id or "")).lower()):
        resolved = _email_scheduler_resolve_schedule_attachment_data(
            plant_name=normalized_plant_code,
            template_id=str(template_id or "").strip(),
            date_key=str(date or "").strip(),
        )
        schedule_name = str(resolved.get("file_name") or "").strip() or None
        schedule_bytes = bytes(resolved.get("file_bytes") or b"")

    if schedule_bytes:
        conversion_schedule_type = "intraday" if "intra" in _email_scheduler_template_category(str(template_id or "")).lower() else "dayahead"
        converted = maybe_convert_for_auto_email(
            plant_code=normalized_plant_code,
            template_id=str(template_id or "").strip(),
            schedule_type=conversion_schedule_type,
            file_name=schedule_name or "schedule.csv",
            file_bytes=schedule_bytes,
            report_date=str(date or "").strip() if conversion_schedule_type == "dayahead" else "",
            source_key=schedule_name or "",
        )
        if converted:
            schedule_name = _email_scheduler_attachment_display_name(
                plant_code=normalized_plant_code,
                template_id=str(template_id or "").strip(),
                schedule_type="intraday" if "intra" in _email_scheduler_template_category(str(template_id or "")).lower() else "dayahead",
                source_key=schedule_name or converted.filename,
                original_name=converted.filename,
                report_date=str(date or "").strip(),
                date_already_day_ahead=day_ahead_date_already_adjusted_flag,
            )
            schedule_bytes = converted.content_bytes
        else:
            schedule_name = _email_scheduler_attachment_display_name(
                plant_code=normalized_plant_code,
                template_id=str(template_id or "").strip(),
                schedule_type=conversion_schedule_type,
                source_key=schedule_name or "",
                original_name=schedule_name or "schedule.csv",
                report_date=str(date or "").strip(),
                date_already_day_ahead=day_ahead_date_already_adjusted_flag,
            )
    dsm_payload = _email_scheduler_parse_json_payload(dsm_summary_payload)
    if dsm_payload is None and "dsm" in _email_scheduler_template_category(str(template_id or "")).lower():
        dsm_payload = await asyncio.to_thread(
            _email_scheduler_build_dsm_payload_from_s3_for_email,
            plant_code=normalized_plant_code,
            report_date=str(date or "").strip(),
        )

    send_subject = str(subject or "").strip()
    send_body = str(body or "").strip()
    if _email_scheduler_is_gsnp_intraday(plant_code=normalized_plant_code, template_id=str(template_id or "")):
        send_subject = _email_scheduler_gsnp_intraday_subject(str(date or "").strip())
        send_body = _email_scheduler_gsnp_intraday_body(str(date or "").strip())
    elif _email_scheduler_is_ilios_pv_intraday(plant_code=normalized_plant_code, template_id=str(template_id or "")):
        send_subject = _email_scheduler_ilios_pv_intraday_subject(str(date or "").strip())
        send_body = _email_scheduler_ilios_pv_intraday_body(str(date or "").strip())
    elif _email_scheduler_is_sirmour_intraday(plant_code=normalized_plant_code, template_id=str(template_id or "")):
        send_subject = _email_scheduler_build_report_subject(
            template_id=str(template_id or "").strip(),
            plant_code=normalized_plant_code,
            report_date=str(date or "").strip(),
        ) or send_subject
        send_body = _email_scheduler_sirmour_intraday_body(str(date or "").strip())
    else:
        send_subject = _email_scheduler_subject_day_ahead_date(
            send_subject,
            template_id=str(template_id or "").strip(),
            report_date=str(date or "").strip(),
            date_already_day_ahead=day_ahead_date_already_adjusted_flag,
        )
        send_body = _email_scheduler_day_ahead_body_date(
            send_body,
            plant_code=normalized_plant_code,
            template_id=str(template_id or "").strip(),
            report_date=str(date or "").strip(),
            date_already_day_ahead=day_ahead_date_already_adjusted_flag,
        )

    sent_at = datetime.now(timezone.utc)
    try:
        await asyncio.to_thread(
            _email_scheduler_send_now,
            template_id=template_id,
            role=role,
            from_email=from_email,
            to_email=to_email,
            cc_email=cc_email,
            subject=send_subject,
            body=send_body,
            employee_name=employee_name,
            dsm_payload=dsm_payload,
            schedule_attachment=(schedule_name or "", schedule_bytes or b"") if schedule_bytes else None,
            attachment=(att_name or "", att_bytes or b"", att_type or "application/octet-stream") if att_bytes else None,
        )
        _email_scheduler_log_event(
            requested_by=user or "",
            employee_name=employee_name or "",
            role=role,
            template_id=template_id,
            plant_code=plant,
            mode="SEND_NOW",
            status="SENT",
            from_email=from_email,
            to_email=to_email,
            cc_email=cc_email,
            subject=send_subject,
            scheduled_at=scheduled_at,
            sent_at=sent_at,
            error_message=None,
        )
    except Exception as exc:
        _email_scheduler_log_event(
            requested_by=user or "",
            employee_name=employee_name or "",
            role=role,
            template_id=template_id,
            plant_code=plant,
            mode="SEND_NOW",
            status="FAILED",
            from_email=from_email,
            to_email=to_email,
            cc_email=cc_email,
            subject=send_subject,
            scheduled_at=scheduled_at,
            sent_at=sent_at,
            error_message=str(exc),
        )
        raise

    return {"ok": True, "status": "sent", "template_id": template_id, "plant_code": plant, "role": role}


@app.post("/email-scheduler/schedule")
async def email_scheduler_schedule(
    template_id: str = Form(...),
    plant_code: str = Form(...),
    date: str = Form(...),
    time: str = Form(...),
    am_pm: str = Form("AM"),
    from_email: str = Form(...),
    to_email: str = Form(...),
    cc_email: str = Form(""),
    employee_name: str = Form(""),
    subject: str = Form(...),
    body: str = Form(...),
    auto_send: str = Form("0"),
    portal_issue: str = Form("0"),
    portal_issue_plants: str = Form(""),
    dsm_summary_payload: str = Form(""),
    schedule_attachment: Optional[UploadFile] = File(None),
    attachment: Optional[UploadFile] = File(None),
    x_user_role: Optional[str] = Header(None, alias="X-User-Role"),
    x_user_name: Optional[str] = Header(None, alias="X-User-Name"),
):
    role = _email_scheduler_normalize_role(x_user_role)
    user = _email_scheduler_normalize_user(x_user_name)
    _email_scheduler_role_guard(role=role, admin_only=False)
    to_email = str(to_email or "").strip()
    cc_email = str(cc_email or "").strip()

    scheduled_at = _email_scheduler_parse_scheduled_at_utc(date_str=date, time_str=time, am_pm=am_pm)

    schedule_name, schedule_bytes, _ = await _email_scheduler_read_upload_bytes(schedule_attachment)
    att_name, att_bytes, att_type = await _email_scheduler_read_upload_bytes(attachment)
    normalized_plant_code = _normalize_plant_code(plant_code)
    is_dsm_template = "dsm" in _email_scheduler_template_category(str(template_id or "")).lower()
    portal_issue_flag = str(portal_issue or "").strip().lower() in {"1", "true", "yes", "y"}

    # For INTRADAY templates, auto-resolve schedule CSV from S3 if caller did not upload one.
    if (not schedule_bytes) and ("intra" in _email_scheduler_template_category(str(template_id or "")).lower()):
        resolved = _email_scheduler_resolve_schedule_attachment_data(
            plant_name=normalized_plant_code,
            template_id=str(template_id or "").strip(),
            date_key=str(date or "").strip(),
        )
        schedule_name = str(resolved.get("file_name") or "").strip() or None
        schedule_bytes = bytes(resolved.get("file_bytes") or b"")

    if schedule_bytes:
        conversion_schedule_type = "intraday" if "intra" in _email_scheduler_template_category(str(template_id or "")).lower() else "dayahead"
        converted = maybe_convert_for_auto_email(
            plant_code=normalized_plant_code,
            template_id=str(template_id or "").strip(),
            schedule_type=conversion_schedule_type,
            file_name=schedule_name or "schedule.csv",
            file_bytes=schedule_bytes,
            report_date=str(date or "").strip() if conversion_schedule_type == "dayahead" else "",
            source_key=schedule_name or "",
        )
        if converted:
            schedule_name = _email_scheduler_attachment_display_name(
                plant_code=normalized_plant_code,
                template_id=str(template_id or "").strip(),
                schedule_type="intraday" if "intra" in _email_scheduler_template_category(str(template_id or "")).lower() else "dayahead",
                source_key=schedule_name or converted.filename,
                original_name=converted.filename,
                report_date=str(date or "").strip(),
            )
            schedule_bytes = converted.content_bytes
        else:
            schedule_name = _email_scheduler_attachment_display_name(
                plant_code=normalized_plant_code,
                template_id=str(template_id or "").strip(),
                schedule_type=conversion_schedule_type,
                source_key=schedule_name or "",
                original_name=schedule_name or "schedule.csv",
                report_date=str(date or "").strip(),
            )

    normalized_dsm_payload = str(dsm_summary_payload or "").strip() or None
    if portal_issue_flag and portal_issue_plants and not normalized_dsm_payload:
        try:
            parsed = json.loads(str(portal_issue_plants or "").strip())
            if isinstance(parsed, list):
                parsed = [str(x).strip() for x in parsed if str(x).strip()]
            else:
                parsed = []
        except Exception:
            parsed = []
        normalized_dsm_payload = json.dumps({"portal_issue_plants": parsed}, ensure_ascii=True, separators=(",", ":"))

    schedule_subject = str(subject or "").strip()
    schedule_body = normalize_day_ahead_body(str(body or ""), str(template_id or ""))
    if _email_scheduler_is_gsnp_intraday(plant_code=normalized_plant_code, template_id=str(template_id or "")):
        schedule_subject = _email_scheduler_gsnp_intraday_subject(str(date or "").strip())
        schedule_body = _email_scheduler_gsnp_intraday_body(str(date or "").strip())
    elif _email_scheduler_is_ilios_pv_intraday(plant_code=normalized_plant_code, template_id=str(template_id or "")):
        schedule_subject = _email_scheduler_ilios_pv_intraday_subject(str(date or "").strip())
        schedule_body = _email_scheduler_ilios_pv_intraday_body(str(date or "").strip())
    elif _email_scheduler_is_sirmour_intraday(plant_code=normalized_plant_code, template_id=str(template_id or "")):
        schedule_subject = _email_scheduler_build_report_subject(
            template_id=str(template_id or "").strip(),
            plant_code=normalized_plant_code,
            report_date=str(date or "").strip(),
        ) or schedule_subject
        schedule_body = _email_scheduler_sirmour_intraday_body(str(date or "").strip())
    else:
        schedule_subject = _email_scheduler_subject_day_ahead_date(
            schedule_subject,
            template_id=str(template_id or "").strip(),
            report_date=str(date or "").strip(),
        )
        schedule_body = _email_scheduler_day_ahead_body_date(
            schedule_body,
            plant_code=normalized_plant_code,
            template_id=str(template_id or "").strip(),
            report_date=str(date or "").strip(),
        )

    db = SessionLocal()
    try:
        to_email, cc_email = _email_scheduler_apply_saved_recipients(
            db,
            plant_code=normalized_plant_code,
            template_id=str(template_id or "").strip(),
            to_email=to_email,
            cc_email=cc_email,
        )
        job = EmailSchedulerJob(
            requested_by=user or None,
            role=role,
            template_id=str(template_id or "").strip(),
            plant_code=normalized_plant_code,
            scheduled_at=scheduled_at,
            auto_send=str(auto_send or "").strip().lower() in {"1", "true", "yes", "y"},
            from_email=str(from_email or "").strip(),
            to_email=str(to_email or "").strip(),
            cc_email=str(cc_email or "").strip() or None,
            employee_name=str(employee_name or "").strip() or None,
            subject=schedule_subject,
            body=schedule_body,
            portal_issue=portal_issue_flag,
            dsm_summary_payload=normalized_dsm_payload,
            schedule_attachment_name=schedule_name,
            schedule_attachment_bytes=schedule_bytes,
            attachment_name=att_name,
            attachment_bytes=att_bytes,
            attachment_content_type=att_type,
            status="SCHEDULED",
        )
        db.add(job)
        db.commit()
        db.refresh(job)
        _email_scheduler_log_event(
            requested_by=user or "",
            employee_name=employee_name or "",
            role=role,
            template_id=str(template_id or "").strip(),
            plant_code=str(job.plant_code or "").strip(),
            mode="SCHEDULE",
            status="SCHEDULED",
            from_email=str(from_email or "").strip(),
            to_email=str(to_email or "").strip(),
            cc_email=str(cc_email or "").strip(),
            subject=schedule_subject,
            scheduled_at=job.scheduled_at,
            sent_at=None,
            error_message=None,
        )
        return {"ok": True, "id": job.id, "status": job.status}
    except HTTPException:
        raise
    except Exception as exc:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Failed to schedule job: {exc}") from exc
    finally:
        db.close()


@app.get("/email-scheduler/send-logs")
def email_scheduler_send_logs(
    x_user_role: Optional[str] = Header(None, alias="X-User-Role"),
    x_user_name: Optional[str] = Header(None, alias="X-User-Name"),
    limit: int = Query(100, ge=1, le=500),
):
    role = _email_scheduler_normalize_role(x_user_role)
    _user = _email_scheduler_normalize_user(x_user_name)
    _email_scheduler_role_guard(role=role, admin_only=True)

    db = SessionLocal()
    try:
        rows = (
            db.query(EmailSendLog)
            .order_by(EmailSendLog.created_at.desc(), EmailSendLog.id.desc())
            .limit(int(limit))
            .all()
        )
        items = []
        for r in rows:
            items.append(
                {
                    "id": r.id,
                    "requested_by": r.requested_by,
                    "employee_name": r.employee_name,
                    "role": r.role,
                    "template_id": r.template_id,
                    "plant_code": r.plant_code,
                    "category": r.category,
                    "mode": r.mode,
                    "from_email": r.from_email,
                    "to_email": r.to_email,
                    "cc_email": r.cc_email,
                    "subject": r.subject,
                    "scheduled_at": r.scheduled_at.isoformat() if r.scheduled_at else None,
                    "sent_at": r.sent_at.isoformat() if r.sent_at else None,
                    "status": r.status,
                    "error_message": r.error_message,
                    "created_at": r.created_at.isoformat() if r.created_at else None,
                }
            )
        return {"ok": True, "items": items}
    finally:
        db.close()


class EmailSchedulerScheduleAllRequest(BaseModel):
    template_id: str
    date: str
    time: str
    am_pm: Optional[str] = "AM"
    from_email: str
    employee_name: Optional[str] = ""
    auto_send: Optional[bool] = False


class EmailSchedulerDailyDsmRunRequest(BaseModel):
    template_id: Optional[str] = None
    from_email: Optional[str] = None
    employee_name: Optional[str] = None
    auto_send: Optional[bool] = True
    dry_run: Optional[bool] = False
    force_repeat: Optional[bool] = False


class EmailSchedulerDailyDaRunRequest(BaseModel):
    template_id: Optional[str] = None
    from_email: Optional[str] = None
    employee_name: Optional[str] = None
    auto_send: Optional[bool] = True
    dry_run: Optional[bool] = False
    force_repeat: Optional[bool] = False


class EmailSchedulerDailyIntradayRunRequest(BaseModel):
    template_id: Optional[str] = None
    from_email: Optional[str] = None
    employee_name: Optional[str] = None
    auto_send: Optional[bool] = True
    dry_run: Optional[bool] = False
    force_repeat: Optional[bool] = False


EMAIL_SCHEDULER_SETTING_DAILY_DSM_ENABLED = "daily_dsm_enabled"
EMAIL_SCHEDULER_SETTING_DAILY_DA_ENABLED = "daily_da_enabled"
EMAIL_SCHEDULER_SETTING_PLANT_AUTO_EMAIL_ENABLED = "plant_auto_email_enabled"
EMAIL_SCHEDULER_SETTING_RECIPIENT_DEFAULTS = "recipient_defaults"


def _email_scheduler_settings_get_bool(db: Session, key: str, default: bool) -> bool:
    if not db or not key:
        return default
    try:
        row = db.query(EmailSchedulerSetting).filter(EmailSchedulerSetting.key == key).first()
    except Exception:
        return default
    if not row:
        return default
    raw = str(getattr(row, "value_text", "") or "").strip().lower()
    if raw in {"1", "true", "yes", "y", "on"}:
        return True
    if raw in {"0", "false", "no", "n", "off"}:
        return False
    return default


def _email_scheduler_settings_set_bool(db: Session, key: str, value: bool) -> None:
    if not db or not key:
        return
    text_value = "true" if bool(value) else "false"
    row = db.query(EmailSchedulerSetting).filter(EmailSchedulerSetting.key == key).first()
    if row:
        row.value_text = text_value
        db.add(row)
        return
    db.add(EmailSchedulerSetting(key=key, value_text=text_value))


def _email_scheduler_settings_get_json_dict(db: Session, key: str) -> Dict[str, Any]:
    if not db or not key:
        return {}
    try:
        row = db.query(EmailSchedulerSetting).filter(EmailSchedulerSetting.key == key).first()
    except Exception:
        return {}
    if not row:
        return {}
    try:
        parsed = json.loads(str(getattr(row, "value_text", "") or "{}"))
    except Exception:
        return {}
    return parsed if isinstance(parsed, dict) else {}


def _email_scheduler_settings_set_json_dict(db: Session, key: str, value: Dict[str, Any]) -> None:
    if not db or not key:
        return
    text_value = json.dumps(value or {}, ensure_ascii=True, separators=(",", ":"))
    row = db.query(EmailSchedulerSetting).filter(EmailSchedulerSetting.key == key).first()
    if row:
        row.value_text = text_value
        db.add(row)
        return
    db.add(EmailSchedulerSetting(key=key, value_text=text_value))


def _email_scheduler_normalize_plant_auto_email_map(raw: Dict[str, Any]) -> Dict[str, bool]:
    out: Dict[str, bool] = {}
    for key, value in (raw or {}).items():
        plant_code = _normalize_plant_code(str(key or "").strip())
        if not plant_code:
            continue
        if isinstance(value, str):
            out[plant_code] = value.strip().lower() in {"1", "true", "yes", "y", "on"}
        else:
            out[plant_code] = bool(value)
    return out


def _email_scheduler_get_plant_auto_email_map(db: Session) -> Dict[str, bool]:
    return _email_scheduler_normalize_plant_auto_email_map(
        _email_scheduler_settings_get_json_dict(db, EMAIL_SCHEDULER_SETTING_PLANT_AUTO_EMAIL_ENABLED)
    )


def _email_scheduler_is_plant_auto_email_enabled(settings: Dict[str, bool], plant_code: str) -> bool:
    plant = _normalize_plant_code(str(plant_code or "").strip())
    if not plant:
        return True
    return bool((settings or {}).get(plant, True))


def _email_scheduler_normalize_recipient_defaults(raw: Dict[str, Any]) -> Dict[str, Dict[str, Dict[str, str]]]:
    out: Dict[str, Dict[str, Dict[str, str]]] = {}
    for plant_key, template_map in (raw or {}).items():
        plant = _normalize_plant_code(str(plant_key or "").strip())
        if not plant or not isinstance(template_map, dict):
            continue
        for template_key, value in (template_map or {}).items():
            template_id = str(template_key or "").strip()
            if not template_id or not isinstance(value, dict):
                continue
            to_email = str(value.get("to_email") or value.get("to") or "").strip()
            cc_email = str(value.get("cc_email") or value.get("cc") or "").strip()
            if not to_email and not cc_email:
                continue
            out.setdefault(plant, {})[template_id] = {
                "to_email": to_email,
                "cc_email": cc_email,
            }
    return out


def _email_scheduler_get_recipient_defaults(db: Session) -> Dict[str, Dict[str, Dict[str, str]]]:
    return _email_scheduler_normalize_recipient_defaults(
        _email_scheduler_settings_get_json_dict(db, EMAIL_SCHEDULER_SETTING_RECIPIENT_DEFAULTS)
    )


def _email_scheduler_get_recipient_default(
    settings: Dict[str, Dict[str, Dict[str, str]]],
    plant_code: str,
    template_id: str,
) -> Dict[str, str]:
    plant = _normalize_plant_code(str(plant_code or "").strip())
    template = str(template_id or "").strip()
    if not plant or not template:
        return {}
    value = ((settings or {}).get(plant) or {}).get(template) or {}
    return value if isinstance(value, dict) else {}


EMAIL_SCHEDULER_PLANT_CAPACITY_MW: Dict[str, float] = {
    "BHUPALPALLY": 10.0,
    "KASIPET": 15.0,
    "KOTHAGUDEM": 37.0,
    "OSEPL": 20.0,
    "ANDAD": 7.5,
    "BALAKWADA": 7.5,
    "GUGARIYAKHEDI": 7.5,
    "NANDGAON": 7.5,
    "BAMKHAL": 5.0,
    "SIRMOUR": 5.1,
    "SAWDA": 7.5,
    "ZETRIC": 25.0,
    "ANJANGAON": 7.5,
    "ILIOS_PV": 50.0,
}


def _email_scheduler_format_subject_capacity(capacity: float) -> str:
    try:
        value = float(capacity or 0.0)
    except Exception:
        value = 0.0
    if not math.isfinite(value) or value <= 0:
        return "0"
    if float(value).is_integer():
        return str(int(value))
    return (f"{value:.3f}").rstrip("0").rstrip(".")


def _email_scheduler_format_subject_date(date_value: Any) -> str:
    raw = str(date_value or "").strip()
    if not raw:
        return raw
    try:
        if isinstance(date_value, date):
            parsed = date_value
        else:
            parsed = datetime.strptime(raw[:10], "%Y-%m-%d").date()
        return parsed.strftime("%d-%m-%Y")
    except Exception:
        return raw


def _email_scheduler_format_dotted_date(date_value: Any) -> str:
    raw = str(date_value or "").strip()
    if not raw:
        return raw
    try:
        if isinstance(date_value, date):
            parsed = date_value
        else:
            parsed = datetime.strptime(raw[:10], "%Y-%m-%d").date()
        return parsed.strftime("%d.%m.%Y")
    except Exception:
        return raw


def _email_scheduler_report_subject_prefix(template_id: str, template: Optional[Dict[str, Any]] = None) -> str:
    tpl_id = str(template_id or (template or {}).get("id") or "").strip().lower()
    category = _email_scheduler_template_category(tpl_id).strip().lower() if tpl_id else ""
    hay = f"{tpl_id} {category}"
    if "dsm" in hay:
        return "DSM Report"
    if "intra" in hay:
        return "Intraday Schedule"
    if "day" in hay or tpl_id.endswith("_da0") or tpl_id.endswith("_da1") or "da0" in tpl_id or "da1" in tpl_id:
        return "Dayahead Schedule"
    return ""


def _email_scheduler_build_report_subject(
    *,
    template_id: str,
    plant_code: str,
    report_date: Any,
    template: Optional[Dict[str, Any]] = None,
) -> str:
    prefix = _email_scheduler_report_subject_prefix(template_id, template)
    plant = _normalize_plant_code(str(plant_code or "").strip())
    if not prefix or not plant:
        return ""
    subject_date = report_date
    template_key = str(template_id or (template or {}).get("id") or "").strip().lower()
    if prefix == "Dayahead Schedule" and _email_scheduler_is_day_ahead_template(template_key):
        subject_date = _email_scheduler_add_days_to_date_key(report_date, 1)
    date_label = _email_scheduler_format_subject_date(subject_date)
    if _email_scheduler_is_gsnp_intraday(plant_code=plant, template_id=template_key):
        return _email_scheduler_gsnp_intraday_subject(subject_date)
    if _email_scheduler_is_ilios_pv_intraday(plant_code=plant, template_id=template_key):
        return _email_scheduler_ilios_pv_intraday_subject(subject_date)
    if plant == "ILIOS_PV" and prefix == "Dayahead Schedule":
        return f"Dayahead Schedule Ilios_PV (50MW) for {date_label}"
    if _email_scheduler_is_sirmour_intraday(plant_code=plant, template_id=template_key):
        capacity = _email_scheduler_format_subject_capacity(EMAIL_SCHEDULER_PLANT_CAPACITY_MW.get(plant, 0.0))
        return f"Final Intraday Schedule {plant} ({capacity} MW) for {date_label}"
    if prefix == "DSM Report" and plant in {"TELANGANA", "BHUPALPALLY", "KASIPET", "KOTHAGUDEM"}:
        return f"DSM Report Telangana State Plants for {date_label}"
    capacity = _email_scheduler_format_subject_capacity(EMAIL_SCHEDULER_PLANT_CAPACITY_MW.get(plant, 0.0))
    return f"{prefix} {plant} ({capacity} MW) for {date_label}"


def _email_scheduler_attachment_revision_label(
    *,
    template_id: str,
    schedule_type: str,
    source_key: str,
) -> str:
    template_key = str(template_id or "").strip().lower()
    type_key = str(schedule_type or "").strip().lower()
    if type_key == "dayahead":
        if "da1" in template_key:
            return "DA1"
        return "DA0"
    if type_key == "intraday":
        source_name = os.path.basename(str(source_key or "").strip())
        explicit = re.search(r"(?:^|[_-])(?:id[_-]?)?r(?:ev(?:ision)?)?[_-]?(\d+)(?:\D|$)", source_name, flags=re.IGNORECASE)
        revision = explicit.group(1) if explicit else None
        if not revision:
            extracted = _extract_schedule_revision_from_key(source_name)
            revision = str(extracted) if extracted is not None else "1"
        return f"ID_R{revision}"
    return "SCHEDULE"


def _email_scheduler_attachment_display_name(
    *,
    plant_code: str,
    template_id: str,
    schedule_type: str,
    source_key: str,
    original_name: str = "",
    report_date: Any = None,
    date_already_day_ahead: bool = False,
) -> str:
    plant = _normalize_plant_code(str(plant_code or "").strip())
    template_key = str(template_id or "").strip().lower()
    type_key = str(schedule_type or "").strip().lower()
    label = _email_scheduler_attachment_revision_label(
        template_id=template_id,
        schedule_type=schedule_type,
        source_key=source_key or original_name,
    )
    ext = os.path.splitext(str(original_name or source_key or "").strip())[1] or ".csv"
    if ext.lower() not in {".csv", ".xlsx", ".xlsm", ".xls"}:
        ext = ".csv"
    date_suffix = _email_scheduler_attachment_date_suffix(
        template_id=template_id,
        schedule_type=schedule_type,
        report_date=report_date,
        date_already_day_ahead=date_already_day_ahead,
    )
    if plant == "SIRMOUR" and ("intra" in template_key or type_key == "intraday"):
        return f"Final_Schedule-Sirmour{date_suffix}{ext}"
    return f"{plant}_{label}{date_suffix}{ext}"


def _email_scheduler_attachment_date_suffix(
    *,
    template_id: str,
    schedule_type: str,
    report_date: Any,
    date_already_day_ahead: bool = False,
) -> str:
    parsed = _email_scheduler_parse_report_date(report_date)
    if not parsed:
        return ""
    type_key = str(schedule_type or "").strip().lower()
    template_key = str(template_id or "").strip().lower()
    if type_key == "dayahead" or "da0" in template_key or "da1" in template_key:
        schedule_date = parsed if date_already_day_ahead else parsed + timedelta(days=1)
    elif type_key == "intraday" or "intra" in template_key:
        schedule_date = parsed
    else:
        return ""
    return f"_{schedule_date.strftime('%d-%m-%Y')}"


def _email_scheduler_parse_report_date(value: Any) -> Optional[date]:
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    text = str(value or "").strip()
    if not text:
        return None
    for fmt in ("%Y-%m-%d", "%d-%m-%Y", "%d/%m/%Y"):
        try:
            return datetime.strptime(text[:10], fmt).date()
        except Exception:
            continue
    return None


class EmailSchedulerSettingsUpdateRequest(BaseModel):
    daily_dsm_enabled: Optional[bool] = None
    daily_da_enabled: Optional[bool] = None
    plant_auto_email_enabled: Optional[Dict[str, bool]] = None
    recipient_defaults: Optional[Dict[str, Any]] = None


@app.get("/email-scheduler/settings")
def email_scheduler_get_settings(
    x_user_role: Optional[str] = Header(None, alias="X-User-Role"),
):
    role = _email_scheduler_normalize_role(x_user_role)
    _email_scheduler_role_guard(role=role, admin_only=False)

    db = SessionLocal()
    try:
        dsm_enabled = _email_scheduler_settings_get_bool(db, EMAIL_SCHEDULER_SETTING_DAILY_DSM_ENABLED, True)
        da_enabled = _email_scheduler_settings_get_bool(db, EMAIL_SCHEDULER_SETTING_DAILY_DA_ENABLED, True)
        plant_auto_email_enabled = _email_scheduler_get_plant_auto_email_map(db)
        recipient_defaults = _email_scheduler_get_recipient_defaults(db)
        return {
            "ok": True,
            "daily_dsm_enabled": bool(dsm_enabled),
            "daily_da_enabled": bool(da_enabled),
            "plant_auto_email_enabled": plant_auto_email_enabled,
            "recipient_defaults": recipient_defaults,
        }
    finally:
        db.close()


@app.post("/email-scheduler/settings")
def email_scheduler_update_settings(
    payload: EmailSchedulerSettingsUpdateRequest,
    x_user_role: Optional[str] = Header(None, alias="X-User-Role"),
):
    role = _email_scheduler_normalize_role(x_user_role)
    _email_scheduler_role_guard(role=role, admin_only=False)
    if payload.recipient_defaults is not None:
        _email_scheduler_role_guard(role=role, admin_only=True)

    db = SessionLocal()
    try:
        if payload.daily_dsm_enabled is not None:
            _email_scheduler_settings_set_bool(db, EMAIL_SCHEDULER_SETTING_DAILY_DSM_ENABLED, bool(payload.daily_dsm_enabled))
        if payload.daily_da_enabled is not None:
            _email_scheduler_settings_set_bool(db, EMAIL_SCHEDULER_SETTING_DAILY_DA_ENABLED, bool(payload.daily_da_enabled))
        if payload.plant_auto_email_enabled is not None:
            _email_scheduler_settings_set_json_dict(
                db,
                EMAIL_SCHEDULER_SETTING_PLANT_AUTO_EMAIL_ENABLED,
                _email_scheduler_normalize_plant_auto_email_map(payload.plant_auto_email_enabled),
            )
        if payload.recipient_defaults is not None:
            _email_scheduler_settings_set_json_dict(
                db,
                EMAIL_SCHEDULER_SETTING_RECIPIENT_DEFAULTS,
                _email_scheduler_normalize_recipient_defaults(payload.recipient_defaults),
            )
        db.commit()
        dsm_enabled = _email_scheduler_settings_get_bool(db, EMAIL_SCHEDULER_SETTING_DAILY_DSM_ENABLED, True)
        da_enabled = _email_scheduler_settings_get_bool(db, EMAIL_SCHEDULER_SETTING_DAILY_DA_ENABLED, True)
        plant_auto_email_enabled = _email_scheduler_get_plant_auto_email_map(db)
        recipient_defaults = _email_scheduler_get_recipient_defaults(db)
        return {
            "ok": True,
            "daily_dsm_enabled": bool(dsm_enabled),
            "daily_da_enabled": bool(da_enabled),
            "plant_auto_email_enabled": plant_auto_email_enabled,
            "recipient_defaults": recipient_defaults,
        }
    except Exception as exc:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Failed to update settings: {exc}") from exc
    finally:
        db.close()


def _email_scheduler_ist_day_utc_bounds(target_day_ist: date) -> Tuple[datetime, datetime]:
    ist = ZoneInfo("Asia/Kolkata")
    start_ist = datetime(target_day_ist.year, target_day_ist.month, target_day_ist.day, 0, 0, 0, tzinfo=ist)
    end_ist = start_ist + timedelta(days=1)
    return start_ist.astimezone(timezone.utc), end_ist.astimezone(timezone.utc)


def _email_scheduler_exists_daily_send_for_plant(
    db: Session,
    *,
    plant_code: str,
    template_id: str,
    day_start_utc: datetime,
    day_end_utc: datetime,
) -> bool:
    plant = _normalize_plant_code(plant_code)
    template = str(template_id or "").strip()
    if not plant or not template:
        return False
    exists_job = (
        db.query(EmailSchedulerJob.id)
        .filter(EmailSchedulerJob.plant_code == plant)
        .filter(EmailSchedulerJob.template_id == template)
        .filter(EmailSchedulerJob.scheduled_at >= day_start_utc)
        .filter(EmailSchedulerJob.scheduled_at < day_end_utc)
        .first()
    )
    if exists_job:
        return True
    exists_log = (
        db.query(EmailSendLog.id)
        .filter(EmailSendLog.plant_code == plant)
        .filter(EmailSendLog.template_id == template)
        .filter(EmailSendLog.created_at >= day_start_utc)
        .filter(EmailSendLog.created_at < day_end_utc)
        .first()
    )
    return bool(exists_log)


def _email_scheduler_guess_attachment_content_type(file_name: str) -> str:
    lower = str(file_name or "").strip().lower()
    if lower.endswith(".xlsx"):
        return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    if lower.endswith(".xls"):
        return "application/vnd.ms-excel"
    if lower.endswith(".csv"):
        return "text/csv"
    if lower.endswith(".pdf"):
        return "application/pdf"
    return "application/octet-stream"


def _email_scheduler_merge_cc(existing_cc: str, extra_cc: str) -> str:
    existing = [x.strip() for x in str(existing_cc or "").split(",") if str(x or "").strip()]
    extra = [x.strip() for x in str(extra_cc or "").split(",") if str(x or "").strip()]
    merged: List[str] = []
    seen = set()
    for addr in existing + extra:
        key = addr.lower()
        if key in seen:
            continue
        seen.add(key)
        merged.append(addr)
    return ",".join(merged)


def _email_scheduler_apply_saved_recipients(
    db: Session,
    *,
    plant_code: str,
    template_id: str,
    to_email: str,
    cc_email: str,
    merge_cc: bool = False,
) -> Tuple[str, str]:
    effective_to = str(to_email or "").strip()
    effective_cc = str(cc_email or "").strip()
    recipient_default = _email_scheduler_get_recipient_default(
        _email_scheduler_get_recipient_defaults(db),
        plant_code,
        str(template_id or "").strip(),
    )
    if not recipient_default:
        return effective_to, effective_cc
    saved_to = str(recipient_default.get("to_email") or "").strip()
    saved_cc = str(recipient_default.get("cc_email") or "").strip()
    if saved_to:
        effective_to = saved_to
    if saved_cc:
        effective_cc = _email_scheduler_merge_cc(effective_cc, saved_cc) if merge_cc else saved_cc
    return effective_to, effective_cc


def _email_scheduler_ensure_intraday_cc(*, plant_code: str, template_id: str, cc_email: str) -> str:
    plant = _normalize_plant_code(str(plant_code or "").strip())
    template_key = str(template_id or "").strip().lower()
    if plant == "SIRMOUR" and "intra" in template_key:
        return _email_scheduler_merge_cc(cc_email, "forecasting.vppl@gmail.com")
    return str(cc_email or "").strip()


def _email_scheduler_is_sirmour_intraday(*, plant_code: str, template_id: str) -> bool:
    plant = _normalize_plant_code(str(plant_code or "").strip())
    template_key = str(template_id or "").strip().lower()
    return plant == "SIRMOUR" and "intra" in template_key


def _email_scheduler_is_gsnp_intraday(*, plant_code: str, template_id: str) -> bool:
    plant = _normalize_plant_code(str(plant_code or "").strip())
    template_key = str(template_id or "").strip().lower()
    return plant == "GSNP" and "intra" in template_key


def _email_scheduler_is_ilios_pv_intraday(*, plant_code: str, template_id: str) -> bool:
    plant = _normalize_plant_code(str(plant_code or "").strip())
    template_key = str(template_id or "").strip().lower()
    return plant == "ILIOS_PV" and "intra" in template_key


def _email_scheduler_gsnp_intraday_subject(report_date: Any) -> str:
    context = _email_scheduler_build_template_context(str(report_date or "")[:10])
    return f"Globus Steel N Power Intraday for {context.get('month_full', '')}-{context.get('year_full', '')}"


def _email_scheduler_gsnp_intraday_body(report_date: Any) -> str:
    return (
        "Dear Sir/mam,\n\n"
        f"Please Find the attached Intraday Forecast of \"Globus Steel N Power\" for Date {_email_scheduler_format_dotted_date(report_date)}"
    )


def _email_scheduler_ilios_pv_intraday_subject(report_date: Any) -> str:
    context = _email_scheduler_build_template_context(str(report_date or "")[:10])
    return f"Ilios_PV Intraday Schedule for the Month of {context.get('month_full', '')}_{context.get('year_full', '')}"


def _email_scheduler_ilios_pv_intraday_body(report_date: Any) -> str:
    return (
        "Dear Sir/Mam,\n\n"
        f"Please find attached the Intraday Schedule ILIOS_PV for Date {_email_scheduler_format_dotted_date(report_date)}"
    )


def _email_scheduler_sirmour_intraday_body(report_date: Any) -> str:
    return (
        "Dear Sir/Mam,\n"
        f"Please find attached Final Intraday Schedule SIRMOUR_PV for Date {_email_scheduler_format_dotted_date(report_date)}."
    )


EMAIL_SCHEDULER_TELANGANA_DA1_BODY_PLANTS = {"BHUPALPALLY", "KASIPET", "KOTHAGUDEM"}


def _email_scheduler_is_day_ahead_template(template_id: str) -> bool:
    template_key = str(template_id or "").strip().lower()
    return "da0" in template_key or "da1" in template_key


def _email_scheduler_is_telangana_da1_body(*, plant_code: str, template_id: str) -> bool:
    plant = _normalize_plant_code(str(plant_code or "").strip())
    template_key = str(template_id or "").strip().lower()
    return plant in EMAIL_SCHEDULER_TELANGANA_DA1_BODY_PLANTS and "da1" in template_key


def _email_scheduler_add_days_to_date_key(date_value: Any, days: int) -> str:
    try:
        base_day = date_value if isinstance(date_value, date) else datetime.strptime(str(date_value or "")[:10], "%Y-%m-%d").date()
        return (base_day + timedelta(days=int(days or 0))).isoformat()
    except Exception:
        return str(date_value or "").strip()


def _email_scheduler_body_template_date_key(*, plant_code: str, template_id: str, report_date: Any) -> str:
    if _email_scheduler_is_day_ahead_template(template_id):
        return _email_scheduler_add_days_to_date_key(report_date, 1)
    return str(report_date or "").strip()


def _email_scheduler_day_ahead_body_date(
    body: str,
    *,
    plant_code: str,
    template_id: str,
    report_date: Any,
    date_already_day_ahead: bool = False,
) -> str:
    normalized = normalize_day_ahead_body(str(body or ""), str(template_id or ""))
    if not _email_scheduler_is_day_ahead_template(template_id):
        return normalized
    template_date_key = str(report_date or "").strip() if date_already_day_ahead else _email_scheduler_body_template_date_key(
        plant_code=plant_code,
        template_id=template_id,
        report_date=report_date,
    )
    context = _email_scheduler_build_template_context(
        template_date_key
    )
    date_label = str(context.get("date_dotted") or "").strip()
    if not date_label:
        return normalized
    return re.sub(
        r"\bDate\s+\d{1,2}[./-]\d{1,2}[./-]\d{2,4}",
        f"Date {date_label}",
        normalized,
        flags=re.IGNORECASE,
    )


def _email_scheduler_subject_day_ahead_date(subject: str, *, template_id: str, report_date: Any, date_already_day_ahead: bool = False) -> str:
    text = str(subject or "")
    if not text or not _email_scheduler_is_day_ahead_template(template_id):
        return text
    date_key = str(report_date or "").strip() if date_already_day_ahead else _email_scheduler_add_days_to_date_key(report_date, 1)
    try:
        day = datetime.strptime(str(date_key or "")[:10], "%Y-%m-%d").date()
    except Exception:
        return text

    def replace_date(match: re.Match) -> str:
        separator = str(match.group("sep") or "-")
        if separator == ".":
            date_label = day.strftime("%d.%m.%Y")
        elif separator == "/":
            date_label = day.strftime("%d/%m/%Y")
        else:
            date_label = day.strftime("%d-%m-%Y")
        return f"{match.group('prefix')}{date_label}"

    return re.sub(
        r"(?P<prefix>\bfor\s+)\d{1,2}(?P<sep>[./-])\d{1,2}(?P=sep)\d{2,4}",
        replace_date,
        text,
        flags=re.IGNORECASE,
    )


def _email_scheduler_telangana_da1_body_date(body: str, *, plant_code: str, template_id: str, report_date: Any) -> str:
    return _email_scheduler_day_ahead_body_date(
        body,
        plant_code=plant_code,
        template_id=template_id,
        report_date=report_date,
    )


def _email_scheduler_build_template_context(date_str: str) -> Dict[str, str]:
    raw = str(date_str or "").strip()
    try:
        dt = datetime.strptime(raw, "%Y-%m-%d")
    except Exception:
        dt = datetime.now()
    next_month_dt = (dt.replace(day=1) + timedelta(days=32)).replace(day=1)
    return {
        "date_dashed": dt.strftime("%Y-%m-%d"),
        "date_dotted": dt.strftime("%d.%m.%Y"),
        "month_full": dt.strftime("%B"),
        "month_short": dt.strftime("%b"),
        "year_full": dt.strftime("%Y"),
        "year_short": dt.strftime("%y"),
        "next_month_short": next_month_dt.strftime("%b"),
    }


def _email_scheduler_render_template_vars(text_value: str, context: Dict[str, str]) -> str:
    """
    Replace `{var}` placeholders in the email subject/body.

    The UI formats these placeholders client-side for "send now", but cron-driven
    jobs store the raw template strings. Keep this replacement minimal and safe:
    unknown placeholders are left as-is.
    """
    out = str(text_value or "")
    for k, v in (context or {}).items():
        token = "{" + str(k) + "}"
        out = out.replace(token, str(v))
    return out


def _email_scheduler_build_tabular_dsm_attachment_bytes(
    *,
    template_id: str,
    date_str: str,
    plant_name: str,
) -> Tuple[Optional[str], Optional[bytes]]:
    tid = str(template_id or "").strip().lower()
    context = _email_scheduler_build_template_context(date_str)
    pname = str(plant_name or "").strip() or "Plant Report"

    schemas: Dict[str, Dict[str, Any]] = {
        "sirmour_dsm": {
            "file_name": "sirmour-dsm-report.csv",
            "columns": ["From", "To", "Project", "Installed Capacity (Mw)", "Generation (Kwh)", "DSM Penalty (Rs.)", "Paisa /Kwh", "Net Revenue", "%Impact"],
            "row": [context["date_dashed"], context["date_dashed"], pname, "5", "12,480", "4,850", "0.39", "1,24,560", "3.89%"],
        },
        "bhupalpally_dsm": {
            "file_name": "bhupalpally-dsm-report.csv",
            "columns": ["To", "Month", "Project", "Installed Capacity (Mw)", "Generation (Kwh)", "DSM Penalty(Rs.) As per Scada Availability", "DSM Penalty As Maintenance Information", "Paisa/Kwh Scada Availability", "Paisa/Kwh Maintenance Information", "Scada Availability(%)"],
            "row": [context["date_dashed"], context["month_full"], pname, "10", "24,920", "8,640", "1,250", "0.35", "0.05", "98.7%"],
        },
        "kasipet_dsm": {
            "file_name": "kasipet-dsm-report.csv",
            "columns": ["To", "Month", "Project", "Installed Capacity (Mw)", "Generation (Kwh)", "DSM Penalty(Rs.) As per Scada Availability", "DSM Penalty As Maintenance Information", "Paisa/Kwh Scada Availability", "Paisa/Kwh Maintenance Information", "Scada Availability(%)"],
            "row": [context["date_dashed"], context["month_full"], pname, "15", "31,250", "9,250", "1,450", "0.30", "0.05", "98.9%"],
        },
        "kothagudem_dsm": {
            "file_name": "kothagudem-dsm-report.csv",
            "columns": ["To", "Month", "Project", "Installed Capacity (Mw)", "Generation (Kwh)", "DSM Penalty(Rs.) As per Scada Availability", "DSM Penalty As Maintenance Information", "Paisa/Kwh Scada Availability", "Paisa/Kwh Maintenance Information", "Scada Availability(%)"],
            "row": [context["date_dashed"], context["month_full"], pname, "37", "74,880", "21,640", "3,120", "0.29", "0.04", "99.0%"],
        },
        "osepl_dsm": {
            "file_name": "osepl-dsm-report.csv",
            "columns": ["From", "Month", "Project", "Installed Capacity", "SCADA availability", "Generation(Kwh)", "Scheduled unit*PPA", "Payable", "Receivable", "DSM Penalty(Rs.)"],
            "row": [context["date_dashed"], context["month_full"], pname, "20", "99.1%", "48,750", "47,900", "12,500", "2,800", "9,700"],
        },
    }

    schema = schemas.get(tid)
    if not schema:
        return None, None
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(schema["columns"])
    writer.writerow(schema["row"])
    return str(schema["file_name"]), output.getvalue().encode("utf-8")


def _email_scheduler_build_simple_daily_dsm_attachment(
    *,
    plant_code: str,
    plant_name: str,
    report_date: str,
) -> Dict[str, Any]:
    """
    Always generate a lightweight DSM CSV attachment for cron auto-send.
    This avoids S3/runtime dependency and guarantees an attachment exists.
    """
    pcode = _normalize_plant_code(plant_code)
    pname = str(plant_name or pcode).strip() or pcode
    day = str(report_date or "").strip()
    if not re.fullmatch(r"\d{4}-\d{2}-\d{2}", day):
        day = datetime.now().strftime("%Y-%m-%d")

    capacity_map = {
        "ANJANGAON": "7.5",
        "ANDAD": "7.5",
        "BALAKWADA": "7.5",
        "GUGARIYAKHEDI": "7.5",
        "NANDGAON": "7.5",
        "SIRMOUR": "5.1",
        "BHUPALPALLY": "10",
        "KASIPET": "15",
        "KOTHAGUDEM": "37",
        "OSEPL": "20",
        "BAMKHAL": "5",
    }
    installed_capacity = str(capacity_map.get(pcode) or "0")

    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow([
        "FROM",
        "TO",
        "PROJECT",
        "INSTALLED CAPACITY (MW)",
        "GENERATION (KWH)",
        "DSM PENALTY (RS.)",
        "PAISA / KWH",
        "NET REVENUE",
        "%IMPACT",
    ])
    writer.writerow([
        day,
        day,
        f"{pname}_Schedule",
        installed_capacity,
        "0",
        "0",
        "--",
        "--",
        "--",
    ])

    file_name = f"{pcode.lower()}-dsm-report-{day}.csv"
    data = buf.getvalue().encode("utf-8")
    return {
        "file_name": file_name,
        "bytes": data,
        "content_type": "text/csv",
        "s3_key": "",
        "generated": True,
    }


def _email_scheduler_build_simple_daily_dsm_table_payload(
    *,
    plant_code: str,
    plant_name: str,
    report_date: str,
) -> Dict[str, Any]:
    """
    Build a minimal DSM payload that renders as an HTML table in the email body.

    This intentionally mirrors the "simple daily DSM" attachment path (zeros/placeholder values),
    but delivers it as structured data for `render_dsm_table_html(...)`.
    """
    pcode = _normalize_plant_code(plant_code)
    pname = str(plant_name or pcode).strip() or pcode
    day = str(report_date or "").strip()
    if not re.fullmatch(r"\d{4}-\d{2}-\d{2}", day):
        day = datetime.now().strftime("%Y-%m-%d")

    context = _email_scheduler_build_template_context(day)
    month_label = f"{context['month_short']}-{context['year_short']}"

    capacity_map = {
        "ANJANGAON": "7.5",
        "ANDAD": "7.5",
        "BALAKWADA": "7.5",
        "GUGARIYAKHEDI": "7.5",
        "NANDGAON": "7.5",
        "SIRMOUR": "5.1",
        "BHUPALPALLY": "10",
        "KASIPET": "15",
        "KOTHAGUDEM": "37",
        "OSEPL": "20",
        "BAMKHAL": "5",
    }
    installed_capacity = str(capacity_map.get(pcode) or "0")

    columns = [
        "DATE",
        "TO",
        "MONTH",
        "PROJECT",
        "INSTALLED CAPACITY (MW)",
        "GENERATION (KWH)",
        "DSM PENALTY (RS.), AS PER SCADA AVAILABILITY",
        "DSM PENALTY (RS.), AS MAINTENANCE INFORMATION",
        "PAISA/KWH SCADA AVAILABILITY",
        "PAISA/KWH MAINTENANCE INFORMATION",
        "SCADA AVAILABILITY(%)",
    ]
    row = {
        "DATE": day,
        "TO": day,
        "MONTH": month_label,
        "PROJECT": str(pname).strip().upper(),
        "INSTALLED CAPACITY (MW)": installed_capacity,
        "GENERATION (KWH)": "0",
        "DSM PENALTY (RS.), AS PER SCADA AVAILABILITY": "0",
        "DSM PENALTY (RS.), AS MAINTENANCE INFORMATION": "0",
        "PAISA/KWH SCADA AVAILABILITY": "0.00",
        "PAISA/KWH MAINTENANCE INFORMATION": "0.00",
        "SCADA AVAILABILITY(%)": "100%",
    }

    # Use the existing "sirmour/multi" palette (green header) so the email resembles the UI preview.
    return {"variant": "sirmour", "columns": columns, "rows": [row]}


def _email_scheduler_build_simple_daily_dsm_table_payload_multi(
    *,
    plants: List[Tuple[str, str]],
    report_date: str,
) -> Dict[str, Any]:
    """
    Build a multi-row DSM payload (e.g. Telangana combined daily summary) that renders
    as a single HTML table in the email body.
    """
    day = str(report_date or "").strip()
    if not re.fullmatch(r"\d{4}-\d{2}-\d{2}", day):
        day = datetime.now().strftime("%Y-%m-%d")

    context = _email_scheduler_build_template_context(day)
    month_label = f"{context['month_short']}-{context['year_short']}"

    capacity_map = {
        "ANJANGAON": "7.5",
        "ANDAD": "7.5",
        "BALAKWADA": "7.5",
        "GUGARIYAKHEDI": "7.5",
        "NANDGAON": "7.5",
        "SIRMOUR": "5.1",
        "BHUPALPALLY": "10",
        "KASIPET": "15",
        "KOTHAGUDEM": "37",
        "OSEPL": "20",
        "BAMKHAL": "5",
    }

    columns = [
        "DATE",
        "TO",
        "MONTH",
        "PROJECT",
        "INSTALLED CAPACITY (MW)",
        "GENERATION (KWH)",
        "DSM PENALTY (RS.), AS PER SCADA AVAILABILITY",
        "DSM PENALTY (RS.), AS MAINTENANCE INFORMATION",
        "PAISA/KWH SCADA AVAILABILITY",
        "PAISA/KWH MAINTENANCE INFORMATION",
        "SCADA AVAILABILITY(%)",
    ]

    rows: List[Dict[str, Any]] = []
    for code, name in plants:
        pcode = _normalize_plant_code(code)
        pname = str(name or pcode).strip() or pcode
        installed_capacity = str(capacity_map.get(pcode) or "0")
        rows.append(
            {
                "DATE": day,
                "TO": day,
                "MONTH": month_label,
                "PROJECT": str(pname).strip().upper(),
                "INSTALLED CAPACITY (MW)": installed_capacity,
                "GENERATION (KWH)": "0",
                "DSM PENALTY (RS.), AS PER SCADA AVAILABILITY": "0",
                "DSM PENALTY (RS.), AS MAINTENANCE INFORMATION": "0",
                "PAISA/KWH SCADA AVAILABILITY": "0.00",
                "PAISA/KWH MAINTENANCE INFORMATION": "0.00",
                "SCADA AVAILABILITY(%)": "100%",
            }
        )

    return {"variant": "multi", "columns": columns, "rows": rows}


def _email_scheduler_public_dsm_columns(row: Dict[str, Any]) -> List[str]:
    return [str(key) for key in (row or {}).keys() if not str(key).startswith("__")]


def _email_scheduler_sirmour_block_end_timestamp(day: str, block: int) -> str:
    minutes = max(0, min(96, int(block or 0))) * 15
    hh = (minutes // 60) % 24
    mm = minutes % 60
    return f"{str(day or '').strip()} {hh:02d}:{mm:02d}"


def _email_scheduler_block_end_timestamp(day: str, block: int) -> str:
    return _email_scheduler_sirmour_block_end_timestamp(day, block)


def _email_scheduler_round_number(value: Any, ndigits: int = 2) -> float:
    try:
        num = float(value or 0.0)
    except Exception:
        return 0.0
    if not math.isfinite(num):
        return 0.0
    return round(num, int(ndigits))


def _email_scheduler_summary_logo_paths() -> Tuple[str, ...]:
    base = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    return (
        os.path.join(base, "public", "vedanjay logo.png"),
        os.path.join(base, "public", "image.png"),
        os.path.join(base, "image.png"),
    )


def _email_scheduler_add_summary_logo(ws: Any, xl_image_cls: Any) -> None:
    for logo_path in _email_scheduler_summary_logo_paths():
        if not os.path.exists(logo_path):
            continue
        try:
            logo = xl_image_cls(logo_path)
            logo.width = 230
            logo.height = 52
            ws.add_image(logo, "B2")
            ws.row_dimensions[2].height = 34
            ws.row_dimensions[3].height = 28
            break
        except Exception:
            continue


def _email_scheduler_dsm_support_attachment_from_payload(
    *,
    payload: Optional[Dict[str, Any]],
    plant_code: str,
    report_date: str,
    telangana_static_values: bool = False,
) -> Optional[Dict[str, Any]]:
    rows = list((payload or {}).get("rows") or [])
    if not rows:
        return None
    columns = list((payload or {}).get("columns") or [])
    if not columns:
        columns = _email_scheduler_public_dsm_columns(rows[0]) if isinstance(rows[0], dict) else list(rows[0].keys())

    pcode = _normalize_plant_code(plant_code)
    date_key = str(report_date or "").strip()
    try:
        date_label = datetime.strptime(date_key, "%Y-%m-%d").strftime("%d-%m-%Y")
    except Exception:
        date_label = date_key or datetime.now().strftime("%d-%m-%Y")

    if pcode in {"TELANGANA", *_EMAIL_SCHEDULER_TELANGANA_DSM_CODES}:
        file_name = f"Daily_DSM_Penalty_Report_TELANGANA_{date_label}.xlsx"
    elif pcode == "OSEPL":
        file_name = f"Daily_DSM_Penalty_Report_OSEPL_{date_label}.xlsx"
    elif pcode == "SIRMOUR":
        file_name = f"Daily_DSM_Penalty_Report_SIRMOUR_{date_label}.xlsx"
    else:
        file_name = f"Daily_DSM_Penalty_Report_{pcode or 'DSM'}_{date_label}.xlsx"

    try:
        from openpyxl import Workbook  # type: ignore
        from openpyxl.drawing.image import Image as XLImage  # type: ignore
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side  # type: ignore

        wb = Workbook()
        sirmour_details = rows[0].get("__support_details") if pcode == "SIRMOUR" and isinstance(rows[0], dict) else None
        sirmour_summary = rows[0].get("__support_summary") if pcode == "SIRMOUR" and isinstance(rows[0], dict) else None
        if pcode == "SIRMOUR" and isinstance(sirmour_details, list) and sirmour_details:
            ws = wb.active
            ws.title = "Summary"
            ws.merge_cells("B2:E2")
            ws.merge_cells("B3:E3")
            _email_scheduler_add_summary_logo(ws, XLImage)

            summary_headers = [
                "From",
                "To",
                "Project",
                "Installed \nCapacity (Mw)",
                "Generation(Kwh)",
                "DSM Penalty(Rs.)\n",
                "Paisa/Kwh\n",
                "Net Revenue\n",
                "%Impact\n",
            ]
            summary = sirmour_summary if isinstance(sirmour_summary, dict) else {}
            summary_values = [
                summary.get("From", ""),
                summary.get("To", ""),
                summary.get("Project", "Sirmour_Schedule"),
                summary.get("Installed Capacity (Mw)", ""),
                summary.get("Generation(Kwh)", ""),
                summary.get("DSM Penalty(Rs.)", ""),
                summary.get("Paisa/Kwh", ""),
                summary.get("Net Revenue", ""),
                summary.get("%Impact", ""),
            ]
            for col_idx, value in enumerate(summary_headers, start=2):
                ws.cell(row=12, column=col_idx).value = value
            for col_idx, value in enumerate(summary_values, start=2):
                ws.cell(row=13, column=col_idx).value = value

            detail_ws = wb.create_sheet("Sirmour_Schedule")
            detail_ws.cell(row=2, column=1).value = "Deviation_Charges Blocks"
            detail_ws.append([])
            detail_ws.cell(row=3, column=1).value = "From"
            detail_ws.cell(row=3, column=2).value = "Upto"
            detail_ws.cell(row=3, column=3).value = "Deviation_Charges(Rs.)"
            for row_idx, band in enumerate([(0, 10, 0), (10, 15, 0.5), (15, 20, 0.75), (20, "", 1)], start=4):
                detail_ws.cell(row=row_idx, column=1).value = band[0]
                detail_ws.cell(row=row_idx, column=2).value = band[1]
                detail_ws.cell(row=row_idx, column=3).value = band[2]

            detail_headers = ["Datetime(Date+Block endtime)", "Schedule(Kwh)", "Meter data(KWh)", "AvC(Kwh)", "% Error", "DSM Penalty"]
            for col_idx, value in enumerate(detail_headers, start=1):
                detail_ws.cell(row=9, column=col_idx).value = value
            for row_idx, detail in enumerate(sirmour_details, start=10):
                if not isinstance(detail, dict):
                    continue
                detail_ws.cell(row=row_idx, column=1).value = detail.get("Datetime(Date+Block endtime)", "")
                detail_ws.cell(row=row_idx, column=2).value = detail.get("Schedule(Kwh)", 0)
                detail_ws.cell(row=row_idx, column=3).value = detail.get("Meter data(KWh)", 0)
                detail_ws.cell(row=row_idx, column=4).value = detail.get("AvC(Kwh)", 0)
                detail_ws.cell(row=row_idx, column=5).value = detail.get("% Error", 0)
                detail_ws.cell(row=row_idx, column=6).value = detail.get("DSM Penalty", 0)

            header_fill = PatternFill("solid", fgColor="16823A")
            header_font = Font(color="FFFFFF", bold=True)
            thin = Side(style="thin", color="D9D9D9")
            border = Border(left=thin, right=thin, top=thin, bottom=thin)
            for cell in ws[12][1:10]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                cell.border = border
            for cell in ws[13][1:10]:
                cell.alignment = Alignment(vertical="center", wrap_text=True)
                cell.border = border
            for cell in detail_ws[3][0:3]:
                cell.font = Font(bold=True)
                cell.border = border
            for cell in detail_ws[9][0:6]:
                cell.font = Font(bold=True)
                cell.border = border
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            for row_cells in detail_ws.iter_rows(min_row=10, max_row=105, min_col=1, max_col=6):
                for cell in row_cells:
                    cell.border = border
                    cell.alignment = Alignment(vertical="center")

            widths = {"B": 18, "C": 18, "D": 18, "E": 22, "F": 18, "G": 18, "H": 14, "I": 16, "J": 14}
            for col, width in widths.items():
                ws.column_dimensions[col].width = width
            for col, width in {"A": 26, "B": 16, "C": 16, "D": 12, "E": 12, "F": 14}.items():
                detail_ws.column_dimensions[col].width = width

            buf = io.BytesIO()
            wb.save(buf)
            return {
                "file_name": file_name,
                "bytes": buf.getvalue(),
                "content_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            }

        if pcode in {"TELANGANA", *_EMAIL_SCHEDULER_TELANGANA_DSM_CODES}:
            ws = wb.active
            ws.title = "Summary"

            summary_headers = [
                "Date",
                "To",
                "Month",
                "Project",
                "Installed \nCapacity (Mw)",
                "Generation(Kwh)",
                "DSM Penalty(Rs.)\nAs per Scada Availability",
                "DSM Penalty As \nMaintenance Information",
                "Paisa/Kwh\nScada Availability",
                "Paisa/Kwh\nMaintenance Information",
                "Scada Availability(%)",
                "Remarks",
            ]
            for col_idx, value in enumerate(summary_headers, start=1):
                ws.cell(row=2, column=col_idx).value = value

            title_map = {
                "KASIPET": "Kasipet",
                "BHUPALPALLY": "Bhupalpally",
                "KOTHAGUDEM": "Kothagudem",
            }
            office_order = ("KASIPET", "BHUPALPALLY", "KOTHAGUDEM")
            summary_by_plant = {}
            for row in rows:
                if not isinstance(row, dict):
                    continue
                plant = _normalize_plant_code(row.get("PROJECT") or row.get("Project") or "")
                if plant in title_map:
                    summary_by_plant[plant] = row

            try:
                report_dt = datetime.strptime(date_key, "%Y-%m-%d")
            except Exception:
                report_dt = datetime.now()
            month_end = (report_dt.replace(day=28) + timedelta(days=4)).replace(day=1) - timedelta(days=1)
            days_in_month = int(month_end.day)

            def _summary_number(row_obj: Dict[str, Any], fallback: Any, *keys: str) -> float:
                for key in keys:
                    raw = row_obj.get(key)
                    if raw not in (None, ""):
                        try:
                            return float(str(raw).replace(",", "").replace("%", "").strip())
                        except Exception:
                            break
                try:
                    return float(fallback)
                except Exception:
                    return 0.0

            for idx, plant in enumerate(office_order, start=3):
                source = summary_by_plant.get(plant, {})
                sheet_name = title_map[plant]
                generation = _summary_number(source, 0, "GENERATION (KWH)", "Generation (kWh)", "Generation(Kwh)")
                dsm_penalty = _summary_number(source, 0, "DSM PENALTY (RS.), AS PER SCADA AVAILABILITY", "DSM Penalty (Rs.) As per SCADA Availability", "DSM Penalty(Rs.)\nAs per Scada Availability")
                maint_penalty = _summary_number(source, dsm_penalty, "DSM PENALTY (RS.), AS MAINTENANCE INFORMATION", "DSM Penalty (Rs.) As Maintenance Information", "DSM Penalty As \nMaintenance Information")
                ws.cell(row=idx, column=1).value = report_dt.replace(hour=0, minute=15, second=0, microsecond=0)
                ws.cell(row=idx, column=2).value = report_dt.replace(hour=0, minute=15, second=0, microsecond=0)
                ws.cell(row=idx, column=3).value = month_end
                ws.cell(row=idx, column=4).value = sheet_name
                ws.cell(row=idx, column=5).value = _summary_number(source, {"KASIPET": 15, "BHUPALPALLY": 10, "KOTHAGUDEM": 37}.get(plant, 0), "INSTALLED CAPACITY (MW)", "Installed Capacity (MW)", "Installed \nCapacity (Mw)")
                if telangana_static_values:
                    ws.cell(row=idx, column=6).value = generation
                    ws.cell(row=idx, column=7).value = dsm_penalty
                    ws.cell(row=idx, column=8).value = maint_penalty
                    ws.cell(row=idx, column=9).value = (dsm_penalty / generation * 100.0) if generation > 0 else 0
                    ws.cell(row=idx, column=10).value = (maint_penalty / generation * 100.0) if generation > 0 else 0
                    scada_availability = _summary_number(source, 100, "SCADA AVAILABILITY(%)", "Scada Availability(%)")
                    ws.cell(row=idx, column=11).value = (scada_availability / 100.0) if scada_availability > 1 else scada_availability
                else:
                    ws.cell(row=idx, column=6).value = f'=SUMIFS({sheet_name}!C:C,{sheet_name}!C:C,"<>#N/A",{sheet_name}!$A:$A,">="&Summary!$A{idx},{sheet_name}!$A:$A,"<="&Summary!$B{idx}+1)'
                    ws.cell(row=idx, column=7).value = f'=SUMIFS({sheet_name}!F:F,{sheet_name}!F:F,"<>#N/A",{sheet_name}!$A:$A,">="&Summary!$A{idx},{sheet_name}!$A:$A,"<="&Summary!$B{idx}+1)'
                    ws.cell(row=idx, column=8).value = f'=SUMIFS({sheet_name}!H:H,{sheet_name}!H:H,"<>#N/A",{sheet_name}!$A:$A,">="&Summary!$A{idx},{sheet_name}!$A:$A,"<="&Summary!$B{idx}+1)'
                    ws.cell(row=idx, column=9).value = f'=G{idx}/F{idx}*100'
                    ws.cell(row=idx, column=10).value = f'=(H{idx}/F{idx}*100)'
                    ws.cell(row=idx, column=11).value = f'=COUNTIFS(INDIRECT($D{idx}&"!$A:$A"),">="&$A{idx},INDIRECT($D{idx}&"!$A:$A"),"<"&$B{idx}+1,INDIRECT($D{idx}&"!$C:$C"),"<>#N/A")/((B{idx}-A{idx}+1)*96)'

            ws.cell(row=9, column=1).value = "*#N/A Means Meter Data Not Available"
            ws.merge_cells("A9:F9")

            for plant in office_order:
                source = summary_by_plant.get(plant, {})
                details = source.get("__support_details") if isinstance(source, dict) else None
                if not isinstance(details, list):
                    details = []
                detail_ws = wb.create_sheet(title_map[plant])
                detail_ws.merge_cells("B8:E8")
                detail_ws.cell(row=2, column=1).value = "Deviation_Charges Blocks"
                detail_ws.cell(row=3, column=1).value = "From"
                detail_ws.cell(row=3, column=2).value = "Upto"
                detail_ws.cell(row=3, column=3).value = "Deviation_Charges(Rs.)"
                for row_idx, band in enumerate([(0, 15, 0), (15, 25, 0.5), (25, 35, 1), (35, "", 1.5)], start=4):
                    detail_ws.cell(row=row_idx, column=1).value = band[0]
                    detail_ws.cell(row=row_idx, column=2).value = band[1]
                    detail_ws.cell(row=row_idx, column=3).value = band[2]
                detail_ws.cell(row=8, column=2).value = title_map[plant]
                detail_headers = [
                    "Datetime(Date+Block endtime)",
                    "Schedule(Kwh)",
                    "Meter data(KWh)",
                    "AvC(Kwh)",
                    "% Error",
                    "DSM penalty",
                    "Maintenance/Scada Update" if plant == "KASIPET" else "Maintenance Update",
                    "DSM penalty as per Maintenance Updates",
                ]
                for col_idx, value in enumerate(detail_headers, start=1):
                    detail_ws.cell(row=9, column=col_idx).value = value
                details_by_block = {}
                for detail in details:
                    if not isinstance(detail, dict):
                        continue
                    try:
                        block = int(detail.get("block") or 0)
                    except Exception:
                        block = 0
                    if 1 <= block <= 96:
                        details_by_block[block] = detail
                avc_default = {"KASIPET": 3750, "BHUPALPALLY": 2500, "KOTHAGUDEM": 9250}.get(plant, 0)
                for day in range(1, days_in_month + 1):
                    for block in range(1, 97):
                        row_idx = 10 + ((day - 1) * 96) + (block - 1)
                        detail = details_by_block.get(block) if day == report_dt.day else None
                        if telangana_static_values:
                            detail_ws.cell(row=row_idx, column=1).value = report_dt.replace(day=day, hour=0, minute=0, second=0, microsecond=0) + timedelta(minutes=block * 15)
                        elif row_idx == 10:
                            detail_ws.cell(row=row_idx, column=1).value = report_dt.replace(day=1, hour=0, minute=15, second=0, microsecond=0)
                        else:
                            detail_ws.cell(row=row_idx, column=1).value = f"=A{row_idx - 1}+TIME(0,15,0)"
                        detail_ws.cell(row=row_idx, column=2).value = _summary_number(detail or {}, 0, "Schedule(Kwh)")
                        detail_ws.cell(row=row_idx, column=3).value = _summary_number(detail or {}, 0, "Meter data(KWh)")
                        detail_ws.cell(row=row_idx, column=4).value = _summary_number(detail or {}, avc_default, "AvC(Kwh)")
                        if telangana_static_values:
                            detail_ws.cell(row=row_idx, column=5).value = _summary_number(detail or {}, 0, "% Error")
                            detail_ws.cell(row=row_idx, column=6).value = _summary_number(detail or {}, 0, "DSM penalty", "DSM Penalty")
                        else:
                            detail_ws.cell(row=row_idx, column=5).value = f"=IF(D{row_idx}=0,0,100*ABS(B{row_idx}-C{row_idx})/D{row_idx})"
                            detail_ws.cell(row=row_idx, column=6).value = f"=IF(AND(E{row_idx}>$A$5,E{row_idx}<=$B$5),(E{row_idx}-$A$5)*(D{row_idx}*$A$5%*$C$5)/$A$5,IF(AND(E{row_idx}>$A$6,E{row_idx}<=$B$6),(D{row_idx}*($B$5-$A$5)%*$C$5)+((E{row_idx}-$A$6)*(D{row_idx}*$A$6%*$C$6)/$A$6),IF(E{row_idx}>$A$7,(D{row_idx}*($B$5-$A$5)%*$C$5)+(D{row_idx}*($B$6-$A$6)%*$C$6)+((E{row_idx}-$A$7)*(D{row_idx}*$A$7%*$C$7)/$A$7),0)))"
                        detail_ws.cell(row=row_idx, column=7).value = _summary_number(detail or {}, 0, "Maintenance Update", "Maintenance/Scada Update")
                        if telangana_static_values:
                            detail_ws.cell(row=row_idx, column=8).value = _summary_number(detail or {}, 0, "DSM penalty as per Maintenance Updates")
                        else:
                            detail_ws.cell(row=row_idx, column=8).value = f"=IF(G{row_idx}=0,F{row_idx},0)"

            header_fill = PatternFill("solid", fgColor="BDD7EE")
            header_font = Font(color="000000", bold=True)
            thin = Side(style="thin", color="D9D9D9")
            border = Border(left=thin, right=thin, top=thin, bottom=thin)
            for cell in ws[2][0:12]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                cell.border = border
            for row_cells in ws.iter_rows(min_row=3, max_row=5, min_col=1, max_col=12):
                for cell in row_cells:
                    cell.alignment = Alignment(vertical="center", wrap_text=True)
                    cell.border = border
            for cell in ws[9][0:6]:
                cell.border = border
            for detail_ws in wb.worksheets[1:]:
                for cell in detail_ws[3][0:3]:
                    cell.font = Font(bold=True)
                    cell.border = border
                for cell in detail_ws[9][0:8]:
                    cell.font = Font(bold=True)
                    cell.border = border
                    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                for row_cells in detail_ws.iter_rows(min_row=10, max_row=9 + (days_in_month * 96), min_col=1, max_col=8):
                    for cell in row_cells:
                        cell.border = border
                        cell.alignment = Alignment(vertical="center")
                detail_ws.column_dimensions["A"].number_format = "dd\\.mm\\.yyyy\\ h:mm"
                for row_idx in range(10, 10 + (days_in_month * 96)):
                    detail_ws.cell(row=row_idx, column=1).number_format = "dd\\.mm\\.yyyy\\ h:mm"
                    for col_idx in (2, 3, 4, 5, 7, 8):
                        detail_ws.cell(row=row_idx, column=col_idx).number_format = "0.00"
                    detail_ws.cell(row=row_idx, column=6).number_format = "0.000"
                for col, width in {"A": 26, "B": 16, "C": 16, "D": 12, "E": 12, "F": 14, "G": 18, "H": 24}.items():
                    detail_ws.column_dimensions[col].width = width
            for row_idx in range(3, 6):
                ws.cell(row=row_idx, column=1).number_format = "dd\\.mm\\.yyyy"
                ws.cell(row=row_idx, column=2).number_format = "dd\\.mm\\.yyyy"
                ws.cell(row=row_idx, column=3).number_format = "mmm-yy"
                ws.cell(row=row_idx, column=6).number_format = "0"
                ws.cell(row=row_idx, column=7).number_format = "0"
                ws.cell(row=row_idx, column=8).number_format = "0"
                ws.cell(row=row_idx, column=9).number_format = "0.00"
                ws.cell(row=row_idx, column=10).number_format = "0.00"
                ws.cell(row=row_idx, column=11).number_format = "0%"
            for col, width in {"A": 18, "B": 18, "C": 12, "D": 18, "E": 22, "F": 18, "G": 28, "H": 30, "I": 22, "J": 24, "K": 20, "L": 16}.items():
                ws.column_dimensions[col].width = width

            buf = io.BytesIO()
            wb.save(buf)
            return {
                "file_name": file_name,
                "bytes": buf.getvalue(),
                "content_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            }

        if pcode == "OSEPL":
            ws = wb.active
            ws.title = "Summary"
            ws.merge_cells("B2:E2")
            ws.merge_cells("B3:E3")
            _email_scheduler_add_summary_logo(ws, XLImage)

            summary_headers = [
                "From",
                "To",
                "Month",
                "Project",
                "Installed Capacity",
                "SCADA availability",
                "Generation(Kwh)",
                "Scheduled unit*PPA",
                "Payable ",
                "Receivable",
                "DSM Penalty(Rs.)",
                "DSM penalty as per SCADA Availability/Maint information",
            ]
            row = rows[0] if isinstance(rows[0], dict) else {}
            summary_values = [
                row.get("From", ""),
                row.get("From", ""),
                row.get("Month", ""),
                row.get("Project", "ESSEL"),
                row.get("Installed Capacity", ""),
                row.get("SCADA availability %", ""),
                row.get("Generation(kWh)", ""),
                row.get("Scheduled unit*PPA", ""),
                row.get("Payable", ""),
                row.get("Receivable", ""),
                row.get("DSM Penalty (Rs.)", ""),
                row.get("SCADA Adjusted DSM", row.get("DSM Penalty (Rs.)", "")),
            ]
            for col_idx, value in enumerate(summary_headers, start=1):
                ws.cell(row=12, column=col_idx).value = value
            for col_idx, value in enumerate(summary_values, start=1):
                ws.cell(row=13, column=col_idx).value = value

            detail_ws = wb.create_sheet("ESSEL")
            detail_ws.cell(row=2, column=1).value = "PPA Rate"
            detail_ws.cell(row=2, column=3).value = row.get("PPA", "9.27")
            detail_ws.cell(row=3, column=1).value = "Error Blocks"
            detail_ws.cell(row=3, column=11).value = 1
            detail_ws.cell(row=4, column=1).value = "From"
            detail_ws.cell(row=4, column=2).value = "Upto"
            detail_ws.cell(row=4, column=3).value = "UnderInjection Penalty(Rs.)"
            detail_ws.cell(row=4, column=4).value = "Overinjection penalty (Rs.)"
            for row_idx, band in enumerate([(0, 10, 9.27, 9.27), (10, 12, 10.197, 8.343), (12, 15, 11.124, 7.416), (15, "", 13.905, 0)], start=5):
                for col_idx, value in enumerate(band, start=1):
                    detail_ws.cell(row=row_idx, column=col_idx).value = value
            detail_ws.cell(row=9, column=10).value = "A-B+C"
            for col_idx, value in enumerate(["A", "B", "C", "D=(B-C)*100/C", "E=A*PPA rate", "F", "G", "H=E-F+G"], start=1):
                detail_ws.cell(row=10, column=col_idx).value = value
            detail_ws.cell(row=11, column=10).value = "Under/Overinjection"
            detail_headers = [
                "TimeSlots(date+endtime)",
                "Forecast(kwH)",
                "Actual(Kwh)",
                "AvC(kwh)",
                "% Error",
                "Scheduled unit*PPA",
                "Payable",
                "Receivable",
                "Total",
                "generator-end penalty(Rs.)",
                "Scada/Maintenance Information Availability",
                "Penalty As per Scada Information",
            ]
            for col_idx, value in enumerate(detail_headers, start=1):
                detail_ws.cell(row=12, column=col_idx).value = value
            details = row.get("__support_details") if isinstance(row, dict) else []
            if not isinstance(details, list):
                details = []
            for row_idx, detail in enumerate(details, start=13):
                if not isinstance(detail, dict):
                    continue
                for col_idx, key in enumerate(detail_headers, start=1):
                    detail_ws.cell(row=row_idx, column=col_idx).value = detail.get(key, "")

            header_fill = PatternFill("solid", fgColor="16823A")
            header_font = Font(color="FFFFFF", bold=True)
            thin = Side(style="thin", color="D9D9D9")
            border = Border(left=thin, right=thin, top=thin, bottom=thin)
            for cell in ws[12][0:12]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                cell.border = border
            for cell in ws[13][0:12]:
                cell.alignment = Alignment(vertical="center", wrap_text=True)
                cell.border = border
            for row_idx in (4, 10, 12):
                for cell in detail_ws[row_idx][0:12]:
                    cell.font = Font(bold=True)
                    cell.border = border
                    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            for row_cells in detail_ws.iter_rows(min_row=13, max_row=max(13, 12 + len(details)), min_col=1, max_col=12):
                for cell in row_cells:
                    cell.border = border
                    cell.alignment = Alignment(vertical="center")
            for col, width in {"A": 26, "B": 16, "C": 16, "D": 12, "E": 12, "F": 18, "G": 14, "H": 14, "I": 14, "J": 22, "K": 28, "L": 24}.items():
                detail_ws.column_dimensions[col].width = width
            for col, width in {"A": 18, "B": 18, "C": 12, "D": 16, "E": 18, "F": 18, "G": 18, "H": 20, "I": 14, "J": 14, "K": 18, "L": 34}.items():
                ws.column_dimensions[col].width = width

            buf = io.BytesIO()
            wb.save(buf)
            return {
                "file_name": file_name,
                "bytes": buf.getvalue(),
                "content_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            }

        ws = wb.active
        ws.title = "DSM Report Preview"
        ws.append(columns)
        for row in rows:
            ws.append([row.get(col, "") for col in columns])

        header_fill = PatternFill("solid", fgColor="16823A")
        header_font = Font(color="FFFFFF", bold=True)
        thin = Side(style="thin", color="808080")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = border
        for row_cells in ws.iter_rows(min_row=2):
            for cell in row_cells:
                cell.alignment = Alignment(vertical="center", wrap_text=True)
                cell.border = border

        for idx, column in enumerate(columns, start=1):
            max_len = max(
                len(str(column or "")),
                *[len(str(row.get(column, "") or "")) for row in rows],
            )
            ws.column_dimensions[ws.cell(row=1, column=idx).column_letter].width = min(max(max_len + 2, 12), 36)

        buf = io.BytesIO()
        wb.save(buf)
        return {
            "file_name": file_name,
            "bytes": buf.getvalue(),
            "content_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        }
    except Exception:
        buf = io.StringIO()
        writer = csv.writer(buf)
        writer.writerow(columns)
        for row in rows:
            writer.writerow([row.get(col, "") for col in columns])
        return {
            "file_name": re.sub(r"\.xlsx$", ".csv", file_name, flags=re.IGNORECASE),
            "bytes": buf.getvalue().encode("utf-8"),
            "content_type": "text/csv",
        }


_DSM_PENALTY_CONFIG_BY_STATE: Dict[str, Any] = {
    "Telangana": {
        "state": "Telangana",
        "byType": {
            "Solar": {
                "baseBand": 15,
                "bands": [
                    {"min": 0, "max": 15, "rate": 0},
                    {"min": 15, "max": 25, "rate": 0.5},
                    {"min": 25, "max": 35, "rate": 1.0},
                    {"min": 35, "max": float("inf"), "rate": 1.5},
                ],
            },
            "Wind": {
                "baseBand": 15,
                "bands": [
                    {"min": 0, "max": 15, "rate": 0},
                    {"min": 15, "max": 25, "rate": 0.5},
                    {"min": 25, "max": 35, "rate": 1.0},
                    {"min": 35, "max": float("inf"), "rate": 1.5},
                ],
            },
        },
    },
    "Madhya Pradesh": {
        "state": "Madhya Pradesh",
        "byType": {
            "Solar": {
                "baseBand": 10,
                "bands": [
                    {"min": 0, "max": 10, "rate": 0},
                    {"min": 10, "max": 15, "rate": 0.5},
                    {"min": 15, "max": 20, "rate": 0.75},
                    {"min": 20, "max": float("inf"), "rate": 1.0},
                ],
            },
            "Wind": {
                "baseBand": 15,
                "bands": [
                    {"min": 0, "max": 15, "rate": 0},
                    {"min": 15, "max": 20, "rate": 0.5},
                    {"min": 20, "max": 25, "rate": 0.75},
                    {"min": 25, "max": float("inf"), "rate": 1.0},
                ],
            },
        },
    },
}

_DSM_DEFAULT_PENALTY_CONFIG: Dict[str, Any] = {
    "state": "Default",
    "byType": {
        "Solar": {
            "baseBand": 10,
            "bands": [
                {"min": 0, "max": 10, "rate": 0},
                {"min": 10, "max": 12, "rate": 0.25},
                {"min": 12, "max": 15, "rate": 0.5},
                {"min": 15, "max": 25, "rate": 0.75},
                {"min": 25, "max": float("inf"), "rate": 1.0},
            ],
        },
        "Wind": {
            "baseBand": 12,
            "bands": [
                {"min": 0, "max": 12, "rate": 0},
                {"min": 12, "max": 15, "rate": 0.25},
                {"min": 15, "max": 20, "rate": 0.5},
                {"min": 20, "max": float("inf"), "rate": 1.0},
            ],
        },
    },
}


def _normalize_state_name(raw: str) -> str:
    return " ".join([p[:1].upper() + p[1:].lower() for p in str(raw or "").strip().split() if p])


def _dsm_get_allowed_band_percent(*, plant_state: str, plant_type: str) -> float:
    config = _DSM_PENALTY_CONFIG_BY_STATE.get(_normalize_state_name(plant_state)) or _DSM_DEFAULT_PENALTY_CONFIG
    by_type = (config.get("byType") or {}) if isinstance(config, dict) else {}
    type_cfg = by_type.get(plant_type) or by_type.get("Solar") or {}
    try:
        return float(type_cfg.get("baseBand"))
    except Exception:
        return 10.0


def _dsm_get_penalty_bands(*, plant_state: str, plant_type: str) -> List[Dict[str, Any]]:
    config = _DSM_PENALTY_CONFIG_BY_STATE.get(_normalize_state_name(plant_state)) or _DSM_DEFAULT_PENALTY_CONFIG
    by_type = (config.get("byType") or {}) if isinstance(config, dict) else {}
    type_cfg = by_type.get(plant_type) or by_type.get("Solar") or {}
    bands = type_cfg.get("bands")
    return bands if isinstance(bands, list) else []


def _dsm_calculate_penalty_rs(
    *,
    scheduled_mw: float,
    actual_mw: float,
    capacity_mw: float,
    plant_state: str,
    plant_type: str,
) -> float:
    eps = 1e-6
    if not isinstance(scheduled_mw, (int, float)) or not isinstance(actual_mw, (int, float)):
        return 0.0
    capacity = max(abs(float(capacity_mw or 0.0)), eps)
    deviation = float(actual_mw) - float(scheduled_mw)
    abs_dev_pct = abs((deviation / capacity) * 100.0)
    if not math.isfinite(abs_dev_pct) or abs_dev_pct <= 0:
        return 0.0

    band_pct = _dsm_get_allowed_band_percent(plant_state=plant_state, plant_type=plant_type)
    allowed_mw = (capacity * float(band_pct)) / 100.0
    if abs(deviation) <= allowed_mw + 1e-9:
        return 0.0

    lower = float(scheduled_mw) - allowed_mw
    upper = float(scheduled_mw) + allowed_mw
    under = (lower - float(actual_mw)) if float(actual_mw) < lower else 0.0
    over = (float(actual_mw) - upper) if float(actual_mw) > upper else 0.0
    excess = max(under, over, 0.0)
    if excess <= eps:
        return 0.0

    deviation_energy_kwh = abs(deviation) * 0.25 * 1000.0
    total = 0.0
    for band in _dsm_get_penalty_bands(plant_state=plant_state, plant_type=plant_type):
        try:
            bmin = float(band.get("min"))
            bmax = float(band.get("max"))
            rate = float(band.get("rate"))
        except Exception:
            continue
        span = min(abs_dev_pct, bmax) - bmin
        if span <= 0:
            continue
        band_energy = deviation_energy_kwh * (span / abs_dev_pct)
        total += band_energy * rate
    return float(total or 0.0)


def _osepl_mw_to_block_energy_kwh(mw: float) -> float:
    return float(mw) * 0.25 * 1000.0


def _osepl_calculate_payable_receivable_by_bands(
    *,
    deviation_kwh: float,
    avc_kwh: float,
    bands: List[Dict[str, float]],
) -> Dict[str, float]:
    error_pct_signed = (deviation_kwh / avc_kwh) * 100.0
    error_pct = abs(error_pct_signed)
    direction = "NONE"
    if deviation_kwh < 0:
        direction = "UNDER"
    elif deviation_kwh > 0:
        direction = "OVER"

    payable_rs = 0.0
    receivable_rs = 0.0
    band_min = 0.0
    for band in bands:
        band_max = float(band.get("maxErrorPercent", float("inf")))
        clamped_upper = min(error_pct, band_max)
        span = clamped_upper - band_min
        if span > 0:
            energy_slice_kwh = avc_kwh * (span / 100.0)
            if direction == "UNDER":
                payable_rs += energy_slice_kwh * float(band.get("underRate", 0.0))
            elif direction == "OVER":
                receivable_rs += energy_slice_kwh * float(band.get("overRate", 0.0))
        band_min = band_max
        if error_pct <= band_max:
            break

    return {
        "errorPctSigned": error_pct_signed,
        "errorPct": error_pct,
        "payableRs": payable_rs,
        "receivableRs": receivable_rs,
        "direction": direction,
    }


def _osepl_calculate_settlement(
    *,
    scheduled_mw: float,
    actual_mw: float,
    capacity_mw: float,
    ppa_rate: float = 9.27,
) -> Optional[Dict[str, float]]:
    if not (math.isfinite(scheduled_mw) and math.isfinite(actual_mw) and math.isfinite(capacity_mw)):
        return None
    if capacity_mw <= 0:
        return None

    # Match ESSEL/OSEPL workbook behavior: negative actual treated as zero penalty.
    if actual_mw < 0:
        scheduled_energy_kwh = _osepl_mw_to_block_energy_kwh(scheduled_mw)
        actual_energy_kwh = _osepl_mw_to_block_energy_kwh(actual_mw)
        avc_kwh = _osepl_mw_to_block_energy_kwh(capacity_mw)
        if not (avc_kwh > 0):
            return None
        deviation_kwh = actual_energy_kwh - scheduled_energy_kwh
        error_pct_signed = (deviation_kwh / avc_kwh) * 100.0
        return {
            "scheduledEnergyKwh": scheduled_energy_kwh,
            "actualEnergyKwh": actual_energy_kwh,
            "avcKwh": avc_kwh,
            "deviationKwh": deviation_kwh,
            "errorPctSigned": error_pct_signed,
            "errorPct": abs(error_pct_signed),
            "payableRs": 0.0,
            "receivableRs": 0.0,
            "finalPenaltyRs": 0.0,
        }

    scheduled_energy_kwh = _osepl_mw_to_block_energy_kwh(scheduled_mw)
    actual_energy_kwh = _osepl_mw_to_block_energy_kwh(actual_mw)
    deviation_kwh = actual_energy_kwh - scheduled_energy_kwh
    avc_kwh = _osepl_mw_to_block_energy_kwh(capacity_mw)
    if not (avc_kwh > 0):
        return None

    bands = [
        {"maxErrorPercent": 10.0, "underRate": ppa_rate, "overRate": ppa_rate},
        {"maxErrorPercent": 12.0, "underRate": ppa_rate * 1.1, "overRate": ppa_rate * 0.9},
        {"maxErrorPercent": 15.0, "underRate": ppa_rate * 1.2, "overRate": ppa_rate * 0.8},
        {"maxErrorPercent": float("inf"), "underRate": ppa_rate * 1.5, "overRate": 0.0},
    ]
    slab = _osepl_calculate_payable_receivable_by_bands(deviation_kwh=deviation_kwh, avc_kwh=avc_kwh, bands=bands)
    payable_rs = slab["payableRs"]
    receivable_rs = slab["receivableRs"]

    schedule_value_rs = scheduled_energy_kwh * ppa_rate
    actual_value_rs = actual_energy_kwh * ppa_rate
    final_penalty_rs = actual_value_rs - (schedule_value_rs + receivable_rs - payable_rs)
    return {
        "scheduledEnergyKwh": scheduled_energy_kwh,
        "actualEnergyKwh": actual_energy_kwh,
        "avcKwh": avc_kwh,
        "deviationKwh": deviation_kwh,
        "errorPctSigned": slab["errorPctSigned"],
        "errorPct": slab["errorPct"],
        "payableRs": payable_rs,
        "receivableRs": receivable_rs,
        "finalPenaltyRs": final_penalty_rs,
    }


def _osepl_calculate_office_payable_receivable(
    *,
    scheduled_mw: float,
    actual_mw: float,
    capacity_mw: float,
    ppa_rate: float = 9.27,
) -> Optional[Dict[str, float]]:
    if not (math.isfinite(scheduled_mw) and math.isfinite(actual_mw) and math.isfinite(capacity_mw)):
        return None
    if capacity_mw <= 0:
        return None

    # Office report treats negative actual as zero payable/receivable.
    if actual_mw < 0:
        return {"payableRs": 0.0, "receivableRs": 0.0}

    scheduled_energy_kwh = _osepl_mw_to_block_energy_kwh(scheduled_mw)
    actual_energy_kwh = _osepl_mw_to_block_energy_kwh(max(0.0, actual_mw))
    deviation_kwh = actual_energy_kwh - scheduled_energy_kwh
    avc_kwh = _osepl_mw_to_block_energy_kwh(capacity_mw)
    if not (avc_kwh > 0):
        return None

    bands = [
        {"maxErrorPercent": 10.0, "underRate": ppa_rate, "overRate": ppa_rate},
        {"maxErrorPercent": 12.0, "underRate": ppa_rate * 1.1, "overRate": ppa_rate * 0.9},
        {"maxErrorPercent": 15.0, "underRate": ppa_rate * 1.2, "overRate": ppa_rate * 0.8},
        {"maxErrorPercent": float("inf"), "underRate": ppa_rate * 1.5, "overRate": 0.0},
    ]
    slab = _osepl_calculate_payable_receivable_by_bands(deviation_kwh=deviation_kwh, avc_kwh=avc_kwh, bands=bands)
    return {"payableRs": slab["payableRs"], "receivableRs": slab["receivableRs"]}


def _parse_block_series_csv(text_value: Optional[str]) -> Dict[int, float]:
    raw = str(text_value or "")
    if not raw.strip():
        return {}
    reader = csv.reader(io.StringIO(raw))
    rows = list(reader)
    if not rows:
        return {}
    header = [str(c or "").strip() for c in (rows[0] or [])]
    start_idx = 1 if any(h for h in header) else 0
    header_lower = [h.lower() for h in header]

    def _find_col(pred) -> Optional[int]:
        for idx, h in enumerate(header_lower):
            if pred(h):
                return idx
        return None

    block_col = _find_col(lambda h: "block" in h)  # type: ignore[arg-type]
    value_col = _find_col(lambda h: "mw" in h or "value" in h or "scheduled" in h or "actual" in h or "meter" in h)  # type: ignore[arg-type]
    if block_col is None:
        block_col = 0
    if value_col is None:
        value_col = 1 if len(header) > 1 else 0

    out: Dict[int, float] = {}
    for r in rows[start_idx:]:
        if not r:
            continue
        try:
            block_raw = r[block_col] if block_col < len(r) else ""
            block = int(float(str(block_raw or "").strip()))
        except Exception:
            continue
        if block < 1 or block > 96:
            continue
        try:
            val_raw = r[value_col] if value_col < len(r) else ""
            val = float(str(val_raw or "").replace(",", "").strip())
        except Exception:
            continue
        if not math.isfinite(val):
            continue
        out[block] = val
    return out


def _to_header_key(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", str(value or "").strip().lower())


def _detect_csv_delimiter(sample: str) -> str:
    s = str(sample or "")
    comma = s.count(",")
    semi = s.count(";")
    tab = s.count("\t")
    # Prefer the most common delimiter in the header row.
    if semi > comma and semi >= tab and semi > 0:
        return ";"
    if tab > comma and tab > semi and tab > 0:
        return "\t"
    return ","


def _parse_csv_with_header_detection(text_value: Optional[str]) -> Tuple[List[str], List[List[str]]]:
    raw = str(text_value or "")
    if not raw.strip():
        return [], []
    raw_lines = [line for line in raw.splitlines() if str(line or "").strip()]
    if not raw_lines:
        return [], []

    def score_header_line(line: str) -> int:
        lowered = str(line or "").lower()
        if not any(d in lowered for d in [",", ";", "\t"]):
            return -1
        score = 0
        if re.search(r"\bblock\b|\bblk\b|\bs\.?\s*no\b|\bsno\b", lowered):
            score += 5
        if re.search(r"\btime\b|\btimestamp\b|\bdate\b", lowered):
            score += 4
        if re.search(r"meter|actual|forecast|sch[^a-z0-9]*mw|schedule", lowered):
            score += 6
        if re.search(r"mw|kw|power|generation", lowered):
            score += 2
        return score

    best_idx = 0
    best_score = -1
    for idx, line in enumerate(raw_lines[:25]):
        score = score_header_line(line)
        if score > best_score:
            best_idx = idx
            best_score = score

    delim = _detect_csv_delimiter(raw_lines[best_idx] if best_idx < len(raw_lines) else raw_lines[0])

    def parse_line(line: str) -> List[str]:
        try:
            return [str(c or "").strip() for c in next(csv.reader([line], delimiter=delim))]
        except Exception:
            return [str(line or "").strip()]

    lines = [parse_line(line) for line in raw_lines]
    headers = [str(c or "").replace("\ufeff", "").strip() for c in (lines[best_idx] or [])]
    header2 = [str(c or "").replace("\ufeff", "").strip() for c in (lines[best_idx + 1] or [])] if len(lines) > best_idx + 1 else []
    use_second = any(re.search(r"(forecast|availability)", str(h or ""), flags=re.IGNORECASE) for h in header2)

    max_cols = max(len(headers), len(header2))
    merged: List[str] = []
    for i in range(max_cols):
        h1 = headers[i] if i < len(headers) else ""
        h2 = header2[i] if i < len(header2) else ""
        if use_second and h1 and h2:
            merged.append(f"{h1} {h2}".strip())
        else:
            merged.append((h1 or h2).strip())

    data_start = best_idx + (2 if use_second else 1)
    rows = [[str(c or "").strip() for c in (r or [])] for r in lines[data_start:]]
    return merged, rows


def _parse_block_number(raw: Any) -> Optional[int]:
    if raw is None:
        return None
    text = str(raw).strip()
    if not text:
        return None
    try:
        direct = int(text)
        return direct
    except Exception:
        pass
    m = re.search(r"([0-9]{1,3})", text)
    if not m:
        return None
    try:
        return int(m.group(1))
    except Exception:
        return None


def _parse_schedule_series_map(text_value: Optional[str], plant_code: str = "") -> Dict[int, float]:
    headers, rows = _parse_csv_with_header_detection(text_value)
    normalized = [_to_header_key(h) for h in (headers or [])]
    block_idx = next((i for i, h in enumerate(normalized) if ("block" in h or "blk" in h or h == "sno" or "srno" in h)), -1)
    if block_idx == -1:
        return {}

    def find_idx(pred) -> int:
        for i, h in enumerate(normalized):
            if pred(h):
                return i
        return -1

    site_code = _normalize_plant_code(plant_code)
    schedule_idx = -1
    if site_code == "OSEPL":
        schedule_idx = find_idx(lambda h: ("declared" in h and "forecast" in h))
    if schedule_idx == -1:
        schedule_idx = find_idx(lambda h: ("stationschedule" in h and "availability" not in h and "capacity" not in h))
    if schedule_idx == -1:
        schedule_idx = find_idx(lambda h: ("schedule" in h and "mw" in h))
    if schedule_idx == -1:
        schedule_idx = find_idx(lambda h: ("schedule" in h or "schmw" in h or ("sch" in h and "mw" in h)))
    if schedule_idx == -1:
        schedule_idx = find_idx(lambda h: ("forecast" in h and "forcastavailability" not in h))
    if schedule_idx == -1:
        schedule_idx = find_idx(lambda h: ("mw" in h and "meter" not in h and "actual" not in h and "kw" not in h))
    if schedule_idx == -1:
        return {}

    is_osepl_end_block_template = (
        site_code == "OSEPL"
        and not any(("time" in h or "from" in h or "to" in h) for h in normalized)
        and "declaredforecast" in normalized
        and "interavc" in normalized
        and "schedule" in normalized
    )

    out: Dict[int, float] = {}
    for cols in (rows or []):
        parsed_block = _parse_block_number(cols[block_idx] if block_idx < len(cols) else None)
        block = (
            parsed_block + 1
            if (is_osepl_end_block_template and isinstance(parsed_block, int) and parsed_block >= 1)
            else parsed_block
        )
        if not isinstance(block, int) or block < 1 or block > 96:
            continue
        try:
            raw_val = cols[schedule_idx] if schedule_idx < len(cols) else ""
            val = float(str(raw_val or "").replace(",", "").strip())
        except Exception:
            continue
        if not math.isfinite(val):
            continue
        out[int(block)] = float(val)
    return out


def _meter_distance_to_block_start_seconds(raw_time: Any, block: int) -> Optional[float]:
    m = re.search(r"(\d{1,2}):(\d{2})(?::(\d{2}))?", str(raw_time or ""))
    if not m:
        return None
    try:
        hh = int(m.group(1))
        mm = int(m.group(2))
        ss = int(m.group(3) or "0")
    except Exception:
        return None
    minutes = (hh * 60) + mm + (ss / 60.0)
    block_start_minutes = (max(1, int(block)) - 1) * 15
    return abs((minutes - block_start_minutes) * 60.0)


def _meter_is_midnight_carry_row(raw_time: Any) -> bool:
    t = str(raw_time or "").strip()
    if not t:
        return False
    return bool(re.match(r"^0?0:0?0(?::0?0)?$", t) or re.match(r"^24:00(?::00)?$", t))


def _meter_parse_block_from_start_timestamp(raw: Any, *, total_blocks: int = 96) -> Optional[int]:
    m = re.search(r"(\d{1,2}):(\d{2})(?::(\d{2}))?(?:\.(\d{1,3}))?", str(raw or "").strip())
    if not m:
        return None
    try:
        hh = int(m.group(1))
        mm = int(m.group(2))
        ss = int(m.group(3) or "0")
        ms = int((m.group(4) or "0").ljust(3, "0")[:3])
    except Exception:
        return None
    if hh == 24:
        if mm == 0 and ss == 0 and ms == 0:
            return total_blocks
        return None
    if not (0 <= hh <= 23 and 0 <= mm <= 59 and 0 <= ss <= 59 and 0 <= ms <= 999):
        return None
    total_minutes = (hh * 60) + mm
    block = int(total_minutes // 15) + 1
    return block if 1 <= block <= total_blocks else None


def _meter_parse_block_from_end_timestamp(raw: Any, *, total_blocks: int = 96) -> Optional[int]:
    """
    Match frontend `parseBlockFromTimestamp`: treat timestamp as END of 15-min block
    (00:15 => block 1). Seconds/millis ignored to avoid accidental shifts.
    """
    m = re.search(r"(\d{1,2}):(\d{2})(?::(\d{2}))?(?:\.(\d{1,3}))?", str(raw or "").strip())
    if not m:
        return None
    try:
        hh = int(m.group(1))
        mm = int(m.group(2))
        ss = int(m.group(3) or "0")
        ms = int((m.group(4) or "0").ljust(3, "0")[:3])
    except Exception:
        return None
    if hh == 24:
        if mm == 0 and ss == 0 and ms == 0:
            return total_blocks
        return None
    if not (0 <= hh <= 23 and 0 <= mm <= 59):
        return None
    total_minutes = (hh * 60) + mm
    if total_minutes <= 0:
        return None
    block = int(math.ceil(total_minutes / 15.0))
    return block if 1 <= block <= total_blocks else None


def _meter_parse_block_from_nearest_quarter_start(raw: Any, *, total_blocks: int = 96) -> Optional[int]:
    """
    Match frontend `parseBlockFromNearestQuarterStart`.
    """
    m = re.search(r"(\d{1,2}):(\d{2})(?::(\d{2}))?(?:\.(\d{1,3}))?", str(raw or "").strip())
    if not m:
        return None
    try:
        hh = int(m.group(1))
        mm = int(m.group(2))
        ss = int(m.group(3) or "0")
        ms = int((m.group(4) or "0").ljust(3, "0")[:3])
    except Exception:
        return None
    if hh == 24:
        if mm == 0 and ss == 0 and ms == 0:
            return total_blocks
        return None
    if not (0 <= hh <= 23 and 0 <= mm <= 59 and 0 <= ss <= 59 and 0 <= ms <= 999):
        return None
    total_minutes = (hh * 60) + mm + (ss / 60.0) + (ms / 60000.0)
    rounded_minutes = math.floor((total_minutes + 7.5) / 15.0) * 15.0
    block = int(math.floor(rounded_minutes / 15.0) + 1)
    return block if 1 <= block <= total_blocks else None


def _meter_build_time_block_resolver(rows: List[List[str]], time_idx: int) -> Any:
    """
    Choose the same time->block convention as the UI by scoring 3 strategies:
    - start of interval
    - end of interval
    - nearest quarter start
    """
    if time_idx == -1:
        return lambda _raw: None

    def _score(resolver) -> int:
        seen: set[int] = set()
        parsed = 0
        duplicates = 0
        sample = rows[: min(len(rows), 200)]
        for cols in sample:
            raw = cols[time_idx] if time_idx < len(cols) else None
            val = str(raw or "").strip()
            if not val:
                continue
            range_match = re.search(r"(\d{1,2}:\d{2})(?:\s*[-\u2013\u2014]\s*)(\d{1,2}:\d{2})", val)
            probe = range_match.group(1) if range_match else val
            block = resolver(probe)
            if not isinstance(block, int) or block < 1 or block > 96:
                continue
            parsed += 1
            if block in seen:
                duplicates += 1
            seen.add(block)
        unique = len(seen)
        missing = max(0, 96 - unique)
        return (unique * 100) - (duplicates * 25) - (missing * 10) + parsed

    start_res = lambda t: _meter_parse_block_from_start_timestamp(t, total_blocks=96)
    end_res = lambda t: _meter_parse_block_from_end_timestamp(t, total_blocks=96)
    near_res = lambda t: _meter_parse_block_from_nearest_quarter_start(t, total_blocks=96)

    rank = {"end": 3, "nearest": 2, "start": 1}
    candidates = [("start", start_res), ("end", end_res), ("nearest", near_res)]
    best_name, best_res = max(candidates, key=lambda item: (_score(item[1]), rank.get(item[0], 0)))

    def resolve(raw_time: Any) -> Optional[int]:
        text_val = str(raw_time or "").strip()
        if not text_val:
            return None
        range_match = re.search(r"(\d{1,2}:\d{2})(?:\s*[-\u2013\u2014]\s*)(\d{1,2}:\d{2})", text_val)
        if range_match:
            start_block = _meter_parse_block_from_start_timestamp(range_match.group(1), total_blocks=96)
            if isinstance(start_block, int):
                return start_block
            end_block = _meter_parse_block_from_end_timestamp(range_match.group(1), total_blocks=96)
            return end_block if isinstance(end_block, int) else None
        if best_name == "start":
            start_block = _meter_parse_block_from_start_timestamp(text_val, total_blocks=96)
            if isinstance(start_block, int):
                return start_block
            nearest_block = _meter_parse_block_from_nearest_quarter_start(text_val, total_blocks=96)
            if isinstance(nearest_block, int):
                return nearest_block
            return _meter_parse_block_from_end_timestamp(text_val, total_blocks=96)
        if best_name == "end":
            end_block = _meter_parse_block_from_end_timestamp(text_val, total_blocks=96)
            if isinstance(end_block, int):
                return end_block
            nearest_block = _meter_parse_block_from_nearest_quarter_start(text_val, total_blocks=96)
            if isinstance(nearest_block, int):
                return nearest_block
            return _meter_parse_block_from_start_timestamp(text_val, total_blocks=96)
        nearest = _meter_parse_block_from_nearest_quarter_start(text_val, total_blocks=96)
        if isinstance(nearest, int):
            return nearest
        start_block = _meter_parse_block_from_start_timestamp(text_val, total_blocks=96)
        if isinstance(start_block, int):
            return start_block
        return _meter_parse_block_from_end_timestamp(text_val, total_blocks=96)

    return resolve


def _parse_meter_series_map(text_value: Optional[str]) -> Dict[int, float]:
    headers, rows = _parse_csv_with_header_detection(text_value)
    normalized = [_to_header_key(h) for h in (headers or [])]
    block_idx = next((i for i, h in enumerate(normalized) if ("block" in h or "blk" in h or h == "sno" or "srno" in h)), -1)
    time_idx = next((i for i, h in enumerate(normalized) if (("time" in h) or ("timestamp" in h) or (h == "date") or ("date" in h) or ("from" in h) or ("to" in h))), -1)

    def find_idx(pred) -> int:
        for i, h in enumerate(normalized):
            if pred(h):
                return i
        return -1

    power_idx = find_idx(lambda h: ("meter" in h and ("mw" in h or "kw" in h or "power" in h)))
    if power_idx == -1:
        power_idx = find_idx(lambda h: ("meterpower" in h or ("meter" in h and "power" in h)))
    if power_idx == -1:
        power_idx = find_idx(lambda h: ("actual" in h and ("mw" in h or "kw" in h or "power" in h)))
    if power_idx == -1:
        power_idx = find_idx(lambda h: (("mw" in h or "kw" in h) and "schedule" not in h and "sch" not in h))
    if power_idx == -1:
        return {}

    power_header = str(normalized[power_idx] or "")
    explicit_kw = ("kw" in power_header) and ("mw" not in power_header)
    explicit_mw = ("mw" in power_header)

    get_block_from_time_text = _meter_build_time_block_resolver(rows or [], time_idx)

    points: List[Dict[str, Any]] = []
    for idx, cols in enumerate(rows or []):
        block_from_col = _parse_block_number(cols[block_idx]) if (block_idx != -1 and block_idx < len(cols)) else None
        time_raw = cols[time_idx] if (time_idx != -1 and time_idx < len(cols)) else None
        block_from_time = get_block_from_time_text(time_raw) if time_idx != -1 else None
        fallback_block = idx + 1

        block: Optional[int] = None
        if isinstance(block_from_col, int) and 1 <= block_from_col <= 96:
            block = int(block_from_col)
        elif isinstance(block_from_time, int) and 1 <= block_from_time <= 96:
            block = int(block_from_time)
        elif block_idx == -1 and time_idx == -1 and 1 <= fallback_block <= 96:
            block = int(fallback_block)

        if not isinstance(block, int):
            continue
        try:
            raw_val = cols[power_idx] if power_idx < len(cols) else ""
            value = float(str(raw_val or "").replace(",", "").strip())
        except Exception:
            continue
        if not math.isfinite(value):
            continue
        points.append(
            {
                "block": block,
                "value": value,
                "idx": idx,
                "timeRaw": time_raw,
                "dist": _meter_distance_to_block_start_seconds(time_raw, block),
            }
        )

    if not points:
        return {}

    should_drop_midnight = block_idx == -1 and time_idx != -1 and len(points) > 96
    norm_points = [p for p in points if not (should_drop_midnight and _meter_is_midnight_carry_row(p.get("timeRaw")))]

    non_zero = [p["value"] for p in points if isinstance(p.get("value"), (int, float)) and p["value"] > 0]
    avg = (sum(non_zero) / len(non_zero)) if non_zero else 0.0
    assume_kw = bool(explicit_kw or (not explicit_mw and avg > 200))
    factor = (1.0 / 1000.0) if assume_kw else 1.0

    best_by_block: Dict[int, Dict[str, Any]] = {}

    def prefer(curr: Optional[Dict[str, Any]], inc: Dict[str, Any]) -> bool:
        if curr is None:
            return True
        a = curr.get("dist")
        b = inc.get("dist")
        a_has = isinstance(a, (int, float)) and math.isfinite(float(a))
        b_has = isinstance(b, (int, float)) and math.isfinite(float(b))
        if a_has and b_has and float(a) != float(b):
            return float(b) < float(a)
        if b_has and not a_has:
            return True
        if a_has and not b_has:
            return False
        return int(inc.get("idx") or 0) > int(curr.get("idx") or 0)

    for p in norm_points:
        b = int(p["block"])
        if prefer(best_by_block.get(b), p):
            best_by_block[b] = p

    out: Dict[int, float] = {}
    for b, p in best_by_block.items():
        out[int(b)] = float(p["value"]) * factor
    return out

def _email_scheduler_build_daily_dsm_row_from_s3(
    *,
    s3_client: Any,
    bucket: str,
    plant_code: str,
    plant_name: str,
    report_date: str,
) -> Optional[Dict[str, Any]]:
    pcode = _normalize_plant_code(plant_code)
    if not pcode or not re.fullmatch(r"\d{4}-\d{2}-\d{2}", str(report_date or "").strip()):
        return None
    day = str(report_date).strip()
    context = _email_scheduler_build_template_context(day)
    month_label = f"{context['month_short']}-{context['year_short']}"

    schedule_text = None
    schedule_prefix = _email_scheduler_sldc_schedule_prefix(pcode, day)
    if schedule_prefix and _s3_proxy_is_allowed_path(schedule_prefix):
        try:
            schedule_objects = _list_s3_objects_paginated(
                s3_client=s3_client,
                bucket=bucket,
                prefix=schedule_prefix,
                max_items=2000,
            )
            schedule_pick = _email_scheduler_pick_latest_sldc_schedule(schedule_objects)
            schedule_key = str((schedule_pick or {}).get("key") or "").strip()
            if schedule_key and _s3_proxy_is_allowed_path(schedule_key):
                obj = s3_client.get_object(Bucket=bucket, Key=schedule_key)
                body = obj.get("Body")
                data = body.read() if body is not None else b""
                schedule_text = _email_scheduler_schedule_bytes_to_csv_text(os.path.basename(schedule_key), data)
        except Exception:
            schedule_text = None
    schedule_map = _parse_schedule_series_map(schedule_text, pcode)
    if not schedule_map:
        return None

    meter_text = None
    # Scan known prefixes and pick the latest CSV (match UI DSM preview behavior).
    if not meter_text:
        meter_prefixes = (
            [f"raw/vedanjay/multiple_generator/ZTRIC/{day}/metered_data/"]
            if _normalize_plant_code(pcode) == "ZETRIC"
            else [
                *[f"raw/vedanjay/{folder}/{day}/metered_data/" for folder in _raw_plant_folder_aliases(pcode)],
                f"generated/vedanjay/{pcode}/outputs/{day}/meter/",
                f"outputs/{day}/meter/",
                f"{day}/meter/",
            ]
        )
        meter_objects: List[Dict[str, str]] = []
        for prefix in meter_prefixes:
            if not prefix or not _s3_proxy_is_allowed_path(prefix):
                continue
            try:
                meter_objects.extend(_list_s3_objects_paginated(s3_client=s3_client, bucket=bucket, prefix=prefix, max_items=2000))
            except Exception:
                continue
        meter_pick = _pick_latest_csv(meter_objects, prefer_suffix=".csv")
        if meter_pick:
            key = str(meter_pick.get("key") or "").strip()
            if key and _s3_proxy_is_allowed_path(key):
                meter_text = _read_s3_text_safe(s3_client, bucket, key)
    meter_map = _parse_meter_series_map(meter_text)
    if not meter_map:
        return None

    capacity_map = {
        "ANJANGAON": 7.5,
        "SIRMOUR": 5.1,
        "BHUPALPALLY": 10.0,
        "KASIPET": 15.0,
        "KOTHAGUDEM": 37.0,
        "OSEPL": 20.0,
        "ANDAD": 7.5,
        "BALAKWADA": 7.5,
        "GUGARIYAKHEDI": 7.5,
        "NANDGAON": 7.5,
        "BAMKHAL": 5.0,
    }
    capacity = float(capacity_map.get(pcode) or 0.0)
    plant_state_map = {
        "ANJANGAON": "Madhya Pradesh",
        "BHUPALPALLY": "Telangana",
        "KASIPET": "Telangana",
        "KOTHAGUDEM": "Telangana",
        "OSEPL": "Maharashtra",
        "ANDAD": "Madhya Pradesh",
        "BALAKWADA": "Madhya Pradesh",
        "GUGARIYAKHEDI": "Madhya Pradesh",
        "NANDGAON": "Madhya Pradesh",
        "BAMKHAL": "Madhya Pradesh",
        "SIRMOUR": "Madhya Pradesh",
        "SAWDA": "Madhya Pradesh",
        "ZETRIC": "Maharashtra",
    }
    plant_state = str(plant_state_map.get(pcode) or "").strip()
    plant_type = "Solar"

    def _current_ist_block() -> int:
        now_ist = datetime.now(timezone.utc).astimezone(ZoneInfo("Asia/Kolkata"))
        total_minutes = (now_ist.hour * 60) + now_ist.minute
        block = (total_minutes // 15) + 1
        return min(max(int(block), 1), 96)

    is_today_ist = day == datetime.now(timezone.utc).astimezone(ZoneInfo("Asia/Kolkata")).date().isoformat()
    is_bhupalpally_dsm = pcode == "BHUPALPALLY"
    dsm_block_limit = _current_ist_block() if is_bhupalpally_dsm and is_today_ist else 96

    def _schedule_mw_for_dsm(block: int) -> Optional[float]:
        sched = schedule_map.get(block)
        if sched is None:
            return None
        value = float(sched)
        if is_bhupalpally_dsm:
            return round(value + 1e-12, 2)
        return value

    # Match UI: generation is derived purely from meter values (independent of schedule presence).
    generation_kwh = 0.0
    for block in range(1, dsm_block_limit + 1):
        act = meter_map.get(block)
        if act is None:
            continue
        generation_kwh += float(act) * 0.25 * 1000.0

    penalty_rs = 0.0
    for block in range(1, dsm_block_limit + 1):
        sched = _schedule_mw_for_dsm(block)
        act = meter_map.get(block)
        if sched is None or act is None:
            continue
        penalty_rs += _dsm_calculate_penalty_rs(
            scheduled_mw=float(sched),
            actual_mw=float(act),
            capacity_mw=capacity,
            plant_state=plant_state,
            plant_type=plant_type,
        )

    paisa_per_kwh = "--"
    if generation_kwh > 0:
        paisa_per_kwh = f"{(penalty_rs / generation_kwh) * 100.0:.2f}"

    if pcode == "SIRMOUR":
        # Match UI's SIRMOUR DSM summary format (EmailScheduler.jsx).
        ppa_rate = 2.94
        net_revenue = generation_kwh * ppa_rate if generation_kwh > 0 else 0.0
        impact_pct = (penalty_rs / net_revenue) * 100.0 if net_revenue > 0 else 0.0
        support_details: List[Dict[str, Any]] = []
        block_energy_factor = 0.25 * 1000.0
        avc_kwh_active = capacity * block_energy_factor if capacity > 0 else 0.0
        for block in range(1, 97):
            sched = _schedule_mw_for_dsm(block)
            act = meter_map.get(block)
            scheduled_mw = float(sched) if sched is not None else 0.0
            actual_mw = float(act) if act is not None else 0.0
            schedule_kwh = scheduled_mw * block_energy_factor
            meter_kwh = actual_mw * block_energy_factor
            avc_kwh = avc_kwh_active if schedule_kwh > 0 else 0.0
            pct_error = (abs(schedule_kwh - meter_kwh) / avc_kwh * 100.0) if avc_kwh > 0 else 0.0
            block_penalty = (
                _dsm_calculate_penalty_rs(
                    scheduled_mw=scheduled_mw,
                    actual_mw=actual_mw,
                    capacity_mw=capacity,
                    plant_state=plant_state,
                    plant_type=plant_type,
                )
                if schedule_kwh > 0 and act is not None
                else 0.0
            )
            support_details.append(
                {
                    "Datetime(Date+Block endtime)": _email_scheduler_sirmour_block_end_timestamp(day, block),
                    "Schedule(Kwh)": _email_scheduler_round_number(schedule_kwh, 2),
                    "Meter data(KWh)": _email_scheduler_round_number(meter_kwh, 2),
                    "AvC(Kwh)": _email_scheduler_round_number(avc_kwh, 2),
                    "% Error": _email_scheduler_round_number(pct_error, 2),
                    "DSM Penalty": _email_scheduler_round_number(block_penalty, 2),
                }
            )
        return {
            "From": day,
            "To": day,
            "Project": "Sirmour_Schedule",
            "Installed Capacity (MW)": f"{capacity:.1f}" if capacity else "0",
            "Generation (kWh)": f"{generation_kwh:.0f}",
            "DSM Penalty (Rs.)": f"{penalty_rs:.0f}",
            "Paisa / kWh": paisa_per_kwh,
            "Net Revenue": f"{net_revenue:.2f}",
            "%Impact": f"{impact_pct:.2f}%",
            "__support_summary": {
                "From": _email_scheduler_sirmour_block_end_timestamp(day, 1),
                "To": _email_scheduler_sirmour_block_end_timestamp(day, 95),
                "Project": "Sirmour_Schedule",
                "Installed Capacity (Mw)": f"{capacity:.1f}" if capacity else "0",
                "Generation(Kwh)": _email_scheduler_round_number(generation_kwh, 2),
                "DSM Penalty(Rs.)": _email_scheduler_round_number(penalty_rs, 2),
                "Paisa/Kwh": _email_scheduler_round_number((penalty_rs / generation_kwh) * 100.0, 2) if generation_kwh > 0 else 0,
                "Net Revenue": _email_scheduler_round_number(net_revenue, 2),
                "%Impact": f"{impact_pct:.2f}%",
            },
            "__support_details": support_details,
        }

    if pcode == "OSEPL":
        # Match UI's OSEPL DSM daily summary format (EmailScheduler.jsx).
        ppa_rate = 9.27
        scheduled_unit_ppa_block_limit = _current_ist_block() if is_today_ist else 96
        scheduled_kwh = 0.0
        for block in range(1, scheduled_unit_ppa_block_limit + 1):
            sched = schedule_map.get(block)
            if sched is None:
                continue
            scheduled_kwh += round(float(sched) + 1e-12, 2) * 0.25 * 1000.0

        totals_payable = 0.0
        totals_receivable = 0.0
        totals_final = 0.0
        support_details: List[Dict[str, Any]] = []
        block_energy_factor = 0.25 * 1000.0
        avc_kwh = capacity * block_energy_factor if capacity > 0 else 0.0
        for block in range(1, dsm_block_limit + 1):
            sched = schedule_map.get(block)
            act = meter_map.get(block)
            if sched is None or act is None:
                continue
            cap_mw = capacity if capacity > 0 else 0.0
            settlement = _osepl_calculate_settlement(
                scheduled_mw=float(sched),
                actual_mw=float(act),
                capacity_mw=cap_mw,
                ppa_rate=ppa_rate,
            )
            office = _osepl_calculate_office_payable_receivable(
                scheduled_mw=float(sched),
                actual_mw=float(act),
                capacity_mw=cap_mw,
                ppa_rate=ppa_rate,
            )
            if office:
                totals_payable += float(office.get("payableRs") or 0.0)
                totals_receivable += float(office.get("receivableRs") or 0.0)
            if settlement:
                totals_final += float(settlement.get("finalPenaltyRs") or 0.0)

        def _osepl_maps_for_day(detail_day: str) -> Tuple[Dict[int, float], Dict[int, float]]:
            if detail_day == day:
                return schedule_map, meter_map
            detail_schedule_text = None
            detail_schedule_prefix = _email_scheduler_sldc_schedule_prefix(pcode, detail_day)
            if detail_schedule_prefix and _s3_proxy_is_allowed_path(detail_schedule_prefix):
                try:
                    detail_schedule_objects = _list_s3_objects_paginated(
                        s3_client=s3_client,
                        bucket=bucket,
                        prefix=detail_schedule_prefix,
                        max_items=2000,
                    )
                    detail_schedule_pick = _email_scheduler_pick_latest_sldc_schedule(detail_schedule_objects)
                    detail_schedule_key = str((detail_schedule_pick or {}).get("key") or "").strip()
                    if detail_schedule_key and _s3_proxy_is_allowed_path(detail_schedule_key):
                        obj = s3_client.get_object(Bucket=bucket, Key=detail_schedule_key)
                        body = obj.get("Body")
                        data = body.read() if body is not None else b""
                        detail_schedule_text = _email_scheduler_schedule_bytes_to_csv_text(os.path.basename(detail_schedule_key), data)
                except Exception:
                    detail_schedule_text = None
            detail_schedule_map = _parse_schedule_series_map(detail_schedule_text, pcode)

            detail_meter_text = None
            detail_meter_prefixes = [
                *[f"raw/vedanjay/{folder}/{detail_day}/metered_data/" for folder in _raw_plant_folder_aliases(pcode)],
                f"generated/vedanjay/{pcode}/outputs/{detail_day}/meter/",
                f"outputs/{detail_day}/meter/",
                f"{detail_day}/meter/",
            ]
            detail_meter_objects: List[Dict[str, str]] = []
            for prefix in detail_meter_prefixes:
                if not prefix or not _s3_proxy_is_allowed_path(prefix):
                    continue
                try:
                    detail_meter_objects.extend(_list_s3_objects_paginated(s3_client=s3_client, bucket=bucket, prefix=prefix, max_items=2000))
                except Exception:
                    continue
            detail_meter_pick = _pick_latest_csv(detail_meter_objects, prefer_suffix=".csv")
            if detail_meter_pick:
                key = str(detail_meter_pick.get("key") or "").strip()
                if key and _s3_proxy_is_allowed_path(key):
                    detail_meter_text = _read_s3_text_safe(s3_client, bucket, key)
            detail_meter_map = _parse_meter_series_map(detail_meter_text)
            return detail_schedule_map, detail_meter_map

        try:
            parsed_day = datetime.strptime(day, "%Y-%m-%d").date()
            detail_days = [date(parsed_day.year, parsed_day.month, d).isoformat() for d in range(1, parsed_day.day + 1)]
        except Exception:
            detail_days = [day]

        for detail_day in detail_days:
            detail_schedule_map, detail_meter_map = _osepl_maps_for_day(detail_day)
            for block in range(1, 97):
                sched = detail_schedule_map.get(block)
                act = detail_meter_map.get(block)
                scheduled_mw = round(float(sched) + 1e-12, 2) if sched is not None else 0.0
                actual_mw = float(act) if act is not None else 0.0
                forecast_kwh = scheduled_mw * block_energy_factor
                actual_kwh = actual_mw * block_energy_factor
                pct_error = ((forecast_kwh - actual_kwh) / avc_kwh * 100.0) if avc_kwh > 0 else 0.0
                cap_mw = capacity if capacity > 0 else 0.0
                office = _osepl_calculate_office_payable_receivable(
                    scheduled_mw=scheduled_mw,
                    actual_mw=actual_mw,
                    capacity_mw=cap_mw,
                    ppa_rate=ppa_rate,
                )
                settlement = _osepl_calculate_settlement(
                    scheduled_mw=scheduled_mw,
                    actual_mw=actual_mw,
                    capacity_mw=cap_mw,
                    ppa_rate=ppa_rate,
                )
                payable = float((office or {}).get("payableRs") or 0.0)
                receivable = float((office or {}).get("receivableRs") or 0.0)
                final_penalty = float((settlement or {}).get("finalPenaltyRs") or 0.0)
                support_details.append(
                    {
                        "TimeSlots(date+endtime)": _email_scheduler_block_end_timestamp(detail_day, block),
                        "Forecast(kwH)": _email_scheduler_round_number(forecast_kwh, 2),
                        "Actual(Kwh)": _email_scheduler_round_number(actual_kwh, 2),
                        "AvC(kwh)": _email_scheduler_round_number(avc_kwh, 2),
                        "% Error": _email_scheduler_round_number(pct_error, 2),
                        "Scheduled unit*PPA": _email_scheduler_round_number(forecast_kwh * ppa_rate, 2),
                        "Payable": _email_scheduler_round_number(payable, 2),
                        "Receivable": _email_scheduler_round_number(receivable, 2),
                        "Total": _email_scheduler_round_number((forecast_kwh * ppa_rate) - payable + receivable, 2),
                        "generator-end penalty(Rs.)": _email_scheduler_round_number(final_penalty, 2),
                        "Scada/Maintenance Information Availability": 0,
                        "Penalty As per Scada Information": _email_scheduler_round_number(final_penalty, 2),
                    }
                )

        return {
            "From": day,
            "Month": month_label,
            "Project": "ESSEL",
            "Installed Capacity": f"{capacity:.0f}" if capacity else "0",
            "SCADA availability %": "100%",
            "Generation(kWh)": f"{generation_kwh:.0f}",
            "Scheduled unit*PPA": f"{(scheduled_kwh * ppa_rate):.0f}",
            "Payable": f"{totals_payable:.0f}",
            "Receivable": f"{totals_receivable:.0f}",
            "DSM Penalty (Rs.)": f"{totals_final:.0f}",
            "SCADA Adjusted DSM": f"{totals_final:.0f}",
            "PPA": f"{ppa_rate:.2f}",
            "__support_details": support_details,
        }

    support_details: List[Dict[str, Any]] = []
    block_energy_factor = 0.25 * 1000.0
    avc_kwh = capacity * block_energy_factor if capacity > 0 else 0.0
    for block in range(1, 97):
        sched = _schedule_mw_for_dsm(block)
        act = meter_map.get(block)
        scheduled_mw = float(sched) if sched is not None else 0.0
        actual_mw = float(act) if act is not None else 0.0
        schedule_kwh = scheduled_mw * block_energy_factor
        meter_kwh = actual_mw * block_energy_factor
        pct_error = (abs(schedule_kwh - meter_kwh) / avc_kwh * 100.0) if avc_kwh > 0 else 0.0
        block_penalty = (
            _dsm_calculate_penalty_rs(
                scheduled_mw=scheduled_mw,
                actual_mw=actual_mw,
                capacity_mw=capacity,
                plant_state=plant_state,
                plant_type=plant_type,
            )
            if sched is not None and act is not None
            else 0.0
        )
        support_details.append(
            {
                "Datetime(Date+Block endtime)": _email_scheduler_block_end_timestamp(day, block),
                "Schedule(Kwh)": _email_scheduler_round_number(schedule_kwh, 2),
                "Meter data(KWh)": _email_scheduler_round_number(meter_kwh, 2),
                "AvC(Kwh)": _email_scheduler_round_number(avc_kwh, 2),
                "% Error": _email_scheduler_round_number(pct_error, 2),
                "DSM penalty": _email_scheduler_round_number(block_penalty, 2),
                "Maintenance Update": 0,
                "DSM penalty as per Maintenance Updates": _email_scheduler_round_number(block_penalty, 2),
            }
        )

    return {
        "DATE": day,
        "TO": day,
        "MONTH": month_label,
        "PROJECT": str(plant_name or pcode).strip().upper() or pcode,
        "INSTALLED CAPACITY (MW)": str(int(round(capacity))) if capacity else "0",
        "GENERATION (KWH)": f"{generation_kwh:.0f}",
        "DSM PENALTY (RS.), AS PER SCADA AVAILABILITY": f"{penalty_rs:.0f}",
        "DSM PENALTY (RS.), AS MAINTENANCE INFORMATION": f"{penalty_rs:.0f}",
        "PAISA/KWH SCADA AVAILABILITY": paisa_per_kwh,
        "PAISA/KWH MAINTENANCE INFORMATION": paisa_per_kwh,
        "SCADA AVAILABILITY(%)": "100%",
        "__support_details": support_details,
    }


_EMAIL_SCHEDULER_TELANGANA_DSM_CODES: Tuple[str, ...] = ("KASIPET", "BHUPALPALLY", "KOTHAGUDEM")


def _email_scheduler_report_date_from_job(job: Any, fallback_now_utc: datetime) -> str:
    raw_dt = getattr(job, "scheduled_at", None) or fallback_now_utc
    try:
        if raw_dt.tzinfo is None:
            raw_dt = raw_dt.replace(tzinfo=timezone.utc)
        return raw_dt.astimezone(ZoneInfo("Asia/Kolkata")).date().isoformat()
    except Exception:
        return fallback_now_utc.astimezone(ZoneInfo("Asia/Kolkata")).date().isoformat()


def _email_scheduler_rebuild_auto_dsm_payload_for_dispatch(
    *,
    job: Any,
    now_utc: datetime,
) -> Optional[Dict[str, Any]]:
    """
    Rebuild cron DSM table just before dispatch so auto-sent values match the Email Scheduler screen.

    Existing queued jobs may have stale DSM JSON; this recalculates from Vedanjay SLDC schedule + meter CSV
    and falls back to stored payload if S3 inputs are unavailable.
    """
    stored_payload = _email_scheduler_parse_json_payload(getattr(job, "dsm_summary_payload", None))
    requested_by = str(getattr(job, "requested_by", "") or "").strip()
    if requested_by and requested_by != _email_scheduler_system_user():
        return stored_payload

    template_id = str(getattr(job, "template_id", "") or "").strip()
    template_category = _email_scheduler_template_category(template_id).strip().lower() if template_id else ""
    if "dsm" not in template_category:
        return stored_payload

    plant_code = _normalize_plant_code(str(getattr(job, "plant_code", "") or "").strip())
    report_date = _email_scheduler_report_date_from_job(job, now_utc)
    try:
        computed = _email_scheduler_build_dsm_payload_from_s3_for_email(
            plant_code=plant_code,
            report_date=report_date,
        )
    except Exception:
        computed = None
    return computed or stored_payload


def _email_scheduler_dsm_inputs_ready(
    *,
    s3_client: Any,
    bucket: str,
    plant_code: str,
    report_date: str,
) -> bool:
    """
    Return True only when the required S3 inputs exist for DSM calculation:
    - latest Vedanjay SLDC schedule for the report date
    - at least one meter CSV for the day

    This is used to gate SYSTEM_CRON daily DSM auto-send so we don't email placeholder zeros.
    """
    pcode = _normalize_plant_code(plant_code)
    day = str(report_date or "").strip()
    if not pcode or not re.fullmatch(r"\d{4}-\d{2}-\d{2}", day):
        return False
    if not s3_client or not bucket:
        return False

    schedule_ok = False
    schedule_prefix = _email_scheduler_sldc_schedule_prefix(pcode, day)
    if schedule_prefix and _s3_proxy_is_allowed_path(schedule_prefix):
        try:
            schedule_objects = _list_s3_objects_paginated(
                s3_client=s3_client,
                bucket=bucket,
                prefix=schedule_prefix,
                max_items=2000,
            )
            schedule_ok = bool(_email_scheduler_pick_latest_sldc_schedule(schedule_objects))
        except Exception:
            schedule_ok = False
    if not schedule_ok:
        return False

    meter_prefixes = (
        [f"raw/vedanjay/multiple_generator/ZTRIC/{day}/metered_data/"]
        if _normalize_plant_code(pcode) == "ZETRIC"
        else [
            *[f"raw/vedanjay/{folder}/{day}/metered_data/" for folder in _raw_plant_folder_aliases(pcode)],
            f"generated/vedanjay/{pcode}/outputs/{day}/meter/",
            f"outputs/{day}/meter/",
            f"{day}/meter/",
        ]
    )
    for prefix in meter_prefixes:
        if not prefix or not _s3_proxy_is_allowed_path(prefix):
            continue
        try:
            keys = _list_s3_keys_safe(s3_client, bucket, prefix, max_keys=200)
        except Exception:
            keys = []
        if any(str(k or "").lower().endswith(".csv") for k in (keys or [])):
            return True
    return False

def _email_scheduler_resolve_daily_dsm_attachment(
    *,
    plant_code: str,
    report_date: str,
    template_id: str,
    from_email: str,
    to_email: str,
    cc_email: str,
    plant_name: str,
) -> Optional[Dict[str, Any]]:
    """
    Best-effort lookup for plant/date DSM report attachment from S3.

    Expected filename hints: dsm / penalty / report with extension csv/xlsx/xls/pdf.
    Prefixes are configurable via EMAIL_SCHEDULER_DSM_ATTACHMENT_PREFIX_TEMPLATES.
    """
    plant = _normalize_plant_code(plant_code)
    date_key = str(report_date or "").strip()
    if not plant or not re.fullmatch(r"\d{4}-\d{2}-\d{2}", date_key):
        return None

    raw_prefixes = str(
        os.getenv(
            "EMAIL_SCHEDULER_DSM_ATTACHMENT_PREFIX_TEMPLATES",
            "raw/vedanjay/{plant}/{date}/,generated/vedanjay/{plant}/reports/{date}/,uploads/vedanjay/{plant}/{date}/",
        )
        or ""
    )
    prefixes = []
    for item in raw_prefixes.split(","):
        p = str(item or "").strip()
        if not p:
            continue
        p = p.format(plant=plant, date=date_key).strip().lstrip("/")
        if p and not p.endswith("/"):
            p = f"{p}/"
        if p and _s3_proxy_is_allowed_path(p):
            prefixes.append(p)
    if not prefixes:
        return None

    file_pattern = str(os.getenv("EMAIL_SCHEDULER_DSM_ATTACHMENT_REGEX") or "").strip()
    if not file_pattern:
        file_pattern = r"(dsm|penalty|report).*\.(csv|xlsx|xls|pdf)$"
    try:
        file_re = re.compile(file_pattern, flags=re.IGNORECASE)
    except Exception:
        file_re = re.compile(r"(dsm|penalty|report).*\.(csv|xlsx|xls|pdf)$", flags=re.IGNORECASE)

    bucket = _derive_s3_bucket_name()
    region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"
    if not bucket:
        return None

    s3 = None
    try:
        import boto3  # type: ignore

        s3 = boto3.client("s3", region_name=region)
    except Exception:
        s3 = None

    candidate_keys: List[str] = []
    for prefix in prefixes:
        try:
            keys = _list_s3_keys_safe(s3, bucket, prefix, max_keys=2000)
        except Exception:
            keys = []
        for key in keys or []:
            k = str(key or "").strip()
            if not k:
                continue
            base = os.path.basename(k)
            if not file_re.search(base):
                continue
            candidate_keys.append(k)
    if not candidate_keys:
        return None

    # Prefer files that mention plant/date and then lexical latest.
    def _score_key(k: str) -> Tuple[int, str]:
        base = os.path.basename(k).lower()
        score = 0
        if plant.lower() in base:
            score += 2
        if date_key in base:
            score += 1
        return (score, k)

    candidate_keys = sorted(set(candidate_keys), key=_score_key, reverse=True)
    target_key = candidate_keys[0]

    data: bytes = b""
    if s3 is not None and bucket:
        try:
            obj = s3.get_object(Bucket=bucket, Key=target_key)
            body = obj.get("Body")
            data = body.read() if body is not None else b""
        except Exception:
            data = b""

    if not data:
        try:
            encoded_key = "/".join(quote(segment) for segment in str(target_key or "").split("/"))
            url = f"{DEFAULT_TEMPLATE_S3_BASE_URL.rstrip('/')}/{encoded_key}"
            with urlopen(url, timeout=30) as resp:
                data = resp.read()
        except Exception:
            data = b""

    if not data:
        out_name, out_bytes = _email_scheduler_build_tabular_dsm_attachment_bytes(
            template_id=str(template_id or "").strip(),
            date_str=str(date_key or "").strip(),
            plant_name=str(plant_name or plant),
        )
        if out_name and out_bytes:
            return {
                "file_name": out_name,
                "bytes": out_bytes,
                "content_type": _email_scheduler_guess_attachment_content_type(out_name),
                "s3_key": "",
                "generated": True,
            }
        return None

    name = os.path.basename(target_key) or f"{plant}_{date_key}_dsm_report.csv"
    return {
        "file_name": name,
        "bytes": data,
        "content_type": _email_scheduler_guess_attachment_content_type(name),
        "s3_key": target_key,
        "generated": False,
    }


@app.post("/email-scheduler/schedule-all")
def email_scheduler_schedule_all(
    payload: EmailSchedulerScheduleAllRequest,
    x_user_role: Optional[str] = Header(None, alias="X-User-Role"),
    x_user_name: Optional[str] = Header(None, alias="X-User-Name"),
):
    role = _email_scheduler_normalize_role(x_user_role)
    user = _email_scheduler_normalize_user(x_user_name)
    _email_scheduler_role_guard(role=role, admin_only=True)

    plants, templates_by_plant, _meta = load_email_scheduler_metadata()
    active_plants = [p for p in (plants or []) if bool(p.get("active")) and str(p.get("plant_code") or "").strip()]
    if not active_plants:
        raise HTTPException(status_code=400, detail="No active plants available in scheduler metadata.")

    scheduled_at = _email_scheduler_parse_scheduled_at_utc(date_str=payload.date, time_str=payload.time, am_pm=str(payload.am_pm or "AM"))

    created = 0
    skipped_auto_disabled = 0
    skipped_auto_disabled_plants: List[str] = []
    db = SessionLocal()
    try:
        plant_auto_email_enabled = _email_scheduler_get_plant_auto_email_map(db)
        recipient_defaults = _email_scheduler_get_recipient_defaults(db)
        for plant in active_plants:
            plant_code = _normalize_plant_code(str(plant.get("plant_code") or "").strip())
            if not _email_scheduler_is_plant_auto_email_enabled(plant_auto_email_enabled, plant_code):
                skipped_auto_disabled += 1
                skipped_auto_disabled_plants.append(plant_code)
                continue
            # Pick defaults from the per-plant template list if available.
            defaults = None
            for tpl in (templates_by_plant or {}).get(plant_code, []) or []:
                if str(tpl.get("id") or "").strip() == str(payload.template_id or "").strip():
                    defaults = tpl
                    break
            recipient_default = _email_scheduler_get_recipient_default(recipient_defaults, plant_code, str(payload.template_id or "").strip())
            to_email = str(recipient_default.get("to_email") or (defaults or {}).get("default_to") or "").strip()
            cc_email = str(recipient_default.get("cc_email") or (defaults or {}).get("default_cc") or "").strip()
            subject = _email_scheduler_build_report_subject(
                template_id=str(payload.template_id or "").strip(),
                plant_code=plant_code,
                report_date=payload.date,
                template=defaults,
            ) or str((defaults or {}).get("subject") or str(payload.template_id or "")).strip()
            body_context = _email_scheduler_build_template_context(
                _email_scheduler_body_template_date_key(
                    plant_code=plant_code,
                    template_id=str(payload.template_id or "").strip(),
                    report_date=payload.date,
                )
            )
            body = _email_scheduler_render_template_vars(
                str((defaults or {}).get("body") or "").strip(),
                body_context,
            )
            body = _email_scheduler_day_ahead_body_date(
                body,
                plant_code=plant_code,
                template_id=str(payload.template_id or "").strip(),
                report_date=payload.date,
            )
            if not to_email:
                # Skip plants without recipients configured.
                continue

            job = EmailSchedulerJob(
                requested_by=user or None,
                role=role,
                template_id=str(payload.template_id or "").strip(),
                plant_code=plant_code,
                scheduled_at=scheduled_at,
                auto_send=bool(payload.auto_send),
                from_email=str(payload.from_email or "").strip(),
                to_email=to_email,
                cc_email=cc_email or None,
                employee_name=str(payload.employee_name or "").strip() or None,
                subject=subject,
                body=body,
                status="SCHEDULED",
            )
            db.add(job)
            created += 1

        db.commit()
        return {
            "ok": True,
            "created": created,
            "skipped_auto_disabled": skipped_auto_disabled,
            "skipped_auto_disabled_plants": skipped_auto_disabled_plants,
        }
    except HTTPException:
        raise
    except Exception as exc:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Schedule-all failed: {exc}") from exc
    finally:
        db.close()


@app.post("/email-scheduler/daily-dsm-run")
def email_scheduler_daily_dsm_run(
    payload: EmailSchedulerDailyDsmRunRequest,
    x_scheduler_secret: Optional[str] = Header(None, alias="X-Scheduler-Secret"),
):
    expected_secret = str(os.getenv("EMAIL_SCHEDULER_DAILY_RUN_SECRET") or "").strip()
    if not expected_secret:
        raise HTTPException(status_code=503, detail="Daily DSM run secret is not configured.")
    if str(x_scheduler_secret or "").strip() != expected_secret:
        raise HTTPException(status_code=403, detail="Forbidden")

    # UI-controlled master switch to pause/resume cron-triggered DSM sends.
    db_settings = SessionLocal()
    try:
        enabled = _email_scheduler_settings_get_bool(db_settings, EMAIL_SCHEDULER_SETTING_DAILY_DSM_ENABLED, True)
    finally:
        db_settings.close()
    if not enabled:
        return {"ok": False, "disabled": True, "reason": "Daily DSM auto email is OFF"}

    template_id = str(payload.template_id or os.getenv("EMAIL_SCHEDULER_DAILY_TEMPLATE_ID") or "DSM").strip()
    template_selector = template_id.strip().lower()
    from_email = str(payload.from_email or os.getenv("EMAIL_SCHEDULER_DAILY_FROM_EMAIL") or EMAIL_SCHEDULER_DEFAULT_FROM_EMAIL).strip()
    forced_to_email = str(os.getenv("EMAIL_SCHEDULER_DAILY_TO_EMAIL") or "").strip()
    employee_name = str(payload.employee_name or os.getenv("EMAIL_SCHEDULER_DAILY_EMPLOYEE_NAME") or _email_scheduler_signature_name()).strip()
    auto_send = bool(payload.auto_send if payload.auto_send is not None else True)
    dry_run = bool(payload.dry_run if payload.dry_run is not None else False)
    force_repeat = bool(payload.force_repeat if payload.force_repeat is not None else False)
    repeat_mode_env = str(os.getenv("EMAIL_SCHEDULER_REPEAT_MODE", "0")).strip().lower() in {"1", "true", "yes", "y", "on"}
    repeat_mode = bool(force_repeat or repeat_mode_env)
    if not from_email:
        raise HTTPException(status_code=400, detail="from_email is required (payload or EMAIL_SCHEDULER_DAILY_FROM_EMAIL).")
    if not template_id:
        raise HTTPException(status_code=400, detail="template_id is required.")

    plants, templates_by_plant, _meta = load_email_scheduler_metadata()
    active_plants = [p for p in (plants or []) if bool(p.get("active")) and str(p.get("plant_code") or "").strip()]
    if not active_plants:
        raise HTTPException(status_code=400, detail="No active plants available in scheduler metadata.")

    now_utc = datetime.now(timezone.utc)
    now_ist = now_utc.astimezone(ZoneInfo("Asia/Kolkata"))
    day_start_utc, day_end_utc = _email_scheduler_ist_day_utc_bounds(now_ist.date())
    template_context = _email_scheduler_build_template_context(now_ist.date().isoformat())

    # Optional: gate DSM auto-send until S3 inputs exist (edited_frozen + meter CSV).
    bucket = _derive_s3_bucket_name()
    region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"
    s3 = None
    if bucket:
        try:
            import boto3  # type: ignore

            s3 = boto3.client("s3", region_name=region)
        except Exception:
            s3 = None

    created = 0
    skipped_existing = 0
    skipped_no_recipients = 0
    attached_count = 0
    missing_attachment_count = 0
    processed_plants: List[str] = []
    skipped_existing_plants: List[str] = []
    skipped_no_recipients_plants: List[str] = []
    missing_attachment_plants: List[str] = []
    not_ready_plants: List[str] = []
    skipped_auto_disabled_plants: List[str] = []
    telangana_codes = {"KASIPET", "BHUPALPALLY", "KOTHAGUDEM"}
    telangana_rows: List[Dict[str, Any]] = []
    telangana_row_columns: List[str] = []
    telangana_processed_plants: List[str] = []
    telangana_to_email = ""
    telangana_cc_email = ""
    telangana_subject = ""
    telangana_body = ""
    telangana_template_id = template_id
    db = SessionLocal()
    try:
        plant_auto_email_enabled = _email_scheduler_get_plant_auto_email_map(db)
        recipient_defaults = _email_scheduler_get_recipient_defaults(db)
        for plant in active_plants:
            plant_code = _normalize_plant_code(str(plant.get("plant_code") or "").strip())
            if not plant_code:
                continue
            if not _email_scheduler_is_plant_auto_email_enabled(plant_auto_email_enabled, plant_code):
                skipped_auto_disabled_plants.append(plant_code)
                continue

            # Gate per-plant DSM auto-send until S3 inputs exist (edited_frozen + meter CSV).
            if s3 and bucket:
                if not _email_scheduler_dsm_inputs_ready(
                    s3_client=s3,
                    bucket=bucket,
                    plant_code=plant_code,
                    report_date=now_ist.date().isoformat(),
                ):
                    not_ready_plants.append(plant_code)
                    continue

            defaults = None
            resolved_template_id = template_id
            for tpl in (templates_by_plant or {}).get(plant_code, []) or []:
                tpl_id = str(tpl.get("id") or "").strip()
                tpl_label = str(tpl.get("label") or "").strip().lower()
                tpl_category = str(tpl.get("category") or "").strip().lower()
                if tpl_id == template_id or tpl_label == template_selector or tpl_category == template_selector:
                    defaults = tpl
                    resolved_template_id = tpl_id or template_id
                    break

            if not repeat_mode:
                if _email_scheduler_exists_daily_send_for_plant(
                    db,
                    plant_code=plant_code,
                    template_id=resolved_template_id,
                    day_start_utc=day_start_utc,
                    day_end_utc=day_end_utc,
                ):
                    skipped_existing += 1
                    skipped_existing_plants.append(plant_code)
                    continue

            recipient_default = _email_scheduler_get_recipient_default(recipient_defaults, plant_code, resolved_template_id)
            to_email = forced_to_email or str(recipient_default.get("to_email") or (defaults or {}).get("default_to") or "").strip()
            cc_email = str(recipient_default.get("cc_email") or (defaults or {}).get("default_cc") or "").strip()
            mandatory_cc = str(os.getenv("EMAIL_SCHEDULER_DAILY_MANDATORY_CC") or "").strip()
            # Always CC forecasting on cron auto-emails (per ops request).
            cc_email = _email_scheduler_merge_cc(cc_email, "forecasting.vppl@gmail.com")
            if mandatory_cc:
                cc_email = _email_scheduler_merge_cc(cc_email, mandatory_cc)
            if not to_email:
                skipped_no_recipients += 1
                skipped_no_recipients_plants.append(plant_code)
                continue

            subject = _email_scheduler_build_report_subject(
                template_id=resolved_template_id,
                plant_code=plant_code,
                report_date=now_ist.date().isoformat(),
                template=defaults,
            ) or _email_scheduler_render_template_vars(
                str((defaults or {}).get("subject") or resolved_template_id).strip(),
                template_context,
            )
            body_context = _email_scheduler_build_template_context(
                _email_scheduler_body_template_date_key(
                    plant_code=plant_code,
                    template_id=resolved_template_id,
                    report_date=now_ist.date().isoformat(),
                )
            )
            body = _email_scheduler_render_template_vars(
                str((defaults or {}).get("body") or "").strip(),
                body_context,
            )
            if _email_scheduler_is_gsnp_intraday(plant_code=plant_code, template_id=resolved_template_id):
                body = _email_scheduler_gsnp_intraday_body(now_ist.date().isoformat())
            body = normalize_day_ahead_body(
                body,
                resolved_template_id,
                str((defaults or {}).get("label") or ""),
            )
            # BHUPALPALLY DSM template text historically says "selected date" (no placeholder).
            # For cron auto-send, replace that phrase with the actual IST report date.
            if plant_code == "BHUPALPALLY":
                try:
                    date_label = str(template_context.get("date_dotted") or template_context.get("date_dashed") or "").strip()
                    if date_label and re.search(r"\bfor\s+the\s+selected\s+date\b", body, flags=re.IGNORECASE):
                        body = re.sub(r"\bfor\s+the\s+selected\s+date\b", f"for Date {date_label}", body, flags=re.IGNORECASE)
                except Exception:
                    pass
            if not to_email:
                skipped_no_recipients += 1
                skipped_no_recipients_plants.append(plant_code)
                continue

            # Attempt S3-based calculation (same as UI) for correct values.
            resolved_row = None
            if s3 and bucket:
                resolved_row = _email_scheduler_build_daily_dsm_row_from_s3(
                    s3_client=s3,
                    bucket=bucket,
                    plant_code=plant_code,
                    plant_name=str(plant.get("plant_name") or plant_code),
                    report_date=now_ist.date().isoformat(),
                )
            if resolved_row:
                variant = "default"
                if plant_code == "OSEPL":
                    variant = "osepl"
                elif plant_code == "SIRMOUR":
                    variant = "sirmour"
                elif plant_code in {"KASIPET", "BHUPALPALLY", "KOTHAGUDEM"}:
                    variant = "multi"
                dsm_table_payload = {"variant": variant, "columns": _email_scheduler_public_dsm_columns(resolved_row), "rows": [resolved_row]}
            else:
                dsm_table_payload = _email_scheduler_build_simple_daily_dsm_table_payload(
                    plant_code=plant_code,
                    plant_name=str(plant.get("plant_name") or plant_code),
                    report_date=now_ist.date().isoformat(),
                )
            attachment_payload = _email_scheduler_dsm_support_attachment_from_payload(
                payload=dsm_table_payload,
                plant_code=plant_code,
                report_date=now_ist.date().isoformat(),
            )

            # Consolidate Telangana plants (Kasipet/Bhupalpally/Kothagudem) into a single email.
            if plant_code in telangana_codes:
                rows = list(dsm_table_payload.get("rows") or [])
                computed_rows = [
                    row
                    for row in rows
                    if isinstance(row, dict)
                    and isinstance(row.get("__support_details"), list)
                    and row.get("__support_details")
                ]
                if computed_rows:
                    if not telangana_row_columns:
                        telangana_row_columns = _email_scheduler_public_dsm_columns(computed_rows[0])
                    telangana_rows.extend(computed_rows)
                else:
                    missing_attachment_count += 1
                    missing_attachment_plants.append(plant_code)
                telangana_processed_plants.append(plant_code)
                telangana_template_id = resolved_template_id or telangana_template_id
                telangana_to_email = _email_scheduler_merge_cc(telangana_to_email, to_email)
                telangana_cc_email = _email_scheduler_merge_cc(telangana_cc_email, cc_email)
                if not telangana_subject:
                    telangana_subject = subject
                if not telangana_body:
                    telangana_body = body
                continue

            dsm_table_payload_json = json.dumps(dsm_table_payload, ensure_ascii=True, separators=(",", ":"))

            if dry_run:
                created += 1
                processed_plants.append(plant_code)
                if attachment_payload:
                    attached_count += 1
                continue

            job = EmailSchedulerJob(
                requested_by=_email_scheduler_system_user(),
                role="admin",
                template_id=resolved_template_id,
                plant_code=plant_code,
                scheduled_at=now_utc,
                auto_send=auto_send,
                from_email=from_email,
                to_email=to_email,
                cc_email=cc_email or None,
                employee_name=employee_name or None,
                subject=subject,
                body=body,
                dsm_summary_payload=dsm_table_payload_json,
                attachment_name=str((attachment_payload or {}).get("file_name") or "") or None,
                attachment_bytes=(attachment_payload or {}).get("bytes") or None,
                attachment_content_type=str((attachment_payload or {}).get("content_type") or "") or None,
                status="SCHEDULED",
            )
            db.add(job)
            created += 1
            processed_plants.append(plant_code)
            if attachment_payload:
                attached_count += 1

        # Emit one consolidated Telangana job (single email with multi-row table).
        if telangana_rows:
            if not telangana_to_email:
                skipped_no_recipients += len(telangana_processed_plants)
                skipped_no_recipients_plants.extend(telangana_processed_plants)
            else:
                telangana_payload = {
                    "variant": "multi",
                    "columns": telangana_row_columns or list(telangana_rows[0].keys()),
                    "rows": telangana_rows,
                }
                telangana_payload_json = json.dumps(telangana_payload, ensure_ascii=True, separators=(",", ":"))
                telangana_attachment_payload = _email_scheduler_dsm_support_attachment_from_payload(
                    payload=telangana_payload,
                    plant_code="TELANGANA",
                    report_date=now_ist.date().isoformat(),
                    telangana_static_values=True,
                )

                if dry_run:
                    created += 1
                    processed_plants.extend(telangana_processed_plants)
                    if telangana_attachment_payload:
                        attached_count += 1
                else:
                    job = EmailSchedulerJob(
                        requested_by=_email_scheduler_system_user(),
                        role="admin",
                        template_id=telangana_template_id,
                        plant_code="TELANGANA",
                        scheduled_at=now_utc,
                        auto_send=auto_send,
                        from_email=from_email,
                        to_email=telangana_to_email,
                        cc_email=telangana_cc_email or None,
                        employee_name=employee_name or None,
                        subject=_email_scheduler_build_report_subject(
                            template_id=telangana_template_id,
                            plant_code="TELANGANA",
                            report_date=now_ist.date().isoformat(),
                        ) or telangana_subject or f"Telangana DSM Summary - {now_ist.date().isoformat()}",
                        body=telangana_body,
                        dsm_summary_payload=telangana_payload_json,
                        attachment_name=str((telangana_attachment_payload or {}).get("file_name") or "") or None,
                        attachment_bytes=(telangana_attachment_payload or {}).get("bytes") or None,
                        attachment_content_type=str((telangana_attachment_payload or {}).get("content_type") or "") or None,
                        status="SCHEDULED",
                    )
                    db.add(job)
                    created += 1
                    processed_plants.extend(telangana_processed_plants)
                    if telangana_attachment_payload:
                        attached_count += 1

        if not dry_run:
            db.commit()
        return {
            "ok": True,
            "date_ist": now_ist.date().isoformat(),
            "scheduled_at_utc": now_utc.isoformat(),
            "template_id": template_id,
            "auto_send": auto_send,
            "dry_run": dry_run,
            "repeat_mode": repeat_mode,
            "active_plants": len(active_plants),
            "created": created,
            "skipped_existing": skipped_existing,
            "skipped_no_recipients": skipped_no_recipients,
            "attached": attached_count,
            "missing_attachment": missing_attachment_count,
            "skipped_not_ready": len(not_ready_plants),
            "processed_plants": processed_plants,
            "skipped_existing_plants": skipped_existing_plants,
            "skipped_no_recipients_plants": skipped_no_recipients_plants,
            "missing_attachment_plants": missing_attachment_plants,
            "skipped_not_ready_plants": not_ready_plants,
            "skipped_auto_disabled": len(skipped_auto_disabled_plants),
            "skipped_auto_disabled_plants": skipped_auto_disabled_plants,
        }
    except HTTPException:
        raise
    except Exception as exc:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Daily DSM run failed: {exc}") from exc
    finally:
        db.close()


@app.post("/email-scheduler/daily-dayahead-run")
def email_scheduler_daily_dayahead_run(
    payload: EmailSchedulerDailyDaRunRequest,
    x_scheduler_secret: Optional[str] = Header(None, alias="X-Scheduler-Secret"),
):
    expected_secret = str(os.getenv("EMAIL_SCHEDULER_DAILY_DA_RUN_SECRET") or "").strip()
    if not expected_secret:
        raise HTTPException(status_code=503, detail="Daily Day-Ahead run secret is not configured.")
    if str(x_scheduler_secret or "").strip() != expected_secret:
        raise HTTPException(status_code=403, detail="Forbidden")

    # UI-controlled master switch to pause/resume cron-triggered DA sends.
    db_settings = SessionLocal()
    try:
        enabled = _email_scheduler_settings_get_bool(db_settings, EMAIL_SCHEDULER_SETTING_DAILY_DA_ENABLED, True)
    finally:
        db_settings.close()
    if not enabled:
        return {"ok": False, "disabled": True, "reason": "Daily Day-Ahead auto email is OFF"}

    template_id = str(payload.template_id or os.getenv("EMAIL_SCHEDULER_DAILY_DA_TEMPLATE_ID") or "DA0").strip()
    template_selector = template_id.strip().lower()
    from_email = str(payload.from_email or os.getenv("EMAIL_SCHEDULER_DAILY_DA_FROM_EMAIL") or EMAIL_SCHEDULER_DEFAULT_FROM_EMAIL).strip()
    forced_to_email = str(os.getenv("EMAIL_SCHEDULER_DAILY_DA_TO_EMAIL") or "").strip()
    employee_name = str(payload.employee_name or os.getenv("EMAIL_SCHEDULER_DAILY_DA_EMPLOYEE_NAME") or _email_scheduler_signature_name()).strip()
    auto_send = bool(payload.auto_send if payload.auto_send is not None else True)
    dry_run = bool(payload.dry_run if payload.dry_run is not None else False)
    force_repeat = bool(payload.force_repeat if payload.force_repeat is not None else False)
    repeat_mode_env = str(os.getenv("EMAIL_SCHEDULER_REPEAT_MODE", "0")).strip().lower() in {"1", "true", "yes", "y", "on"}
    repeat_mode = bool(force_repeat or repeat_mode_env)
    if not from_email:
        raise HTTPException(status_code=400, detail="from_email is required (payload or EMAIL_SCHEDULER_DAILY_DA_FROM_EMAIL).")
    if not template_id:
        raise HTTPException(status_code=400, detail="template_id is required.")

    plants, templates_by_plant, _meta = load_email_scheduler_metadata()
    active_plants = [p for p in (plants or []) if bool(p.get("active")) and str(p.get("plant_code") or "").strip()]
    if not active_plants:
        raise HTTPException(status_code=400, detail="No active plants available in scheduler metadata.")

    now_utc = datetime.now(timezone.utc)
    now_ist = now_utc.astimezone(ZoneInfo("Asia/Kolkata"))
    day_start_utc, day_end_utc = _email_scheduler_ist_day_utc_bounds(now_ist.date())
    template_context = _email_scheduler_build_template_context(now_ist.date().isoformat())

    created = 0
    skipped_existing = 0
    skipped_no_recipients = 0
    skipped_missing_attachment = 0
    processed_plants: List[str] = []
    skipped_existing_plants: List[str] = []
    skipped_no_recipients_plants: List[str] = []
    skipped_missing_attachment_plants: List[str] = []
    skipped_auto_disabled_plants: List[str] = []
    db = SessionLocal()
    try:
        plant_auto_email_enabled = _email_scheduler_get_plant_auto_email_map(db)
        recipient_defaults = _email_scheduler_get_recipient_defaults(db)
        for plant in active_plants:
            plant_code = _normalize_plant_code(str(plant.get("plant_code") or "").strip())
            if not plant_code:
                continue
            if not _email_scheduler_is_plant_auto_email_enabled(plant_auto_email_enabled, plant_code):
                skipped_auto_disabled_plants.append(plant_code)
                continue

            defaults, resolved_template_id = _email_scheduler_pick_template_for_plant(
                templates_for_plant=(templates_by_plant or {}).get(plant_code, []) or [],
                template_id=template_id,
            )
            resolved_template_id = resolved_template_id or template_id
            if not _email_scheduler_is_auto_schedule_window_open(
                plant_code=plant_code,
                template_id=resolved_template_id,
                now_ist=now_ist,
            ):
                continue

            if not repeat_mode:
                if _email_scheduler_exists_daily_send_for_plant(
                    db,
                    plant_code=plant_code,
                    template_id=resolved_template_id,
                    day_start_utc=day_start_utc,
                    day_end_utc=day_end_utc,
                ):
                    skipped_existing += 1
                    skipped_existing_plants.append(plant_code)
                    continue

            recipient_default = _email_scheduler_get_recipient_default(recipient_defaults, plant_code, resolved_template_id)
            to_email = forced_to_email or str(recipient_default.get("to_email") or (defaults or {}).get("default_to") or "").strip()
            cc_email = str(recipient_default.get("cc_email") or (defaults or {}).get("default_cc") or "").strip()
            mandatory_cc = str(os.getenv("EMAIL_SCHEDULER_DAILY_MANDATORY_CC") or "").strip()
            # Always CC forecasting on cron auto-emails (per ops request).
            cc_email = _email_scheduler_merge_cc(cc_email, "forecasting.vppl@gmail.com")
            if mandatory_cc:
                cc_email = _email_scheduler_merge_cc(cc_email, mandatory_cc)
            if not to_email:
                skipped_no_recipients += 1
                skipped_no_recipients_plants.append(plant_code)
                continue

            try:
                resolved = _email_scheduler_resolve_schedule_attachment_data(
                    plant_name=plant_code,
                    template_id=resolved_template_id,
                    date_key=now_ist.date().isoformat(),
                )
            except HTTPException:
                skipped_missing_attachment += 1
                skipped_missing_attachment_plants.append(plant_code)
                continue

            schedule_name = str(resolved.get("file_name") or "schedule.csv")
            schedule_bytes = bytes(resolved.get("file_bytes") or b"")
            if not schedule_bytes:
                skipped_missing_attachment += 1
                skipped_missing_attachment_plants.append(plant_code)
                continue

            subject = _email_scheduler_build_report_subject(
                template_id=resolved_template_id,
                plant_code=plant_code,
                report_date=now_ist.date().isoformat(),
                template=defaults,
            ) or _email_scheduler_render_template_vars(
                str((defaults or {}).get("subject") or resolved_template_id).strip(),
                template_context,
            )
            body_context = _email_scheduler_build_template_context(
                _email_scheduler_body_template_date_key(
                    plant_code=plant_code,
                    template_id=resolved_template_id,
                    report_date=now_ist.date().isoformat(),
                )
            )
            body = _email_scheduler_render_template_vars(
                str((defaults or {}).get("body") or "").strip(),
                body_context,
            )
            body = _email_scheduler_day_ahead_body_date(
                body,
                plant_code=plant_code,
                template_id=resolved_template_id,
                report_date=now_ist.date().isoformat(),
            )

            # Cron-driven auto emails should match the Email Scheduler screen attachments:
            # - Telangana plants: XLSX SLDC template
            # - SIRMOUR: GSNP XLSX template
            # - OSEPL: keep CSV as-is
            converted = maybe_convert_for_auto_email(
                plant_code=plant_code,
                template_id=resolved_template_id,
                schedule_type=str(resolved.get("schedule_type") or ""),
                file_name=schedule_name,
                file_bytes=schedule_bytes,
                report_date=str(resolved.get("lookup_date") or now_ist.date().isoformat()),
                source_key=str(resolved.get("attachment_revision_source_key") or resolved.get("s3_key") or ""),
            )
            if converted:
                schedule_name = _email_scheduler_attachment_display_name(
                    plant_code=plant_code,
                    template_id=resolved_template_id,
                    schedule_type=str(resolved.get("schedule_type") or ""),
                    source_key=str(resolved.get("attachment_revision_source_key") or resolved.get("s3_key") or ""),
                    original_name=converted.filename,
                    report_date=str(resolved.get("lookup_date") or now_ist.date().isoformat()),
                    date_already_day_ahead=str(resolved.get("schedule_type") or "").strip().lower() == "dayahead",
                )
                schedule_bytes = converted.content_bytes

            if dry_run:
                created += 1
                processed_plants.append(plant_code)
                continue

            job = EmailSchedulerJob(
                requested_by=_email_scheduler_system_user(),
                role="admin",
                template_id=resolved_template_id,
                plant_code=plant_code,
                scheduled_at=now_utc,
                auto_send=auto_send,
                from_email=from_email,
                to_email=to_email,
                cc_email=cc_email or None,
                employee_name=employee_name or None,
                subject=subject,
                body=body,
                schedule_attachment_name=schedule_name,
                schedule_attachment_bytes=schedule_bytes,
                status="SCHEDULED",
            )
            db.add(job)
            created += 1
            processed_plants.append(plant_code)

        if not dry_run:
            db.commit()
        return {
            "ok": True,
            "date_ist": now_ist.date().isoformat(),
            "scheduled_at_utc": now_utc.isoformat(),
            "template_id": template_id,
            "auto_send": auto_send,
            "dry_run": dry_run,
            "repeat_mode": repeat_mode,
            "active_plants": len(active_plants),
            "created": created,
            "skipped_existing": skipped_existing,
            "skipped_no_recipients": skipped_no_recipients,
            "skipped_missing_attachment": skipped_missing_attachment,
            "processed_plants": processed_plants,
            "skipped_existing_plants": skipped_existing_plants,
            "skipped_no_recipients_plants": skipped_no_recipients_plants,
            "skipped_missing_attachment_plants": skipped_missing_attachment_plants,
            "skipped_auto_disabled": len(skipped_auto_disabled_plants),
            "skipped_auto_disabled_plants": skipped_auto_disabled_plants,
        }
    except HTTPException:
        raise
    except Exception as exc:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Daily Day-Ahead run failed: {exc}") from exc
    finally:
        db.close()


@app.post("/email-scheduler/daily-intraday-run")
def email_scheduler_daily_intraday_run(
    payload: EmailSchedulerDailyIntradayRunRequest,
    x_scheduler_secret: Optional[str] = Header(None, alias="X-Scheduler-Secret"),
):
    """
    Cron-triggered intraday auto email run.

    Current scope: SIRMOUR, GSNP, and ILIOS_PV by template id. Attachments are resolved
    through the same path used by the Email Scheduler UI.
    """
    expected_secret = str(os.getenv("EMAIL_SCHEDULER_DAILY_INTRADAY_RUN_SECRET") or "").strip()
    if not expected_secret:
        raise HTTPException(status_code=503, detail="Daily Intraday run secret is not configured.")
    if str(x_scheduler_secret or "").strip() != expected_secret:
        raise HTTPException(status_code=403, detail="Forbidden")

    enabled = str(os.getenv("EMAIL_SCHEDULER_DAILY_INTRADAY_ENABLED", "true")).strip().lower() in {"1", "true", "yes", "y", "on"}
    if not enabled:
        return {"ok": False, "disabled": True, "reason": "Daily Intraday auto email is OFF"}

    template_id = str(payload.template_id or os.getenv("EMAIL_SCHEDULER_DAILY_INTRADAY_TEMPLATE_ID") or "sirmour_intraday").strip()
    template_selector = template_id.strip().lower()
    from_email = str(payload.from_email or os.getenv("EMAIL_SCHEDULER_DAILY_INTRADAY_FROM_EMAIL") or EMAIL_SCHEDULER_DEFAULT_FROM_EMAIL).strip()
    forced_to_email = str(os.getenv("EMAIL_SCHEDULER_DAILY_INTRADAY_TO_EMAIL") or "").strip()
    employee_name = str(payload.employee_name or os.getenv("EMAIL_SCHEDULER_DAILY_INTRADAY_EMPLOYEE_NAME") or _email_scheduler_signature_name()).strip()
    auto_send = bool(payload.auto_send if payload.auto_send is not None else True)
    dry_run = bool(payload.dry_run if payload.dry_run is not None else False)
    force_repeat = bool(payload.force_repeat if payload.force_repeat is not None else False)
    repeat_mode_env = str(os.getenv("EMAIL_SCHEDULER_REPEAT_MODE", "0")).strip().lower() in {"1", "true", "yes", "y", "on"}
    repeat_mode = bool(force_repeat or repeat_mode_env)
    if not from_email:
        raise HTTPException(status_code=400, detail="from_email is required (payload or EMAIL_SCHEDULER_DAILY_INTRADAY_FROM_EMAIL).")
    if not template_id:
        raise HTTPException(status_code=400, detail="template_id is required.")

    plants, templates_by_plant, _meta = load_email_scheduler_metadata()
    requested_template_key = template_id.strip().lower()
    if "ilios" in requested_template_key:
        cron_intraday_codes = {"ILIOS_PV"}
    elif "gsnp" in requested_template_key:
        cron_intraday_codes = {"GSNP"}
    else:
        cron_intraday_codes = {"SIRMOUR"}
    # SIRMOUR, GSNP, and ILIOS_PV use the same intraday cron timing path.
    active_plants = [
        p
        for p in (plants or [])
        if bool(p.get("active")) and str(p.get("plant_code") or "").strip().upper() in cron_intraday_codes
    ]
    if not active_plants:
        raise HTTPException(status_code=400, detail="SIRMOUR/GSNP/ILIOS_PV is not active/available in scheduler metadata.")

    now_utc = datetime.now(timezone.utc)
    now_ist = now_utc.astimezone(ZoneInfo("Asia/Kolkata"))
    day_start_utc, day_end_utc = _email_scheduler_ist_day_utc_bounds(now_ist.date())
    template_context = _email_scheduler_build_template_context(now_ist.date().isoformat())

    created = 0
    skipped_existing = 0
    skipped_no_recipients = 0
    skipped_missing_attachment = 0
    processed_plants: List[str] = []
    skipped_existing_plants: List[str] = []
    skipped_no_recipients_plants: List[str] = []
    skipped_missing_attachment_plants: List[str] = []
    skipped_auto_disabled_plants: List[str] = []
    db = SessionLocal()
    try:
        plant_auto_email_enabled = _email_scheduler_get_plant_auto_email_map(db)
        recipient_defaults = _email_scheduler_get_recipient_defaults(db)
        for plant in active_plants:
            plant_code = _normalize_plant_code(str(plant.get("plant_code") or "").strip())
            if not plant_code:
                continue
            if not _email_scheduler_is_plant_auto_email_enabled(plant_auto_email_enabled, plant_code):
                skipped_auto_disabled_plants.append(plant_code)
                continue

            defaults, resolved_template_id = _email_scheduler_pick_template_for_plant(
                templates_for_plant=(templates_by_plant or {}).get(plant_code, []) or [],
                template_id=template_id,
            )
            resolved_template_id = resolved_template_id or template_id
            if not _email_scheduler_is_auto_schedule_window_open(
                plant_code=plant_code,
                template_id=resolved_template_id,
                now_ist=now_ist,
            ):
                continue

            if not repeat_mode:
                if _email_scheduler_exists_daily_send_for_plant(
                    db,
                    plant_code=plant_code,
                    template_id=resolved_template_id,
                    day_start_utc=day_start_utc,
                    day_end_utc=day_end_utc,
                ):
                    skipped_existing += 1
                    skipped_existing_plants.append(plant_code)
                    continue

            recipient_default = _email_scheduler_get_recipient_default(recipient_defaults, plant_code, resolved_template_id)
            to_email = forced_to_email or str(recipient_default.get("to_email") or (defaults or {}).get("default_to") or "").strip()
            cc_email = str(recipient_default.get("cc_email") or (defaults or {}).get("default_cc") or "").strip()
            mandatory_cc = str(os.getenv("EMAIL_SCHEDULER_DAILY_MANDATORY_CC") or "").strip()
            # Always CC forecasting on cron auto-emails (per ops request).
            cc_email = _email_scheduler_merge_cc(cc_email, "forecasting.vppl@gmail.com")
            if mandatory_cc:
                cc_email = _email_scheduler_merge_cc(cc_email, mandatory_cc)
            cc_email = _email_scheduler_ensure_intraday_cc(
                plant_code=plant_code,
                template_id=resolved_template_id,
                cc_email=cc_email,
            )
            subject = _email_scheduler_build_report_subject(
                template_id=resolved_template_id,
                plant_code=plant_code,
                report_date=now_ist.date().isoformat(),
                template=defaults,
            ) or _email_scheduler_render_template_vars(
                str((defaults or {}).get("subject") or resolved_template_id).strip(),
                template_context,
            )
            body = _email_scheduler_render_template_vars(
                str((defaults or {}).get("body") or "").strip(),
                template_context,
            )
            if _email_scheduler_is_gsnp_intraday(plant_code=plant_code, template_id=resolved_template_id):
                body = _email_scheduler_gsnp_intraday_body(now_ist.date())
            elif _email_scheduler_is_ilios_pv_intraday(plant_code=plant_code, template_id=resolved_template_id):
                subject = _email_scheduler_ilios_pv_intraday_subject(now_ist.date())
                body = _email_scheduler_ilios_pv_intraday_body(now_ist.date())
            elif _email_scheduler_is_sirmour_intraday(plant_code=plant_code, template_id=resolved_template_id):
                body = _email_scheduler_sirmour_intraday_body(now_ist.date())
            if not to_email:
                skipped_no_recipients += 1
                skipped_no_recipients_plants.append(plant_code)
                continue

            try:
                resolved = _email_scheduler_resolve_schedule_attachment_data(
                    plant_name=plant_code,
                    template_id=resolved_template_id,
                    date_key=now_ist.date().isoformat(),
                )
            except HTTPException:
                skipped_missing_attachment += 1
                skipped_missing_attachment_plants.append(plant_code)
                continue

            schedule_name = str(resolved.get("file_name") or "schedule.csv")
            schedule_bytes = bytes(resolved.get("file_bytes") or b"")
            if not schedule_bytes:
                skipped_missing_attachment += 1
                skipped_missing_attachment_plants.append(plant_code)
                continue

            converted = maybe_convert_for_auto_email(
                plant_code=plant_code,
                template_id=resolved_template_id,
                schedule_type=str(resolved.get("schedule_type") or ""),
                file_name=schedule_name,
                file_bytes=schedule_bytes,
                report_date=now_ist.date().isoformat(),
                source_key=str(resolved.get("attachment_revision_source_key") or resolved.get("s3_key") or ""),
            )
            if converted:
                schedule_name = _email_scheduler_attachment_display_name(
                    plant_code=plant_code,
                    template_id=resolved_template_id,
                    schedule_type=str(resolved.get("schedule_type") or ""),
                    source_key=str(resolved.get("attachment_revision_source_key") or resolved.get("s3_key") or ""),
                    original_name=converted.filename,
                    report_date=str(resolved.get("lookup_date") or now_ist.date().isoformat()),
                    date_already_day_ahead=str(resolved.get("schedule_type") or "").strip().lower() == "dayahead",
                )
                schedule_bytes = converted.content_bytes

            if dry_run:
                created += 1
                processed_plants.append(plant_code)
                continue

            job = EmailSchedulerJob(
                requested_by=_email_scheduler_system_user(),
                role="admin",
                template_id=resolved_template_id,
                plant_code=plant_code,
                scheduled_at=now_utc,
                auto_send=auto_send,
                from_email=from_email,
                to_email=to_email,
                cc_email=cc_email or None,
                employee_name=employee_name or None,
                subject=subject,
                body=body,
                schedule_attachment_name=schedule_name,
                schedule_attachment_bytes=schedule_bytes,
                status="SCHEDULED",
            )
            db.add(job)
            created += 1
            processed_plants.append(plant_code)

        if not dry_run:
            db.commit()
        return {
            "ok": True,
            "date_ist": now_ist.date().isoformat(),
            "scheduled_at_utc": now_utc.isoformat(),
            "template_id": template_id,
            "auto_send": auto_send,
            "dry_run": dry_run,
            "repeat_mode": repeat_mode,
            "active_plants": len(active_plants),
            "created": created,
            "skipped_existing": skipped_existing,
            "skipped_no_recipients": skipped_no_recipients,
            "skipped_missing_attachment": skipped_missing_attachment,
            "processed_plants": processed_plants,
            "skipped_existing_plants": skipped_existing_plants,
            "skipped_no_recipients_plants": skipped_no_recipients_plants,
            "skipped_missing_attachment_plants": skipped_missing_attachment_plants,
            "skipped_auto_disabled": len(skipped_auto_disabled_plants),
            "skipped_auto_disabled_plants": skipped_auto_disabled_plants,
        }
    except HTTPException:
        raise
    except Exception as exc:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Daily Intraday run failed: {exc}") from exc
    finally:
        db.close()


@app.get("/email-scheduler/jobs")
def email_scheduler_list_jobs(
    x_user_role: Optional[str] = Header(None, alias="X-User-Role"),
    x_user_name: Optional[str] = Header(None, alias="X-User-Name"),
    limit: int = Query(100, ge=1, le=500),
):
    role = _email_scheduler_normalize_role(x_user_role)
    user = _email_scheduler_normalize_user(x_user_name)

    db = SessionLocal()
    try:
        q = db.query(EmailSchedulerJob)
        if role != "admin":
            q = q.filter(EmailSchedulerJob.requested_by == (user or None))
        rows = (
            q.order_by(EmailSchedulerJob.scheduled_at.desc(), EmailSchedulerJob.id.desc())
            .limit(int(limit))
            .all()
        )
        items = []
        for r in rows:
            items.append(
                {
                    "id": r.id,
                    "requested_by": r.requested_by,
                    "role": r.role,
                    "template_id": r.template_id,
                    "plant_code": r.plant_code,
                    "scheduled_at": r.scheduled_at.isoformat() if r.scheduled_at else None,
                    "auto_send": bool(r.auto_send),
                    "status": r.status,
                    "sent_at": r.sent_at.isoformat() if r.sent_at else None,
                    "error_message": r.error_message,
                    "created_at": r.created_at.isoformat() if r.created_at else None,
                }
            )
        return {"ok": True, "items": items}
    finally:
        db.close()


@app.delete("/email-scheduler/jobs/{job_id}")
def email_scheduler_delete_job(
    job_id: int,
    x_user_role: Optional[str] = Header(None, alias="X-User-Role"),
    x_user_name: Optional[str] = Header(None, alias="X-User-Name"),
):
    role = _email_scheduler_normalize_role(x_user_role)
    user = _email_scheduler_normalize_user(x_user_name)

    db = SessionLocal()
    try:
        job = db.query(EmailSchedulerJob).filter(EmailSchedulerJob.id == int(job_id)).first()
        if not job:
            raise HTTPException(status_code=404, detail="Job not found")
        if role != "admin" and (job.requested_by or "") != (user or ""):
            raise HTTPException(status_code=403, detail="Not allowed to delete this job")
        job.status = "CANCELED"
        db.delete(job)
        db.commit()
        return {"ok": True}
    except HTTPException:
        raise
    except Exception as exc:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Delete failed: {exc}") from exc
    finally:
        db.close()


@app.post("/email-scheduler/jobs/{job_id}/retry")
def email_scheduler_retry_job(
    job_id: int,
    x_user_role: Optional[str] = Header(None, alias="X-User-Role"),
    x_user_name: Optional[str] = Header(None, alias="X-User-Name"),
):
    role = _email_scheduler_normalize_role(x_user_role)
    user = _email_scheduler_normalize_user(x_user_name)

    db = SessionLocal()
    try:
        job = db.query(EmailSchedulerJob).filter(EmailSchedulerJob.id == job_id).first()
        if not job:
            raise HTTPException(status_code=404, detail="Job not found")
        if role != "admin" and job.requested_by != (user or None):
            raise HTTPException(status_code=403, detail="Not allowed")

        # Re-queue for immediate dispatch if auto-send is enabled.
        ist = ZoneInfo("Asia/Kolkata")
        now_ist = datetime.now(timezone.utc).astimezone(ist)
        scheduled_at = now_ist.astimezone(timezone.utc)

        job.scheduled_at = scheduled_at
        job.status = "SCHEDULED"
        job.sent_at = None
        job.error_message = None
        job.auto_send = True
        db.add(job)
        db.commit()
        db.refresh(job)
        return {"ok": True, "id": job.id, "status": job.status, "scheduled_at": job.scheduled_at.isoformat()}
    finally:
        db.close()


async def _email_scheduler_dispatch_due_jobs_loop() -> None:
    # Single-process dispatcher (use a dedicated worker container if you need HA).
    interval = int(os.getenv("EMAIL_SCHEDULER_DISPATCH_INTERVAL_SECONDS") or "20")
    while True:
        try:
            now_utc = datetime.now(timezone.utc)
            db = SessionLocal()
            try:
                due = (
                    db.query(EmailSchedulerJob)
                    .filter(EmailSchedulerJob.status == "SCHEDULED")
                    .filter(EmailSchedulerJob.auto_send == True)  # noqa: E712
                    .filter(EmailSchedulerJob.scheduled_at <= now_utc)
                    .order_by(EmailSchedulerJob.scheduled_at.asc(), EmailSchedulerJob.id.asc())
                    .limit(10)
                    .all()
                )
                for job in due:
                    try:
                        plant_auto_email_enabled = _email_scheduler_get_plant_auto_email_map(db)
                        if not _email_scheduler_is_plant_auto_email_enabled(plant_auto_email_enabled, job.plant_code):
                            job.status = "CANCELED"
                            job.error_message = "Cron auto email disabled for plant"
                            db.add(job)
                            db.commit()
                            _email_scheduler_log_event(
                                requested_by=job.requested_by or "",
                                employee_name=job.employee_name or "",
                                role=job.role or "",
                                template_id=job.template_id,
                                plant_code=job.plant_code,
                                mode="DISPATCHER",
                                status="CANCELED",
                                from_email=job.from_email,
                                to_email=job.to_email,
                                cc_email=job.cc_email or "",
                                subject=job.subject,
                                scheduled_at=job.scheduled_at,
                                sent_at=None,
                                error_message="Cron auto email disabled for plant",
                            )
                            continue
                        dsm_payload = (
                            _email_scheduler_rebuild_auto_dsm_payload_for_dispatch(job=job, now_utc=now_utc)
                            or _email_scheduler_parse_json_payload(job.dsm_summary_payload)
                        )
                        schedule_att = None
                        if job.schedule_attachment_bytes:
                            schedule_att = (job.schedule_attachment_name or "schedule.csv", bytes(job.schedule_attachment_bytes))
                        att = None
                        is_system_telangana_dsm = (
                            str(getattr(job, "requested_by", "") or "").strip() == _email_scheduler_system_user()
                            and _normalize_plant_code(str(job.plant_code or "")) == "TELANGANA"
                            and "dsm" in _email_scheduler_template_category(str(job.template_id or "")).lower()
                        )
                        if job.attachment_bytes:
                            att = (job.attachment_name or "attachment.bin", bytes(job.attachment_bytes), job.attachment_content_type or "application/octet-stream")
                        if is_system_telangana_dsm and dsm_payload:
                            att = None
                        if dsm_payload and (att is None) and "dsm" in _email_scheduler_template_category(str(job.template_id or "")).lower():
                            generated_att = _email_scheduler_dsm_support_attachment_from_payload(
                                payload=dsm_payload,
                                plant_code=str(job.plant_code or ""),
                                report_date=_email_scheduler_report_date_from_job(job, now_utc),
                                telangana_static_values=is_system_telangana_dsm,
                            )
                            if generated_att and generated_att.get("bytes"):
                                att = (
                                    str(generated_att.get("file_name") or "support-file.xlsx"),
                                    bytes(generated_att.get("bytes") or b""),
                                    str(generated_att.get("content_type") or "application/octet-stream"),
                                )
                                job.attachment_name = att[0]
                                job.attachment_bytes = att[1]
                                job.attachment_content_type = att[2]
                        dispatch_to_email, dispatch_cc_base = _email_scheduler_apply_saved_recipients(
                            db,
                            plant_code=str(job.plant_code or ""),
                            template_id=str(job.template_id or "").strip(),
                            to_email=job.to_email or "",
                            cc_email=job.cc_email or "",
                            merge_cc=True,
                        )
                        dispatch_cc_email = _email_scheduler_ensure_intraday_cc(
                            plant_code=job.plant_code,
                            template_id=job.template_id,
                            cc_email=dispatch_cc_base,
                        )
                        dispatch_body = (
                            _email_scheduler_gsnp_intraday_body(
                                (job.scheduled_at or now_utc).astimezone(ZoneInfo("Asia/Kolkata")).date()
                            )
                            if _email_scheduler_is_gsnp_intraday(plant_code=job.plant_code, template_id=job.template_id)
                            else
                            _email_scheduler_ilios_pv_intraday_body(
                                (job.scheduled_at or now_utc).astimezone(ZoneInfo("Asia/Kolkata")).date()
                            )
                            if _email_scheduler_is_ilios_pv_intraday(plant_code=job.plant_code, template_id=job.template_id)
                            else
                            _email_scheduler_sirmour_intraday_body(
                                (job.scheduled_at or now_utc).astimezone(ZoneInfo("Asia/Kolkata")).date()
                            )
                            if _email_scheduler_is_sirmour_intraday(plant_code=job.plant_code, template_id=job.template_id)
                            else job.body
                        )
                        dispatch_subject = (
                            _email_scheduler_build_report_subject(
                                template_id=job.template_id,
                                plant_code=job.plant_code,
                                report_date=(job.scheduled_at or now_utc).astimezone(ZoneInfo("Asia/Kolkata")).date(),
                            )
                            if (
                                _email_scheduler_is_gsnp_intraday(plant_code=job.plant_code, template_id=job.template_id)
                                or
                                _email_scheduler_is_ilios_pv_intraday(plant_code=job.plant_code, template_id=job.template_id)
                                or
                                _email_scheduler_is_sirmour_intraday(plant_code=job.plant_code, template_id=job.template_id)
                                or _normalize_plant_code(str(job.plant_code or "")) == "TELANGANA"
                            )
                            else job.subject
                        )

                        _email_scheduler_send_now(
                            template_id=str(job.template_id or ""),
                            role=job.role or "testing",
                            from_email=job.from_email,
                            to_email=dispatch_to_email,
                            cc_email=dispatch_cc_email,
                            subject=dispatch_subject,
                            body=dispatch_body,
                            employee_name=_email_scheduler_normalize_signature_name(job.employee_name),
                            dsm_payload=dsm_payload,
                            schedule_attachment=schedule_att,
                            attachment=att,
                            include_employee_mobile=False,
                        )
                        job.status = "SENT"
                        job.sent_at = datetime.now(timezone.utc)
                        job.error_message = None
                        db.add(job)
                        db.commit()
                        _email_scheduler_log_event(
                            requested_by=job.requested_by or "",
                            employee_name=job.employee_name or "",
                            role=job.role or "",
                            template_id=job.template_id,
                            plant_code=job.plant_code,
                            mode="DISPATCHER",
                            status="SENT",
                            from_email=job.from_email,
                            to_email=dispatch_to_email,
                            cc_email=dispatch_cc_email,
                            subject=dispatch_subject,
                            scheduled_at=job.scheduled_at,
                            sent_at=job.sent_at,
                            error_message=None,
                        )
                    except Exception as exc:
                        job.status = "FAILED"
                        job.error_message = str(exc)
                        db.add(job)
                        db.commit()
                        _email_scheduler_log_event(
                            requested_by=job.requested_by or "",
                            employee_name=job.employee_name or "",
                            role=job.role or "",
                            template_id=job.template_id,
                            plant_code=job.plant_code,
                            mode="DISPATCHER",
                            status="FAILED",
                            from_email=job.from_email,
                            to_email=job.to_email,
                            cc_email=job.cc_email or "",
                            subject=job.subject,
                            scheduled_at=job.scheduled_at,
                            sent_at=datetime.now(timezone.utc),
                            error_message=str(exc),
                        )
            finally:
                db.close()
        except Exception:
            # Keep loop alive; errors visible via job error_message when possible.
            pass
        await asyncio.sleep(max(5, interval))


@app.on_event("startup")
async def _email_scheduler_start_dispatcher() -> None:
    enabled = str(os.getenv("EMAIL_SCHEDULER_DISPATCH_ENABLED", "true")).strip().lower() in {"1", "true", "yes", "y"}
    if not enabled:
        return
    asyncio.create_task(_email_scheduler_dispatch_due_jobs_loop())


_EMAIL_SCHEDULER_AUTO_SCHEDULE_WINDOWS: Dict[Tuple[str, str], List[Tuple[str, str]]] = {
    ("SIRMOUR", "DA0"): [],
    ("SIRMOUR", "DA1"): [("20:00", "21:00")],
    ("SIRMOUR", "INTRADAY"): [("08:00", "09:00")],
    ("BHUPALPALLY", "DA0"): [("05:00", "06:00")],
    ("BHUPALPALLY", "DA1"): [("22:45", "23:45")],
    ("KASIPET", "DA0"): [("05:00", "06:00")],
    ("KASIPET", "DA1"): [("22:45", "23:45")],
    ("KASIPETH", "DA0"): [("05:00", "06:00")],
    ("KASIPETH", "DA1"): [("22:45", "23:45")],
    ("KOTHAGUDEM", "DA0"): [("05:00", "06:00")],
    ("KOTHAGUDEM", "DA1"): [("22:45", "23:45")],
    ("KOTHAGUDAM", "DA0"): [("05:00", "06:00")],
    ("KOTHAGUDAM", "DA1"): [("22:45", "23:45")],
    ("OSEPL", "DA0"): [("05:00", "06:00")],
    ("OSEPL", "DA1"): [],
}


def _email_scheduler_auto_schedule_type(template_id: str) -> str:
    key = str(template_id or "").strip().lower()
    if "intraday" in key or "intra" in key:
        return "INTRADAY"
    if "da1" in key:
        return "DA1"
    if "da0" in key:
        return "DA0"
    return ""


def _email_scheduler_auto_time_minutes(value: str) -> int:
    match = re.fullmatch(r"\s*(\d{1,2}):(\d{2})\s*", str(value or ""))
    if not match:
        return -1
    hour = int(match.group(1))
    minute = int(match.group(2))
    if hour < 0 or hour > 23 or minute < 0 or minute > 59:
        return -1
    return hour * 60 + minute


def _email_scheduler_is_auto_schedule_window_open(
    *,
    plant_code: str,
    template_id: str,
    now_ist: datetime,
) -> bool:
    plant = _normalize_plant_code(str(plant_code or "").strip())
    schedule_type = _email_scheduler_auto_schedule_type(template_id)
    if not plant or not schedule_type:
        return True
    # DA0 and intraday are now cron-driven directly, so don't block them behind
    # legacy backend time windows.
    if schedule_type in {"DA0", "INTRADAY"}:
        return True
    windows = _EMAIL_SCHEDULER_AUTO_SCHEDULE_WINDOWS.get((plant, schedule_type))
    if windows is None:
        return True
    if not windows:
        return False
    current = int(now_ist.hour) * 60 + int(now_ist.minute)
    for start_raw, end_raw in windows:
        start = _email_scheduler_auto_time_minutes(start_raw)
        end = _email_scheduler_auto_time_minutes(end_raw)
        if start < 0 or end < 0:
            continue
        if start <= end:
            if start <= current <= end:
                return True
        elif current >= start or current <= end:
            return True
    return False


async def _email_scheduler_internal_poll_loop() -> None:
    """
    Optional internal poller to trigger cron-style auto email runs periodically.

    Behavior:
    - Runs the same "daily-*-run" flows (DSM, DA0, DA1, Intraday) every N seconds.
    - De-dupe/guard rails remain enforced by existing per-day job checks.
    - DA0/DA1/Intraday only create a job when the required attachment exists.

    This is disabled by default to avoid double-running alongside external cron.
    """
    interval = int(os.getenv("EMAIL_SCHEDULER_INTERNAL_POLL_INTERVAL_SECONDS") or "1800")  # 30 min default
    while True:
        try:
            # DSM
            try:
                dsm_secret = str(os.getenv("EMAIL_SCHEDULER_DAILY_RUN_SECRET") or "").strip()
                if dsm_secret:
                    email_scheduler_daily_dsm_run(
                        EmailSchedulerDailyDsmRunRequest(auto_send=True, dry_run=False, force_repeat=False),
                        x_scheduler_secret=dsm_secret,
                    )
            except Exception:
                pass

            # Day-ahead (try DA0 then DA1)
            try:
                da_secret = str(os.getenv("EMAIL_SCHEDULER_DAILY_DA_RUN_SECRET") or "").strip()
                if da_secret:
                    email_scheduler_daily_dayahead_run(
                        EmailSchedulerDailyDaRunRequest(template_id="DA0", auto_send=True, dry_run=False, force_repeat=False),
                        x_scheduler_secret=da_secret,
                    )
                    email_scheduler_daily_dayahead_run(
                        EmailSchedulerDailyDaRunRequest(template_id="DA1", auto_send=True, dry_run=False, force_repeat=False),
                        x_scheduler_secret=da_secret,
                    )
            except Exception:
                pass

            # Intraday (SIRMOUR + GSNP + ILIOS_PV)
            try:
                intra_secret = str(os.getenv("EMAIL_SCHEDULER_DAILY_INTRADAY_RUN_SECRET") or "").strip()
                if intra_secret:
                    email_scheduler_daily_intraday_run(
                        EmailSchedulerDailyIntradayRunRequest(template_id="sirmour_intraday", auto_send=True, dry_run=False, force_repeat=False),
                        x_scheduler_secret=intra_secret,
                    )
                    email_scheduler_daily_intraday_run(
                        EmailSchedulerDailyIntradayRunRequest(template_id="gsnp_intraday", auto_send=True, dry_run=False, force_repeat=False),
                        x_scheduler_secret=intra_secret,
                    )
                    email_scheduler_daily_intraday_run(
                        EmailSchedulerDailyIntradayRunRequest(template_id="ilios_pv_intraday", auto_send=True, dry_run=False, force_repeat=False),
                        x_scheduler_secret=intra_secret,
                    )
            except Exception:
                pass
        finally:
            await asyncio.sleep(max(60, interval))


@app.on_event("startup")
async def _email_scheduler_start_internal_poller() -> None:
    enabled = str(os.getenv("EMAIL_SCHEDULER_INTERNAL_POLL_ENABLED", "false")).strip().lower() in {"1", "true", "yes", "y", "on"}
    if not enabled:
        return
    asyncio.create_task(_email_scheduler_internal_poll_loop())


@app.get("/email-scheduler/dispatcher-status")
def email_scheduler_dispatcher_status():
    enabled = str(os.getenv("EMAIL_SCHEDULER_DISPATCH_ENABLED", "true")).strip().lower() in {"1", "true", "yes", "y"}
    interval = int(os.getenv("EMAIL_SCHEDULER_DISPATCH_INTERVAL_SECONDS") or "20")
    now_utc = datetime.now(timezone.utc)
    ist = ZoneInfo("Asia/Kolkata")
    now_ist = now_utc.astimezone(ist)
    return {
        "ok": True,
        "enabled": bool(enabled),
        "interval_seconds": max(5, interval),
        "now_utc": now_utc.isoformat(),
        "now_ist": now_ist.isoformat(),
    }


@app.get("/api/s3/bytes")
async def s3_proxy_get_bytes(key: str = Query(..., min_length=1, max_length=1024)):
    """Fetch an S3 object as bytes via backend (works even when S3 CORS blocks browser)."""
    key = str(key or "").strip()
    if not _s3_proxy_is_allowed_path(key):
        raise HTTPException(status_code=400, detail="Key not allowed")

    bucket = _derive_s3_bucket_name()
    region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"

    try:
        import boto3  # type: ignore
        from botocore.exceptions import ClientError  # type: ignore
        if bucket:
            s3 = boto3.client("s3", region_name=region)
            try:
                obj = s3.get_object(Bucket=bucket, Key=key)
            except ClientError as exc:
                err = (exc.response or {}).get("Error", {}) or {}
                code = str(err.get("Code", "")).strip()
                if code in {"NoSuchKey", "NotFound", "404"}:
                    raise HTTPException(status_code=404, detail="S3 object not found") from exc
                if code in {"AccessDenied", "403"}:
                    raise HTTPException(status_code=403, detail="S3 access denied") from exc
                raise
            body = obj.get("Body")
            data = body.read() if body is not None else b""
            return StreamingResponse(io.BytesIO(data), media_type="application/octet-stream")
    except HTTPException:
        raise
    except Exception:
        pass

    try:
        url = f"{DEFAULT_TEMPLATE_S3_BASE_URL.rstrip('/')}/{quote(key)}"
        with urlopen(url, timeout=30) as resp:
            data = resp.read()
        return StreamingResponse(io.BytesIO(data), media_type="application/octet-stream")
    except HTTPError as e:
        if getattr(e, "code", None) == 404:
            raise HTTPException(status_code=404, detail="S3 object not found") from e
        if getattr(e, "code", None) == 403:
            raise HTTPException(status_code=403, detail="S3 access denied") from e
        raise HTTPException(status_code=502, detail=f"Failed to fetch S3 object: HTTP {e.code}") from e
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Failed to fetch S3 object: {e}")


# ==================== HEALTH CHECK ====================
@app.get("/api/health")
async def health_check_v1():
    """Health check endpoint"""
    return {"status": "ok", "message": "Server is running"}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=3001)


