"""Backend-only all-plant penalty calculation and report generation."""
from __future__ import annotations

import csv
import hashlib
import io
import json
import os
import re
from dataclasses import dataclass
from datetime import date, datetime, timedelta, timezone
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

import boto3
from botocore.exceptions import ClientError
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import load_workbook
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from sqlalchemy import func
from sqlalchemy.orm import Session

from models import (
    BlockPenaltyResult,
    DailyPenaltySummary,
    GeneratedPenaltyReport,
    MeterData,
    Plant,
    VedanjayScheduleUpload,
)


CALCULATION_VERSION = "all-plant-penalty-v1"
COMPARISON_CALCULATION_VERSION = "comparison-screen-v1"
SOURCES = ("SYSTEM", "MANUAL", "ENERCAST", "VEDANJAY")
COMPARISON_SOURCES = SOURCES + ("TESTENV",)
SOURCE_LABELS = {
    "SYSTEM": "System",
    "MANUAL": "Manual",
    "ENERCAST": "Enercast",
    "VEDANJAY": "Vedanjay",
    "TESTENV": "TestEnv",
}
SOURCE_FILES = {
    "SYSTEM": "system_frozen.csv",
    "MANUAL": "edited_frozen.csv",
    "ENERCAST": "enercast_edited_frozen.csv",
}
SOURCE_MISSING_MESSAGES = {
    "SYSTEM": "System schedule not available.",
    "MANUAL": "Manual edited schedule not available.",
    "ENERCAST": "Enercast schedule not available.",
}
VALID_STATUSES = {
    "Calculated",
    "Partially Calculated",
    "Pending",
    "Not Calculated",
    "Zero Penalty",
    "Failed",
}
BLOCK_HOURS = 0.25
KWH_PER_MWH = 1000.0
EPSILON = 1e-9
DEFAULT_BUCKET = "vedanjay-schedules1"
DEFAULT_REGION = "ap-south-1"
PENALTY_REPORT_PLANT_CODES = (
    "SIRMOUR",
    "BHUPALPALLY",
    "KOTHAGUDEM",
    "KASIPET",
    "OSEPL",
    "ANDAD",
    "BALAKWADA",
    "GUGARIYAKHEDI",
    "NANDGAON",
    "SAWDA",
    "ANJANGAON",
    "BAMKHAL",
)
REQUIRED_PLANT_FALLBACKS: Dict[str, Dict[str, Any]] = {
    "SAWDA": {
        "code": "SAWDA",
        "name": "SAWDA",
        "state": "Madhya Pradesh",
        "type": "Solar",
        "capacity": 7.5,
    },
}

STATE_RULES: Dict[str, Dict[str, List[Tuple[float, float, float]]]] = {
    "Telangana": {
        "Solar": [(0, 15, 0), (15, 25, 0.5), (25, 35, 1.0), (35, float("inf"), 1.5)],
        "Wind": [(0, 15, 0), (15, 25, 0.5), (25, 35, 1.0), (35, float("inf"), 1.5)],
    },
    "Maharashtra": {
        "Solar": [(0, 10, 0), (10, 12, 0.25), (12, 15, 0.5), (15, 25, 0.75), (25, float("inf"), 1.0)],
        "Wind": [(0, 12, 0), (12, 15, 0.25), (15, 20, 0.5), (20, float("inf"), 1.0)],
    },
    "Madhya Pradesh": {
        "Solar": [(0, 10, 0), (10, 15, 0.5), (15, 20, 0.75), (20, float("inf"), 1.0)],
        "Wind": [(0, 15, 0), (15, 20, 0.5), (20, 25, 0.75), (25, float("inf"), 1.0)],
    },
}
DEFAULT_RULES = {
    "Solar": [(0, 10, 0), (10, 12, 0.25), (12, 15, 0.5), (15, 25, 0.75), (25, float("inf"), 1.0)],
    "Wind": [(0, 12, 0), (12, 15, 0.25), (15, 20, 0.5), (20, float("inf"), 1.0)],
}


@dataclass
class SourceData:
    values: Dict[int, float]
    file_name: str
    file_hash: str


def normalize_plant_code(value: Any) -> str:
    code = re.sub(r"[^A-Za-z0-9_-]", "", str(value or "").strip()).upper()
    aliases = {
        "OSEL": "OSEPL",
        "SHRIMOUR": "SIRMOUR",
        "SHROMOUR": "SIRMOUR",
        "ANJANGOAN": "ANJANGAON",
        "KOTHAGUDAM": "KOTHAGUDEM",
        "KASIEPTH": "KASIPET",
    }
    return aliases.get(code, code)


def special_s3_plant_folder(value: Any) -> str:
    code = normalize_plant_code(value)
    if code == "ANJANGAON":
        return "ANJANGOAN"
    return code


def special_s3_plant_folder_aliases(value: Any) -> List[str]:
    code = normalize_plant_code(value)
    preferred = special_s3_plant_folder(code)
    aliases: List[str] = []
    for item in [preferred, code]:
        if item and item not in aliases:
            aliases.append(item)
    return aliases


def normalize_state(value: Any) -> str:
    raw = str(value or "").strip()
    aliases = {"MH": "Maharashtra", "TL": "Telangana", "MP": "Madhya Pradesh"}
    if raw.upper() in aliases:
        return aliases[raw.upper()]
    return " ".join(part.capitalize() for part in raw.split())


def sha256_bytes(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _normalize_header(value: Any) -> str:
    return re.sub(r"[^a-z0-9]+", "", str(value or "").strip().lower())


def _as_float(value: Any) -> Optional[float]:
    if value is None or isinstance(value, bool):
        return None
    raw = str(value).strip().replace(",", "")
    if not raw or raw.lower() in {"nan", "none", "null", "-", "--"}:
        return None
    try:
        number = float(raw)
    except (TypeError, ValueError):
        return None
    return number if number == number and abs(number) != float("inf") else None


def _as_block(value: Any) -> Optional[int]:
    number = _as_float(value)
    if number is None:
        return None
    block = int(round(number))
    return block if 1 <= block <= 96 else None


def _rows_from_upload(filename: str, content: bytes) -> List[List[Any]]:
    lower = str(filename or "").lower()
    if lower.endswith((".xlsx", ".xlsm")):
        workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        sheet = workbook.active
        return [list(row) for row in sheet.iter_rows(values_only=True)]
    text = _decode_text(content)
    sample = text[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
    except csv.Error:
        dialect = csv.excel
    return [list(row) for row in csv.reader(io.StringIO(text), dialect)]


def _is_schedule_value_header(header: str) -> bool:
    h = str(header or "")
    if not h:
        return False
    excluded = (
        "actual",
        "meter",
        "capacity",
        "availability",
        "avc",
        "interavc",
        "error",
        "date",
        "time",
        "from",
        "to",
        "block",
    )
    if any(token in h for token in excluded):
        return False
    preferred = (
        "stationschedule",
        "scheduledmw",
        "scheduled",
        "schedule",
        "declaredforecast",
        "forecastmw",
        "forecast",
        "forcast",
        "quantum",
        "sirmour",
        "anjangaon",
        "anjangoan",
    )
    return any(token in h for token in preferred) or h in {"mw", "pv", "plant"}


def _find_header_row(rows: Sequence[Sequence[Any]]) -> Tuple[int, List[str]]:
    best_index = 0
    best_headers = [_normalize_header(cell) for cell in rows[0]] if rows else []
    best_score = -1
    for index, row in enumerate(rows[:80]):
        normalized = [_normalize_header(cell) for cell in row]
        if not any(normalized):
            continue
        has_block = any("block" in header or header in {"srno", "sno", "serialno"} for header in normalized)
        has_schedule = any(_is_schedule_value_header(header) for header in normalized)
        non_empty = sum(1 for header in normalized if header)
        score = (100 if has_block else 0) + (80 if has_schedule else 0) + min(non_empty, 20)
        if score > best_score:
            best_index = index
            best_headers = normalized
            best_score = score
        if has_block and has_schedule:
            return index, normalized
    return best_index, best_headers


def _merge_schedule_headers(
    headers: Sequence[str],
    next_row: Optional[Sequence[Any]],
) -> Tuple[List[str], bool]:
    if not next_row:
        return list(headers or []), False
    next_headers = [_normalize_header(cell) for cell in next_row]
    use_second = any(header in {"forecast", "forcast", "availability", "avc"} for header in next_headers)
    if not use_second:
        return list(headers or []), False

    max_cols = max(len(headers or []), len(next_headers))
    merged: List[str] = []
    parent = ""
    for index in range(max_cols):
        first = str(headers[index] if index < len(headers or []) else "")
        second = str(next_headers[index] if index < len(next_headers) else "")
        if first:
            parent = first
        effective_parent = first or parent
        if effective_parent and second:
            merged.append(f"{effective_parent}{second}")
        else:
            merged.append(effective_parent or second)
    return merged, True


def _pick_schedule_value_column(headers: Sequence[str], rows: Sequence[Sequence[Any]], header_index: int, block_index: int) -> int:
    normalized = list(headers or [])
    exact_preferred = (
        "stationschedule",
        "scheduledmw",
        "schedule",
        "declaredforecast",
        "forecastmw",
        "forecast",
        "forcast",
        "quantum",
        "mw",
    )
    for key in exact_preferred:
        found = next(
            (
                i
                for i, h in enumerate(normalized)
                if (h == key or h.endswith(key))
                and not any(token in h for token in ("availability", "capacity", "avc", "interavc"))
            ),
            -1,
        )
        if found >= 0:
            return found
    found = next((i for i, h in enumerate(normalized) if _is_schedule_value_header(h)), -1)
    if found >= 0:
        return found

    ignored = ("actual", "meter", "capacity", "availability", "avc", "error", "date", "time", "from", "to", "block")
    best_index = -1
    best_score = -1
    for col in range(max((len(row) for row in rows), default=len(normalized))):
        if col == block_index:
            continue
        header = normalized[col] if col < len(normalized) else ""
        if any(token in header for token in ignored):
            continue
        numeric_count = 0
        magnitude = 0.0
        for row in rows[header_index + 1: header_index + 121]:
            value = _as_float(row[col] if col < len(row) else None)
            if value is None:
                continue
            numeric_count += 1
            magnitude += abs(value)
        score = numeric_count * 1000 + magnitude
        if score > best_score:
            best_index = col
            best_score = score
    if best_index >= 0:
        return best_index

    candidates = [
        i for i, h in enumerate(normalized)
        if i != block_index and not any(token in h for token in ignored)
    ]
    return candidates[-1] if candidates else max(0, len(normalized) - 1)


def parse_schedule_upload(filename: str, content: bytes) -> Dict[int, float]:
    rows = _rows_from_upload(filename, content)
    if not rows:
        raise ValueError("Uploaded schedule is empty.")
    header_index, headers = _find_header_row(rows)
    headers, used_second_header = _merge_schedule_headers(
        headers,
        rows[header_index + 1] if header_index + 1 < len(rows) else None,
    )
    data_start = header_index + (2 if used_second_header else 1)
    block_index = next(
        (
            i
            for i, h in enumerate(headers)
            if h in {"block", "blockno", "blocknumber", "srno", "sno", "serialno"} or "block" in h
        ),
        -1,
    )
    value_index = _pick_schedule_value_column(headers, rows, header_index, block_index)

    values: Dict[int, float] = {}
    for row in rows[data_start:]:
        block = _as_block(row[block_index] if 0 <= block_index < len(row) else None)
        value = _as_float(row[value_index] if value_index < len(row) else None)
        if block is not None and value is not None:
            values[block] = value
    if not values and value_index >= 0:
        next_block = 1
        for row in rows[data_start:]:
            value = _as_float(row[value_index] if value_index < len(row) else None)
            if value is None:
                continue
            values[next_block] = value
            next_block += 1
            if next_block > 96:
                break
    if not values:
        raise ValueError("No valid schedule blocks were found in the uploaded file.")
    return values


def parse_schedule_text(content: bytes) -> Dict[int, float]:
    return parse_schedule_upload("schedule.csv", content)


def _time_to_block(value: Any) -> Optional[int]:
    raw = str(value or "").strip()
    if not raw:
        return None
    match = re.search(r"(?:T|\s|^)(\d{1,2}):(\d{2})", raw)
    if not match:
        return None
    hour, minute = int(match.group(1)), int(match.group(2))
    if hour == 24 and minute == 0:
        return 96
    total = max(0, min((hour * 60) + minute, 1439))
    return min(96, (total // 15) + 1)


def parse_meter_content(filename: str, content: bytes) -> Dict[int, float]:
    rows = _rows_from_upload(filename, content)
    if not rows:
        return {}
    header_index, headers = _find_header_row(rows)
    block_index = next((i for i, h in enumerate(headers) if h in {"block", "blockno", "blocknumber"} or h.startswith("block")), -1)
    time_index = next((i for i, h in enumerate(headers) if any(token in h for token in ("timestamp", "datetime", "time"))), -1)
    preferred = ("actualmw", "metermw", "generationmw", "actual", "meter", "generation", "mw")
    value_index = next((i for key in preferred for i, h in enumerate(headers) if h == key or h.endswith(key)), -1)
    if value_index < 0:
        candidates = [i for i in range(len(headers)) if i not in {block_index, time_index}]
        value_index = candidates[-1] if candidates else -1

    values: Dict[int, float] = {}
    for position, row in enumerate(rows[header_index + 1:], start=1):
        block = _as_block(row[block_index] if 0 <= block_index < len(row) else None)
        if block is None and 0 <= time_index < len(row):
            block = _time_to_block(row[time_index])
        if block is None and position <= 96:
            block = position
        actual = _as_float(row[value_index] if 0 <= value_index < len(row) else None)
        if block is not None and actual is not None:
            values[block] = actual
    return values


def _rule_for(state: str, plant_type: str) -> List[Tuple[float, float, float]]:
    by_type = STATE_RULES.get(normalize_state(state), DEFAULT_RULES)
    return by_type.get(str(plant_type or "Solar").title(), by_type.get("Solar", DEFAULT_RULES["Solar"]))


def penalty_rule_label(state: str, plant_type: str) -> str:
    bands = _rule_for(state, plant_type)
    free_limit = next((upper for lower, upper, rate in bands if lower == 0 and rate == 0), 0)
    return f"{normalize_state(state) or 'Default'} {plant_type or 'Solar'} DSM, free band {free_limit:g}%"


def _penalty_context_for_plant(plant_code: Any, state: Any, plant_type: Any) -> Tuple[str, str]:
    if normalize_plant_code(plant_code) == "BAMKHAL":
        return "Madhya Pradesh", "Solar"
    return normalize_state(state), str(plant_type or "Solar").title()


def calculate_standard_penalty(
    scheduled_mw: float,
    actual_mw: float,
    capacity_mw: float,
    state: str,
    plant_type: str,
) -> Dict[str, float]:
    capacity = max(abs(float(capacity_mw or 0)), EPSILON)
    deviation = actual_mw - scheduled_mw
    deviation_percent = (deviation / capacity) * 100.0
    absolute_percent = abs(deviation_percent)
    energy_kwh = abs(deviation) * BLOCK_HOURS * KWH_PER_MWH
    penalty = 0.0
    for lower, upper, rate in _rule_for(state, plant_type):
        span = min(absolute_percent, upper) - lower
        if span > 0:
            penalty += energy_kwh * (span / absolute_percent) * rate
    return {
        "deviation_mw": deviation,
        "deviation_percent": deviation_percent,
        "penalty_amount": penalty,
        "payable_amount": 0.0,
        "receivable_amount": 0.0,
        "net_settlement": -penalty,
        "ppa_amount": scheduled_mw * BLOCK_HOURS * KWH_PER_MWH,
    }


def calculate_osepl_penalty(scheduled_mw: float, actual_mw: float, capacity_mw: float) -> Dict[str, float]:
    ppa_rate = 9.27
    scheduled_kwh = scheduled_mw * BLOCK_HOURS * KWH_PER_MWH
    actual_kwh = actual_mw * BLOCK_HOURS * KWH_PER_MWH
    capacity_kwh = max(capacity_mw * BLOCK_HOURS * KWH_PER_MWH, EPSILON)
    deviation_kwh = actual_kwh - scheduled_kwh
    error_signed = (deviation_kwh / capacity_kwh) * 100.0
    error = abs(error_signed)
    payable = 0.0
    receivable = 0.0
    if actual_mw >= 0:
        bands = (
            (10.0, 9.27, 9.27),
            (12.0, 10.197, 8.343),
            (15.0, 11.124, 7.416),
            (float("inf"), 13.905, 0.0),
        )
        lower = 0.0
        for upper, under_rate, over_rate in bands:
            span = min(error, upper) - lower
            if span > 0:
                slice_kwh = capacity_kwh * (span / 100.0)
                if deviation_kwh < 0:
                    payable += slice_kwh * under_rate
                elif deviation_kwh > 0:
                    receivable += slice_kwh * over_rate
            lower = upper
            if error <= upper:
                break
    ppa_amount = scheduled_kwh * ppa_rate
    net = (actual_kwh * ppa_rate) - (ppa_amount + receivable - payable)
    return {
        "deviation_mw": actual_mw - scheduled_mw,
        "deviation_percent": ((actual_mw - scheduled_mw) / max(capacity_mw, EPSILON)) * 100.0,
        "penalty_amount": net,
        "payable_amount": payable,
        "receivable_amount": receivable,
        "net_settlement": net,
        "ppa_amount": ppa_amount,
    }


def calculate_daily_penalty(
    *,
    schedule: Dict[int, float],
    meter: Dict[int, float],
    capacity_mw: float,
    state: str,
    plant_type: str,
    plant_code: str,
) -> Dict[str, Any]:
    details: List[Dict[str, Any]] = []
    missing_meter: List[int] = []
    missing_schedule: List[int] = []
    for block in range(1, 97):
        scheduled = schedule.get(block)
        actual = meter.get(block)
        if scheduled is None:
            missing_schedule.append(block)
        if actual is None:
            missing_meter.append(block)
        if scheduled is None or actual is None:
            details.append({
                "block_number": block,
                "scheduled_mw": scheduled,
                "actual_meter_mw": actual,
                "deviation_mw": None,
                "deviation_percent": None,
                "penalty_amount": None,
                "payable_amount": None,
                "receivable_amount": None,
                "net_settlement": None,
                "ppa_amount": None,
                "status": "Pending",
                "missing_data_reason": "Schedule data not available." if scheduled is None else "Meter data not available.",
            })
            continue
        result = (
            calculate_osepl_penalty(scheduled, actual, capacity_mw)
            if normalize_plant_code(plant_code) == "OSEPL"
            else calculate_standard_penalty(
                scheduled,
                actual,
                capacity_mw,
                *_penalty_context_for_plant(plant_code, state, plant_type),
            )
        )
        details.append({
            "block_number": block,
            "scheduled_mw": scheduled,
            "actual_meter_mw": actual,
            **result,
            "status": "Calculated",
            "missing_data_reason": None,
        })

    calculated = [row for row in details if row["penalty_amount"] is not None]
    total = sum(float(row["penalty_amount"]) for row in calculated) if calculated else None
    highest = max(calculated, key=lambda row: abs(float(row["penalty_amount"])), default=None)
    if not calculated:
        status = "Pending"
    elif len(calculated) < 96:
        status = "Partially Calculated"
    elif abs(float(total or 0)) <= EPSILON:
        status = "Zero Penalty"
    else:
        status = "Calculated"

    reasons = []
    if missing_meter:
        reasons.append(f"Meter data was missing for Blocks {_format_blocks(missing_meter)}.")
    if missing_schedule:
        reasons.append(f"Schedule data was missing for Blocks {_format_blocks(missing_schedule)}.")
    reason = " ".join(reasons) or None
    if status == "Partially Calculated":
        reason = f"Calculated using {len(calculated)} of 96 blocks. {reason or ''}".strip()
    return {
        "status": status,
        "total_penalty": total,
        "calculated_blocks": len(calculated),
        "highest_penalty_block": highest["block_number"] if highest else None,
        "highest_penalty_amount": abs(float(highest["penalty_amount"])) if highest else None,
        "missing_data_reason": reason,
        "blocks": details,
    }


def _format_blocks(blocks: Sequence[int]) -> str:
    if not blocks:
        return ""
    ranges: List[str] = []
    start = previous = blocks[0]
    for block in blocks[1:]:
        if block == previous + 1:
            previous = block
            continue
        ranges.append(str(start) if start == previous else f"{start}-{previous}")
        start = previous = block
    ranges.append(str(start) if start == previous else f"{start}-{previous}")
    return " and ".join(ranges)


class ReadOnlyS3Source:
    """Reads calculation inputs from S3. This class intentionally has no write method."""

    def __init__(self) -> None:
        self.bucket = str(os.getenv("S3_BUCKET") or os.getenv("VITE_S3_BUCKET") or DEFAULT_BUCKET).strip()
        self.region = str(os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or DEFAULT_REGION).strip()
        self.client = boto3.client("s3", region_name=self.region)

    def _get(self, key: str) -> Optional[bytes]:
        try:
            response = self.client.get_object(Bucket=self.bucket, Key=key)
            return response["Body"].read()
        except ClientError as exc:
            code = str(exc.response.get("Error", {}).get("Code", ""))
            if code in {"NoSuchKey", "404", "AccessDenied"}:
                return None
            raise

    def _latest_under(self, prefixes: Iterable[str], filename: Optional[str] = None) -> Optional[Tuple[str, bytes]]:
        candidates: List[Tuple[datetime, str]] = []
        for prefix in prefixes:
            token = None
            while True:
                kwargs: Dict[str, Any] = {"Bucket": self.bucket, "Prefix": prefix, "MaxKeys": 1000}
                if token:
                    kwargs["ContinuationToken"] = token
                response = self.client.list_objects_v2(**kwargs)
                for item in response.get("Contents", []):
                    key = str(item.get("Key") or "")
                    if not key.lower().endswith((".csv", ".xlsx", ".xlsm")):
                        continue
                    if filename and os.path.basename(key).lower() != filename.lower():
                        continue
                    candidates.append((item.get("LastModified") or datetime.min.replace(tzinfo=timezone.utc), key))
                token = response.get("NextContinuationToken")
                if not token:
                    break
        for _, key in sorted(candidates, reverse=True):
            content = self._get(key)
            if content is not None:
                return key, content
        return None

    def schedule(self, plant_code: str, schedule_date: date, source: str) -> Optional[SourceData]:
        code = normalize_plant_code(plant_code)
        day = schedule_date.isoformat()
        filename = SOURCE_FILES[source]
        direct = [
            *[
                f"frozenschedules/vedanjay/{folder}/{day}/{filename}"
                for folder in special_s3_plant_folder_aliases(code)
            ],
            f"generated/vedanjay/{code}/outputs/{day}/frozen/{filename}",
            f"generated/{code}/{code.lower()}/outputs/{day}/frozen/{filename}",
            f"outputs/{day}/frozen/{filename}",
        ]
        for key in direct:
            content = self._get(key)
            if content:
                return SourceData(parse_schedule_text(content), key, sha256_bytes(content))
        found = self._latest_under(
            [
                *[
                    f"frozenschedules/vedanjay/{folder}/{day}/"
                    for folder in special_s3_plant_folder_aliases(code)
                ],
                f"generated/vedanjay/{code}/outputs/{day}/frozen/",
            ],
            filename,
        )
        if not found:
            return None
        key, content = found
        return SourceData(parse_schedule_text(content), key, sha256_bytes(content))

    def meter(self, plant_code: str, schedule_date: date) -> Optional[SourceData]:
        code = normalize_plant_code(plant_code)
        day = schedule_date.isoformat()
        aliases = [code] + (["ANJANGOAN"] if code == "ANJANGAON" else [])
        prefixes: List[str] = []
        for alias in aliases:
            prefixes.extend([
                f"raw/vedanjay/{alias}/{day}/metered_data/",
                f"generated/vedanjay/{alias}/outputs/{day}/meter/",
            ])
        prefixes.extend([f"outputs/{day}/meter/", f"{day}/meter/"])
        found = self._latest_under(prefixes)
        if not found:
            return None
        key, content = found
        values = parse_meter_content(os.path.basename(key), content)
        return SourceData(values, key, sha256_bytes(content)) if values else None


def _plant_code_from_name(name: str) -> str:
    match = re.search(r"\(([A-Za-z0-9_-]+)\)", str(name or ""))
    return normalize_plant_code(match.group(1) if match else name)


def configured_plants(db: Session) -> List[Dict[str, Any]]:
    plants = db.query(Plant).order_by(Plant.name.asc()).all()
    configured = [
        {
            "code": _plant_code_from_name(str(plant.name)),
            "name": str(plant.name),
            "state": normalize_state(plant.state),
            "type": str(plant.type or "Solar").title(),
            "capacity": float(plant.capacity or 0),
        }
        for plant in plants
        if str(getattr(plant, "status", "Active") or "Active").lower() != "inactive"
    ]
    by_code = {plant["code"]: plant for plant in configured}
    for code, fallback in REQUIRED_PLANT_FALLBACKS.items():
        by_code.setdefault(code, fallback)
    return [
        by_code[code]
        for code in PENALTY_REPORT_PLANT_CODES
        if code in by_code
    ]


def active_upload(db: Session, plant_code: str, schedule_date: date) -> Optional[VedanjayScheduleUpload]:
    return (
        db.query(VedanjayScheduleUpload)
        .filter(VedanjayScheduleUpload.plant_code == normalize_plant_code(plant_code))
        .filter(VedanjayScheduleUpload.schedule_date == schedule_date)
        .filter(VedanjayScheduleUpload.is_active.is_(True))
        .order_by(VedanjayScheduleUpload.uploaded_at.desc(), VedanjayScheduleUpload.id.desc())
        .first()
    )


def store_vedanjay_upload(
    db: Session,
    *,
    plant: Dict[str, Any],
    schedule_date: date,
    filename: str,
    content_type: str,
    content: bytes,
    uploader: str,
) -> VedanjayScheduleUpload:
    blocks = parse_schedule_upload(filename, content)
    code = normalize_plant_code(plant["code"])
    file_hash = sha256_bytes(content)
    db.query(VedanjayScheduleUpload).filter(
        VedanjayScheduleUpload.plant_code == code,
        VedanjayScheduleUpload.schedule_date == schedule_date,
        VedanjayScheduleUpload.is_active.is_(True),
    ).update({"is_active": False}, synchronize_session=False)
    existing = (
        db.query(VedanjayScheduleUpload)
        .filter(VedanjayScheduleUpload.plant_code == code)
        .filter(VedanjayScheduleUpload.schedule_date == schedule_date)
        .filter(VedanjayScheduleUpload.file_hash == file_hash)
        .first()
    )
    if existing:
        existing.is_active = True
        existing.uploader = uploader or existing.uploader
        existing.uploaded_at = datetime.now(timezone.utc)
        upload = existing
    else:
        upload = VedanjayScheduleUpload(
            plant_code=code,
            plant_name=plant["name"],
            schedule_date=schedule_date,
            filename=filename,
            storage_key=f"postgresql://vedanjay_schedule_uploads/{code}/{schedule_date.isoformat()}/{file_hash}",
            uploader=uploader,
            file_hash=file_hash,
            original_content_type=content_type,
            original_file=content,
            normalized_blocks_json=json.dumps({str(k): v for k, v in blocks.items()}, separators=(",", ":")),
            is_active=True,
            validation_status="VALID",
        )
        db.add(upload)
    db.commit()
    db.refresh(upload)
    return upload


def _meter_from_database(db: Session, plant: Dict[str, Any], schedule_date: date) -> Optional[SourceData]:
    candidates = (
        db.query(MeterData)
        .filter(MeterData.dataDate == schedule_date)
        .order_by(MeterData.updatedAt.desc(), MeterData.id.desc())
        .all()
    )
    code = normalize_plant_code(plant["code"])
    row = next((item for item in candidates if _plant_code_from_name(str(item.plantName)) == code), None)
    if not row:
        return None
    try:
        raw = json.loads(str(row.blockData or "{}"))
    except json.JSONDecodeError:
        return None
    values: Dict[int, float] = {}
    iterable = raw.items() if isinstance(raw, dict) else enumerate(raw, start=1)
    for key, value in iterable:
        block = _as_block(key if isinstance(raw, dict) else key)
        actual = _as_float(value.get("actualMw") if isinstance(value, dict) else value)
        if block and actual is not None:
            values[block] = actual
    encoded = str(row.blockData or "").encode("utf-8")
    return SourceData(values, f"postgresql://meter_data/{row.id}", sha256_bytes(encoded)) if values else None


def _missing_summary(source: str, reason: str, status: str = "Pending") -> Dict[str, Any]:
    return {
        "schedule_source": source,
        "total_penalty": None,
        "status": status,
        "missing_data_reason": reason,
        "calculated_blocks": 0,
        "highest_penalty_block": None,
        "highest_penalty_amount": None,
        "blocks": [],
    }


def calculate_and_store_daily(
    db: Session,
    *,
    plant: Dict[str, Any],
    schedule_date: date,
    source: str,
    s3: Optional[ReadOnlyS3Source] = None,
    force: bool = False,
) -> DailyPenaltySummary:
    source = str(source).upper()
    if source not in SOURCES:
        raise ValueError(f"Unsupported schedule source: {source}")
    code = normalize_plant_code(plant["code"])
    s3_reader = s3 or ReadOnlyS3Source()
    upload = active_upload(db, code, schedule_date) if source == "VEDANJAY" else None
    if source == "VEDANJAY":
        if not upload:
            calculation = _missing_summary(source, "Vedanjay schedule not uploaded.", "Pending")
            schedule_data = None
        else:
            values = {int(k): float(v) for k, v in json.loads(upload.normalized_blocks_json).items()}
            schedule_data = SourceData(values, upload.storage_key, upload.file_hash)
            calculation = None
    else:
        schedule_data = s3_reader.schedule(code, schedule_date, source)
        calculation = None if schedule_data else _missing_summary(
            source,
            SOURCE_MISSING_MESSAGES[source],
            "Not Calculated",
        )

    meter_data = None
    if schedule_data:
        try:
            meter_data = s3_reader.meter(code, schedule_date)
        except Exception:
            meter_data = None
        meter_data = meter_data or _meter_from_database(db, plant, schedule_date)
        if not meter_data:
            calculation = _missing_summary(source, "Meter data not available.", "Pending")

    current = (
        db.query(DailyPenaltySummary)
        .filter(DailyPenaltySummary.plant_code == code)
        .filter(DailyPenaltySummary.schedule_date == schedule_date)
        .filter(DailyPenaltySummary.schedule_source == source)
        .first()
    )
    if (
        not force
        and current
        and schedule_data
        and meter_data
        and current.schedule_hash == schedule_data.file_hash
        and current.meter_hash == meter_data.file_hash
        and current.calculation_version == CALCULATION_VERSION
    ):
        return current

    if calculation is None and schedule_data and meter_data:
        calculation = calculate_daily_penalty(
            schedule=schedule_data.values,
            meter=meter_data.values,
            capacity_mw=float(plant["capacity"]),
            state=str(plant["state"]),
            plant_type=str(plant["type"]),
            plant_code=code,
        )
        calculation["schedule_source"] = source

    assert calculation is not None
    summary = current or DailyPenaltySummary(
        plant_code=code,
        schedule_date=schedule_date,
        schedule_source=source,
        plant_name=plant["name"],
        calculation_version=CALCULATION_VERSION,
        status="Pending",
    )
    summary.plant_name = plant["name"]
    summary.state = plant["state"]
    summary.capacity_mw = plant["capacity"]
    summary.total_penalty = calculation["total_penalty"]
    summary.status = calculation["status"]
    summary.missing_data_reason = calculation["missing_data_reason"]
    summary.calculated_blocks = calculation["calculated_blocks"]
    summary.highest_penalty_block = calculation["highest_penalty_block"]
    summary.highest_penalty_amount = calculation["highest_penalty_amount"]
    summary.schedule_file = schedule_data.file_name if schedule_data else None
    summary.schedule_hash = schedule_data.file_hash if schedule_data else None
    summary.meter_file = meter_data.file_name if meter_data else None
    summary.meter_hash = meter_data.file_hash if meter_data else None
    summary.calculation_version = CALCULATION_VERSION
    summary.upload_id = upload.id if upload else None
    summary.calculated_at = datetime.now(timezone.utc)
    summary.observation = _daily_observation(source, calculation)
    if current is None:
        db.add(summary)
    db.flush()
    db.query(BlockPenaltyResult).filter(BlockPenaltyResult.summary_id == summary.id).delete(synchronize_session=False)
    for row in calculation.get("blocks", []):
        db.add(BlockPenaltyResult(
            summary_id=summary.id,
            plant_code=code,
            schedule_date=schedule_date,
            schedule_source=source,
            **row,
        ))
    db.commit()
    db.refresh(summary)
    return summary


def _daily_observation(source: str, calculation: Dict[str, Any]) -> str:
    if calculation["status"] in {"Pending", "Not Calculated", "Failed"}:
        return str(calculation.get("missing_data_reason") or calculation["status"])
    total = float(calculation.get("total_penalty") or 0)
    block = calculation.get("highest_penalty_block")
    amount = calculation.get("highest_penalty_amount")
    label = SOURCE_LABELS.get(source, source.title())
    text = f"{label} penalty was Rs {total:,.2f}."
    if block and amount is not None:
        text += f" Highest penalty was Rs {float(amount):,.2f} at Block {block}."
    if calculation.get("missing_data_reason"):
        text += f" {calculation['missing_data_reason']}"
    return text


def summary_dict(summary: DailyPenaltySummary) -> Dict[str, Any]:
    return {
        "id": summary.id,
        "plant_code": summary.plant_code,
        "plant_name": summary.plant_name,
        "state": summary.state,
        "capacity_mw": summary.capacity_mw,
        "schedule_date": summary.schedule_date.isoformat(),
        "schedule_source": summary.schedule_source,
        "total_penalty": summary.total_penalty,
        "status": summary.status,
        "missing_data_reason": summary.missing_data_reason,
        "calculated_blocks": summary.calculated_blocks,
        "highest_penalty_block": summary.highest_penalty_block,
        "highest_penalty_amount": summary.highest_penalty_amount,
        "schedule_file": summary.schedule_file,
        "meter_file": summary.meter_file,
        "calculation_version": summary.calculation_version,
        "observation": summary.observation,
    }


def store_comparison_results(
    db: Session,
    *,
    plant: Dict[str, Any],
    schedule_date: date,
    sources: Sequence[Dict[str, Any]],
) -> List[DailyPenaltySummary]:
    """Persist the exact results already calculated and displayed by Comparison."""
    code = normalize_plant_code(plant["code"])
    submitted = {str(item.get("source") or "").upper(): item for item in sources}
    stored: List[DailyPenaltySummary] = []
    for source in COMPARISON_SOURCES:
        item = submitted.get(source) or {}
        block_rows = item.get("blocks") if isinstance(item.get("blocks"), list) else []
        valid_blocks = []
        for raw in block_rows:
            block = _as_block(raw.get("block_number"))
            penalty = _as_float(raw.get("penalty_amount"))
            scheduled = _as_float(raw.get("scheduled_mw"))
            actual = _as_float(raw.get("actual_meter_mw"))
            if block is None:
                continue
            valid_blocks.append({
                "block_number": block,
                "scheduled_mw": scheduled,
                "actual_meter_mw": actual,
                "deviation_mw": _as_float(raw.get("deviation_mw")),
                "deviation_percent": _as_float(raw.get("deviation_percent")),
                "penalty_amount": penalty,
                "payable_amount": _as_float(raw.get("payable_amount")),
                "receivable_amount": _as_float(raw.get("receivable_amount")),
                "net_settlement": _as_float(raw.get("net_settlement")),
                "ppa_amount": _as_float(raw.get("ppa_amount")),
                "status": "Calculated" if penalty is not None else "Not Calculated",
                "missing_data_reason": None,
            })

        calculated = [row for row in valid_blocks if row["penalty_amount"] is not None]
        has_schedule = any(row["scheduled_mw"] is not None for row in valid_blocks)
        has_meter = any(row["actual_meter_mw"] is not None for row in valid_blocks)
        total_penalty = sum(float(row["penalty_amount"]) for row in calculated) if calculated else None
        highest = max(calculated, key=lambda row: abs(float(row["penalty_amount"])), default=None)
        if calculated:
            status = "Zero Penalty" if abs(float(total_penalty or 0)) <= EPSILON else (
                "Calculated" if len(calculated) == 96 else "Partially Calculated"
            )
        elif not has_schedule:
            status = "Not Calculated"
        elif not has_meter:
            status = "Pending"
        else:
            status = "Not Calculated"

        schedule_values = [
            [row["block_number"], row["scheduled_mw"]]
            for row in valid_blocks
            if row["scheduled_mw"] is not None
        ]
        meter_values = [
            [row["block_number"], row["actual_meter_mw"]]
            for row in valid_blocks
            if row["actual_meter_mw"] is not None
        ]
        current = (
            db.query(DailyPenaltySummary)
            .filter_by(plant_code=code, schedule_date=schedule_date, schedule_source=source)
            .first()
        )
        summary = current or DailyPenaltySummary(
            plant_code=code,
            schedule_date=schedule_date,
            schedule_source=source,
            plant_name=plant["name"],
            calculation_version=COMPARISON_CALCULATION_VERSION,
            status=status,
        )
        summary.plant_name = plant["name"]
        summary.state = plant["state"]
        summary.capacity_mw = plant["capacity"]
        summary.total_penalty = total_penalty
        summary.status = status
        summary.missing_data_reason = None
        summary.calculated_blocks = len(calculated)
        summary.highest_penalty_block = highest["block_number"] if highest else None
        summary.highest_penalty_amount = abs(float(highest["penalty_amount"])) if highest else None
        summary.schedule_file = str(item.get("schedule_file") or "") or None
        summary.schedule_hash = (
            sha256_bytes(json.dumps(schedule_values, separators=(",", ":")).encode("utf-8"))
            if schedule_values else None
        )
        summary.meter_file = str(item.get("meter_file") or "") or None
        summary.meter_hash = (
            sha256_bytes(json.dumps(meter_values, separators=(",", ":")).encode("utf-8"))
            if meter_values else None
        )
        summary.calculation_version = COMPARISON_CALCULATION_VERSION
        summary.calculated_at = datetime.now(timezone.utc)
        summary.observation = _daily_observation(source, {
            "status": status,
            "total_penalty": total_penalty,
            "highest_penalty_block": summary.highest_penalty_block,
            "highest_penalty_amount": summary.highest_penalty_amount,
            "missing_data_reason": None,
        }) if calculated else None
        if current is None:
            db.add(summary)
        db.flush()
        db.query(BlockPenaltyResult).filter_by(summary_id=summary.id).delete(synchronize_session=False)
        for row in valid_blocks:
            db.add(BlockPenaltyResult(
                summary_id=summary.id,
                plant_code=code,
                schedule_date=schedule_date,
                schedule_source=source,
                **row,
            ))
        stored.append(summary)
    db.commit()
    for summary in stored:
        db.refresh(summary)
    return stored


def comparison_readiness(
    db: Session,
    *,
    start_date: date,
    end_date: date,
) -> Dict[str, Any]:
    plants = configured_plants(db)
    required = [
        (plant["code"], day)
        for plant in plants
        for day in _date_range(start_date, end_date)
    ]
    loaded = set()
    if required:
        rows = (
            db.query(DailyPenaltySummary.plant_code, DailyPenaltySummary.schedule_date)
            .filter(DailyPenaltySummary.schedule_date >= start_date)
            .filter(DailyPenaltySummary.schedule_date <= end_date)
            .filter(DailyPenaltySummary.calculation_version == COMPARISON_CALCULATION_VERSION)
            .group_by(DailyPenaltySummary.plant_code, DailyPenaltySummary.schedule_date)
            .having(func.count(DailyPenaltySummary.id) >= len(SOURCES))
            .all()
        )
        loaded = {(normalize_plant_code(code), day) for code, day in rows}
    missing = [
        {"plant_code": code, "schedule_date": day.isoformat()}
        for code, day in required
        if (code, day) not in loaded
    ]
    loaded_items = [
        {"plant_code": code, "schedule_date": day.isoformat()}
        for code, day in required
        if (code, day) in loaded
    ]
    return {
        "ready": bool(required) and not missing,
        "required_count": len(required),
        "loaded_count": len(required) - len(missing),
        "loaded": loaded_items,
        "missing": missing,
    }


def _date_range(start_date: date, end_date: date) -> Iterable[date]:
    current = start_date
    while current <= end_date:
        yield current
        current += timedelta(days=1)


def _blockwise_observation(
    db: Session,
    *,
    summaries: Dict[str, DailyPenaltySummary],
    values: Dict[str, Optional[float]],
    best_source: Optional[str],
) -> str:
    available = [(source, float(value)) for source, value in values.items() if value is not None]
    if not available:
        return ""

    sentences = []
    if all(abs(value) <= EPSILON for _, value in available):
        sentences.append("All available schedules recorded zero penalty.")
    elif best_source:
        best_value = float(values[best_source] or 0)
        best_label = SOURCE_LABELS.get(best_source, best_source.title())
        sentence = f"{best_label} had the lowest total penalty at Rs {best_value:,.2f}"
        system_value = values.get("SYSTEM")
        if best_source != "SYSTEM" and system_value is not None:
            saving = float(system_value) - best_value
            if saving > EPSILON:
                sentence += f", saving Rs {saving:,.2f} compared with System"
            elif saving < -EPSILON:
                sentence += f", Rs {abs(saving):,.2f} higher than System"
        sentences.append(f"{sentence}.")

    summary_ids = [summary.id for summary in summaries.values() if summary.id is not None]
    block_rows = (
        db.query(BlockPenaltyResult)
        .filter(BlockPenaltyResult.summary_id.in_(summary_ids))
        .filter(BlockPenaltyResult.penalty_amount.isnot(None))
        .all()
        if summary_ids else []
    )
    highest = max(
        block_rows,
        key=lambda row: abs(float(row.penalty_amount or 0)),
        default=None,
    )
    if highest is not None:
        source_label = SOURCE_LABELS.get(highest.schedule_source, highest.schedule_source.title())
        amount = abs(float(highest.penalty_amount or 0))
        block_text = f"Block {highest.block_number} ({_block_interval(highest.block_number)})"
        scheduled = highest.scheduled_mw
        actual = highest.actual_meter_mw
        if scheduled is not None and actual is not None:
            gap = float(actual) - float(scheduled)
            if abs(gap) <= EPSILON:
                reason = (
                    f"actual generation matched the {float(scheduled):,.2f} MW schedule"
                )
            else:
                direction = "below" if gap < 0 else "above"
                reason = (
                    f"actual generation was {abs(gap):,.2f} MW {direction} schedule "
                    f"(actual {float(actual):,.2f} MW, scheduled {float(scheduled):,.2f} MW)"
                )
            sentences.append(
                f"The highest block penalty was Rs {amount:,.2f} for {source_label} at "
                f"{block_text} because {reason}."
            )
        else:
            sentences.append(
                f"The highest block penalty was Rs {amount:,.2f} for {source_label} at {block_text}."
            )

    if best_source and best_source != "SYSTEM":
        best_blocks = [
            row for row in block_rows
            if row.schedule_source == best_source and abs(float(row.penalty_amount or 0)) > EPSILON
        ]
        system_blocks = {
            row.block_number: row for row in block_rows
            if row.schedule_source == "SYSTEM"
        }
        improved = sum(
            1 for row in best_blocks
            if row.block_number in system_blocks
            and abs(float(row.penalty_amount or 0))
            < abs(float(system_blocks[row.block_number].penalty_amount or 0))
        )
        compared = sum(1 for row in best_blocks if row.block_number in system_blocks)
        if compared:
            sentences.append(
                f"{SOURCE_LABELS.get(best_source, best_source.title())} had a lower block penalty "
                f"than System in {improved} of {compared} comparable penalty blocks."
            )
    return " ".join(sentences)


def _osepl_month_key(day: date) -> str:
    return day.strftime("%b-%y")


def _osepl_source_rows_for_day(
    db: Session,
    *,
    plant_code: str,
    plant_name: str,
    plant_capacity: float,
    day: date,
) -> List[Dict[str, Any]]:
    if normalize_plant_code(plant_code) != "OSEPL":
        return []

    block_rows = (
        db.query(BlockPenaltyResult)
        .filter(BlockPenaltyResult.plant_code == normalize_plant_code(plant_code))
        .filter(BlockPenaltyResult.schedule_date == day)
        .all()
    )
    grouped: Dict[str, List[BlockPenaltyResult]] = {}
    for row in block_rows:
        grouped.setdefault(str(row.schedule_source or "").upper(), []).append(row)

    def aggregate(source: Optional[str], label: str) -> Dict[str, Any]:
        rows = grouped.get(str(source or "").upper(), []) if source else []
        if not rows:
            return {
                "Type": label,
                "Project Details": "--",
                "Net Settlement (Rs)": "--",
                "Installed Capacity": f"{float(plant_capacity or 0):.0f}",
                "SCADA Availability": "--",
                "Generation (kWh)": "--",
                "Scheduled Units": "--",
                "DSM Penalty (Rs)": "--",
                "Payable (Rs)": "--",
                "Receivable (Rs)": "--",
            }

        scheduled_kwh = sum(
            float(row.scheduled_mw or 0.0) * BLOCK_HOURS * KWH_PER_MWH
            for row in rows
            if row.scheduled_mw is not None
        )
        generation_kwh = sum(
            float(row.actual_meter_mw or 0.0) * BLOCK_HOURS * KWH_PER_MWH
            for row in rows
            if row.actual_meter_mw is not None
        )
        payable_rs = sum(float(row.payable_amount or 0.0) for row in rows if row.payable_amount is not None)
        receivable_rs = sum(float(row.receivable_amount or 0.0) for row in rows if row.receivable_amount is not None)
        penalty_rs = sum(float(row.penalty_amount or 0.0) for row in rows if row.penalty_amount is not None)
        net_settlement = receivable_rs - payable_rs - penalty_rs
        return {
            "Type": label,
            "Project Details": f"{day.isoformat()} / {_osepl_month_key(day)} / {plant_name}",
            "Net Settlement (Rs)": f"{round(net_settlement):.0f}",
            "Installed Capacity": f"{float(plant_capacity or 0):.0f}",
            "SCADA Availability": "100%",
            "Generation (kWh)": f"{round(generation_kwh):.0f}",
            "Scheduled Units": f"{round(scheduled_kwh):.0f}",
            "DSM Penalty (Rs)": f"{round(penalty_rs):.0f}",
            "Payable (Rs)": f"{round(payable_rs):.0f}",
            "Receivable (Rs)": f"{round(receivable_rs):.0f}",
            "TestEnv": "--",
        }

    return [
        aggregate("SYSTEM", "System (Auto)"),
        aggregate("MANUAL", "Manual"),
        aggregate("VEDANJAY", "Vedanjay (UI)"),
        aggregate("TESTENV", "Testing Env"),
    ]


def _word_add_osepl_report_sections(document: Document, plant: Dict[str, Any], sections: Sequence[Dict[str, Any]]) -> None:
    if not sections:
        return

    headers = [
        "Type",
        "Project Details",
        "Net Settlement (Rs)",
        "Installed Capacity",
        "SCADA Availability",
        "Generation (kWh)",
        "Scheduled Units",
        "DSM Penalty (Rs)",
        "Payable (Rs)",
        "Receivable (Rs)",
        "TestEnv",
    ]

    for index, section in enumerate(sections):
        if index > 0:
            document.add_paragraph()
        day_text = datetime.strptime(str(section.get("date") or ""), "%Y-%m-%d").strftime("%d-%m-%y")
        title = document.add_paragraph()
        title_run = title.add_run(f"DATE {day_text}")
        title_run.bold = True
        title_run.font.name = "Times New Roman"
        title_run.font.size = Pt(12)

        note = document.add_paragraph()
        note_run = note.add_run("Net Settlement = Receivable - Payable - DSM Penalty")
        note_run.bold = True
        note_run.font.name = "Times New Roman"
        note_run.font.size = Pt(10)
        note_run.font.color.rgb = RGBColor(200, 0, 0)

        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.autofit = False
        widths = [Inches(1.0), Inches(2.2), Inches(1.2), Inches(1.2), Inches(1.1), Inches(1.2), Inches(1.2), Inches(1.2), Inches(1.1), Inches(1.1), Inches(0.9)]
        for col_index, width in enumerate(widths):
            table.columns[col_index].width = width
        header_row = table.rows[0]
        for col_index, header in enumerate(headers):
            _word_set_cell_text(header_row.cells[col_index], header, bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
            _word_set_cell_shading(header_row.cells[col_index], "2F855A")

        for row_data in section.get("rows", []):
            row = table.add_row()
            net_text = str(row_data.get("Net Settlement (Rs)") or "")
            try:
                net_value = float(net_text.replace(",", "").strip() or "nan")
            except Exception:
                net_value = float("nan")
            for col_index, header in enumerate(headers):
                value = row_data.get(header, "--")
                _word_set_cell_text(row.cells[col_index], str(value), bold=col_index == 0, size=8)
            if net_value == net_value:
                _word_set_cell_shading(row.cells[2], "C6F6D5" if net_value >= 0 else "FED7D7")


def _pdf_add_osepl_report_sections(story: List[Any], plant: Dict[str, Any], sections: Sequence[Dict[str, Any]]) -> None:
    if not sections:
        return

    styles = getSampleStyleSheet()
    heading_style = ParagraphStyle(
        "OseplHeading",
        parent=styles["Heading2"],
        fontName="Times-Bold",
        fontSize=11,
        textColor=colors.HexColor("#111827"),
        spaceBefore=4,
        spaceAfter=3,
    )
    note_style = ParagraphStyle(
        "OseplNote",
        parent=styles["Normal"],
        fontName="Times-Bold",
        fontSize=8,
        textColor=colors.HexColor("#B91C1C"),
        spaceAfter=4,
    )
    cell_style = ParagraphStyle(
        "OseplCell",
        parent=styles["Normal"],
        fontName="Times-Roman",
        fontSize=6.8,
        leading=8,
    )
    bold_style = ParagraphStyle(
        "OseplCellBold",
        parent=cell_style,
        fontName="Times-Bold",
    )
    headers = [
        "Type",
        "Project Details",
        "Net Settlement (Rs)",
        "Installed Capacity",
        "SCADA Availability",
        "Generation (kWh)",
        "Scheduled Units",
        "DSM Penalty (Rs)",
        "Payable (Rs)",
        "Receivable (Rs)",
        "TestEnv",
    ]

    for index, section in enumerate(sections):
        if index > 0:
            story.append(Spacer(1, 4))
        day_text = datetime.strptime(str(section.get("date") or ""), "%Y-%m-%d").strftime("%d-%m-%y")
        story.append(Paragraph(f"DATE {day_text}", heading_style))
        story.append(Paragraph("Net Settlement = Receivable - Payable - DSM Penalty", note_style))

        rows = [[_pdf_text(header, bold_style) for header in headers]]
        for row_data in section.get("rows", []):
            rows.append([_pdf_text(str(row_data.get(header, "--")), cell_style) for header in headers])

        table = Table(
            rows,
            repeatRows=1,
            colWidths=[16 * mm, 40 * mm, 20 * mm, 18 * mm, 18 * mm, 18 * mm, 18 * mm, 18 * mm, 18 * mm, 18 * mm, 14 * mm],
        )
        commands = [
            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#94a3b8")),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2F855A")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 3),
            ("RIGHTPADDING", (0, 0), (-1, -1), 3),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]
        for row_index, row_data in enumerate(section.get("rows", []), start=1):
            net_text = str(row_data.get("Net Settlement (Rs)") or "")
            try:
                net_value = float(net_text.replace(",", "").strip() or "nan")
            except Exception:
                net_value = float("nan")
            if net_value == net_value:
                commands.append((
                    "BACKGROUND",
                    (2, row_index),
                    (2, row_index),
                    colors.HexColor("#C6F6D5" if net_value >= 0 else "#FED7D7"),
                ))
        table.setStyle(TableStyle(commands))
        story.append(table)


def build_report_data(
    db: Session,
    *,
    start_date: date,
    end_date: date,
    include_block_details: bool,
    plant_codes: Optional[Sequence[str]] = None,
    s3: Optional[ReadOnlyS3Source] = None,
) -> Dict[str, Any]:
    allowed_codes = None
    if plant_codes:
        allowed_codes = {normalize_plant_code(code) for code in plant_codes if normalize_plant_code(code)}
    report_plants = []
    for plant in configured_plants(db):
        if allowed_codes is not None and plant["code"] not in allowed_codes:
            continue
        daily_rows = []
        osepl_report_rows: List[Dict[str, Any]] = []
        source_totals = {source: 0.0 for source in SOURCES}
        source_has_values = {source: False for source in SOURCES}
        missing_days = 0
        for day in _date_range(start_date, end_date):
            summaries: Dict[str, DailyPenaltySummary] = {}
            for source in SOURCES:
                summary = (
                    db.query(DailyPenaltySummary)
                    .filter_by(
                        plant_code=plant["code"],
                        schedule_date=day,
                        schedule_source=source,
                        calculation_version=COMPARISON_CALCULATION_VERSION,
                    )
                    .first()
                )
                if summary is None:
                    raise ValueError(
                        f"Comparison data has not been loaded for {plant['name']} on {day.isoformat()}."
                    )
                summaries[source] = summary
            testenv_summary = (
                db.query(DailyPenaltySummary)
                .filter_by(
                    plant_code=plant["code"],
                    schedule_date=day,
                    schedule_source="TESTENV",
                    calculation_version=COMPARISON_CALCULATION_VERSION,
                )
                .first()
            )
            values = {
                source: summaries[source].total_penalty
                if summaries[source].status in {"Calculated", "Partially Calculated", "Zero Penalty"}
                else None
                for source in SOURCES
            }
            for source, value in values.items():
                if value is not None:
                    source_totals[source] += float(value)
                    source_has_values[source] = True
            valid = [(source, value) for source, value in values.items() if value is not None]
            best = min(valid, key=lambda item: item[1])[0] if valid else None
            if any(value is None for value in values.values()):
                missing_days += 1
            observation = _blockwise_observation(
                db,
                summaries=summaries,
                values=values,
                best_source=best,
            )
            daily_rows.append({
                "date": day.isoformat(),
                "sources": {
                    **{source: summary_dict(summaries[source]) for source in SOURCES},
                    "TESTENV": summary_dict(testenv_summary) if testenv_summary is not None else {
                        "schedule_source": "TESTENV",
                        "total_penalty": None,
                        "status": "Not Applicable",
                        "missing_data_reason": None,
                        "calculated_blocks": 0,
                        "highest_penalty_block": None,
                        "highest_penalty_amount": None,
                    },
                },
                "best": SOURCE_LABELS.get(best, "--") if best else "--",
                "observation": observation,
            })
            if normalize_plant_code(plant["code"]) == "OSEPL":
                osepl_report_rows.append({
                    "date": day.isoformat(),
                    "month_key": _osepl_month_key(day),
                    "rows": _osepl_source_rows_for_day(
                        db,
                        plant_code=plant["code"],
                        plant_name=str(plant["name"] or "OSEPL"),
                        plant_capacity=float(plant["capacity"] or 0.0),
                        day=day,
                    ),
                })
        system_total = source_totals["SYSTEM"] if source_has_values["SYSTEM"] else None
        vedanjay_total = source_totals["VEDANJAY"] if source_has_values["VEDANJAY"] else None
        penalty_days = [
            (row["date"], source, detail["total_penalty"])
            for row in daily_rows
            for source, detail in row["sources"].items()
            if detail["total_penalty"] is not None
        ]
        highest_day = max(penalty_days, key=lambda item: abs(float(item[2])), default=None)
        block_candidates = [
            (
                row["date"],
                source,
                detail["highest_penalty_block"],
                detail["highest_penalty_amount"],
            )
            for row in daily_rows
            for source, detail in row["sources"].items()
            if detail["highest_penalty_amount"] is not None
        ]
        highest_block = max(block_candidates, key=lambda item: float(item[3]), default=None)
        report_plants.append({
            **plant,
            "penalty_rule": penalty_rule_label(plant["state"], plant["type"]),
            "daily": daily_rows,
            "osepl_report_rows": osepl_report_rows,
            "totals": {
                source: source_totals[source] if source_has_values[source] else None
                for source in SOURCES
            },
            "vedanjay_savings_vs_system": (
                system_total - vedanjay_total
                if system_total is not None and vedanjay_total is not None
                else None
            ),
            "missing_data_days": missing_days,
            "highest_penalty_day": {
                "date": highest_day[0],
                "source": highest_day[1],
                "amount": highest_day[2],
            } if highest_day else None,
            "highest_penalty_block": {
                "date": highest_block[0],
                "source": highest_block[1],
                "block": highest_block[2],
                "amount": highest_block[3],
            } if highest_block else None,
        })
    data = {
        "title": "All Plant Penalty Report",
        "calculation_version": COMPARISON_CALCULATION_VERSION,
        "start_date": start_date.isoformat(),
        "end_date": end_date.isoformat(),
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "include_block_details": include_block_details,
        "plants": report_plants,
    }
    if include_block_details:
        data["block_details"] = _report_block_details(db, start_date, end_date)
    return data


def _store_failed_summary(
    db: Session,
    plant: Dict[str, Any],
    schedule_date: date,
    source: str,
    message: str,
) -> DailyPenaltySummary:
    summary = (
        db.query(DailyPenaltySummary)
        .filter_by(plant_code=plant["code"], schedule_date=schedule_date, schedule_source=source)
        .first()
        or DailyPenaltySummary(
            plant_code=plant["code"],
            schedule_date=schedule_date,
            schedule_source=source,
            plant_name=plant["name"],
            calculation_version=CALCULATION_VERSION,
        )
    )
    summary.plant_name = plant["name"]
    summary.state = plant["state"]
    summary.capacity_mw = plant["capacity"]
    summary.total_penalty = None
    summary.status = "Failed"
    summary.missing_data_reason = message[:2000]
    summary.calculated_blocks = 0
    summary.calculation_version = CALCULATION_VERSION
    db.add(summary)
    db.commit()
    db.refresh(summary)
    return summary


def _report_block_details(db: Session, start_date: date, end_date: date) -> List[Dict[str, Any]]:
    rows = (
        db.query(BlockPenaltyResult)
        .filter(BlockPenaltyResult.schedule_date >= start_date)
        .filter(BlockPenaltyResult.schedule_date <= end_date)
        .order_by(
            BlockPenaltyResult.plant_code,
            BlockPenaltyResult.schedule_date,
            BlockPenaltyResult.schedule_source,
            BlockPenaltyResult.block_number,
        )
        .all()
    )
    return [{
        "plant_code": row.plant_code,
        "date": row.schedule_date.isoformat(),
        "source": row.schedule_source,
        "block": row.block_number,
        "scheduled_mw": row.scheduled_mw,
        "actual_meter_mw": row.actual_meter_mw,
        "deviation_mw": row.deviation_mw,
        "deviation_percent": row.deviation_percent,
        "penalty_amount": row.penalty_amount,
        "status": row.status,
        "reason": row.missing_data_reason,
    } for row in rows]


def _money(value: Optional[float]) -> str:
    return "" if value is None else f"Rs {float(value):,.2f}"


REPORT_SOURCE_ORDER = ("VEDANJAY", "SYSTEM", "MANUAL", "ENERCAST", "TESTENV")
REPORT_SOURCE_HEADERS = {
    "VEDANJAY": "Vedanjay",
    "SYSTEM": "System\nProduction",
    "MANUAL": "Manual\nedited",
    "ENERCAST": "Enercast",
    "TESTENV": "TestEnv",
}
PENALTY_COLOR_LESS = "00FF00"
PENALTY_COLOR_MARGINAL = "FFFF00"
PENALTY_COLOR_HIGH = "00D9D9"


def _report_date(value: str) -> str:
    try:
        return datetime.strptime(str(value), "%Y-%m-%d").strftime("%d-%m-%y")
    except ValueError:
        return str(value)


def _block_interval(block: Optional[int]) -> str:
    if not block or block < 1 or block > 96:
        return "--"
    start = (block - 1) * 15
    end = start + 15
    return f"{start // 60:02d}:{start % 60:02d}-{end // 60:02d}:{end % 60:02d}"


def _penalty_display(summary: Dict[str, Any]) -> str:
    if not summary:
        return "--"
    status = str(summary.get("status") or "")
    value = summary.get("total_penalty")
    if status in {"Calculated", "Partially Calculated", "Zero Penalty"} and value is not None:
        return f"{float(value):,.2f}"
    return ""


def _highest_block_display(summary: Dict[str, Any]) -> str:
    if not summary:
        return "--"
    block = summary.get("highest_penalty_block")
    amount = summary.get("highest_penalty_amount")
    if block and amount is not None:
        return f"B{block} ({_block_interval(int(block))})\nRs: {float(amount):,.2f}"
    return ""


def _missing_day_message(day: Dict[str, Any]) -> Optional[str]:
    return None


def _penalty_rank_colors(day: Dict[str, Any]) -> Dict[str, str]:
    values = []
    for source in REPORT_SOURCE_ORDER:
        summary = (day.get("sources") or {}).get(source) or {}
        value = summary.get("total_penalty")
        if value is not None and summary.get("status") in {"Calculated", "Partially Calculated", "Zero Penalty"}:
            values.append((source, float(value)))
    if not values:
        return {}
    distinct = sorted({value for _, value in values})
    if len(distinct) == 1:
        return {source: PENALTY_COLOR_LESS for source, _ in values}
    minimum, maximum = distinct[0], distinct[-1]
    colors_by_source = {}
    for source, value in values:
        if value == minimum:
            colors_by_source[source] = PENALTY_COLOR_LESS
        elif value == maximum:
            colors_by_source[source] = PENALTY_COLOR_HIGH
        else:
            colors_by_source[source] = PENALTY_COLOR_MARGINAL
    return colors_by_source


def _word_set_cell_shading(cell: Any, fill: str) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = cell_properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        cell_properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _word_set_cell_text(
    cell: Any,
    text: str,
    *,
    bold: bool = False,
    color: Optional[RGBColor] = None,
    size: float = 8,
    align: Any = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _word_add_metadata_line(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(f"{label} – {value}")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(47, 84, 150)


def _pdf_text(value: Any, style: ParagraphStyle) -> Paragraph:
    return Paragraph(str(value).replace("\n", "<br/>"), style)


def _reference_penalty_rule(plant: Dict[str, Any]) -> str:
    state = str(plant.get("state") or "Applicable State")
    threshold = 15 if state == "Telangana" else 10
    return f"DSM as per {state} SERC bands (≤{threshold}% no penalty, >{threshold}% slab-based charges)"


def generate_word_report(data: Dict[str, Any]) -> bytes:
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)

    for plant_index, plant in enumerate(data["plants"]):
        header = document.add_table(rows=1, cols=6)
        header.autofit = False
        header.columns[0].width = Inches(1.8)
        header.columns[1].width = Inches(1.0)
        header.columns[2].width = Inches(1.5)
        header.columns[3].width = Inches(1.3)
        header.columns[4].width = Inches(1.9)
        header.columns[5].width = Inches(1.3)
        values = (
            "Site Name",
            f"– {plant['name'].title()}",
            "Color Code =>",
            "Less Penalty",
            "Marginally high penalty",
            "High Penalty",
        )
        for index, value in enumerate(values):
            _word_set_cell_text(header.cell(0, index), value, bold=True, size=9)
        _word_set_cell_shading(header.cell(0, 3), PENALTY_COLOR_LESS)
        _word_set_cell_shading(header.cell(0, 4), PENALTY_COLOR_MARGINAL)
        _word_set_cell_shading(header.cell(0, 5), PENALTY_COLOR_HIGH)

        _word_add_metadata_line(document, "State", str(plant["state"]))
        _word_add_metadata_line(document, "Plant Capacity", f"{float(plant['capacity']):g} MW")
        _word_add_metadata_line(document, "Schedule Type", str(plant["type"]))
        _word_add_metadata_line(document, "Penalty Rule", _reference_penalty_rule(plant))
        if normalize_plant_code(plant.get("code")) == "OSEPL":
            _word_add_osepl_report_sections(document, plant, plant.get("osepl_report_rows") or [])
            document.add_page_break()

        table = document.add_table(rows=3, cols=12)
        table.style = "Table Grid"
        table.autofit = False
        table.cell(0, 0).merge(table.cell(2, 0))
        table.cell(0, 1).merge(table.cell(0, 10))
        table.cell(1, 1).merge(table.cell(1, 5))
        table.cell(1, 6).merge(table.cell(1, 10))
        table.cell(1, 11).merge(table.cell(2, 11))
        _word_set_cell_text(table.cell(0, 0), "Date", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        _word_set_cell_text(table.cell(0, 1), f"{plant['name'].title()} Site", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        _word_set_cell_text(table.cell(1, 1), "Penalties", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        _word_set_cell_text(
            table.cell(1, 6),
            "High Penalty Blocks\n(Block No & Time & in Rs penalty)",
            bold=True,
            size=8,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _word_set_cell_text(table.cell(1, 11), "Observation", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        for index, source in enumerate(REPORT_SOURCE_ORDER):
            _word_set_cell_text(table.cell(2, 1 + index), REPORT_SOURCE_HEADERS[source], bold=True, size=8)
            _word_set_cell_text(table.cell(2, 6 + index), REPORT_SOURCE_HEADERS[source], bold=True, size=8)

        for day in reversed(plant["daily"]):
            row = table.add_row()
            _word_set_cell_text(row.cells[0], _report_date(day["date"]), bold=True, size=8)
            missing_message = _missing_day_message(day)
            if missing_message:
                merged = row.cells[1].merge(row.cells[10])
                _word_set_cell_text(
                    merged,
                    missing_message,
                    bold=True,
                    color=RGBColor(255, 0, 0),
                    size=9,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                )
                _word_set_cell_text(row.cells[11], day.get("observation") or "", size=8)
                continue
            rank_colors = _penalty_rank_colors(day)
            for index, source in enumerate(REPORT_SOURCE_ORDER):
                summary = day["sources"][source]
                penalty_cell = row.cells[1 + index]
                _word_set_cell_text(penalty_cell, _penalty_display(summary), bold=True, size=8)
                if source in rank_colors:
                    _word_set_cell_shading(penalty_cell, rank_colors[source])
                _word_set_cell_text(row.cells[6 + index], _highest_block_display(summary), bold=True, size=8)
            _word_set_cell_text(row.cells[11], day.get("observation") or "", size=8)
        if plant_index < len(data["plants"]) - 1:
            document.add_page_break()

    if data.get("include_block_details"):
        document.add_page_break()
        document.add_heading("96-Block Details", level=1)
        detail_table = document.add_table(rows=1, cols=8)
        detail_table.style = "Table Grid"
        headings = ("Plant", "Date", "Source", "Block", "Schedule", "Meter", "Deviation %", "Penalty")
        for index, value in enumerate(headings):
            detail_table.rows[0].cells[index].text = value
        for row in data.get("block_details", []):
            cells = detail_table.add_row().cells
            values = (
                row["plant_code"], row["date"], row["source"], row["block"],
                row["scheduled_mw"], row["actual_meter_mw"], row["deviation_percent"], row["penalty_amount"],
            )
            for index, value in enumerate(values):
                cells[index].text = "" if value is None else str(round(value, 3) if isinstance(value, float) else value)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def generate_pdf_report(data: Dict[str, Any]) -> bytes:
    output = io.BytesIO()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("ReportTitle", parent=styles["Title"], alignment=TA_CENTER, textColor=colors.HexColor("#1f2937"))
    document = SimpleDocTemplate(
        output,
        pagesize=landscape(A4),
        rightMargin=12 * mm,
        leftMargin=12 * mm,
        topMargin=12 * mm,
        bottomMargin=12 * mm,
    )
    metadata_style = ParagraphStyle(
        "Metadata",
        parent=styles["Normal"],
        fontName="Times-Bold",
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#2F5496"),
    )
    cell_style = ParagraphStyle(
        "PenaltyCell",
        parent=styles["Normal"],
        fontName="Times-Roman",
        fontSize=7,
        leading=9,
    )
    bold_cell_style = ParagraphStyle(
        "PenaltyCellBold",
        parent=cell_style,
        fontName="Times-Bold",
    )
    center_style = ParagraphStyle(
        "PenaltyCenter",
        parent=bold_cell_style,
        alignment=TA_CENTER,
    )
    story: List[Any] = []
    for plant_index, plant in enumerate(data["plants"]):
        legend = Table(
            [[
                _pdf_text(f"Site Name – {plant['name'].title()}", bold_cell_style),
                _pdf_text("Color Code =>", bold_cell_style),
                _pdf_text("Less Penalty", bold_cell_style),
                _pdf_text("Marginally high penalty", bold_cell_style),
                _pdf_text("High Penalty", bold_cell_style),
            ]],
            colWidths=[55 * mm, 28 * mm, 32 * mm, 50 * mm, 32 * mm],
        )
        legend.setStyle(TableStyle([
            ("BACKGROUND", (2, 0), (2, 0), colors.HexColor(f"#{PENALTY_COLOR_LESS}")),
            ("BACKGROUND", (3, 0), (3, 0), colors.HexColor(f"#{PENALTY_COLOR_MARGINAL}")),
            ("BACKGROUND", (4, 0), (4, 0), colors.HexColor(f"#{PENALTY_COLOR_HIGH}")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 2),
            ("RIGHTPADDING", (0, 0), (-1, -1), 2),
        ]))
        story.extend([
            legend,
            Spacer(1, 6),
            Paragraph(f"State – {plant['state']}", metadata_style),
            Paragraph(f"Plant Capacity – {float(plant['capacity']):g} MW", metadata_style),
            Paragraph(f"Schedule Type – {plant['type']}", metadata_style),
            Paragraph(f"Penalty Rule – {_reference_penalty_rule(plant)}", metadata_style),
            Spacer(1, 7),
        ])
        if normalize_plant_code(plant.get("code")) == "OSEPL":
            _pdf_add_osepl_report_sections(story, plant, plant.get("osepl_report_rows") or [])
            story.append(PageBreak())
        rows = [
            [
                _pdf_text("Date", center_style),
                _pdf_text(f"{plant['name'].title()} Site", center_style),
                "", "", "", "", "", "", "", "", "", "",
            ],
            [
                "",
                _pdf_text("Penalties", center_style), "", "", "", "",
                _pdf_text("High Penalty Blocks<br/>(Block No &amp; Time &amp; in Rs penalty)", center_style), "", "", "", "",
                _pdf_text("Observation", center_style),
            ],
            [
                "",
                *[_pdf_text(REPORT_SOURCE_HEADERS[source], bold_cell_style) for source in REPORT_SOURCE_ORDER],
                *[_pdf_text(REPORT_SOURCE_HEADERS[source], bold_cell_style) for source in REPORT_SOURCE_ORDER],
                "",
            ],
        ]
        row_styles: List[Tuple[str, Tuple[int, int], Tuple[int, int], Any]] = []
        for day in reversed(plant["daily"]):
            row_index = len(rows)
            missing_message = _missing_day_message(day)
            if missing_message:
                rows.append([
                    _pdf_text(_report_date(day["date"]), bold_cell_style),
                    _pdf_text(f'<font color="red"><b>{missing_message}</b></font>', center_style),
                    "", "", "", "", "", "", "", "", "",
                    _pdf_text(day.get("observation") or "", cell_style),
                ])
                row_styles.append(("SPAN", (1, row_index), (10, row_index), None))
                continue
            rank_colors = _penalty_rank_colors(day)
            row = [_pdf_text(_report_date(day["date"]), bold_cell_style)]
            row.extend(
                _pdf_text(_penalty_display(day["sources"][source]), bold_cell_style)
                for source in REPORT_SOURCE_ORDER
            )
            row.extend(
                _pdf_text(_highest_block_display(day["sources"][source]), bold_cell_style)
                for source in REPORT_SOURCE_ORDER
            )
            row.append(_pdf_text(day.get("observation") or "", cell_style))
            rows.append(row)
            for source_index, source in enumerate(REPORT_SOURCE_ORDER, start=1):
                if source in rank_colors:
                    row_styles.append((
                        "BACKGROUND",
                        (source_index, row_index),
                        (source_index, row_index),
                        colors.HexColor(f"#{rank_colors[source]}"),
                    ))
        table = Table(
            rows,
            repeatRows=3,
            colWidths=[18 * mm, 18 * mm, 18 * mm, 18 * mm, 18 * mm, 18 * mm, 18 * mm, 18 * mm, 18 * mm, 18 * mm, 18 * mm, 62 * mm],
        )
        table_commands = [
            ("SPAN", (0, 0), (0, 2)),
            ("SPAN", (1, 0), (10, 0)),
            ("SPAN", (1, 1), (5, 1)),
            ("SPAN", (6, 1), (10, 1)),
            ("SPAN", (11, 1), (11, 2)),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#94a3b8")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ALIGN", (0, 0), (10, 2), "CENTER"),
            ("LEFTPADDING", (0, 0), (-1, -1), 3),
            ("RIGHTPADDING", (0, 0), (-1, -1), 3),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]
        for command, start, end, value in row_styles:
            table_commands.append((command, start, end) if value is None else (command, start, end, value))
        table.setStyle(TableStyle(table_commands))
        story.append(table)
        if plant_index < len(data["plants"]) - 1:
            story.append(PageBreak())
    if data.get("include_block_details"):
        story.extend([PageBreak(), Paragraph("96-Block Details", styles["Heading1"])])
        detail_rows = [["Plant", "Date", "Source", "Block", "Schedule", "Meter", "Deviation %", "Penalty"]]
        for row in data.get("block_details", []):
            detail_rows.append([
                row["plant_code"], row["date"], row["source"], row["block"],
                "" if row["scheduled_mw"] is None else f"{row['scheduled_mw']:.3f}",
                "" if row["actual_meter_mw"] is None else f"{row['actual_meter_mw']:.3f}",
                "" if row["deviation_percent"] is None else f"{row['deviation_percent']:.2f}",
                "" if row["penalty_amount"] is None else f"{row['penalty_amount']:.2f}",
            ])
        detail_table = Table(detail_rows, repeatRows=1)
        detail_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e2e8f0")),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#94a3b8")),
            ("FONTSIZE", (0, 0), (-1, -1), 6),
        ]))
        story.append(detail_table)
    document.build(story)
    return output.getvalue()


def generate_and_store_report(
    db: Session,
    *,
    report_type: str,
    start_date: date,
    end_date: date,
    formats: Sequence[str],
    include_block_details: bool,
    requested_by: str,
    s3: Optional[ReadOnlyS3Source] = None,
    plant_codes: Optional[Sequence[str]] = None,
) -> GeneratedPenaltyReport:
    normalized_formats = sorted({str(value).upper() for value in formats if str(value).upper() in {"WORD", "PDF"}})
    if not normalized_formats:
        raise ValueError("At least one report format is required.")
    history = GeneratedPenaltyReport(
        report_type=report_type.title(),
        start_date=start_date,
        end_date=end_date,
        requested_formats=",".join(normalized_formats),
        include_block_details=include_block_details,
        requested_by=requested_by,
        status="Generating",
    )
    db.add(history)
    db.commit()
    db.refresh(history)
    try:
        data = build_report_data(
            db,
            start_date=start_date,
            end_date=end_date,
            include_block_details=include_block_details,
            plant_codes=plant_codes,
            s3=s3,
        )
        base = f"all-plant-penalty-{report_type.lower()}-{start_date.isoformat()}-{end_date.isoformat()}"
        history.report_data_json = json.dumps(data, separators=(",", ":"), default=str)
        if "WORD" in normalized_formats:
            history.word_filename = f"{base}.docx"
            history.word_content = generate_word_report(data)
        if "PDF" in normalized_formats:
            history.pdf_filename = f"{base}.pdf"
            history.pdf_content = generate_pdf_report(data)
        history.status = "Ready"
        history.completed_at = datetime.now(timezone.utc)
    except Exception as exc:
        history.status = "Failed"
        history.error_message = str(exc)[:4000]
        db.commit()
        raise
    db.commit()
    db.refresh(history)
    return history
